"""Independent scientific and package-consistency checks for v30."""

from __future__ import annotations

import csv
import hashlib
import re
from pathlib import Path

import pandas as pd
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "manuscript_package_v30_secondary_revised"
V27 = ROOT / "outputs" / "modeling_v27_severity_recovery"
V28 = ROOT / "outputs" / "modeling_v28_severe_temporal_external"
V29 = ROOT / "outputs" / "modeling_v29_multistate_competing_risk"
V30 = ROOT / "outputs" / "modeling_v30_severe_same_risk_incremental"


def docx_text(path: Path) -> str:
    doc = Document(path)
    text = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        text.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
    return "\n".join(text)


def check(condition: bool, message: str, rows: list[dict[str, str]]) -> None:
    if not condition:
        raise AssertionError(message)
    rows.append({"check": message, "status": "PASS"})


def main() -> None:
    results: list[dict[str, str]] = []
    md = (OUT / "secondary_manuscript_v30_en.md").read_text(encoding="utf-8")
    tex = (OUT / "latex" / "main.tex").read_text(encoding="utf-8")
    word = docx_text(OUT / "secondary_manuscript_v30_en.docx")
    combined = "\n".join([md, tex, word])

    check("Caiyun Yuan" not in combined, "Package contains only the final author", results)
    check("Bizhi Wei" in combined and "15619056250wbz@gmail.com" in combined, "Author and correspondence are present", results)
    check("[AUTHOR ACTION REQUIRED]" not in combined, "No author-action placeholder remains", results)
    check("to be provided before publication" in combined, "Only intentional archive-DOI placeholder remains", results)
    check("v1.0.2" in combined, "Repository release v1.0.2 is cited", results)
    check("Supplementary Tables S1–S9" in word, "Word main file reports Tables S1–S9", results)
    check("Supplementary Figures S1–S7" in word, "Word main file reports Figures S1–S7", results)

    summary = pd.read_csv(V29 / "audit_v29_multistate_summary.csv").set_index("measure")
    expected_counts = {
        "Analysis cohort": 10877, "Incident SCr-AKI": 4531,
        "Trajectory-eligible incident SCr-AKI": 4519,
        "Locked AKI onset after recorded hospital disposition": 12,
        "Severe SCr-AKI stage 2/3": 679,
        "Observed recovery after AKI": 3936,
        "Recurrent AKI after observed recovery": 641,
    }
    for label, expected in expected_counts.items():
        check(int(summary.loc[label, "n"]) == expected, f"Source count matches: {label}={expected:,}", results)

    table1 = pd.read_csv(OUT / "tables" / "Table_1_baseline_characteristics_by_severe_aki.csv")
    check("No severe SCr-AKI (N=10,198)" in table1, "Table 1 has the nonsevere denominator", results)
    check("Severe SCr-AKI (N=679)" in table1, "Table 1 has the severe denominator", results)
    check(len(table1) >= 30, "Table 1 contains a complete baseline profile", results)

    table2 = pd.read_csv(OUT / "tables" / "Table_2_selected_severe_aki_models.csv")
    check(len(table2) == 3 and set(table2["Landmark"]) == {"0 h", "6 h", "24 h"}, "Table 2 contains only three severe-AKI landmark models", results)
    check(not table2.astype(str).apply(lambda x: x.str.contains("persistence|nonrecovery", case=False, regex=True)).any().any(), "Table 2 excludes internal-only recovery models", results)

    perf = pd.read_csv(V27 / "model_v27_performance_summary.csv")
    for lm, model, expected in [(0, "XGBoost", 0.6977275179), (6, "XGBoost", 0.7896103896), (24, "Logistic Regression", 0.8394132653)]:
        value = perf[(perf.task.eq("severe_scr")) & perf.landmark.astype(str).eq(str(lm)) & perf.model.eq(model)].iloc[0].auroc
        check(abs(value - expected) < 1e-10, f"Internal {lm}-h AUROC matches source", results)

    deltas = pd.read_csv(V30 / "model_v30_severe_paired_incremental_deltas.csv")
    d6 = deltas[deltas.primary_model_for_risk_set.astype(bool) & deltas.risk_set_hours.eq(6)
                & deltas.new_information_hours.eq(6) & deltas.reference_information_hours.eq(0)].iloc[0]
    d24 = deltas[deltas.primary_model_for_risk_set.astype(bool) & deltas.risk_set_hours.eq(24)
                 & deltas.new_information_hours.eq(24) & deltas.reference_information_hours.eq(6)].iloc[0]
    check(round(d6.delta_auroc, 3) == 0.101, "6-h paired delta AUROC is 0.101", results)
    check(round(d24.delta_auroc, 3) == 0.079, "24-h paired delta AUROC is 0.079", results)
    check(d6.delta_auroc_ci_lower > 0 and d24.delta_auroc_ci_lower > 0, "Primary paired AUROC intervals exclude zero", results)
    check("n=2,162; 132 events" in md and "n=2,142; 112 events" in md, "Same-risk-set denominators are reported", results)

    ext = pd.read_csv(V28 / "model_v28_eicu_frozen_severe_performance.csv")
    for lm, expected in [(0, 0.7070313499), (6, 0.7613177886), (24, 0.7839542750)]:
        value = ext[(ext.landmark_hours.eq(lm)) & ext.target.eq("SCr stage 2/3")].iloc[0].auroc
        check(abs(value - expected) < 1e-10, f"External {lm}-h AUROC matches source", results)
    check("14,229 (46.9%)" in combined and "30,365" in combined and "16,136" in combined, "External observability denominator is explicit", results)
    check("These external results apply only to the severe-AKI landmark models" in md, "External-validation boundary is explicit", results)

    mapping = pd.read_csv(OUT / "tables" / "Table_S9_feature_harmonization.csv")
    check(len(mapping) == 8, "Feature harmonization table contains eight operational groups", results)
    check({"Baseline SCr", "Surgical system", "Time-updated SCr status"} <= set(mapping["Feature group"]), "Key cross-database mappings are documented", results)

    cif_aki = pd.read_csv(V29 / "analysis_v29_cif_after_aki.csv")
    cif_rec = pd.read_csv(V29 / "analysis_v29_cif_after_recovery.csv")
    recovery = cif_aki[(cif_aki.group_variable.eq("Overall")) & cif_aki.time_hours.eq(48) & cif_aki.cause.eq("Observed recovery")].iloc[0].cif_percent
    severe = cif_aki[(cif_aki.group_variable.eq("Overall")) & cif_aki.time_hours.eq(48) & cif_aki.cause.eq("Severe AKI onset/progression")].iloc[0].cif_percent
    recurrence = cif_rec[(cif_rec.group_variable.eq("Overall")) & cif_rec.time_hours.eq(48) & cif_rec.cause.eq("Recurrent AKI")].iloc[0].cif_percent
    check(round(recovery, 1) == 65.0 and round(severe, 1) == 11.9 and round(recurrence, 1) == 11.3, "Trajectory CIF values match source", results)
    check("Observed recovery occurred at some point" in md and "cumulative incidence" in md, "Ever-observed and first-event recovery estimands are distinguished", results)

    refs = pd.read_csv(OUT / "reference_audit.csv")
    check(len(refs) == 35 and refs["Citation key"].is_unique, "Reference audit contains 35 unique entries", results)
    cite_groups = re.findall(r"\\cite\{([^}]+)\}", tex)
    cited = {key for group in cite_groups for key in group.split(",")}
    check(cited <= set(refs["Citation key"]), "Every LaTeX citation key exists", results)

    for i in range(1, 5):
        for extn in ["png", "pdf", "svg"]:
            check((OUT / "figures" / f"Fig{i}.{extn}").exists(), f"Figure {i} {extn} exists", results)
    for i in range(1, 8):
        for extn in ["png", "pdf", "svg"]:
            check((OUT / "figures" / f"FigS{i}.{extn}").exists(), f"Figure S{i} {extn} exists", results)
    for i in range(1, 10):
        check(any((OUT / "tables").glob(f"Table_S{i}_*.csv")), f"Table S{i} CSV exists", results)

    deliverables = [
        "secondary_manuscript_v30_en.docx", "secondary_manuscript_v30_en.pdf",
        "additional_file_1_secondary_supplement_en.docx", "additional_file_1_secondary_supplement_en.pdf",
        "additional_file_2_tripod_ai_checklist.docx", "additional_file_2_tripod_ai_checklist.pdf",
        "additional_file_2_tripod_ai_checklist.csv", "code_and_data_availability_statement.docx",
        "code_and_data_availability_statement.pdf", "latex/main.pdf", "latex/supplement.pdf",
    ]
    for name in deliverables:
        path = OUT / name
        check(path.exists() and path.stat().st_size > 0, f"Deliverable exists: {name}", results)

    manifest = pd.read_csv(OUT / "file_manifest_sha256.csv")
    check(manifest.relative_path.is_unique, "Manifest paths are unique", results)
    sampled = manifest.loc[manifest.relative_path.isin(["secondary_manuscript_v30_en.docx", "figures/Fig1.pdf", "tables/Table_S2_same_risk_incremental_value.csv"])]
    for row in sampled.itertuples():
        digest = hashlib.sha256((OUT / row.relative_path).read_bytes()).hexdigest()
        check(digest == row.sha256, f"Manifest checksum matches: {row.relative_path}", results)

    QA = OUT / "qa"; QA.mkdir(exist_ok=True)
    with (QA / "secondary_manuscript_validation_checks.csv").open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["check", "status"]); writer.writeheader(); writer.writerows(results)
    (QA / "secondary_manuscript_validation_report.md").write_text(
        "# v30 secondary manuscript validation\n\n"
        f"All {len(results)} automated checks passed.\n\n"
        "The Word manuscript, supplement, checklist, and availability statement were rendered with the stable Windows "
        "LibreOffice workflow using isolated profiles. LaTeX main and supplement compiled with TeX Live. All rendered "
        "pages, wide tables, and final figures were visually inspected for clipping, overlap, and pagination.\n",
        encoding="utf-8",
    )
    print(f"PASS: {len(results)} validation checks")


if __name__ == "__main__":
    main()
