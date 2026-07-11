from __future__ import annotations

import csv
import shutil
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
V6 = ROOT / "outputs" / "final_manuscript_v6"
OUT = ROOT / "outputs" / "manuscript_package_v7"
LATEX_DIR = OUT / "english_latex"
WORD_DIR = OUT / "chinese_word"
DATA_DIR = OUT / "source_tables"
FIG_DIR = LATEX_DIR / "figures"
QA_DIR = OUT / "qa"

TITLE_EN = "Dynamic prediction of incident acute kidney injury after major surgery in intensive care: a retrospective MIMIC-IV cohort study"
TITLE_ZH = "外科重症患者术后新发急性肾损伤的动态预测：基于 MIMIC-IV 的回顾性队列研究"
RUNNING_EN = "Dynamic postoperative AKI prediction"
RUNNING_ZH = "术后 AKI 动态预测"

ABSTRACT_EN = """Background: Postoperative acute kidney injury (AKI) evolves over time, but prediction studies commonly use a single baseline assessment or predictors that are not aligned to a clinically defined prediction time. We developed internally validated landmark models for incident AKI after major surgery in intensive care.

Methods: We conducted a retrospective cohort study in MIMIC-IV version 3.1. The analysis unit was the first intensive care unit (ICU) stay per hospital admission in a strict therapeutic surgical cohort. Incident AKI within seven days of ICU admission was recomputed using Kidney Disease: Improving Global Outcomes serum-creatinine criteria; urine output was not used. Separate risk sets were constructed at ICU admission, 6 h and 24 h, excluding AKI that had already occurred by each landmark. Logistic regression, random forest, XGBoost and LightGBM were evaluated using an 80:20 subject-grouped split, patient-level bootstrap confidence intervals, calibration, decision-curve analysis, SHAP attribution and prespecified sensitivity analyses.

Results: Among 11,943 qualifying admissions, 10,877 were evaluable for incident AKI and 4,531 developed AKI (41.7%). Stage 1, 2 and 3 accounted for 84.9%, 9.8% and 5.3% of events, respectively. Risk sets comprised 10,877 admissions at 0 h, 10,624 at 6 h and 9,301 at 24 h. XGBoost achieved the highest test AUROC at 0 h (0.728, 95% CI 0.706-0.748) and 6 h (0.740, 0.719-0.762). At 24 h, logistic regression had the highest AUROC (0.754, 0.729-0.777), AUPRC 0.602 and Brier score 0.176. Removing creatinine-derived predictors reduced 24 h AUROC by 0.009-0.025 across model families, whereas restriction to a pre-ICU seven-day baseline creatinine produced no material change.

Conclusions: Routinely available peri-ICU data supported moderate internal discrimination for postoperative incident AKI across dynamic landmarks. Performance partly depended on creatinine information at 24 h, although non-creatinine models retained useful signal. External and temporal validation are required before clinical use."""

INTRO_EN = [
    "Acute kidney injury (AKI) is a frequent complication of major surgery and is associated with short-term complications, chronic kidney disease, cardiovascular events and death. The postoperative setting is biologically heterogeneous: haemodynamic instability, inflammation, nephrotoxic exposure, anaemia and pre-existing kidney vulnerability may contribute at different times after surgery. Consensus guidance therefore emphasizes early risk recognition, repeated monitoring and avoidance of secondary kidney insults rather than reliance on a single static assessment \\citep{Prowle2021,Wang2017}.",
    "The Kidney Disease: Improving Global Outcomes (KDIGO) definition provides a reproducible framework for AKI ascertainment using changes in serum creatinine and urine output \\citep{KDIGO2012}. However, serum creatinine is both a delayed marker of kidney dysfunction and a dominant input in many prediction systems. Models evaluated after ICU admission may therefore appear accurate partly because early creatinine change already lies on the pathway to the creatinine-defined outcome. Prediction time, predictor availability and exclusion of already established AKI must be aligned to avoid information leakage and to preserve a clinically interpretable target.",
    "Electronic health record models have shown that impending AKI can be predicted from routinely collected longitudinal data \\citep{Koyner2018,Tomasev2019}. Yet model performance may not transport across surgical populations, outcome definitions and landmark times. Postoperative cohorts also require careful separation of therapeutic operations from diagnostic or bedside procedures, and subject-level data splitting is needed when patients contribute more than one hospital admission.",
    "We therefore constructed a strict postoperative surgical ICU cohort in MIMIC-IV, recomputed seven-day incident AKI from timestamped serum-creatinine measurements, and developed separate prediction datasets at ICU admission, 6 h and 24 h. We compared four model families using a common preprocessing and patient-grouped internal-validation strategy. We further examined calibration, clinical net benefit, global feature attribution, subgroup performance, a no-creatinine analysis and restriction to patients with a pre-ICU baseline creatinine."
]

METHODS_EN = {
    "Study design and data source": [
        "We conducted a retrospective cohort study using the de-identified MIMIC-IV version 3.1 database, which contains hospital and ICU data from Beth Israel Deaconess Medical Center \\citep{Johnson2023}. The analysis unit was the first qualifying ICU stay within each hospital admission. The index time was ICU admission (intime); the recorded surgery date was retained for audit but was not treated as a precise operative end time. Reporting was organized with reference to TRIPOD+AI, and risk-of-bias considerations were informed by PROBAST \\citep{Collins2024,Wolff2019}."
    ],
    "Cohort construction": [
        "Adults were eligible for the strict primary cohort when an explicitly therapeutic major surgical procedure was recorded on the day of ICU admission or the preceding day and the first ICU location was surgical, cardiac vascular, trauma surgical, mixed medical-surgical, neuro-surgical or post-anaesthesia care. Eligible operations comprised cardiac, vascular, general gastrointestinal or hepatobiliary, major orthopaedic, neurosurgical, and thoracic or respiratory surgery. Diagnostic imaging, electrocardiography, vascular access, tracheal intubation, enteral nutrition, dialysis, non-operative respiratory measurements and obstetric procedures were excluded from the strict surgical definition.",
        "Patients with AKI already present at or before ICU admission, no usable baseline serum creatinine (SCr), or no post-index SCr measurement within seven days were retained for audit but excluded from incident-AKI modelling. A broad procedure-based cohort was retained during cohort development as a sensitivity cohort but was not used for primary model training."
    ],
    "Baseline kidney function and AKI outcome": [
        "Baseline SCr was the lowest value during the seven days before ICU admission. If unavailable, the earliest SCr during the first 24 h after hospital admission was used, and its source and timestamp were retained. At any prediction landmark, a fallback baseline was exposed to the model only if its timestamp preceded that landmark.",
        "Incident AKI was adjudicated from timestamped SCr values without urine-output criteria. AKI was present when SCr increased by at least 0.3 mg/dL from a prior result within 48 h or reached at least 1.5 times baseline within seven days after ICU admission \\citep{KDIGO2012}. Patients meeting either criterion before or at ICU admission were classified as prevalent AKI. Severity was assigned from the seven-day peak SCr: stage 1, 1.5 to less than 2.0 times baseline or an absolute rise of at least 0.3 mg/dL; stage 2, 2.0 to less than 3.0 times baseline; and stage 3, at least 3.0 times baseline or peak SCr at least 4.0 mg/dL."
    ],
    "Dynamic landmark datasets and predictors": [
        "Prediction datasets were constructed at ICU admission (0 h), 6 h and 24 h. Patients with AKI onset at or before a landmark were excluded from that landmark risk set. The landmark-specific outcome was new AKI after the landmark and within seven days of ICU admission.",
        "Static predictors included demographics, admission characteristics, comorbidities, Charlson score, surgical category and first ICU type. The most recent laboratory values during the 24 h before ICU admission were used as pre-index features. For 6 h and 24 h, minimum, maximum, most recent and measurement-count features were recalculated from timestamped laboratory and vital-sign observations in (0, landmark]. Whole-period post-index summaries, mortality, length of stay, AKI-derived variables and untimed laboratory summaries were excluded."
    ],
    "Model development and internal validation": [
        "Logistic regression, random forest, XGBoost and LightGBM models were fitted separately at each landmark \\citep{Breiman2001,Chen2016,Ke2017}. Continuous variables were median-imputed with missingness indicators. Categorical and binary variables were imputed using the most frequent value, and categorical variables were one-hot encoded. Continuous variables were standardized for logistic regression.",
        "An 80:20 subject-grouped split was selected from 500 candidate GroupShuffleSplit assignments to approximate overall outcome prevalence. The same subject assignment was reused at all landmarks, and no subject appeared in both training and test sets. Random forest used 500 trees and minimum leaf size 5; XGBoost used 500 trees, learning rate 0.03 and maximum depth 4; LightGBM used 500 trees, learning rate 0.03 and 31 leaves. Hyperparameters were fixed before test evaluation."
    ],
    "Performance and statistical analysis": [
        "Discrimination was summarized by the area under the receiver-operating-characteristic curve (AUROC) and area under the precision-recall curve (AUPRC). Accuracy, sensitivity, specificity, precision, F1 score and confusion matrices were calculated at a threshold of 0.5. A secondary threshold maximizing the Youden index was selected in training predictions and applied unchanged to the test set.",
        "Calibration was assessed using the Brier score, calibration intercept, calibration slope and ten equal-frequency calibration bins \\citep{VanCalster2019}. Decision-curve analysis quantified net benefit over threshold probabilities 0.05 to 0.80 relative to treat-all and treat-none strategies \\citep{Vickers2006}. Confidence intervals were obtained from 1,000 subject-level bootstrap resamples of the held-out test set. Prespecified subgroup estimates used 300 subject-level bootstrap resamples and were omitted for groups with fewer than 50 observations or a single outcome class."
    ],
    "Interpretability and sensitivity analyses": [
        "Global SHAP values were calculated in up to 1,000 held-out observations for the tree model with the highest test AUROC at each landmark \\citep{Lundberg2017}. SHAP magnitudes were interpreted as model-attribution measures, not causal effects.",
        "Two sensitivity analyses preserved the original subject assignment. First, all predictors containing creatinine, baseline SCr, or baseline-to-ICU timing information were removed and models were retrained. Second, models were retrained only among patients with a baseline SCr measured during the seven days before ICU admission. Paired subject-level bootstrap resampling compared predictions on identical test patients."
    ],
    "Software": [
        "Analyses were performed in Python 3.14 using pandas 3.0, scikit-learn 1.9, XGBoost 3.3, LightGBM 4.6 and SHAP 0.52 \\citep{Pedregosa2011}. All cohort, outcome, feature and modelling stages were implemented as reproducible scripts."
    ]
}

RESULTS_EN = {
    "Cohort finalization and incident AKI": [
        "The strict postoperative surgical ICU cohort included 11,943 hospital admissions (Fig. 1). We excluded 1,014 admissions with AKI present at or before ICU admission, 50 without a usable baseline SCr and two without a post-index SCr measurement, leaving 10,877 admissions in the incident-AKI analytic cohort. Incident AKI occurred in 4,531 admissions (41.7%). Of these events, 3,847 (84.9%) were stage 1, 445 (9.8%) were stage 2 and 239 (5.3%) were stage 3. Median AKI onset was 31.7 h after ICU admission (interquartile range, 16.9-42.5 h).",
        "Median age was 66 years and was higher among patients who developed AKI than among those who did not (69 versus 64 years; Table 1). Cardiac operations accounted for 74.1% of the cohort, and 77.6% had a baseline SCr measured within seven days before ICU admission. Chronic kidney disease, heart failure, hypertension, diabetes, anaemia and cardiac surgery were more frequent among admissions with AKI."
    ],
    "Dynamic prediction risk sets": [
        "All 10,877 evaluable admissions entered the 0 h dataset, with 4,531 subsequent AKI events (41.7%). Exclusion of 253 events occurring by 6 h yielded 10,624 admissions and 4,278 future events at 6 h. Exclusion of 1,576 cumulative events occurring by 24 h yielded 9,301 admissions and 2,955 future events at 24 h (Fig. 1). Because the populations and outcome windows changed across landmarks, between-landmark metric differences were not treated as paired longitudinal comparisons."
    ],
    "Model discrimination and calibration": [
        "At 0 h, XGBoost provided the highest test discrimination, with AUROC 0.728 (95% CI, 0.706-0.748) and AUPRC 0.665 (0.633-0.696). XGBoost also had the highest AUROC at 6 h (0.740, 0.719-0.762) and AUPRC 0.675 (0.644-0.708). At 24 h, logistic regression had the highest AUROC (0.754, 0.729-0.777), AUPRC 0.602 (0.562-0.643) and Brier score 0.176 (Fig. 2; Table 2). Confidence intervals overlapped substantially across model families and did not establish algorithmic superiority.",
        "Calibration was close to ideal for XGBoost at 0 h (intercept -0.003; slope 0.997) and remained acceptable at 6 h (intercept -0.018; slope 0.935). The 24 h logistic model had an intercept of -0.144 and slope of 0.831, consistent with moderately over-extreme predictions in the held-out sample. The selected models showed greater internal net benefit than treat-all and treat-none strategies across most evaluated thresholds (Fig. 3)."
    ],
    "Model interpretation and subgroup performance": [
        "At 6 h, baseline SCr, age, early minimum haemoglobin, Charlson score and chronic kidney disease were prominent XGBoost features. At 24 h, the most recent creatinine had the largest mean absolute SHAP value, followed by age, minimum haemoglobin, cardiac vascular ICU location, cardiac surgery, minimum creatinine, white-cell count and blood urea nitrogen (Fig. 4).",
        "Discrimination was broadly preserved across sex and age groups, although estimates were less precise in smaller surgical subgroups (Supplementary Fig. S4; Supplementary Table S2). The selected 24 h logistic model had AUROC 0.767 (95% CI, 0.733-0.807) in women and 0.745 (0.715-0.776) in men. Corresponding estimates were 0.732 (0.693-0.767) for patients younger than 65 years and 0.750 (0.719-0.776) for those aged 65 years or older."
    ],
    "Sensitivity analyses": [
        "Removing creatinine-derived predictors reduced 24 h discrimination in every model family (Supplementary Fig. S2; Table 3). Logistic-regression AUROC decreased from 0.754 to 0.729, a paired difference of -0.025 (95% CI, -0.036 to -0.013). Paired confidence intervals also excluded zero for random forest, XGBoost and LightGBM. Nevertheless, the no-creatinine models retained AUROCs of approximately 0.73.",
        "Restricting the cohort to patients with a pre-index seven-day SCr baseline did not materially alter 24 h discrimination (Supplementary Fig. S3). The retrained logistic model achieved AUROC 0.755, compared with 0.757 for the full-cohort model evaluated in the same restricted test patients; the paired difference was -0.002 (95% CI, -0.007 to 0.004). Paired AUROC confidence intervals crossed zero for all four model families."
    ]
}

DISCUSSION_EN = [
    "In this strict postoperative surgical ICU cohort, routinely available data supported moderate prediction of seven-day incident AKI at ICU admission, 6 h and 24 h. XGBoost had the highest AUROC in the 0 h and 6 h test sets, whereas logistic regression performed best at 24 h. Differences between model families were small relative to their bootstrap uncertainty. This pattern argues against assuming that greater algorithmic complexity is necessary for later postoperative risk estimation.",
    "The landmark design is central to interpretation. At 6 h and 24 h, patients whose AKI had already occurred were removed, and predictors were restricted to measurements available by the corresponding time. The later models therefore estimated residual future risk among event-free patients rather than repeatedly scoring a fixed cohort. The modest rise in AUROC across landmarks cannot be interpreted as a within-patient improvement caused by accumulating data because event prevalence, case mix and prediction horizon changed simultaneously.",
    "Creatinine was the dominant 24 h XGBoost attribution and removing all creatinine-related predictors reduced AUROC across model families. This result has two complementary interpretations. First, early postoperative kidney-function trajectories contain genuine predictive information. Second, because the outcome itself was defined by future creatinine change, part of the apparent performance may reflect proximity between a predictor and its outcome definition. The retained AUROC of approximately 0.73 in no-creatinine models indicates that age, comorbidity, surgical context, haemoglobin and other physiological data contributed information beyond direct creatinine measurements.",
    "Calibration and clinical utility deserve equal attention to discrimination. The selected 24 h logistic model had the highest AUROC but showed a calibration slope below one and a negative intercept, indicating that recalibration may be required in a new setting. Decision curves suggested potential net benefit over default strategies across a broad internal range of thresholds, but this analysis does not define an intervention, prove patient benefit or identify an acceptable alert burden. Prospective implementation would require an explicit response pathway, threshold selection with clinicians, calibration monitoring and assessment of workflow consequences.",
    "The study has several strengths. Surgical eligibility excluded diagnostic and bedside procedures that are often misclassified when any procedure code is treated as surgery. AKI was recomputed from timestamped laboratory data after the index time was fixed. Landmark-specific features were reconstructed from their source timestamps, patient identity was respected during data splitting and paired bootstrap analyses were used for sensitivity comparisons. The no-creatinine and pre-index-baseline analyses directly tested two important sources of optimistic performance.",
    "Several limitations constrain interpretation. MIMIC-IV represents one tertiary academic centre and the cohort was dominated by cardiac surgery, limiting transportability to other institutions and surgical mixes. The retrospective design is susceptible to measurement and selection processes embedded in routine care. AKI ascertainment omitted urine output, and admission-based fallback baselines may differ from stable outpatient kidney function. Procedure dates were day-level rather than precise operative end times. Intraoperative exposures, fluid balance, medication timing and several potentially modifiable factors were not comprehensively modelled. The held-out split provides internal rather than temporal or external validation. Finally, SHAP values describe predictive attribution and cannot identify causal or modifiable determinants of AKI."
]

CONCLUSION_EN = "In a strict MIMIC-IV postoperative surgical ICU cohort, dynamic landmark models achieved moderate internal discrimination for seven-day incident AKI. Early models favoured XGBoost, while logistic regression performed best at 24 h; no single algorithm was clearly superior. Creatinine trajectories contributed substantially at 24 h, but non-creatinine information retained predictive signal and restriction to a pre-ICU baseline did not materially change performance. External and temporal validation, urine-output outcome assessment and prospective workflow evaluation are required before clinical deployment."

ABSTRACT_ZH = """背景：术后急性肾损伤（AKI）的风险随时间演变，但既往预测研究常采用单一基线评估，或未将预测变量严格限定在临床预测时点之前。本研究拟建立并内部验证外科重症患者术后新发 AKI 的动态地标预测模型。

方法：基于 MIMIC-IV 3.1 开展回顾性队列研究。严格治疗性外科手术队列以每次住院首次 ICU stay 为分析单位。依据 KDIGO 血清肌酐标准重新判定 ICU 入科后 7 天内新发 AKI，暂不使用尿量标准。在 ICU 入科、6 h 和 24 h 构建独立风险集，并排除各地标时点前已发生 AKI 的患者。采用按患者分组的 80:20 开发/测试划分，比较 Logistic 回归、随机森林、XGBoost 和 LightGBM；评估 bootstrap 置信区间、校准、决策曲线、SHAP 归因以及预设敏感性分析。

结果：11,943 次符合条件的住院中，10,877 次可评价新发 AKI，其中 4,531 次发生 AKI（41.7%）；AKI 1、2、3 期分别占 84.9%、9.8% 和 5.3%。0、6、24 h 风险集分别包含 10,877、10,624 和 9,301 次住院。XGBoost 在 0 h（AUROC 0.728，95% CI 0.706-0.748）和 6 h（0.740，0.719-0.762）表现最佳。24 h 时 Logistic 回归 AUROC 最高（0.754，0.729-0.777），AUPRC 为 0.602，Brier 分数为 0.176。移除肌酐相关预测变量后，各模型 24 h AUROC 下降 0.009-0.025；限定 ICU 前 7 天存在基线肌酐后，模型性能无实质变化。

结论：常规围 ICU 数据可在不同动态地标对术后新发 AKI 提供中等程度的内部区分能力。24 h 预测部分依赖肌酐信息，但非肌酐变量仍保留预测信号。模型临床使用前仍需外部和时间验证。"""

INTRO_ZH = [
    "急性肾损伤（AKI）是大手术后的常见并发症，与短期围术期并发症、慢性肾脏病、心血管事件及死亡风险增加相关[1,2]。术后 AKI 的病理生理具有明显异质性：血流动力学不稳定、炎症、肾毒性暴露、贫血和既往肾脏易感性可能在术后不同阶段发挥作用。因此，围术期共识强调早期风险识别、重复监测和避免继发性肾损伤，而不能仅依赖单次静态风险评估[1]。",
    "KDIGO 标准依据血清肌酐变化和尿量为 AKI 提供了可重复的判定框架[3]。然而，血清肌酐既是肾功能损害的滞后指标，也是许多预测模型中的核心变量。若在 ICU 入科后评估模型，而未明确预测时点、变量可用时间和既存 AKI，模型性能可能部分来自已经发生的早期肌酐变化，甚至产生信息泄漏。",
    "既往电子病历研究表明，利用纵向常规临床数据可以预测即将发生的 AKI[4,5]。但模型性能受到手术人群、结局定义和预测时点影响，未必能够直接迁移。术后队列还必须区分真正的治疗性手术与影像检查、置管等非手术操作；若同一患者存在多次住院，则模型开发与测试亦需按患者分组。",
    "本研究据此构建严格的外科术后 ICU 队列，使用时间戳化血清肌酐数据重新计算 ICU 入科后 7 天内新发 AKI，并在 0、6、24 h 建立独立动态地标数据集。我们在一致的预处理和患者分组内部验证框架下比较四类模型，并进一步评估校准、决策曲线、SHAP 归因、亚组表现、去肌酐模型和 ICU 前基线肌酐限定模型。"
]

METHODS_ZH = {
    "研究设计与数据来源": [
        "本研究为回顾性队列研究，使用去标识化的 MIMIC-IV 3.1 数据库。该数据库包含 Beth Israel Deaconess Medical Center 的住院及 ICU 数据[6]。分析单位为每次住院内首次符合条件的 ICU stay。索引时点定义为 ICU 入科时间（intime）；手术日期仅用于队列审计，不视为精确手术结束时间。报告结构参考 TRIPOD+AI，偏倚风险考虑参考 PROBAST[14,15]。"
    ],
    "队列构建": [
        "当 ICU 入科当日或前一日记录了明确治疗性大手术，且首次 ICU 科室为外科、心血管、创伤外科、综合内外科、神经外科或麻醉后监护单元时，纳入严格主队列。手术类别包括心脏、血管、普外/胃肠肝胆、重大骨科、神经外科以及胸外/呼吸系统手术。单纯影像检查、心电图、血管通路、气管插管、肠内营养、透析、非手术性呼吸测量和产科操作不作为严格外科手术。",
        "索引时点前或当时已满足 AKI、无可用基线肌酐，或索引后 7 天内无肌酐随访的病例保留用于审计，但排除出新发 AKI 建模队列。宽泛 procedure 队列仅用于队列构建敏感性审计，不用于主要模型训练。"
    ],
    "基线肾功能与 AKI 结局": [
        "基线血清肌酐定义为 ICU 入科前 7 天内最低值。若无该值，则使用住院后最初 24 h 内最早的肌酐，并保留其来源与时间戳。在任何地标时点，只有当回退基线的测量时间早于该地标时，才允许模型使用。",
        "依据带时间戳的血清肌酐判定新发 AKI，暂不使用尿量标准。若 48 h 内肌酐较既往值升高至少 0.3 mg/dL，或 ICU 入科后 7 天内达到基线的至少 1.5 倍，则判定 AKI[3]。索引时点前或当时满足任一标准者定义为既存 AKI。分期基于 7 天峰值：1 期为基线的 1.5 至不足 2.0 倍或绝对升高至少 0.3 mg/dL；2 期为 2.0 至不足 3.0 倍；3 期为至少 3.0 倍或峰值至少 4.0 mg/dL。"
    ],
    "动态地标数据与预测变量": [
        "分别在 ICU 入科（0 h）、6 h 和 24 h 构建预测数据。各地标排除在该时点或此前已发生 AKI 的病例；结局为地标之后至 ICU 入科后 7 天内新发 AKI。",
        "静态变量包括人口学特征、入院特征、合并症、Charlson 评分、手术类别及首次 ICU 类型。ICU 入科前 24 h 的实验室指标采用最近一次值。6 h 和 24 h 模型的最小值、最大值、最近值和测量次数均由 (0, 地标] 内带时间戳的实验室和生命体征重新计算。全随访期 post_* 汇总、死亡、住院时长、AKI 衍生变量及无时间信息的实验室汇总均被排除。"
    ],
    "模型开发与内部验证": [
        "分别在每个地标训练 Logistic 回归、随机森林、XGBoost 和 LightGBM[7-10]。连续变量以中位数插补并保留缺失指示；分类和二元变量以众数插补；分类变量进行独热编码。Logistic 回归对连续变量进行标准化。",
        "从 500 个 GroupShuffleSplit 候选方案中选择接近总体结局比例的 80:20 患者分组划分，并在三个地标复用同一患者分配。同一患者不会同时进入训练集和测试集。随机森林设置 500 棵树、最小叶节点样本数 5；XGBoost 设置 500 棵树、学习率 0.03、最大深度 4；LightGBM 设置 500 棵树、学习率 0.03、31 个叶节点。测试集评估前固定超参数。"
    ],
    "性能与统计分析": [
        "区分度采用 AUROC 和 AUPRC。以 0.5 阈值计算准确率、敏感度、特异度、阳性预测值、F1 分数和混淆矩阵。另在训练集预测中确定 Youden 指数最大的阈值，并原样应用于测试集。",
        "校准采用 Brier 分数、校准截距、校准斜率及 10 个等频分箱[12]。决策曲线分析在 0.05-0.80 阈值概率范围内比较模型与全部干预、均不干预策略的净获益[13]。测试集采用 1,000 次患者级 bootstrap 计算置信区间；亚组分析采用 300 次患者级 bootstrap，样本量不足 50 或仅有单一结局类别时不估计。"
    ],
    "可解释性与敏感性分析": [
        "在各地标测试集 AUROC 最高的树模型中，最多抽取 1,000 条测试集记录计算全局 SHAP 值[11]。SHAP 仅解释模型归因，不解释因果作用。",
        "两项敏感性分析均保留原患者划分。第一项删除名称中含 creatinine、baseline SCr 或 baseline-to-ICU 时间信息的全部预测变量后重新训练；第二项仅纳入 ICU 前 7 天存在基线肌酐的患者重新训练。采用配对患者级 bootstrap 在相同测试患者中比较预测。"
    ],
    "软件": [
        "分析使用 Python 3.14、pandas 3.0、scikit-learn 1.9、XGBoost 3.3、LightGBM 4.6 和 SHAP 0.52[7-11]。队列、结局、特征和建模过程均由可复现脚本实现。"
    ]
}

RESULTS_ZH = {
    "队列与新发 AKI": [
        "严格外科术后 ICU 队列共 11,943 次住院（图1）。排除索引时已存在 AKI 的 1,014 次、无可用基线肌酐的 50 次及无索引后肌酐的 2 次后，10,877 次住院进入新发 AKI 分析。4,531 次发生 AKI（41.7%）；其中 1 期 3,847 次（84.9%）、2 期 445 次（9.8%）、3 期 239 次（5.3%）。AKI 中位发生时间为 ICU 入科后 31.7 h（四分位距 16.9-42.5 h）。",
        "总体中位年龄 66 岁，AKI 患者高于非 AKI 患者（69 岁 vs 64 岁；表1）。心脏手术占 74.1%，77.6% 的病例具有 ICU 前 7 天基线肌酐。慢性肾脏病、心力衰竭、高血压、糖尿病、贫血及心脏手术在 AKI 组更常见。"
    ],
    "动态预测风险集": [
        "0 h 数据集纳入全部 10,877 次可评价住院，后续 AKI 4,531 次（41.7%）。排除 6 h 内已发生 AKI 的 253 次后，6 h 风险集为 10,624 次，后续 AKI 4,278 次。累计排除 24 h 内已发生 AKI 的 1,576 次后，24 h 风险集为 9,301 次，后续 AKI 2,955 次（图1）。由于地标间人群与结局时间窗均发生变化，不将地标间指标差异解释为固定队列的配对纵向改善。"
    ],
    "模型区分度、校准与临床净获益": [
        "0 h 时 XGBoost 测试集区分度最高，AUROC 0.728（95% CI 0.706-0.748），AUPRC 0.665（0.633-0.696）。6 h 时 XGBoost AUROC 仍最高，为 0.740（0.719-0.762），AUPRC 0.675（0.644-0.708）。24 h 时 Logistic 回归 AUROC 最高，为 0.754（0.729-0.777），AUPRC 0.602（0.562-0.643），Brier 分数 0.176（图2、表2）。不同模型的置信区间高度重叠，不能据此证明某一算法具有明确优势。",
        "0 h XGBoost 校准接近理想状态（截距 -0.003，斜率 0.997），6 h 仍可接受（截距 -0.018，斜率 0.935）。24 h Logistic 回归截距为 -0.144、斜率为 0.831，提示测试集中预测风险存在一定过度离散。所选模型在大部分评估阈值上较全部干预和均不干预策略具有更高的内部净获益（图3）。"
    ],
    "模型解释与亚组表现": [
        "6 h XGBoost 的主要特征包括基线肌酐、年龄、早期最低血红蛋白、Charlson 评分和慢性肾脏病。24 h 时最近一次肌酐具有最大的平均绝对 SHAP 值，其后依次包括年龄、最低血红蛋白、心血管 ICU、心脏手术、最低肌酐、白细胞和血尿素氮（图4）。",
        "按性别和年龄分层后区分度总体保持稳定，但较小手术亚组的置信区间较宽（补充图 S4、补充表 S2）。24 h Logistic 模型在女性和男性中的 AUROC 分别为 0.767（0.733-0.807）和 0.745（0.715-0.776）；在小于 65 岁和至少 65 岁患者中分别为 0.732（0.693-0.767）和 0.750（0.719-0.776）。"
    ],
    "敏感性分析": [
        "删除全部肌酐相关预测变量后，四类模型 24 h 区分度均下降（补充图 S2、表3）。Logistic 回归 AUROC 从 0.754 降至 0.729，配对差值为 -0.025（95% CI -0.036 至 -0.013）；随机森林、XGBoost 和 LightGBM 的配对置信区间同样不跨越零。尽管如此，去肌酐模型 AUROC 仍约为 0.73。",
        "仅纳入 ICU 前 7 天存在基线肌酐的患者后，24 h 区分度无实质改变（补充图 S3）。重新训练的 Logistic 模型 AUROC 为 0.755，而完整模型在相同限定测试患者中的 AUROC 为 0.757；配对差值 -0.002（95% CI -0.007 至 0.004）。四类模型的配对 AUROC 置信区间均跨越零。"
    ]
}

DISCUSSION_ZH = [
    "在严格外科术后 ICU 队列中，常规临床数据能够在 ICU 入科、6 h 和 24 h 对 7 天内新发 AKI 提供中等程度的预测能力。XGBoost 在 0 h 和 6 h 的 AUROC 最高，而 24 h 时 Logistic 回归表现最佳；不同模型之间的差异小于其 bootstrap 不确定性范围。因此，术后较晚时点并不必然需要更复杂的算法。",
    "动态地标设计是解释结果的关键。6 h 和 24 h 模型均排除了此前已经发生 AKI 的病例，并将预测变量严格限制为对应时点前已获得的信息。因此，后续模型估计的是仍未发生事件患者的剩余未来风险，而不是对固定队列反复评分。随着地标推进，事件率、病例构成和预测时距同时变化，AUROC 的增加不能直接解释为同一患者因数据积累而获得的纵向改善。",
    "肌酐是 24 h XGBoost 中最主要的模型归因，删除所有肌酐相关变量后，四类模型 AUROC 均下降。这一结果具有两方面含义：早期术后肾功能轨迹确实包含预测信息；与此同时，结局本身由未来肌酐变化定义，部分性能可能来自预测变量与结局定义之间的时间邻近性。去肌酐模型仍保持约 0.73 的 AUROC，说明年龄、合并症、手术背景、血红蛋白和其他生理信息提供了独立于直接肌酐测量的预测信号。",
    "预测研究不能只关注区分度。24 h Logistic 模型具有最高 AUROC，但校准斜率低于 1、截距为负，提示在新环境中可能需要重新校准。DCA 显示在内部验证的大部分阈值范围内存在潜在净获益，但这并未定义具体干预措施，也不能证明改善患者结局。未来实施需要明确预警后的响应路径，与临床团队共同选择阈值，并评估警报负担、校准漂移和工作流程影响。",
    "本研究的优势包括：通过可审计规则排除常被误认为手术的诊断和床旁操作；在确定索引时点后依据时间戳化实验室数据重新计算 AKI；按地标重建变量窗口；按患者进行开发/测试划分；并以配对 bootstrap 直接评估去肌酐和术前基线限定的稳健性。",
    "本研究仍有重要局限。MIMIC-IV 来自单个三级学术医疗中心，且队列以心脏手术为主，限制了向其他机构和手术构成的推广。回顾性常规数据受到测量过程和临床选择影响。AKI 判定未纳入尿量，住院后回退基线也不等同于稳定门诊肾功能。手术代码仅有日期级时间精度。术中暴露、液体平衡、药物时间及多项潜在可干预因素尚未完整建模。本研究仅进行内部随机分组验证，尚无时间验证或外部验证。最后，SHAP 仅描述预测归因，不能识别 AKI 的因果或可干预因素。"
]

CONCLUSION_ZH = "在严格的 MIMIC-IV 外科术后 ICU 队列中，动态地标模型对 7 天内新发 AKI 获得了中等程度的内部区分度。早期模型以 XGBoost 表现较好，24 h 时 Logistic 回归表现最佳，但尚无证据证明某一算法明确优于其他算法。24 h 风险预测明显利用了肌酐轨迹，但非肌酐信息仍保留预测信号，且限定 ICU 前基线肌酐后性能无实质变化。临床使用前仍需外部和时间验证、尿量结局评估以及前瞻性工作流程研究。"

REFERENCES = [
    "Prowle JR, Forni LG, Bell M, et al. Postoperative acute kidney injury in adult non-cardiac surgery: joint consensus report of the Acute Disease Quality Initiative and PeriOperative Quality Initiative. Nat Rev Nephrol. 2021;17:605-618. doi:10.1038/s41581-021-00418-2.",
    "Wang Y, Bellomo R. Cardiac surgery-associated acute kidney injury: risk factors, pathophysiology and treatment. Nat Rev Nephrol. 2017;13:697-711. doi:10.1038/nrneph.2017.119.",
    "Kidney Disease: Improving Global Outcomes Acute Kidney Injury Work Group. KDIGO Clinical Practice Guideline for Acute Kidney Injury. Kidney Int Suppl. 2012;2:1-138.",
    "Koyner JL, Carey KA, Edelson DP, Churpek MM. The development of a machine learning inpatient acute kidney injury prediction model. Crit Care Med. 2018;46:1070-1077. doi:10.1097/CCM.0000000000003123.",
    "Tomasev N, Glorot X, Rae JW, et al. A clinically applicable approach to continuous prediction of future acute kidney injury. Nature. 2019;572:116-119. doi:10.1038/s41586-019-1390-1.",
    "Johnson AEW, Bulgarelli L, Shen L, et al. MIMIC-IV, a freely accessible electronic health record dataset. Sci Data. 2023;10:1. doi:10.1038/s41597-022-01899-x.",
    "Chen T, Guestrin C. XGBoost: A scalable tree boosting system. Proc 22nd ACM SIGKDD Int Conf Knowl Discov Data Min. 2016:785-794. doi:10.1145/2939672.2939785.",
    "Ke G, Meng Q, Finley T, et al. LightGBM: A highly efficient gradient boosting decision tree. Adv Neural Inf Process Syst. 2017;30:3146-3154.",
    "Breiman L. Random forests. Mach Learn. 2001;45:5-32. doi:10.1023/A:1010933404324.",
    "Pedregosa F, Varoquaux G, Gramfort A, et al. Scikit-learn: Machine learning in Python. J Mach Learn Res. 2011;12:2825-2830.",
    "Lundberg SM, Lee SI. A unified approach to interpreting model predictions. Adv Neural Inf Process Syst. 2017;30:4765-4774.",
    "Van Calster B, McLernon DJ, van Smeden M, Wynants L, Steyerberg EW. Calibration: the Achilles heel of predictive analytics. BMC Med. 2019;17:230. doi:10.1186/s12916-019-1466-7.",
    "Vickers AJ, Elkin EB. Decision curve analysis: a novel method for evaluating prediction models. Med Decis Making. 2006;26:565-574. doi:10.1177/0272989X06295361.",
    "Collins GS, Moons KGM, Dhiman P, et al. TRIPOD+AI statement: updated guidance for reporting clinical prediction models that use regression or machine learning methods. BMJ. 2024;385:e078378. doi:10.1136/bmj-2023-078378.",
    "Wolff RF, Moons KGM, Riley RD, et al. PROBAST: a tool to assess the risk of bias and applicability of prediction model studies. Ann Intern Med. 2019;170:51-58. doi:10.7326/M18-1376."
]

BIBTEX = r"""@article{Prowle2021,author={Prowle, John R. and Forni, Lui G. and Bell, Max and others},title={Postoperative acute kidney injury in adult non-cardiac surgery: joint consensus report of the Acute Disease Quality Initiative and PeriOperative Quality Initiative},journal={Nature Reviews Nephrology},year={2021},volume={17},pages={605--618},doi={10.1038/s41581-021-00418-2}}
@article{Wang2017,author={Wang, Ying and Bellomo, Rinaldo},title={Cardiac surgery-associated acute kidney injury: risk factors, pathophysiology and treatment},journal={Nature Reviews Nephrology},year={2017},volume={13},pages={697--711},doi={10.1038/nrneph.2017.119}}
@article{KDIGO2012,author={{Kidney Disease: Improving Global Outcomes Acute Kidney Injury Work Group}},title={KDIGO Clinical Practice Guideline for Acute Kidney Injury},journal={Kidney International Supplements},year={2012},volume={2},pages={1--138}}
@article{Koyner2018,author={Koyner, Jay L. and Carey, Kyle A. and Edelson, Dana P. and Churpek, Matthew M.},title={The development of a machine learning inpatient acute kidney injury prediction model},journal={Critical Care Medicine},year={2018},volume={46},pages={1070--1077},doi={10.1097/CCM.0000000000003123}}
@article{Tomasev2019,author={Tomasev, Nenad and Glorot, Xavier and Rae, Jack W. and others},title={A clinically applicable approach to continuous prediction of future acute kidney injury},journal={Nature},year={2019},volume={572},pages={116--119},doi={10.1038/s41586-019-1390-1}}
@article{Johnson2023,author={Johnson, Alistair E. W. and Bulgarelli, Lucas and Shen, Lu and others},title={MIMIC-IV, a freely accessible electronic health record dataset},journal={Scientific Data},year={2023},volume={10},pages={1},doi={10.1038/s41597-022-01899-x}}
@inproceedings{Chen2016,author={Chen, Tianqi and Guestrin, Carlos},title={XGBoost: A scalable tree boosting system},booktitle={Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining},year={2016},pages={785--794},doi={10.1145/2939672.2939785}}
@inproceedings{Ke2017,author={Ke, Guolin and Meng, Qi and Finley, Thomas and others},title={LightGBM: A highly efficient gradient boosting decision tree},booktitle={Advances in Neural Information Processing Systems},year={2017},volume={30},pages={3146--3154}}
@article{Breiman2001,author={Breiman, Leo},title={Random forests},journal={Machine Learning},year={2001},volume={45},pages={5--32},doi={10.1023/A:1010933404324}}
@article{Pedregosa2011,author={Pedregosa, Fabian and Varoquaux, Gael and Gramfort, Alexandre and others},title={Scikit-learn: Machine learning in Python},journal={Journal of Machine Learning Research},year={2011},volume={12},pages={2825--2830}}
@inproceedings{Lundberg2017,author={Lundberg, Scott M. and Lee, Su-In},title={A unified approach to interpreting model predictions},booktitle={Advances in Neural Information Processing Systems},year={2017},volume={30},pages={4765--4774}}
@article{VanCalster2019,author={Van Calster, Ben and McLernon, David J. and van Smeden, Maarten and Wynants, Laure and Steyerberg, Ewout W.},title={Calibration: the Achilles heel of predictive analytics},journal={BMC Medicine},year={2019},volume={17},pages={230},doi={10.1186/s12916-019-1466-7}}
@article{Vickers2006,author={Vickers, Andrew J. and Elkin, Elena B.},title={Decision curve analysis: a novel method for evaluating prediction models},journal={Medical Decision Making},year={2006},volume={26},pages={565--574},doi={10.1177/0272989X06295361}}
@article{Collins2024,author={Collins, Gary S. and Moons, Karel G. M. and Dhiman, Paula and others},title={TRIPOD+AI statement: updated guidance for reporting clinical prediction models that use regression or machine learning methods},journal={BMJ},year={2024},volume={385},pages={e078378},doi={10.1136/bmj-2023-078378}}
@article{Wolff2019,author={Wolff, Robert F. and Moons, Karel G. M. and Riley, Richard D. and others},title={PROBAST: a tool to assess the risk of bias and applicability of prediction model studies},journal={Annals of Internal Medicine},year={2019},volume={170},pages={51--58},doi={10.7326/M18-1376}}
"""

FIGURES = [
    ("Figure_1_cohort_flowchart", "Figure 1. Cohort derivation and dynamic landmark risk sets. The later risk sets exclude AKI with onset at or before the landmark."),
    ("Figure_2_dynamic_ROC", "Figure 2. Receiver-operating-characteristic curves at ICU admission, 6 h and 24 h. Landmark populations and future-event windows differ."),
    ("Figure_3_calibration_DCA", "Figure 3. Calibration and decision-curve analysis. Top, observed versus predicted risk in ten equal-frequency bins; bottom, net benefit relative to treat-all and treat-none strategies."),
    ("Figure_4_SHAP_6h_24h", "Figure 4. Global XGBoost feature importance at 6 h and 24 h. Mean absolute SHAP values for the twelve leading features quantify model attribution rather than causal effects.")
]

FIGURES_ZH = [
    ("Figure_1_cohort_flowchart", "图1. 队列构建与动态地标风险集。6 h 和 24 h 风险集排除了该地标或此前已发生 AKI 的病例。"),
    ("Figure_2_dynamic_ROC", "图2. ICU 入科、6 h 和 24 h 的 ROC 曲线。各地标人群和未来结局窗口不同。"),
    ("Figure_3_calibration_DCA", "图3. 校准与决策曲线。上排为 10 个等频分箱的观察风险与预测风险；下排为相对于全部干预和均不干预策略的净获益。"),
    ("Figure_4_SHAP_6h_24h", "图4. 6 h 与 24 h XGBoost 全局特征重要性。前 12 个特征的平均绝对 SHAP 值仅表示模型归因，不表示因果作用。")
]

SUPP_FIGURES_ZH = [
    ("Figure_S1_precision_recall", "补充图 S1. 四类模型在 0、6、24 h 的精确率-召回率曲线。"),
    ("Figure_S2_no_creatinine_sensitivity", "补充图 S2. 24 h 去肌酐敏感性分析及配对 AUROC 差值。"),
    ("Figure_S3_preindex_baseline_sensitivity", "补充图 S3. ICU 前 7 天基线肌酐限定模型的配对敏感性分析。"),
    ("Figure_S4_subgroup_performance", "补充图 S4. 所选地标模型在预设亚组中的 AUROC 及患者级 bootstrap 95% CI。"),
    ("Figure_S5_predictor_missingness", "补充图 S5. 建模数据中高缺失率术前实验室预测变量。")
]


def read_csv(name: str) -> list[dict[str, str]]:
    with (V6 / "tables" / name).open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def tex_escape(value: object) -> str:
    text = str(value)
    replacements = {
        "\\": r"\textbackslash{}", "&": r"\&", "%": r"\%", "$": r"\$",
        "#": r"\#", "_": r"\_", "{": r"\{", "}": r"\}",
        "~": r"\textasciitilde{}", "^": r"\textasciicircum{}",
        "–": "--", "—": "---", "−": "-", "≥": r"$\geq$", "≤": r"$\leq$", "Δ": r"$\Delta$",
    }
    for source, target in replacements.items():
        text = text.replace(source, target)
    return text


def write_tex_table(rows: list[dict[str, str]], columns: list[tuple[str, str]], caption: str, label: str,
                    path: Path, landscape: bool = False, long: bool = False) -> None:
    n = len(columns)
    alignment = "p{4.0cm}" + "p{2.2cm}" * (n - 1) if long else "l" + "c" * (n - 1)
    env = "longtable" if long else "tabular"
    lines = []
    if landscape:
        lines.append(r"\begin{landscape}")
    lines += [r"\begin{table}[p]" if not long else "", r"\centering", r"\scriptsize"]
    if not long:
        lines.append(r"\caption{" + tex_escape(caption) + r"}\label{" + label + "}")
        lines.append(r"\resizebox{\textwidth}{!}{%")
    lines.append(r"\begin{" + env + "}{" + alignment + "}")
    if long:
        lines.append(r"\caption{" + tex_escape(caption) + r"}\label{" + label + r"}\\")
    lines.append(r"\toprule")
    lines.append(" & ".join(tex_escape(title) for _, title in columns) + r" \\")
    lines.append(r"\midrule")
    if long:
        lines += [r"\endfirsthead", r"\toprule", " & ".join(tex_escape(title) for _, title in columns) + r" \\", r"\midrule", r"\endhead"]
    for row in rows:
        lines.append(" & ".join(tex_escape(row.get(key, "")) for key, _ in columns) + r" \\")
    lines.append(r"\bottomrule")
    lines.append(r"\end{" + env + "}")
    if not long:
        lines += ["}", r"\end{table}"]
    if landscape:
        lines.append(r"\end{landscape}")
    path.write_text("\n".join(line for line in lines if line != "") + "\n", encoding="utf-8")


def combine_s1() -> list[dict[str, str]]:
    output = []
    for row in read_csv("Table_S1_full_model_performance.csv"):
        output.append({
            "time": f"{row['landmark_hours']} h", "model": row["model"], "n": row["test_n"],
            "event": f"{float(row['test_event_rate']) * 100:.1f}%",
            "auroc": f"{float(row['auroc']):.3f} ({float(row['auroc_ci_lower']):.3f}-{float(row['auroc_ci_upper']):.3f})",
            "auprc": f"{float(row['auprc']):.3f} ({float(row['auprc_ci_lower']):.3f}-{float(row['auprc_ci_upper']):.3f})",
            "brier": f"{float(row['brier_score']):.3f}",
            "cal": f"{float(row['calibration_intercept']):.3f} / {float(row['calibration_slope']):.3f}",
            "youden": f"{float(row['youden_threshold']):.3f} ({float(row['threshold_youden_sensitivity']):.3f}/{float(row['threshold_youden_specificity']):.3f})",
        })
    return output


def combine_s2() -> list[dict[str, str]]:
    output = []
    group_names = {
        "baseline_scr_source": "Baseline SCr source", "first_careunit": "First ICU",
        "cardiac_surgery": "Cardiac surgery", "vascular_surgery": "Vascular surgery",
        "general_gi_hepatobiliary_surgery": "GI/hepatobiliary surgery",
        "orthopedic_major_surgery": "Major orthopaedic surgery", "neurosurgery": "Neurosurgery",
        "thoracic_respiratory_surgery": "Thoracic/respiratory surgery", "age_group": "Age group",
        "gender": "Sex", "ckd": "Chronic kidney disease",
    }
    level_names = {
        "lowest_scr_7d_pre_icu": "Pre-ICU 7-day baseline",
        "admission_scr_first_24h_fallback": "Admission SCr fallback",
        "Cardiac Vascular Intensive Care Unit (CVICU)": "CVICU",
        "Surgical Intensive Care Unit (SICU)": "SICU",
        "Trauma SICU (TSICU)": "TSICU",
    }
    for row in read_csv("Table_S2_subgroup_performance.csv"):
        output.append({
            "time": f"{row['landmark_hours']} h", "model": row["model"],
            "group": group_names.get(row["subgroup_dimension"], row["subgroup_dimension"].replace("_", " ").title()),
            "level": level_names.get(row["subgroup_level"], row["subgroup_level"]),
            "n": f"{row['n']} / {row['event_n']}",
            "auroc": f"{float(row['auroc']):.3f} ({float(row['auroc_ci_lower']):.3f}-{float(row['auroc_ci_upper']):.3f})",
            "auprc": f"{float(row['auprc']):.3f} ({float(row['auprc_ci_lower']):.3f}-{float(row['auprc_ci_upper']):.3f})",
            "brier": f"{float(row['brier_score']):.3f}",
        })
    return output


def high_missingness() -> list[dict[str, str]]:
    predictor_names = {
        "lab_pre24h_bicarbonate_last": "Pre-index bicarbonate (last)",
        "lab_pre24h_bun_last": "Pre-index BUN (last)",
        "lab_pre24h_creatinine_last": "Pre-index creatinine (last)",
        "lab_pre24h_hemoglobin_last": "Pre-index haemoglobin (last)",
        "lab_pre24h_inr_last": "Pre-index INR (last)",
        "lab_pre24h_platelet_last": "Pre-index platelet count (last)",
        "lab_pre24h_wbc_last": "Pre-index white-cell count (last)",
    }
    output = []
    for row in read_csv("Table_S3_predictor_missingness.csv"):
        if row["missing_gt40pct"].strip().lower() != "true":
            continue
        row = dict(row)
        row["predictor"] = predictor_names.get(row["predictor"], row["predictor"].replace("_", " "))
        row["variable_type"] = row["variable_type"].replace("continuous_numeric", "Continuous").replace("binary", "Binary").replace("categorical", "Categorical")
        output.append(row)
    return output


def combine_s4() -> list[dict[str, str]]:
    test_n = {"0": "2,175", "6": "2,130", "24": "1,848"}
    output = []
    for row in read_csv("Table_S4_sensitivity_all_landmarks.csv"):
        preindex = row["sensitivity_analysis"] == "preindex_baseline_only"
        reference = row["full_model_auroc"] if preindex else row["full_auroc"]
        sensitivity = row["preindex_model_auroc"] if preindex else row["no_creatinine_auroc"]
        n_value = str(int(float(row["restricted_test_n"]))) if preindex and row["restricted_test_n"] else test_n[row["landmark_hours"]]
        output.append({
            "analysis": "Pre-index baseline only" if preindex else "No creatinine",
            "time": f"{row['landmark_hours']} h", "model": row["model"], "n": n_value,
            "reference": f"{float(reference):.3f}", "sensitivity": f"{float(sensitivity):.3f}",
            "delta": f"{float(row['delta_auroc']):+.3f} ({float(row['delta_auroc_ci_lower']):+.3f} to {float(row['delta_auroc_ci_upper']):+.3f})",
        })
    return output


def build_latex() -> None:
    LATEX_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    (LATEX_DIR / "tables").mkdir(exist_ok=True)
    for file in (V6 / "figures").glob("*.pdf"):
        shutil.copy2(file, FIG_DIR / file.name)

    shutil.copy2(V6 / "tables" / "Table_1_baseline_characteristics.csv", LATEX_DIR / "tables" / "Table_1_baseline_characteristics.csv")
    for file in (V6 / "tables").glob("*.csv"):
        shutil.copy2(file, DATA_DIR / file.name)

    write_tex_table(read_csv("Table_1_baseline_characteristics.csv"), [
        ("Characteristic", "Characteristic"), ("Overall (N=10,877)", "Overall"),
        ("No AKI (N=6,346)", "No AKI"), ("Incident AKI (N=4,531)", "Incident AKI"),
        ("Standardized mean difference", "SMD"), ("Missing, n", "Missing")],
        "Baseline characteristics of the strict incident-AKI evaluable cohort", "tab:baseline",
        LATEX_DIR / "tables" / "table1.tex", landscape=True, long=True)
    write_tex_table(read_csv("Table_2_selected_model_performance.csv"), [
        ("Landmark", "Landmark"), ("Selected model", "Model"), ("Test N", "Test N"),
        ("Event rate", "Event rate"), ("AUROC (95% CI)", "AUROC (95% CI)"),
        ("AUPRC (95% CI)", "AUPRC (95% CI)"), ("Brier score", "Brier"),
        ("Calibration intercept", "Intercept"), ("Calibration slope", "Slope"),
        ("Youden threshold", "Youden"), ("Sensitivity / specificity", "Sensitivity / specificity")],
        "Selected landmark-model performance in held-out test sets", "tab:selected",
        LATEX_DIR / "tables" / "table2.tex", landscape=True)
    write_tex_table(read_csv("Table_3_sensitivity_analyses_24h.csv"), [
        ("Sensitivity analysis", "Sensitivity analysis"), ("Model", "Model"),
        ("Full/reference AUROC", "Reference AUROC"), ("Sensitivity AUROC", "Sensitivity AUROC"),
        ("ΔAUROC (95% paired CI)", "Delta AUROC (95% paired CI)")],
        "Sensitivity analyses at the 24 h landmark", "tab:sensitivity", LATEX_DIR / "tables" / "table3.tex")
    write_tex_table(combine_s1(), [("time", "Landmark"), ("model", "Model"), ("n", "Test N"),
        ("event", "Event rate"), ("auroc", "AUROC (95% CI)"), ("auprc", "AUPRC (95% CI)"),
        ("brier", "Brier"), ("cal", "Calibration I/S"), ("youden", "Youden (sens/spec)")],
        "Complete model performance by landmark", "tab:s1", LATEX_DIR / "tables" / "tableS1.tex", landscape=True)
    write_tex_table(combine_s2(), [("time", "Landmark"), ("model", "Model"), ("group", "Subgroup"),
        ("level", "Level"), ("n", "N/events"), ("auroc", "AUROC (95% CI)"),
        ("auprc", "AUPRC (95% CI)"), ("brier", "Brier")],
        "Prespecified subgroup performance", "tab:s2", LATEX_DIR / "tables" / "tableS2.tex", landscape=True, long=True)
    write_tex_table(high_missingness(), [("landmark_hours", "Landmark (h)"), ("predictor", "Predictor"),
        ("variable_type", "Type"), ("n_missing", "Missing N"), ("missing_percent", "Missing percent"),
        ("n_observed", "Observed N")], "Predictors with more than 40 percent missingness", "tab:s3",
        LATEX_DIR / "tables" / "tableS3.tex", landscape=True, long=True)
    write_tex_table(combine_s4(), [("analysis", "Sensitivity analysis"), ("time", "Landmark"),
        ("model", "Model"), ("n", "Test N"), ("reference", "Reference AUROC"),
        ("sensitivity", "Sensitivity AUROC"), ("delta", "Delta AUROC (95% paired CI)")],
        "Sensitivity analyses across all landmarks", "tab:s4", LATEX_DIR / "tables" / "tableS4.tex",
        landscape=True, long=True)

    methods = "\n".join("\\subsection{" + tex_escape(title) + "}\n" + "\n\n".join(paragraphs) for title, paragraphs in METHODS_EN.items())
    results = "\n".join("\\subsection{" + tex_escape(title) + "}\n" + "\n\n".join(paragraphs) for title, paragraphs in RESULTS_EN.items())
    intro = "\n\n".join(INTRO_EN)
    discussion = "\n\n".join(DISCUSSION_EN)
    main_tex = r"""\documentclass[11pt]{article}
\usepackage[letterpaper,margin=1in]{geometry}
\usepackage[T1]{fontenc}
\usepackage{lmodern}
\usepackage{microtype}
\usepackage{graphicx}
\usepackage{booktabs}
\usepackage{longtable}
\usepackage{pdflscape}
\usepackage{caption}
\usepackage{natbib}
\usepackage{lineno}
\usepackage[colorlinks=true,allcolors=blue]{hyperref}
\usepackage{setspace}
\onehalfspacing
\linenumbers
\title{__TITLE__}
\author{[Author names]\\{}\textit{[Affiliations]}\\{}Correspondence: [email]}
\date{}
\begin{document}
\maketitle
\begin{center}\textit{Running title: __RUNNING__}\end{center}

\begin{abstract}
__ABSTRACT__
\end{abstract}
\noindent\textbf{Keywords:} acute kidney injury; postoperative care; intensive care; dynamic prediction; machine learning; MIMIC-IV

\section{Introduction}
__INTRO__

\section{Methods}
__METHODS__

\section{Results}
__RESULTS__

\section{Discussion}
__DISCUSSION__

\section{Conclusion}
__CONCLUSION__

\section*{Declarations}
\textbf{Ethics approval and consent to participate.} MIMIC-IV contains de-identified data released under its original governance and credentialed PhysioNet access. [AUTHOR ACTION REQUIRED: insert the local institutional determination or exemption and approval number, if applicable.]\\
\textbf{Consent for publication.} Not applicable to the de-identified database analysis; confirm according to the target journal.\\
\textbf{Data availability.} MIMIC-IV is available to credentialed users through PhysioNet. Derived data cannot be redistributed outside the applicable data-use agreement.\\
\textbf{Code availability.} The complete analytic pipeline is retained with the project. [AUTHOR ACTION REQUIRED: insert the public repository and release identifier before submission.]\\
\textbf{Funding.} [AUTHOR ACTION REQUIRED.]\\
\textbf{Competing interests.} [AUTHOR ACTION REQUIRED.]\\
\textbf{Author contributions.} [AUTHOR ACTION REQUIRED: provide a CRediT statement.]\\
\textbf{Acknowledgements.} [AUTHOR ACTION REQUIRED.]

\bibliographystyle{unsrtnat}
\bibliography{references}

\clearpage
\input{tables/table1.tex}
\clearpage
\input{tables/table2.tex}
\clearpage
\input{tables/table3.tex}

\clearpage
__FIGURES__
\end{document}
"""
    abstract_tex = "\n\n".join(tex_escape(p) for p in ABSTRACT_EN.split("\n\n"))
    figures_tex = []
    for name, caption in FIGURES:
        latex_caption = caption.split(". ", 1)[1] if ". " in caption else caption
        figures_tex.append(r"\begin{figure}[p]\centering" + "\n" +
                           r"\includegraphics[width=\textwidth]{figures/" + name + ".pdf}" + "\n" +
                           r"\caption{" + tex_escape(latex_caption) + "}" + "\n" + r"\end{figure}\clearpage")
    replacements = {
        "__TITLE__": tex_escape(TITLE_EN), "__RUNNING__": tex_escape(RUNNING_EN), "__ABSTRACT__": abstract_tex,
        "__INTRO__": intro, "__METHODS__": methods, "__RESULTS__": results,
        "__DISCUSSION__": discussion, "__CONCLUSION__": CONCLUSION_EN,
        "__FIGURES__": "\n".join(figures_tex),
    }
    for key, value in replacements.items():
        main_tex = main_tex.replace(key, value)
    (LATEX_DIR / "main.tex").write_text(main_tex, encoding="utf-8")
    (LATEX_DIR / "references.bib").write_text(BIBTEX, encoding="utf-8")

    supp_figures = [
        ("Figure_S1_precision_recall", "Precision-recall curves for all model families at 0 h, 6 h and 24 h."),
        ("Figure_S2_no_creatinine_sensitivity", "No-creatinine sensitivity analysis at 24 h."),
        ("Figure_S3_preindex_baseline_sensitivity", "Pre-index baseline-only sensitivity analysis at 24 h."),
        ("Figure_S4_subgroup_performance", "Subgroup discrimination for selected landmark models."),
        ("Figure_S5_predictor_missingness", "High-missingness predictors across dynamic datasets."),
    ]
    supp_parts = [r"\documentclass[10pt]{article}", r"\usepackage[letterpaper,margin=0.75in]{geometry}",
                  r"\usepackage[T1]{fontenc}", r"\usepackage{lmodern}", r"\usepackage{graphicx}",
                  r"\usepackage{booktabs}", r"\usepackage{longtable}", r"\usepackage{pdflscape}",
                  r"\usepackage{caption}", r"\usepackage[colorlinks=true,allcolors=blue]{hyperref}",
                  r"\begin{document}", r"\begin{center}\Large\textbf{Supplementary material}\\[6pt]" + tex_escape(TITLE_EN) + r"\end{center}",
                  r"\section*{Supplementary tables}", r"\input{tables/tableS1.tex}", r"\clearpage",
                  r"\input{tables/tableS2.tex}", r"\clearpage", r"\input{tables/tableS3.tex}", r"\clearpage",
                  r"\input{tables/tableS4.tex}",
                  r"\clearpage\section*{Supplementary figures}"]
    for name, caption in supp_figures:
        supp_parts += [r"\begin{figure}[p]\centering", r"\includegraphics[width=\textwidth]{figures/" + name + ".pdf}",
                       r"\caption{" + tex_escape(caption) + "}", r"\end{figure}\clearpage"]
    supp_parts.append(r"\end{document}")
    (LATEX_DIR / "supplement.tex").write_text("\n".join(supp_parts), encoding="utf-8")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader"); node.set(qn("w:val"), "true"); tr_pr.append(node)


def set_table_geometry(table, widths: list[int], indent=120):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout"); tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW"); tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total)); tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd"); tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent)); tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW"); tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width)); tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, latin="Times New Roman", east_asia="SimSun", size=11, bold=None, color=None, italic=None):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color is not None: run.font.color.rgb = RGBColor(*color)


def configure_document(doc: Document, running_title: str):
    section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492); section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"; normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0); normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    for name, size, before, after, color in [
        ("Heading 1", 16, 18, 10, (46,116,181)), ("Heading 2", 13, 12, 6, (46,116,181)),
        ("Heading 3", 12, 8, 4, (31,77,120))]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"; style.font.size = Pt(size); style.font.bold = True
        style.font.color.rgb = RGBColor(*color); style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.add_run(running_title), east_asia="Microsoft YaHei", size=8.5, color=(90,90,90))
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("第 "); set_run_font(run, east_asia="SimSun", size=9, color=(100,100,100))
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); footer._p.append(fld)
    run = footer.add_run(" 页"); set_run_font(run, east_asia="SimSun", size=9, color=(100,100,100))


def add_cover(doc: Document, title: str, subtitle: str):
    doc.add_paragraph().paragraph_format.space_after = Pt(72)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("研究论文稿件"); set_run_font(r, east_asia="Microsoft YaHei", size=11, bold=True, color=(122,90,0))
    p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); set_run_font(r, east_asia="Microsoft YaHei", size=22, bold=True, color=(32,55,72))
    p.paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle); set_run_font(r, east_asia="SimSun", size=11, italic=True, color=(80,80,80))
    p.paragraph_format.space_after = Pt(48)
    for text in ["作者：[请填写]", "单位：[请填写]", "通讯作者：[请填写]", "稿件版本：v7（内部验证稿）"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(text), east_asia="SimSun", size=10.5, color=(70,70,70))
        p.paragraph_format.space_after = Pt(5)
    doc.add_page_break()


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_paragraphs(doc, paragraphs: Iterable[str]):
    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.widow_control = True


def add_abstract(doc: Document):
    add_heading(doc, "摘要", 1)
    for block in ABSTRACT_ZH.split("\n\n"):
        label, body = block.split("：", 1)
        p = doc.add_paragraph()
        r = p.add_run(label + "："); set_run_font(r, east_asia="Microsoft YaHei", bold=True)
        r = p.add_run(body); set_run_font(r, east_asia="SimSun")
    p = doc.add_paragraph()
    r = p.add_run("关键词："); set_run_font(r, east_asia="Microsoft YaHei", bold=True)
    r = p.add_run("急性肾损伤；术后；重症监护；动态预测；机器学习；MIMIC-IV"); set_run_font(r, east_asia="SimSun")


def add_section_dict(doc, heading: str, content: dict[str, list[str]]):
    add_heading(doc, heading, 1)
    for subtitle, paragraphs in content.items():
        add_heading(doc, subtitle, 2); add_paragraphs(doc, paragraphs)


CHAR_ZH = {
    "Age, years":"年龄，岁", "Female sex":"女性", "Race: White":"种族：白人", "Race: Black":"种族：黑人",
    "Race: Asian":"种族：亚裔", "Race: Hispanic/Latino":"种族：西班牙裔/拉丁裔", "Race: Other/unknown":"种族：其他/未知",
    "Admission: Elective/same-day surgical":"入院：择期/当日手术", "Admission: Urgent/emergency":"入院：急诊/紧急",
    "Admission: Observation/other":"入院：观察/其他", "Baseline serum creatinine, mg/dL":"基线血清肌酐，mg/dL",
    "Pre-index 7-day creatinine baseline":"ICU 前 7 天基线肌酐", "Charlson comorbidity score":"Charlson 合并症评分",
    "Congestive heart failure":"充血性心力衰竭", "Hypertension":"高血压", "Diabetes mellitus":"糖尿病",
    "Chronic kidney disease":"慢性肾脏病", "Chronic pulmonary disease":"慢性肺疾病", "Liver disease":"肝病",
    "Cancer":"恶性肿瘤", "Peripheral vascular disease":"外周血管疾病", "Stroke":"卒中", "Myocardial infarction":"心肌梗死",
    "Obesity":"肥胖", "Anaemia":"贫血", "Cardiac surgery":"心脏手术", "Non-cardiac surgery":"非心脏手术",
    "Vascular surgery":"血管手术", "General/GI/hepatobiliary surgery":"普外/胃肠肝胆手术",
    "Major orthopaedic surgery":"重大骨科手术", "Neurosurgery":"神经外科手术", "Thoracic/respiratory surgery":"胸外/呼吸系统手术",
    "First ICU: CVICU":"首次 ICU：CVICU", "First ICU: SICU":"首次 ICU：SICU", "First ICU: TSICU":"首次 ICU：TSICU",
    "First ICU: Other surgical ICU":"首次 ICU：其他外科 ICU"
}


def add_table(doc: Document, rows: list[dict[str, str]], columns: list[tuple[str, str]], widths: list[int],
              font_size=7.5, translate_first=False):
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for idx, (_, title) in enumerate(columns):
        cell = table.rows[0].cells[idx]; cell.text = title; cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for run in cell.paragraphs[0].runs:
            set_run_font(run, east_asia="Microsoft YaHei", size=font_size, bold=True, color=(255,255,255))
        shade = OxmlElement("w:shd"); shade.set(qn("w:fill"), "2E74B5"); cell._tc.get_or_add_tcPr().append(shade)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, (key, _) in enumerate(columns):
            value = row.get(key, "")
            if translate_first and idx == 0: value = CHAR_ZH.get(value, value)
            cells[idx].text = str(value); cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_before = Pt(0); paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs: set_run_font(run, east_asia="SimSun", size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_landscape_section(doc: Document):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11); section.page_height = Inches(8.5)
    section.left_margin = Inches(0.6); section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.65); section.bottom_margin = Inches(0.65)
    section.header_distance = Inches(0.3); section.footer_distance = Inches(0.3)
    return section


def add_portrait_section(doc: Document):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    section.left_margin = Inches(1); section.right_margin = Inches(1)
    section.top_margin = Inches(1); section.bottom_margin = Inches(1)
    section.header_distance = Inches(0.492); section.footer_distance = Inches(0.492)
    return section


def add_figure(doc: Document, filename: str, caption: str):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(V6 / "figures" / f"{filename}.png"), width=Inches(6.25))
    p.paragraph_format.keep_with_next = True; p.paragraph_format.space_after = Pt(4)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption); set_run_font(r, east_asia="SimSun", size=9.5)
    p.paragraph_format.space_after = Pt(6)


def add_references(doc: Document):
    add_heading(doc, "参考文献", 1)
    for i, reference in enumerate(REFERENCES, 1):
        p = doc.add_paragraph(style="List Number")
        p.text = reference
        p.paragraph_format.left_indent = Inches(0.375); p.paragraph_format.first_line_indent = Inches(-0.194)
        p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.208
        for run in p.runs: set_run_font(run, east_asia="SimSun", size=9.5)


def add_declarations(doc: Document):
    add_heading(doc, "声明", 1)
    declarations = [
        ("伦理审批与知情同意", "MIMIC-IV 为按原数据库治理要求发布的去标识化数据，本研究通过 PhysioNet 认证访问。［作者需填写：所在机构的伦理认定/豁免及编号（如适用）。］"),
        ("发表同意", "去标识化数据库分析通常不涉及个体发表同意；请按目标期刊要求确认。"),
        ("数据可用性", "MIMIC-IV 可由完成认证的研究者通过 PhysioNet 获取；衍生数据的共享须遵守数据使用协议。"),
        ("代码可用性", "完整分析流程已保存在项目目录。［作者需填写：投稿前公开代码仓库及版本号。］"),
        ("基金资助", "［作者需填写。］"), ("利益冲突", "［作者需填写。］"),
        ("作者贡献", "［作者需依据 CRediT 分类填写。］"), ("致谢", "［作者需填写。］")]
    for label, value in declarations:
        p = doc.add_paragraph();
        r = p.add_run(label + "："); set_run_font(r, east_asia="Microsoft YaHei", bold=True)
        r = p.add_run(value); set_run_font(r, east_asia="SimSun")


def build_word_main() -> Path:
    doc = Document(); configure_document(doc, RUNNING_ZH); add_cover(doc, TITLE_ZH, "中文完整稿件｜基于 MIMIC-IV 3.1 的内部验证研究")
    add_abstract(doc)
    add_heading(doc, "引言", 1); add_paragraphs(doc, INTRO_ZH)
    add_section_dict(doc, "方法", METHODS_ZH)
    add_section_dict(doc, "结果", RESULTS_ZH)
    add_heading(doc, "讨论", 1); add_paragraphs(doc, DISCUSSION_ZH)
    add_heading(doc, "结论", 1); add_paragraphs(doc, [CONCLUSION_ZH])
    add_declarations(doc); add_references(doc)

    add_landscape_section(doc); add_heading(doc, "表1. 严格新发 AKI 可评价队列的基线特征", 1)
    add_table(doc, read_csv("Table_1_baseline_characteristics.csv"), [
        ("Characteristic","特征"), ("Overall (N=10,877)","总体"), ("No AKI (N=6,346)","无 AKI"),
        ("Incident AKI (N=4,531)","新发 AKI"), ("Standardized mean difference","标准化差异"), ("Missing, n","缺失")],
        [3500,2400,2400,2400,1900,1400], font_size=7.5, translate_first=True)
    add_heading(doc, "表2. 所选地标模型在测试集中的性能", 1)
    add_table(doc, read_csv("Table_2_selected_model_performance.csv"), [
        ("Landmark","地标"), ("Selected model","模型"), ("Test N","测试 N"), ("Event rate","事件率"),
        ("AUROC (95% CI)","AUROC (95% CI)"), ("AUPRC (95% CI)","AUPRC (95% CI)"),
        ("Brier score","Brier"), ("Calibration intercept","校准截距"), ("Calibration slope","校准斜率"),
        ("Youden threshold","Youden 阈值"), ("Sensitivity / specificity","敏感度/特异度")],
        [700,1500,800,900,1650,1650,900,1050,1050,1050,1750], font_size=6.8)
    add_portrait_section(doc); add_heading(doc, "表3. 24 h 地标敏感性分析", 1)
    add_table(doc, read_csv("Table_3_sensitivity_analyses_24h.csv"), [
        ("Sensitivity analysis","敏感性分析"), ("Model","模型"), ("Full/reference AUROC","参照 AUROC"),
        ("Sensitivity AUROC","敏感性 AUROC"), ("ΔAUROC (95% paired CI)","ΔAUROC（配对 95% CI）")],
        [2300,1700,1600,1600,2160], font_size=8.0)
    add_heading(doc, "正文图", 1)
    for index, (name, caption) in enumerate(FIGURES_ZH):
        if index: doc.add_page_break()
        add_figure(doc, name, caption)
    path = WORD_DIR / "manuscript_zh.docx"; doc.save(path); return path


def build_word_supplement() -> Path:
    doc = Document(); configure_document(doc, RUNNING_ZH + "｜补充材料")
    add_cover(doc, "补充材料", TITLE_ZH)
    add_landscape_section(doc); add_heading(doc, "补充表 S1. 各地标四类模型完整性能", 1)
    add_table(doc, combine_s1(), [("time","地标"),("model","模型"),("n","测试 N"),("event","事件率"),
        ("auroc","AUROC (95% CI)"),("auprc","AUPRC (95% CI)"),("brier","Brier"),("cal","校准截距/斜率"),
        ("youden","Youden（敏感度/特异度）")], [700,1500,850,850,1900,1900,850,1500,1950], font_size=7.0)
    add_heading(doc, "补充表 S2. 预设亚组性能", 1)
    add_table(doc, combine_s2(), [("time","地标"),("model","模型"),("group","亚组维度"),("level","水平"),
        ("n","N/事件"),("auroc","AUROC (95% CI)"),("auprc","AUPRC (95% CI)"),("brier","Brier")],
        [700,1500,2000,1600,1000,2300,2300,900], font_size=6.8)
    add_heading(doc, "补充表 S3. 缺失率超过 40% 的预测变量", 1)
    add_table(doc, high_missingness(), [("landmark_hours","地标 h"),("predictor","预测变量"),("variable_type","类型"),
        ("n_missing","缺失 N"),("missing_percent","缺失 %"),("n_observed","观测 N")],
        [900,5200,1600,1400,1400,1400], font_size=7.2)
    p = doc.add_paragraph("注：所有 312 行预测变量缺失率明细随稿件包以 Table_S3_predictor_missingness.csv 提供。")
    for run in p.runs: set_run_font(run, east_asia="SimSun", size=9, italic=True, color=(90,90,90))
    add_heading(doc, "补充表 S4. 全部地标敏感性分析", 1)
    add_table(doc, combine_s4(), [("analysis","敏感性分析"),("time","地标"),("model","模型"),("n","测试 N"),
        ("reference","参照 AUROC"),("sensitivity","敏感性 AUROC"),("delta","ΔAUROC（配对 95% CI）")],
        [2200,900,1900,1000,1600,1700,3000], font_size=7.0)
    add_portrait_section(doc); add_heading(doc, "补充图", 1)
    for index, (name, caption) in enumerate(SUPP_FIGURES_ZH):
        if index: doc.add_page_break()
        add_figure(doc, name, caption)
    path = WORD_DIR / "supplement_zh.docx"; doc.save(path); return path


def write_readmes():
    (OUT / "README.md").write_text("""# Complete manuscript package v7

## English LaTeX
- `english_latex/main.tex` and compiled `manuscript_en.pdf`
- `english_latex/supplement.tex` and compiled `supplement_en.pdf`
- `english_latex/references.bib`, formatted table fragments and vector figures

## Chinese Word
- `chinese_word/manuscript_zh.docx`
- `chinese_word/supplement_zh.docx`

## Author action required
Author names, affiliations, corresponding author, ethics/local exemption wording, funding, conflicts of interest, CRediT contributions, acknowledgements and public code repository remain explicitly marked. A target-journal template has not been imposed.

## Boundaries
This is an internally validated, single-centre, serum-creatinine-only AKI prediction study. SHAP is predictive attribution, not causal or modifiable-factor evidence. External/temporal validation and clinical workflow evaluation remain necessary.
""", encoding="utf-8")
    (OUT / "author_action_required.md").write_text("""# Author action required before submission

1. Replace author, affiliation and correspondence placeholders.
2. Confirm local ethics/exemption wording and any approval identifier.
3. Add funding, conflicts of interest, CRediT contributions and acknowledgements.
4. Add the public code repository and archived release identifier.
5. Select the target journal and apply its title-page, abstract, reference and supplement format.
6. Confirm whether urine-output AKI can be added as a sensitivity analysis.
7. Do not describe SHAP features as causal or modifiable factors.
8. Do not describe the models as externally validated or ready for clinical deployment.
""", encoding="utf-8")


def main():
    if OUT.exists(): shutil.rmtree(OUT)
    for directory in [LATEX_DIR, WORD_DIR, DATA_DIR, FIG_DIR, QA_DIR]: directory.mkdir(parents=True, exist_ok=True)
    build_latex()
    main_docx = build_word_main(); supp_docx = build_word_supplement()
    write_readmes()
    print(f"Package: {OUT}")
    print(f"Word: {main_docx}, {supp_docx}")


if __name__ == "__main__":
    main()
