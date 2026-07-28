"""Independent scientific and package-consistency checks for v31."""

from __future__ import annotations

import csv
import hashlib
import re
from pathlib import Path

import pandas as pd
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "manuscript_package_v31_nested_selection"
V17 = ROOT / "outputs" / "modeling_v17_eicu_recalibration_heterogeneity"
V18 = ROOT / "outputs" / "modeling_v18_eicu_outcome_observability"
V21 = ROOT / "outputs" / "modeling_v21_eicu_selection_bias_sensitivity"
V27 = ROOT / "outputs" / "modeling_v27_severity_recovery"
V28 = ROOT / "outputs" / "modeling_v28_severe_temporal_external"
V29 = ROOT / "outputs" / "modeling_v29_multistate_competing_risk"
V30 = ROOT / "outputs" / "modeling_v30_severe_same_risk_incremental"
V31 = ROOT / "outputs" / "modeling_v31_nested_grouped_cv"


def docx_text(path: Path) -> str:
    doc = Document(path)
    text = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        text.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
    return "\n".join(text)


def check(condition: bool, message: str, rows: list[dict[str, str]]) -> None:
    if not condition:
        raise AssertionError(message)
    rows.append({"check": message, "status": "PASS"})


def main() -> None:
    results: list[dict[str, str]] = []
    md = (OUT / "secondary_manuscript_v31_en.md").read_text(encoding="utf-8")
    tex = (OUT / "latex" / "main.tex").read_text(encoding="utf-8")
    supp_tex = (OUT / "latex" / "supplement.tex").read_text(encoding="utf-8")
    word = docx_text(OUT / "secondary_manuscript_v31_en.docx")
    supplement_word = docx_text(OUT / "additional_file_1_secondary_supplement_en.docx")
    combined = "\n".join([md, tex, word, supplement_word])

    check("Caiyun Yuan" not in combined, "Package contains only the final author", results)
    check("Bizhi Wei" in combined and "15619056250wbz@gmail.com" in combined, "Author and correspondence are present", results)
    check("[AUTHOR ACTION REQUIRED]" not in combined, "No author-action placeholder remains", results)
    check("to be provided before publication" not in combined, "No archive-DOI placeholder remains", results)
    check("10.5281/zenodo.21640763" in combined, "Zenodo version DOI is cited", results)
    check("v1.0.4" in combined and "v1.0.3" not in combined and "v1.0.2" not in combined, "Repository release v1.0.4 is cited consistently", results)
    check("Supplementary Tables S1–S12" in word, "Word main file reports Tables S1–S12", results)
    check("Supplementary Figures S1–S7" in word, "Word main file reports Figures S1–S7", results)
    check("Supplementary Tables S1--S12" in tex, "LaTeX main file reports Tables S1–S12", results)
    for i in range(1, 13):
        check(rf"\input{{tables/tableS{i}.tex}}" in supp_tex, f"LaTeX supplement includes Table S{i}", results)

    summary = pd.read_csv(V29 / "audit_v29_multistate_summary.csv").set_index("measure")
    expected_counts = {
        "Analysis cohort": 10877,
        "Incident SCr-AKI": 4531,
        "Trajectory-eligible incident SCr-AKI": 4519,
        "Locked AKI onset after recorded hospital disposition": 12,
        "Severe SCr-AKI stage 2/3": 679,
        "Observed recovery after AKI": 3936,
        "Recurrent AKI after observed recovery": 641,
    }
    for label, expected in expected_counts.items():
        check(int(summary.loc[label, "n"]) == expected, f"Source count matches: {label}={expected:,}", results)

    table1 = pd.read_csv(OUT / "tables" / "Table_1_baseline_characteristics_by_severe_aki.csv")
    check("No severe SCr-AKI (N=10,198)" in table1, "Table 1 has the nonsevere denominator", results)
    check("Severe SCr-AKI (N=679)" in table1, "Table 1 has the severe denominator", results)
    check(len(table1) >= 30, "Table 1 contains a complete baseline profile", results)

    table2 = pd.read_csv(OUT / "tables" / "Table_2_selected_severe_aki_models.csv")
    check("Reported model" in table2, "Table 2 labels frozen models as reported rather than audit-selected", results)
    check(len(table2) == 3 and set(table2["Landmark"]) == {"0 h", "6 h", "24 h"}, "Table 2 contains three severe-AKI landmark models", results)

    perf = pd.read_csv(V27 / "model_v27_performance_summary.csv")
    for lm, model, expected in [(0, "XGBoost", 0.6977275179), (6, "XGBoost", 0.7896103896), (24, "Logistic Regression", 0.8394132653)]:
        value = perf[(perf.task.eq("severe_scr")) & perf.landmark.astype(str).eq(str(lm)) & perf.model.eq(model)].iloc[0].auroc
        check(abs(value - expected) < 1e-10, f"Internal {lm}-h reported-model AUROC matches source", results)

    nested = pd.read_csv(V31 / "audit_v31_training_only_model_selection_summary.csv")
    check(set(nested["landmark_hours"].astype(int)) == {0, 6, 24}, "Nested-CV summary covers all landmarks", results)
    check(nested["locked_test_rows_used_for_selection"].eq(0).all(), "Nested selection used zero locked-test rows", results)
    check(nested["subject_overlap_any_fold"].eq(0).all(), "Nested folds have zero subject overlap", results)
    check((nested["nested_outer_fold_n"].eq(5) & nested["nested_inner_fold_n"].eq(4)).all(), "Nested audit used 5 outer and 4 inner folds", results)
    expected_lr = {0: 0.7264143485, 6: 0.7954765514, 24: 0.8462628119}
    expected_xgb = {0: 0.7514069169, 6: 0.8085574839, 24: 0.8476092434}
    for row in nested.itertuples(index=False):
        lm = int(row.landmark_hours)
        check(abs(row.training_oof_lr_auroc - expected_lr[lm]) < 1e-9, f"{lm}-h LR training OOF AUROC matches source", results)
        check(abs(row.training_oof_xgb_auroc - expected_xgb[lm]) < 1e-9, f"{lm}-h XGBoost training OOF AUROC matches source", results)
    r24 = nested.loc[nested["landmark_hours"].eq(24)].iloc[0]
    check(r24.training_only_recommended_model == "XGBoost" and r24.previously_reported_model == "Logistic Regression", "24-h selection discordance is retained", results)
    check("At 24 h, the training OOF AUROC difference was only 0.001" in md, "24-h model-selection instability is explained", results)
    s10 = pd.read_csv(OUT / "tables" / "Table_S10_training_only_nested_grouped_cv.csv")
    check(len(s10) == 3 and s10.loc[s10["Landmark"].eq("24 h"), "Concordant"].iloc[0] == "No", "Table S10 records 24-h discordance", results)

    deltas = pd.read_csv(V30 / "model_v30_severe_paired_incremental_deltas.csv")
    d6 = deltas[deltas.primary_model_for_risk_set.astype(bool) & deltas.risk_set_hours.eq(6) & deltas.new_information_hours.eq(6) & deltas.reference_information_hours.eq(0)].iloc[0]
    d24 = deltas[deltas.primary_model_for_risk_set.astype(bool) & deltas.risk_set_hours.eq(24) & deltas.new_information_hours.eq(24) & deltas.reference_information_hours.eq(6)].iloc[0]
    check(round(d6.delta_auroc, 3) == 0.101 and round(d24.delta_auroc, 3) == 0.079, "Same-risk paired AUROC deltas remain unchanged", results)

    ext = pd.read_csv(V28 / "model_v28_eicu_frozen_severe_performance.csv")
    for lm, expected in [(0, 0.7070313499), (6, 0.7613177886), (24, 0.7839542750)]:
        value = ext[(ext.landmark_hours.eq(lm)) & ext.target.eq("SCr stage 2/3")].iloc[0].auroc
        check(abs(value - expected) < 1e-10, f"External severe-AKI {lm}-h AUROC matches source", results)
    check("14,229 (46.9%)" in combined and "30,365" in combined and "16,136" in combined, "External observability denominator is explicit", results)

    s11 = pd.read_csv(OUT / "tables" / "Table_S11_eicu_outcome_observability.csv")
    comparison = pd.read_csv(V18 / "audit_v18_evaluable_vs_unevaluable_comparison.csv")
    comparison = comparison.loc[comparison["variable"].ne("ineligibility_reason")]
    check(len(s11) == len(comparison) + 5, "Table S11 contains all clinically interpretable group comparisons plus denominator/model rows", results)
    check({"Analysis evaluability", "Creatinine-record observability"} <= set(s11["Variable/model"]), "Table S11 documents both observability definitions", results)
    check("ineligibility_reason" not in set(s11["Variable/model"]), "Table S11 excludes mechanical eligibility-reason SMDs", results)
    check({"Age, years", "APACHE IVa score", "Operative system", "ICU type"} <= set(s11["Variable/model"]), "Table S11 uses standardized clinical labels", results)
    obsperf = pd.read_csv(V18 / "model_v18_outcome_observability_performance.csv")
    check(round(obsperf.loc[obsperf.model.eq("clinical_only"), "auroc"].iloc[0], 3) == 0.671, "Clinical observability AUROC matches source", results)
    check(round(obsperf.loc[obsperf.model.eq("clinical_plus_hospital_identifier"), "auroc"].iloc[0], 3) == 0.740, "Hospital-augmented observability AUROC matches source", results)

    s12 = pd.read_csv(OUT / "tables" / "Table_S12_eicu_ipw_strict_baseline_sensitivity.csv")
    check(len(s12) == 21, "Table S12 contains 6 baseline, 9 IPW, and 6 pattern-mixture rows", results)
    check(set(s12["Section"]) == {"Strict baseline", "IPW selection sensitivity", "Pattern-mixture"}, "Table S12 contains all requested sensitivity sections", results)
    ipw = pd.read_csv(V21 / "analysis_v21_ipw_weighted_external_performance.csv")
    for lm, expected in [(0, 0.6932977379), (6, 0.7196549307), (24, 0.7059387728)]:
        value = ipw[(ipw.landmark_hours.eq(lm)) & ipw.method.eq("ipw_clinical")].iloc[0].auroc
        check(abs(value - expected) < 1e-10, f"Clinical-IPW {lm}-h AUROC matches source", results)
    strict = pd.read_csv(V17 / "analysis_v17_strict_pre_icu_baseline_sensitivity.csv")
    for lm, expected in [(0, 0.6728832200), (6, 0.6961448188), (24, 0.6866278421)]:
        value = strict[(strict.landmark_hours.eq(lm)) & strict.sensitivity_cohort.eq("strict_pre_icu_baseline_only")].iloc[0].auroc
        check(abs(value - expected) < 1e-10, f"Strict-baseline {lm}-h AUROC matches source", results)
    check("These results concern the any-stage incident SCr-AKI endpoint" in supplement_word, "Supplement explicitly separates any-stage eICU sensitivities", results)
    check("do not eliminate selection bias" in md, "Selection-bias interpretation remains cautious", results)

    mapping = pd.read_csv(OUT / "tables" / "Table_S9_feature_harmonization.csv")
    check(len(mapping) == 8, "Feature harmonization table contains eight operational groups", results)
    cif_aki = pd.read_csv(V29 / "analysis_v29_cif_after_aki.csv")
    recovery = cif_aki[(cif_aki.group_variable.eq("Overall")) & cif_aki.time_hours.eq(48) & cif_aki.cause.eq("Observed recovery")].iloc[0].cif_percent
    check(round(recovery, 1) == 65.0, "Trajectory 48-h recovery CIF remains unchanged", results)

    refs = pd.read_csv(OUT / "reference_audit.csv")
    check(len(refs) == 35 and refs["Citation key"].is_unique, "Reference audit contains 35 unique entries", results)
    cite_groups = re.findall(r"\\cite\{([^}]+)\}", tex)
    cited = {key for group in cite_groups for key in group.split(",")}
    check(cited <= set(refs["Citation key"]), "Every LaTeX citation key exists", results)

    for i in range(1, 5):
        for extn in ["png", "pdf", "svg"]:
            check((OUT / "figures" / f"Fig{i}.{extn}").exists(), f"Figure {i} {extn} exists", results)
    for i in range(1, 8):
        for extn in ["png", "pdf", "svg"]:
            check((OUT / "figures" / f"FigS{i}.{extn}").exists(), f"Figure S{i} {extn} exists", results)
    for i in range(1, 13):
        check(any((OUT / "tables").glob(f"Table_S{i}_*.csv")), f"Table S{i} CSV exists", results)

    deliverables = [
        "secondary_manuscript_v31_en.docx",
        "secondary_manuscript_v31_en.pdf",
        "additional_file_1_secondary_supplement_en.docx",
        "additional_file_1_secondary_supplement_en.pdf",
        "additional_file_2_tripod_ai_checklist.docx",
        "additional_file_2_tripod_ai_checklist.pdf",
        "additional_file_2_tripod_ai_checklist.csv",
        "code_and_data_availability_statement.docx",
        "code_and_data_availability_statement.pdf",
        "latex/main.pdf",
        "latex/supplement.pdf",
    ]
    for name in deliverables:
        path = OUT / name
        check(path.exists() and path.stat().st_size > 0, f"Deliverable exists: {name}", results)

    manifest = pd.read_csv(OUT / "file_manifest_sha256.csv")
    check(manifest.relative_path.is_unique, "Manifest paths are unique", results)
    sampled = manifest.loc[
        manifest.relative_path.isin(
            [
                "secondary_manuscript_v31_en.docx",
                "figures/Fig1.pdf",
                "tables/Table_S10_training_only_nested_grouped_cv.csv",
                "tables/Table_S11_eicu_outcome_observability.csv",
                "tables/Table_S12_eicu_ipw_strict_baseline_sensitivity.csv",
            ]
        )
    ]
    check(len(sampled) == 5, "Manifest contains all sampled v31 files", results)
    for row in sampled.itertuples():
        digest = hashlib.sha256((OUT / row.relative_path).read_bytes()).hexdigest()
        check(digest == row.sha256, f"Manifest checksum matches: {row.relative_path}", results)

    QA = OUT / "qa"
    QA.mkdir(exist_ok=True)
    with (QA / "secondary_manuscript_validation_checks.csv").open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["check", "status"])
        writer.writeheader()
        writer.writerows(results)
    (QA / "secondary_manuscript_validation_report.md").write_text(
        "# v31 secondary manuscript validation\n\n"
        f"All {len(results)} automated checks passed.\n\n"
        "The audit confirms that nested model-family selection used only the original training partition, with zero subject overlap and zero locked-test rows. The eICU observability, IPW, pattern-mixture, and strict-baseline sensitivities are explicitly labeled as any-stage SCr-AKI analyses. Word and LaTeX render checks are recorded separately.\n",
        encoding="utf-8",
    )
    print(f"PASS: {len(results)} validation checks")


if __name__ == "__main__":
    main()
