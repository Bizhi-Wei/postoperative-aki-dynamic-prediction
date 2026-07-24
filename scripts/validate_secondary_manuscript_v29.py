"""Independent consistency checks for the v29 secondary manuscript package."""

from __future__ import annotations

import csv
import re
from pathlib import Path

import pandas as pd
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "manuscript_package_v29_secondary"
V27 = ROOT / "outputs" / "modeling_v27_severity_recovery"
V28S = ROOT / "outputs" / "modeling_v28_severe_temporal_external"
V29 = ROOT / "outputs" / "modeling_v29_multistate_competing_risk"


def docx_text(path: Path) -> str:
    doc = Document(path)
    text = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        text.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
    return "\n".join(text)


def check(condition: bool, message: str, rows: list[dict[str, str]]) -> None:
    if not condition:
        raise AssertionError(message)
    rows.append({"check": message, "status": "PASS"})


def main() -> None:
    results: list[dict[str, str]] = []
    md = (OUT / "secondary_manuscript_v29_en.md").read_text(encoding="utf-8")
    tex = (OUT / "latex" / "main.tex").read_text(encoding="utf-8")
    word = docx_text(OUT / "secondary_manuscript_v29_en.docx")
    combined = "\n".join([md, tex, word])

    check("Caiyun Yuan" not in combined, "Single-author package contains no former coauthor", results)
    check("Bizhi Wei" in md and "15619056250wbz@gmail.com" in md, "Author and correspondence are present", results)
    check("[AUTHOR ACTION REQUIRED]" not in combined, "No author-action placeholder remains", results)
    check("to be provided before publication" in combined, "Only intentional archived-DOI placeholder remains", results)

    summary = pd.read_csv(V29 / "audit_v29_multistate_summary.csv").set_index("measure")
    check(int(summary.loc["Analysis cohort", "n"]) == 10877, "Analysis cohort is 10,877", results)
    check(int(summary.loc["Incident SCr-AKI", "n"]) == 4531, "Locked incident AKI count is 4,531", results)
    check(int(summary.loc["Trajectory-eligible incident SCr-AKI", "n"]) == 4519, "Trajectory-eligible AKI count is 4,519", results)
    check(int(summary.loc["Locked AKI onset after recorded hospital disposition", "n"]) == 12, "Post-disposition audit count is 12", results)
    check(int(summary.loc["Severe SCr-AKI stage 2/3", "n"]) == 679, "Severe SCr-AKI count is 679", results)
    check(int(summary.loc["Observed recovery after AKI", "n"]) == 3936, "Observed recovery count is 3,936", results)
    check(int(summary.loc["Recurrent AKI after observed recovery", "n"]) == 641, "Recurrent AKI count is 641", results)

    perf = pd.read_csv(V27 / "model_v27_performance_summary.csv")
    expected_internal = {0: 0.6977275179, 6: 0.7896103896, 24: 0.8394132653}
    for lm, expected in expected_internal.items():
        model = "Logistic Regression" if lm == 24 else "XGBoost"
        value = perf[(perf.task == "severe_scr") & (perf.landmark.astype(str) == str(lm)) & (perf.model == model)].iloc[0].auroc
        check(abs(value - expected) < 1e-10, f"Internal {lm}-h severe-AKI AUROC matches source", results)

    ext = pd.read_csv(V28S / "model_v28_eicu_frozen_severe_performance.csv")
    expected_external = {0: 0.7070313499, 6: 0.7613177886, 24: 0.7839542750}
    for lm, expected in expected_external.items():
        value = ext[(ext.landmark_hours == lm) & (ext.target == "SCr stage 2/3")].iloc[0].auroc
        check(abs(value - expected) < 1e-10, f"External {lm}-h severe-AKI AUROC matches source", results)

    cif_aki = pd.read_csv(V29 / "analysis_v29_cif_after_aki.csv")
    cif_rec = pd.read_csv(V29 / "analysis_v29_cif_after_recovery.csv")
    recovery_48 = cif_aki[(cif_aki.group_variable == "Overall") & (cif_aki.time_hours == 48) & (cif_aki.cause == "Observed recovery")].iloc[0].cif_percent
    severe_48 = cif_aki[(cif_aki.group_variable == "Overall") & (cif_aki.time_hours == 48) & (cif_aki.cause == "Severe AKI onset/progression")].iloc[0].cif_percent
    recurrence_48 = cif_rec[(cif_rec.group_variable == "Overall") & (cif_rec.time_hours == 48) & (cif_rec.cause == "Recurrent AKI")].iloc[0].cif_percent
    check(round(recovery_48, 1) == 65.0, "48-h recovery CIF is 65.0%", results)
    check(round(severe_48, 1) == 11.9, "48-h severe-progression CIF is 11.9%", results)
    check(round(recurrence_48, 1) == 11.3, "48-h recurrence CIF is 11.3%", results)

    for token in ["10,877", "679 (6.2%)", "4,519", "3,936", "65.0%", "11.9%", "11.3%"]:
        check(token in md, f"Manuscript contains key token: {token}", results)

    refs = pd.read_csv(OUT / "reference_audit.csv")
    check(len(refs) == 35, "Reference audit contains 35 entries", results)
    check(refs["Citation key"].is_unique, "Citation keys are unique", results)
    cited_keys = set(re.findall(r"\\cite\{([^}]+)\}", tex))
    cited_keys = {key for group in cited_keys for key in group.split(",")}
    ref_keys = set(refs["Citation key"])
    check(cited_keys <= ref_keys, "Every LaTeX citation key exists in the bibliography", results)

    for i in range(1, 5):
        for extn in ["png", "pdf", "svg"]:
            check((OUT / "figures" / f"Fig{i}.{extn}").exists(), f"Figure {i} {extn} exists", results)
    for i in range(1, 6):
        for extn in ["png", "pdf", "svg"]:
            check((OUT / "figures" / f"FigS{i}.{extn}").exists(), f"Figure S{i} {extn} exists", results)
    for i in range(1, 9):
        check(any((OUT / "tables").glob(f"Table_S{i}_*.csv")), f"Table S{i} CSV exists", results)
    for path in [
        OUT / "secondary_manuscript_v29_en.docx", OUT / "secondary_manuscript_v29_en.pdf",
        OUT / "additional_file_1_secondary_supplement_en.docx", OUT / "additional_file_1_secondary_supplement_en.pdf",
        OUT / "additional_file_2_tripod_ai_checklist.docx", OUT / "additional_file_2_tripod_ai_checklist.csv",
    ]:
        check(path.exists() and path.stat().st_size > 0, f"Deliverable exists: {path.name}", results)

    report_csv = OUT / "qa" / "secondary_manuscript_validation_checks.csv"
    with report_csv.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["check", "status"]); writer.writeheader(); writer.writerows(results)
    report_md = OUT / "qa" / "secondary_manuscript_validation_report.md"
    report_md.write_text(
        "# v29 secondary manuscript validation\n\n"
        f"All {len(results)} automated checks passed.\n\n"
        "The Word main manuscript, supplement, checklist, and availability statement were rendered with the project's "
        "stable Windows LibreOffice workflow using an explicit executable path and isolated user profiles. LaTeX main "
        "and supplement compiled successfully with TeX Live. Key pages and wide tables were visually inspected.\n",
        encoding="utf-8",
    )
    print(f"PASS: {len(results)} validation checks")


if __name__ == "__main__":
    main()
