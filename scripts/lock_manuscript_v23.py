"""Create the locked final manuscript package after external validation (v23).

The previous v8 package is intentionally preserved.  This script builds a new,
versioned package that integrates the completed temporal/external validation,
competing-risk, and selection-bias analyses without overstating clinical
readiness.
"""

from __future__ import annotations

import csv
import importlib.util
import shutil
from pathlib import Path

import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "manuscript_package_v23_locked"
V14 = ROOT / "outputs" / "modeling_v14_final_sensitivities"
V16 = ROOT / "outputs" / "modeling_v16_eicu_external_validation"
V17 = ROOT / "outputs" / "modeling_v17_eicu_recalibration_heterogeneity"
V21 = ROOT / "outputs" / "modeling_v21_eicu_selection_bias_sensitivity"

spec = importlib.util.spec_from_file_location("v8", ROOT / "scripts" / "build_manuscript_package_v8.py")
v8 = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(v8)


def configure_paths() -> None:
    v8.ROOT = ROOT
    v8.V6 = ROOT / "outputs" / "final_manuscript_v6"
    v8.V7 = ROOT / "outputs" / "manuscript_package_v7"
    v8.OUT = OUT
    v8.LATEX = OUT / "latex"
    v8.FIGURES = OUT / "figures"
    v8.TABLES = OUT / "tables"
    v8.QA = OUT / "qa"
    v8.ASSETS = ROOT / "scripts" / "assets"
    v8.GRAPHICAL = ROOT / "outputs" / "_v8_figure_stage"


def configure_text() -> None:
    v8.TITLE = (
        "Development and validation of dynamic models for incident postoperative acute kidney injury "
        "in surgical intensive care: a retrospective MIMIC-IV and eICU cohort study"
    )
    v8.RUNNING_TITLE = "Dynamic prediction of postoperative AKI"
    v8.ABSTRACT = {
        "Background": (
            "Postoperative acute kidney injury (AKI) is common after major surgery, but risk evolves during early "
            "critical illness. We developed and validated landmark-specific models for incident AKI in a rigorously "
            "defined surgical intensive care cohort."
        ),
        "Methods": (
            "We used MIMIC-IV version 3.1 for retrospective development. The analysis unit was the first intensive "
            "care unit (ICU) stay per hospital admission. Incident AKI within 7 days of ICU admission was recomputed "
            "using Kidney Disease: Improving Global Outcomes serum creatinine criteria. Separate risk sets were built "
            "at ICU admission, 6 h, and 24 h. Logistic regression, random forest, XGBoost, and LightGBM were evaluated "
            "using an 80:20 patient-grouped split. Temporal validation, eICU feature-harmonized external validation, "
            "hospital-heldout recalibration, competing-risk analyses, and selection-bias sensitivity analyses were performed."
        ),
        "Results": (
            "Among 11,943 qualifying admissions, 10,877 were evaluable for incident AKI and 4,531 developed AKI "
            "(41.7%). Development-benchmark AUROCs were 0.728 at 0 h and 0.740 at 6 h for XGBoost and 0.754 at "
            "24 h for logistic regression. Selected parsimonious models retained AUROCs of 0.726, 0.736, and 0.756, "
            "respectively. In eICU, 14,229 of 30,365 surgical ICU stays were outcome-evaluable; external AUROCs were "
            "0.673, 0.696, and 0.687, with calibration slopes of 0.707, 0.846, and 0.351. Hospital-based logistic "
            "recalibration improved heldout-hospital calibration."
        ),
        "Conclusions": (
            "Dynamic peri-ICU data provided moderate discrimination for incident postoperative AKI. External validation "
            "supported transportable ranking but showed calibration heterogeneity and selective outcome observability. "
            "Local recalibration and prospective silent validation are required before any clinical action threshold is used."
        ),
    }
    v8.INTRODUCTION[-1] = (
        "We therefore developed dynamic models for incident AKI at ICU admission, 6 h, and 24 h in a strict therapeutic "
        "surgical cohort and evaluated calibration, clinical net benefit, attribution, subgroup performance, creatinine and "
        "baseline-creatinine sensitivities, temporal validation, and external validation in eICU [[CITE:Pollard2018]]."
    )
    v8.METHODS["Study design, setting, and reporting"] = [
        "We performed a retrospective prediction-model development study in MIMIC-IV version 3.1, a deidentified electronic "
        "health record database from Beth Israel Deaconess Medical Center in Boston, Massachusetts, United States "
        "[[CITE:Johnson2023,MIMIC2024]]. We then undertook temporal validation and external validation in the eICU "
        "Collaborative Research Database, a multicenter critical care resource [[CITE:Pollard2018]]. The analysis unit was "
        "the first qualifying ICU stay within each hospital admission. Reporting was guided by TRIPOD+AI, and design decisions "
        "were considered in relation to PROBAST [[CITE:Collins2024,Wolff2019]]."
    ]
    v8.METHODS["External validation, recalibration, and selection sensitivity"] = [
        "For eICU validation, we constructed a strict surgical first-ICU-stay cohort using prespecified operative organ-system "
        "diagnoses. Because EHR variables were not identical across databases, externally evaluated models used a feature-harmonized "
        "portable predictor set rather than the full MIMIC-IV models. Baseline creatinine, incident SCr-defined AKI, and landmark "
        "risk sets were rederived using the same temporal logic. Discrimination, calibration, Brier score, and patient-level "
        "bootstrap confidence intervals were calculated separately at 0 h, 6 h, and 24 h.",
        "To assess transportability of probability estimates, hospitals were split into calibration and heldout sets; intercept-only "
        "and logistic recalibration updates learned in calibration hospitals were applied unchanged to heldout hospitals. A competing-risk "
        "sensitivity analysis required patients to remain in ICU at the landmark and counted AKI only before ICU-unit exit, with death "
        "and discharge as competing events. Decision-curve results were interpreted as model utility metrics, not as evidence of clinical benefit.",
        "We audited outcome observability and used five-fold subject-grouped cross-fitting to estimate creatinine-record observability. "
        "Stabilized inverse-probability weights were truncated at the first and 99th percentiles among analysis-evaluable stays. These "
        "analyses assume missing at random conditional on observed covariates. A pattern-mixture analysis varied the unobserved-to-observed "
        "AKI risk ratio from 0.5 to 2.0; it did not impute individual outcomes or identify the true missing-data mechanism."
    ]
    v8.METHODS["Model finalization and deployment boundary"] = [
        "After evaluation, we prespecified a parsimonious translation candidate: XGBoost with 36 predictors at 0 h, XGBoost with 72 "
        "predictors at 6 h, and logistic regression with 72 predictors at 24 h. These specifications were refit on all eligible MIMIC-IV "
        "development data solely to create versioned research-use scoring artifacts. The refit did not replace validation estimates. The "
        "artifacts return a continuous probability and deliberately contain no universal action threshold or automated treatment recommendation."
    ]
    v8.RESULTS["Temporal validation and selected parsimonious models"] = [
        "Rolling temporal validation showed stable discrimination: AUROCs ranged from 0.696 to 0.712 at 0 h, 0.711 to 0.738 at 6 h, "
        "and 0.736 to 0.757 at 24 h across later-year validation blocks (Additional file 1: Table S4). The selected parsimonious "
        "models used 36 predictors at 0 h and 72 predictors at both 6 h and 24 h. Their heldout AUROCs were 0.726 (95% CI, 0.707-0.748), "
        "0.736 (0.713-0.757), and 0.756 (0.732-0.780), respectively (Table 2)."
    ]
    v8.RESULTS["External validation, calibration, and observability"] = [
        "The eICU strict surgical cohort comprised 30,365 first ICU stays from 197 hospitals; 14,229 (46.9%) were evaluable for "
        "incident AKI. The feature-harmonized portable models yielded AUROCs of 0.673 (95% CI, 0.664-0.683), 0.696 (0.686-0.707), "
        "and 0.687 (0.674-0.700) at 0 h, 6 h, and 24 h, respectively (Table S5). Calibration slopes were 0.707, 0.846, and 0.351, "
        "indicating marked calibration deterioration at 24 h.",
        "In heldout eICU hospitals, logistic recalibration improved Brier scores from 0.168 to 0.163 at 0 h, 0.150 to 0.149 at 6 h, "
        "and 0.106 to 0.101 at 24 h; corresponding calibration slopes changed from 0.79 to 1.16, 0.91 to 1.09, and 0.33 to 0.92. "
        "In the active-ICU competing-risk sensitivity analysis, recalibrated AUROCs were 0.706, 0.725, and 0.678 at 0 h, 6 h, and 24 h."
    ]
    v8.RESULTS["Selection-bias sensitivity analysis"] = [
        "Creatinine-record observability was 52.0%, whereas 46.9% of the strict eICU cohort met the full incident-AKI evaluability "
        "definition. In the active-ICU sensitivity analysis, unweighted complete-case AUROCs were 0.674, 0.704, and 0.698; clinical "
        "inverse-probability weighting yielded 0.693, 0.720, and 0.706 at 0 h, 6 h, and 24 h (Table S6). Under pattern-mixture scenarios "
        "in which unobserved AKI risk was 0.5 to 2.0 times observed risk, the implied external 7-day AKI incidence ranged from 17.8% to 35.9%."
    ]
    v8.DISCUSSION.insert(5, (
        "External evaluation provides an important but qualified transportability assessment. The portable eICU models retained moderate "
        "discrimination, but absolute risks were systematically miscalibrated, particularly at 24 h. Hospital-heldout recalibration improved "
        "calibration without changing ranking, supporting the need for site-specific probability updating rather than a single exported "
        "risk scale. The eICU models were feature-harmonized approximations and should not be interpreted as direct external validation of "
        "every predictor in the full MIMIC-IV models."
    ))
    v8.DISCUSSION.insert(6, (
        "Outcome observability was also a substantive external limitation. IPW estimates were broadly consistent with complete-case "
        "discrimination under a conditional missing-at-random assumption, whereas pattern-mixture analyses showed that population incidence "
        "could vary materially under alternative untestable assumptions. These analyses quantify uncertainty; they do not prove that selection "
        "bias has been eliminated. Competing discharge and death further support defining landmark risk sets among patients who remain under "
        "ICU observation."
    ))
    v8.DISCUSSION[-1] = (
        "Several limitations remain. MIMIC-IV represents one tertiary academic center and cardiac surgery dominated the development cohort. "
        "Although eICU supplied multicenter validation, surgical eligibility, feature mapping, and creatinine monitoring differed across data "
        "sources. The retrospective design is vulnerable to selection and measurement processes; AKI ascertainment in the primary model excluded "
        "urine output; procedure timing was date-level; and intraoperative exposures, fluid balance, medication timing, and potentially modifiable "
        "factors such as hypotension were not modeled comprehensively [[CITE:Park2020]]. Finally, SHAP and adjusted associations describe predictive "
        "attribution or prognostic association, not causal or modifiable effects."
    )
    v8.CONCLUSION = (
        "Dynamic landmark models based on routinely available peri-ICU data achieved moderate discrimination for incident postoperative AKI. "
        "External eICU validation supported transportable risk ranking but demonstrated calibration heterogeneity, competing-risk effects, and "
        "selective creatinine observability. The frozen parsimonious models are research-use deployment candidates only; local recalibration, "
        "prospective silent validation, and a clinically governed response pathway are required before implementation."
    )
    v8.FIGURE_INFO[1] = ("Fig2", "Figure_2_dynamic_ROC", "Discrimination across dynamic prediction landmarks",
                          "Receiver-operating-characteristic curves for four development-benchmark model families at ICU admission, 6 h, and 24 h. Landmark populations and remaining outcome windows differ and should not be interpreted as paired longitudinal comparisons.")
    v8.FIGURE_INFO[2] = ("Fig3", "Figure_3_calibration_DCA", "Calibration and clinical net benefit",
                          "Top row: development-benchmark observed versus predicted risk in ten equal-frequency groups. Bottom row: decision curves relative to treat-all and treat-none strategies. DCA, decision-curve analysis.")
    v8.REFS.append(("Pollard2018", "Pollard TJ, Johnson AEW, Raffa JD, Celi LA, Mark RG, Badawi O. The eICU Collaborative Research Database, a freely available multi-center database for critical care research. Sci Data. 2018;5:180178. doi:10.1038/sdata.2018.178.", "10.1038/sdata.2018.178", "30204154", "external validation database"))
    v8.REF_INDEX = {key: i + 1 for i, (key, *_rest) in enumerate(v8.REFS)}
    base_checklist_location = v8.checklist_location
    checklist_overrides = {
        "12f": "Methods—Model finalization and deployment boundary; Results—Temporal validation and selected parsimonious models",
        "12g": "Methods—External validation, recalibration, and selection sensitivity; Results—External validation, calibration, and observability",
        "20c": "Additional file 1: Tables S1-S6 and Figures S1-S5",
        "21": "Results—Dynamic risk sets; Table 2; Additional file 1: Tables S4-S6",
        "23b": "Results—External validation, calibration, and observability; hospital-level calibration heterogeneity analysis",
        "24": "Methods—Model finalization and deployment boundary; v22 versioned research-use artifacts",
        "27a": "Methods—Sample size and missing data; External validation, recalibration, and selection sensitivity; Discussion",
        "27b": "Methods—External validation, recalibration, and selection sensitivity; Results—Selection-bias sensitivity analysis",
        "27c": "Discussion; Conclusions; model card v22",
    }
    v8.checklist_location = lambda item: checklist_overrides.get(item, base_checklist_location(item))


def external_tables() -> tuple[list[dict[str, str]], list[dict[str, str]]]:
    perf = pd.read_csv(V16 / "model_v16_portable_external_validation_performance.csv")
    ext = perf.loc[perf["evaluation_dataset"].eq("eICU external validation")].copy()
    s5 = []
    for r in ext.itertuples():
        s5.append({
            "Landmark": f"{int(r.landmark_hours)} h", "Model": str(r.model_family), "n": f"{int(r.n):,}",
            "Event rate": f"{r.event_rate * 100:.1f}%", "AUROC (95% CI)": f"{r.auroc:.3f} ({r.auroc_ci_lower:.3f}-{r.auroc_ci_upper:.3f})",
            "AUPRC": f"{r.auprc:.3f}", "Brier": f"{r.brier_score:.3f}", "Calibration": f"{r.calibration_intercept:.3f}/{r.calibration_slope:.3f}",
        })
    ipw = pd.read_csv(V21 / "analysis_v21_ipw_weighted_external_performance.csv")
    scenario = pd.read_csv(V21 / "analysis_v21_pattern_mixture_aki_incidence.csv")
    s6 = []
    for lm in [0, 6, 24]:
        d = ipw.loc[ipw["landmark_hours"].eq(lm)].set_index("method")
        s6.append({
            "Landmark": f"{lm} h", "Complete-case AUROC": f"{d.loc['unweighted_complete_case', 'auroc']:.3f}",
            "Clinical-IPW AUROC": f"{d.loc['ipw_clinical', 'auroc']:.3f}", "Clinical+hospital-IPW AUROC": f"{d.loc['ipw_clinical_hospital', 'auroc']:.3f}",
            "Complete-case event rate": f"{d.loc['unweighted_complete_case', 'weighted_event_rate'] * 100:.1f}%",
            "Clinical-IPW event rate": f"{d.loc['ipw_clinical', 'weighted_event_rate'] * 100:.1f}%",
        })
    low, high = scenario["implied_population_aki_incidence"].min(), scenario["implied_population_aki_incidence"].max()
    s6[0]["Pattern-mixture range"] = f"{low * 100:.1f}% to {high * 100:.1f}%"
    for row in s6[1:]: row["Pattern-mixture range"] = "See 0-h scenario"
    return s5, s6


def final_parsimonious_table(tables: tuple) -> tuple:
    t1, t2, t3, s1, s2, s3, s4 = tables
    final = pd.read_csv(V14 / "model_v14_selected_parsimonious_model_finalization.csv").set_index("landmark_hours")
    calibrations = {0: (-0.009, 0.986), 6: (-0.011, 0.931), 24: (-0.094, 0.929)}
    for row in t2:
        lm = int(row["Landmark"].split()[0]); r = final.loc[lm]
        row["Selected model"] = v8.clean_model(str(r["final_model_family"]))
        row["Test N"] = str(int(r["heldout_test_n"]))
        row["Event rate"] = f"{r['heldout_test_event_rate'] * 100:.1f}%"
        row["AUROC (95% CI)"] = f"{r['auroc']:.3f} ({r['auroc_ci_lower']:.3f}-{r['auroc_ci_upper']:.3f})"
        row["AUPRC (95% CI)"] = f"{r['auprc']:.3f} ({r['auprc_ci_lower']:.3f}-{r['auprc_ci_upper']:.3f})"
        row["Brier score"] = f"{r['brier_score']:.3f}"
        row["Calibration intercept"] = f"{calibrations[lm][0]:.3f}"
        row["Calibration slope"] = f"{calibrations[lm][1]:.3f}"
    return t1, t2, t3, s1, s2, s3, s4


def write_extra_supplementary_tables(s5: list[dict[str, str]], s6: list[dict[str, str]]) -> None:
    v8.latex_table(s5, [(k, k) for k in s5[0]], "Feature-harmonized eICU external validation performance", "tab:s5", "tableS5.tex", landscape=True)
    v8.latex_table(s6, [(k, k) for k in s6[0]], "External outcome-observability and selection-bias sensitivity", "tab:s6", "tableS6.tex", landscape=True)
    with (v8.TABLES / "Table_S5_external_validation.csv").open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=list(s5[0])); writer.writeheader(); writer.writerows(s5)
    with (v8.TABLES / "Table_S6_selection_bias_sensitivity.csv").open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=list(s6[0])); writer.writeheader(); writer.writerows(s6)


def extend_supplement_word(s5: list[dict[str, str]], s6: list[dict[str, str]]) -> None:
    path = OUT / "additional_file_1_supplementary_material_en.docx"
    doc = Document(path)
    v8.landscape(doc)
    doc.add_heading("Additional external validation tables", level=1)
    v8.add_table(doc, "Table S5", "Feature-harmonized eICU external validation performance", s5, [(k, k) for k in s5[0]],
                 [700, 1300, 750, 1050, 1800, 900, 800, 1300], "Portable models use harmonized predictors; calibration is intercept/slope.", font=7)
    v8.add_table(doc, "Table S6", "External outcome-observability and selection-bias sensitivity", s6, [(k, k) for k in s6[0]],
                 [700, 1200, 1200, 1400, 1300, 1300, 1900], "IPW assumes missing at random conditional on observed covariates. Pattern-mixture scenarios range from an unobserved-to-observed AKI risk ratio of 0.5 to 2.0.", font=6.6)
    doc.save(path)


def update_supplement_tex() -> None:
    path = v8.LATEX / "supplement.tex"
    text = path.read_text(encoding="utf-8")
    text = text.replace(r"\input{tables/tableS4.tex}\clearpage", r"\input{tables/tableS4.tex}\clearpage\input{tables/tableS5.tex}\clearpage\input{tables/tableS6.tex}\clearpage")
    path.write_text(text, encoding="utf-8")


def write_lock_files() -> None:
    (OUT / "MANUSCRIPT_LOCK.md").write_text(
        "# Manuscript lock v23\n\n"
        "This is the final submission-candidate package. It supersedes v8 for submission preparation while preserving v8 unchanged. "
        "Core MIMIC-IV cohort counts and development-benchmark results are retained. The manuscript now reports temporal validation, "
        "eICU feature-harmonized external validation, hospital-heldout recalibration, competing-risk utility sensitivity, outcome-observability "
        "analysis, and assumption-driven selection-bias sensitivity.\n\n"
        "The v22 artifacts are research-use deployment candidates only. No universal clinical threshold, treatment recommendation, or claim of "
        "prospective clinical effectiveness is included.\n\n"
        "The public code repository and v1.0.1 release are fixed at https://github.com/Bizhi-Wei/postoperative-aki-dynamic-prediction. "
        "Before submission: provide the archived DOI; refresh checklist page locations only after producing the final journal PDF.\n",
        encoding="utf-8",
    )
    (OUT / "README.md").write_text(
        f"# Locked manuscript package v23\n\n"
        f"Target journal: Critical Care (Research article).\n\n"
        f"- English main manuscript: LaTeX and Word\n- Supplementary material: LaTeX and Word\n"
        f"- References: {len(v8.REFS)} Vancouver-style entries\n- Main display items: 3 tables and 4 figures\n"
        f"- Additional file 1: Tables S1-S6 and Figures S1-S5\n- Additional file 2: TRIPOD+AI checklist\n"
        f"- Status: locked research manuscript; public repository and v1.0.1 release fixed; archived DOI remains to be provided before publication.\n",
        encoding="utf-8",
    )
    old = OUT / "author_action_required.md"
    if old.exists(): old.unlink()


def main() -> None:
    configure_paths(); configure_text()
    if OUT.exists(): shutil.rmtree(OUT)
    for path in [OUT, v8.LATEX, v8.FIGURES, v8.TABLES, v8.QA]: path.mkdir(parents=True, exist_ok=True)
    v8.copy_figures()
    tables = final_parsimonious_table(v8.standardize_tables())
    s5, s6 = external_tables()
    v8.build_latex(tables); write_extra_supplementary_tables(s5, s6); update_supplement_tex()
    v8.build_word_main(tables); v8.build_word_supplement(tables); extend_supplement_word(s5, s6)
    v8.build_checklist(); v8.build_availability_doc(); v8.write_audits(); write_lock_files()
    print(f"Locked manuscript package: {OUT}")
    print(f"Abstract words: {v8.abstract_word_count()}; main text words: {v8.main_word_count()}; references: {len(v8.REFS)}")


if __name__ == "__main__":
    main()
