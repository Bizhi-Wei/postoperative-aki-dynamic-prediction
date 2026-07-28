"""Build v31 with training-only nested-CV and eICU selection audits."""

from __future__ import annotations

import csv
import hashlib
import importlib.util
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
spec = importlib.util.spec_from_file_location(
    "v30_builder", ROOT / "scripts" / "build_secondary_manuscript_v30.py"
)
b = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(b)

OUT = ROOT / "outputs" / "manuscript_package_v31_nested_selection"
LATEX = OUT / "latex"
TABLES = OUT / "tables"
FIGURES = OUT / "figures"
QA = OUT / "qa"
SOURCE_DATA = OUT / "source_data"
V31 = ROOT / "outputs" / "modeling_v31_nested_grouped_cv"
V17 = ROOT / "outputs" / "modeling_v17_eicu_recalibration_heterogeneity"
V18 = ROOT / "outputs" / "modeling_v18_eicu_outcome_observability"
V21 = ROOT / "outputs" / "modeling_v21_eicu_selection_bias_sensitivity"

for module in [b, b.m, b.m.v8]:
    for name, value in {
        "OUT": OUT,
        "LATEX": LATEX,
        "TABLES": TABLES,
        "FIGURES": FIGURES,
        "QA": QA,
    }.items():
        setattr(module, name, value)
b.SOURCE_DATA = SOURCE_DATA


def selection_summary() -> pd.DataFrame:
    path = V31 / "audit_v31_training_only_model_selection_summary.csv"
    if not path.exists():
        raise FileNotFoundError(
            "Run training_only_nested_grouped_cv_v31.py before building the manuscript"
        )
    return pd.read_csv(path).sort_values("landmark_hours")


def selection_result_paragraph(summary: pd.DataFrame) -> str:
    clauses = []
    for row in summary.itertuples(index=False):
        clauses.append(
            f"at {int(row.landmark_hours)} h, logistic-regression and XGBoost training OOF AUROCs were "
            f"{row.training_oof_lr_auroc:.3f} and {row.training_oof_xgb_auroc:.3f}, respectively, and the nested "
            f"procedure selected logistic regression in {int(row.outer_fold_lr_selected_n)} and XGBoost in "
            f"{int(row.outer_fold_xgb_selected_n)} outer folds"
        )
    if summary["training_only_recommended_model"].nunique() == 1:
        only = summary["training_only_recommended_model"].iloc[0].replace(
            "Logistic Regression", "logistic regression"
        )
        recommendations = f"{only} at all three landmarks"
    else:
        recommendations = ", ".join(
            f"{row.training_only_recommended_model.replace('Logistic Regression', 'logistic regression')} at {int(row.landmark_hours)} h"
            for row in summary.itertuples(index=False)
        )
    concordant = bool(summary["recommendation_concordant_with_reported"].astype(bool).all())
    final = (
        "The training-only recommendations matched the model families reported in the locked-test analysis."
        if concordant
        else "At least one training-only recommendation differed from the previously reported model family; this instability is retained transparently and the locked test set was not reopened for selection."
    )
    joined = "; ".join(clauses)
    if joined.startswith("at "):
        joined = "At " + joined[3:]
    return (
        "The training-only nested subject-grouped audit did not use any row from the locked 20% test partition. "
        + joined
        + f". The audit's training-only AUROC rule recommended {recommendations}. "
        + final
        + " At 24 h, the training OOF AUROC difference was only 0.001; the previously frozen logistic-regression external analysis was retained to avoid post-audit model switching."
        + " Nested-selection OOF performance and fold-level decisions are reported in Supplementary Table S10."
    )


TITLE = b.TITLE
RUNNING_TITLE = b.RUNNING_TITLE
ABSTRACT = b.ABSTRACT
KEYWORDS = b.KEYWORDS
INTRODUCTION = b.INTRODUCTION
CONCLUSION = b.CONCLUSION
AUTHORS_WORD = b.AUTHORS_WORD
AFFILIATION_WORD = b.AFFILIATION_WORD
CORRESPONDENCE_WORD = b.CORRESPONDENCE_WORD
ETHICS = b.ETHICS
FUNDING = b.FUNDING
COMPETING = b.COMPETING
CONTRIBUTIONS = b.CONTRIBUTIONS
ACKNOWLEDGEMENTS = b.ACKNOWLEDGEMENTS
AI_DISCLOSURE = b.AI_DISCLOSURE
AVAILABILITY = b.AVAILABILITY.replace("v1.0.2", "v1.0.3")
b.AVAILABILITY = AVAILABILITY
b.m.AVAILABILITY = AVAILABILITY
FIGURE_SOURCES = b.FIGURE_SOURCES
FIGURE_LEGENDS = b.FIGURE_LEGENDS
SUPP_FIGURE_LEGENDS = b.SUPP_FIGURE_LEGENDS


METHODS = {key: list(value) for key, value in b.METHODS.items()}
METHODS["Landmark and onset-anchored prediction models"].append(
    "We additionally audited model-family selection entirely within the original 80% training partition. A five-fold outer stratified subject-grouped cross-validation estimated the performance of the selection procedure; within each outer training fold, four-fold stratified subject-grouped cross-validation selected logistic regression or XGBoost by the higher pooled inner out-of-fold AUROC. Preprocessing and hyperparameters remained fixed, with no tuning search. A separate five-fold training-only grouped out-of-fold comparison supplied an AUROC-based family recommendation without using any row or outcome from the locked 20% test partition."
)

methods_reordered: dict[str, list[str]] = {}
for key, value in METHODS.items():
    methods_reordered[key] = value
    if key == "Temporal and external validation":
        methods_reordered["eICU outcome observability and selection-bias sensitivities"] = [
            "For the previously defined any-stage incident SCr-AKI endpoint, we audited external outcome observability separately from model performance. Analysis evaluability required a baseline SCr, absence of AKI before or at ICU admission, and at least one post-index SCr through day 7. A distinct creatinine-record-observability indicator required baseline and post-index SCr records regardless of pre-index AKI. We compared evaluable and nonevaluable stays, summarized hospital-level observability, and predicted record observability in a subject-grouped held-out set using clinical variables with or without hospital identifier.",
            "Selection-bias sensitivity used five-fold subject-grouped cross-fitted observability probabilities. Stabilized inverse-probability weights were truncated at the first and 99th percentiles among evaluable stays. These analyses require missing at random conditional on included covariates and cannot recover unobserved AKI outcomes. Pattern-mixture scenarios varied unobserved relative to observed AKI risk from 0.50 to 2.00. A strict-baseline sensitivity restricted external validation to stays whose baseline SCr was observed strictly before ICU admission within seven days. These any-stage endpoint audits are reported separately from the severe-AKI external models."
        ]
METHODS = methods_reordered


RESULTS = {key: list(value) for key, value in b.RESULTS.items()}
RESULTS["Severe-AKI model performance and incremental information value"][0] = RESULTS[
    "Severe-AKI model performance and incremental information value"
][0].replace("the selected XGBoost models", "the previously reported XGBoost models")
RESULTS["Severe-AKI model performance and incremental information value"].append(
    "__NESTED_SELECTION_RESULT__"
)
RESULTS["Temporal stability, external validation, and recalibration"].append(
    "In the companion any-stage eICU observability audit, 15,781 of 30,365 stays (52.0%) had both baseline and follow-up creatinine records, while 14,229 (46.9%) met the complete analysis-evaluability definition. Held-out AUROC for predicting record observability was 0.671 with clinical variables and 0.740 after adding hospital identifier, confirming institution-dependent measurement selection (Supplementary Table S11). Clinical-only IPW changed complete-case AUROCs from 0.674 to 0.693 at 0 h, 0.704 to 0.720 at 6 h, and 0.698 to 0.706 at 24 h; pattern-mixture assumptions yielded implied population any-stage AKI incidences from 17.8% to 35.9%. Restriction to a strictly pre-ICU seven-day baseline retained 99.6% of evaluable stays and left AUROCs essentially unchanged at 0.673, 0.696, and 0.687 (Supplementary Table S12). These assumption-dependent analyses quantify robustness but do not eliminate selection bias."
)


DISCUSSION = list(b.DISCUSSION)
DISCUSSION[2] = (
    DISCUSSION[2]
    + " In the separate any-stage outcome audit, creatinine-record observability was predicted by clinical features and even more strongly by hospital identity. IPW and pattern-mixture analyses therefore serve as assumption-driven bounds, not as correction proof; strict pre-ICU baseline restriction was reassuring but addressed baseline provenance rather than follow-up selection."
)
DISCUSSION[-1] = DISCUSSION[-1].replace(
    "Model-family choices were made after internal comparisons before external evaluation,",
    "Model-family selection was retrospectively audited with nested grouped cross-validation inside the training partition, but only two fixed candidate families were compared,"
)


def words(text: str) -> int:
    return b.words(text)


def abstract_word_count() -> int:
    return b.abstract_word_count()


def main_word_count() -> int:
    blocks = (
        INTRODUCTION
        + [p for ps in METHODS.values() for p in ps]
        + [p for ps in RESULTS.values() for p in ps]
        + DISCUSSION
        + [CONCLUSION]
    )
    return words(" ".join(blocks))


def write_csv(name: str, rows: list[dict[str, object]]) -> None:
    if not rows:
        raise ValueError(f"No rows for {name}")
    with (TABLES / name).open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0]))
        writer.writeheader()
        writer.writerows(rows)


def build_tables() -> dict[str, list[dict[str, object]]]:
    summary = selection_summary()
    RESULTS["Severe-AKI model performance and incremental information value"][-1] = (
        selection_result_paragraph(summary)
    )
    b.METHODS = METHODS
    b.RESULTS = RESULTS
    b.DISCUSSION = DISCUSSION
    b.TABLES = TABLES
    b.m.TABLES = TABLES
    tables = b.build_tables()
    tables["T2"] = [
        {("Reported model" if key == "Selected model" else key): value for key, value in row.items()}
        for row in tables["T2"]
    ]
    write_csv("Table_2_selected_severe_aki_models.csv", tables["T2"])

    s10: list[dict[str, object]] = []
    for row in summary.itertuples(index=False):
        s10.append(
            {
                "Landmark": f"{int(row.landmark_hours)} h",
                "Training n/events": f"{int(row.training_n):,}/{int(row.training_event_n):,}",
                "Outer/inner folds": f"{int(row.nested_outer_fold_n)}/{int(row.nested_inner_fold_n)}",
                "Nested-selection AUROC (95% CI)": f"{row.nested_selection_oof_auroc:.3f} ({row.nested_selection_oof_auroc_ci95_low:.3f}–{row.nested_selection_oof_auroc_ci95_high:.3f})",
                "AUPRC/Brier": f"{row.nested_selection_oof_auprc:.3f}/{row.nested_selection_oof_brier:.3f}",
                "Outer folds selecting LR/XGBoost": f"{int(row.outer_fold_lr_selected_n)}/{int(row.outer_fold_xgb_selected_n)}",
                "Training OOF AUROC LR/XGBoost": f"{row.training_oof_lr_auroc:.3f}/{row.training_oof_xgb_auroc:.3f}",
                "Training-only recommendation": row.training_only_recommended_model.replace("Regression", "regression"),
                "Previously reported": row.previously_reported_model.replace("Regression", "regression"),
                "Concordant": "Yes" if bool(row.recommendation_concordant_with_reported) else "No",
            }
        )

    comparison = pd.read_csv(V18 / "audit_v18_evaluable_vs_unevaluable_comparison.csv")
    comparison = comparison.loc[comparison["variable"].ne("ineligibility_reason")].copy()
    performance = pd.read_csv(V18 / "model_v18_outcome_observability_performance.csv")
    hospital = pd.read_csv(V18 / "analysis_v18_hospital_observability_summary.csv").iloc[0]
    s11: list[dict[str, object]] = [
        {
            "Section": "Denominator",
            "Variable/model": "Analysis evaluability",
            "Level/scenario": "Strict surgical first ICU stays",
            "Evaluable/result": "14,229 (46.9%)",
            "Unevaluable/comparator": "16,136 (53.1%)",
            "SMD/AUROC": "—",
            "AUPRC/Brier": "—",
            "Note": "Baseline SCr, no pre-index AKI, and follow-up SCr required",
        },
        {
            "Section": "Denominator",
            "Variable/model": "Creatinine-record observability",
            "Level/scenario": "Baseline plus post-index SCr",
            "Evaluable/result": "15,781 (52.0%)",
            "Unevaluable/comparator": "14,584 (48.0%)",
            "SMD/AUROC": "—",
            "AUPRC/Brier": "—",
            "Note": "Defined regardless of pre-index AKI",
        },
        {
            "Section": "Hospital variation",
            "Variable/model": "Creatinine-record observability",
            "Level/scenario": f"All {int(hospital.hospital_n)} hospitals",
            "Evaluable/result": f"Median {100*hospital.observability_median:.1f}% (IQR {100*hospital.observability_iqr_lower:.1f}–{100*hospital.observability_iqr_upper:.1f}%)",
            "Unevaluable/comparator": f"{int(hospital.hospital_n_ge_100)} hospitals with ≥100 stays: {100*hospital.large_hospital_observability_median:.1f}%",
            "SMD/AUROC": "—",
            "AUPRC/Brier": "—",
            "Note": f"Large-hospital range {100*hospital.large_hospital_observability_range_lower:.1f}–{100*hospital.large_hospital_observability_range_upper:.1f}%",
        },
    ]
    for row in performance.itertuples(index=False):
        label = (
            "Clinical variables"
            if row.model == "clinical_only"
            else "Clinical variables plus hospital identifier"
        )
        s11.append(
            {
                "Section": "Observability prediction",
                "Variable/model": label,
                "Level/scenario": row.split,
                "Evaluable/result": f"{int(row.n):,} test stays; {int(row.creatinine_record_observable_n):,} observable",
                "Unevaluable/comparator": f"{100*row.creatinine_record_observable_rate:.1f}% observable",
                "SMD/AUROC": f"AUROC {row.auroc:.3f}",
                "AUPRC/Brier": f"{row.auprc:.3f}/{row.brier_score:.3f}",
                "Note": f"{int(row.predictor_n)} predictors",
            }
        )
    variable_labels = {
        "anchor_age": "Age, years",
        "charlson_score": "Charlson comorbidity proxy",
        "acutephysiologyscore": "APACHE IVa Acute Physiology Score",
        "apachescore": "APACHE IVa score",
        "predictedicumortality": "APACHE IVa predicted ICU mortality",
        "predictedhospitalmortality": "APACHE IVa predicted hospital mortality",
        "unitdischargeoffset": "ICU length of stay, min",
        "hospitaldischargeoffset": "Hospital length of stay, min",
        "chf": "Congestive heart failure",
        "hypertension": "Hypertension",
        "dm": "Diabetes mellitus",
        "ckd": "Chronic kidney disease",
        "copd": "Chronic obstructive pulmonary disease",
        "liver": "Liver disease",
        "cancer": "Cancer",
        "pvd": "Peripheral vascular disease",
        "stroke": "Stroke",
        "mi": "Myocardial infarction",
        "obesity": "Obesity",
        "anemia": "Anemia",
        "icu_death": "ICU mortality",
        "hospital_death": "Hospital mortality",
        "operative_system": "Operative system",
        "unittype": "ICU type",
        "gender": "Sex",
        "unitadmitsource": "ICU admission source",
        "hospitaladmitsource": "Hospital admission source",
    }
    level_labels = {"continuous": "Continuous", "yes": "Yes", "F": "Female", "M": "Male"}
    for row in comparison.itertuples(index=False):
        s11.append(
            {
                "Section": "Evaluable vs unevaluable",
                "Variable/model": variable_labels.get(row.variable, row.variable),
                "Level/scenario": level_labels.get(row.level, row.level),
                "Evaluable/result": row.evaluable_value,
                "Unevaluable/comparator": row.unevaluable_value,
                "SMD/AUROC": f"SMD {float(row.standardized_mean_difference):.3f}",
                "AUPRC/Brier": "—",
                "Note": f"Evaluable n={int(row.evaluable_n):,}; unevaluable n={int(row.unevaluable_n):,}",
            }
        )

    strict = pd.read_csv(V17 / "analysis_v17_strict_pre_icu_baseline_sensitivity.csv")
    ipw = pd.read_csv(V21 / "analysis_v21_ipw_weighted_external_performance.csv")
    pattern = pd.read_csv(V21 / "analysis_v21_pattern_mixture_aki_incidence.csv")
    s12: list[dict[str, object]] = []
    for row in strict.itertuples(index=False):
        cohort = (
            "All eICU evaluable"
            if row.sensitivity_cohort == "all_eicu_evaluable"
            else "Strict pre-ICU baseline only"
        )
        s12.append(
            {
                "Section": "Strict baseline",
                "Landmark/scenario": f"{int(row.landmark_hours)} h",
                "Analysis/cohort": cohort,
                "n/events": f"{int(row.n):,}/{int(row.event_n):,}",
                "Event/incidence": f"{100*row.event_rate:.1f}%",
                "AUROC": f"{row.auroc:.3f}",
                "AUPRC": f"{row.auprc:.3f}",
                "Brier": f"{row.brier_score:.3f}",
                "Additional": f"Mean predicted risk {100*row.mean_predicted_risk:.1f}%",
            }
        )
    method_labels = {
        "unweighted_complete_case": "Complete case",
        "ipw_clinical": "Clinical IPW",
        "ipw_clinical_hospital": "Clinical plus hospital IPW",
    }
    for row in ipw.itertuples(index=False):
        s12.append(
            {
                "Section": "IPW selection sensitivity",
                "Landmark/scenario": f"{int(row.landmark_hours)} h",
                "Analysis/cohort": method_labels[row.method],
                "n/events": f"{int(row.n):,}/{int(row.event_n):,}",
                "Event/incidence": f"{100*row.weighted_event_rate:.1f}%",
                "AUROC": f"{row.auroc:.3f}",
                "AUPRC": f"{row.auprc:.3f}",
                "Brier": f"{row.brier_score:.3f}",
                "Additional": "—" if row.selection_assumption == "none" else row.selection_assumption,
            }
        )
    for row in pattern.itertuples(index=False):
        s12.append(
            {
                "Section": "Pattern-mixture",
                "Landmark/scenario": f"Risk ratio {row.assumed_unobserved_to_observed_aki_risk_ratio:.2f}",
                "Analysis/cohort": "All eligible strict eICU stays",
                "n/events": f"{int(row.analysis_target_n_excluding_known_preindex_aki):,} target stays",
                "Event/incidence": f"{100*row.implied_population_aki_incidence:.1f}% implied",
                "AUROC": "—",
                "AUPRC": "—",
                "Brier": "—",
                "Additional": f"Assumed unobserved risk {100*row.assumed_unobserved_aki_risk:.1f}%",
            }
        )

    tables.update({"S10": s10, "S11": s11, "S12": s12})
    write_csv("Table_S10_training_only_nested_grouped_cv.csv", s10)
    write_csv("Table_S11_eicu_outcome_observability.csv", s11)
    write_csv("Table_S12_eicu_ipw_strict_baseline_sensitivity.csv", s12)
    return tables


TABLE_SPECS = dict(b.TABLE_SPECS)
TABLE_SPECS.update(
    {
        "T2": (
            "Held-out performance of previously reported severe-AKI landmark models",
            "tab:t2",
            [
                ("Landmark", "Landmark"),
                ("Reported model", "Model"),
                ("Test n/events", "Test n/events"),
                ("AUROC (95% CI)", "AUROC (95% CI)"),
                ("AUPRC (95% CI)", "AUPRC (95% CI)"),
                ("Brier", "Brier"),
                ("Calibration intercept/slope", "Intercept/slope"),
            ],
            [1000, 1700, 1300, 2100, 2100, 900, 1600],
            True,
            7.5,
            "These frozen reported models were used for temporal and external validation. The separate training-only model-family audit is reported in Table S10; confidence intervals use 1,000 subject-cluster bootstrap resamples.",
        ),
        "S10": (
            "Training-only nested subject-grouped model-family selection audit",
            "tab:s10",
            [
                ("Landmark", "Landmark"),
                ("Training n/events", "Training n/events"),
                ("Outer/inner folds", "Outer/inner folds"),
                ("Nested-selection AUROC (95% CI)", "Nested AUROC (95% CI)"),
                ("AUPRC/Brier", "AUPRC/Brier"),
                ("Outer folds selecting LR/XGBoost", "LR/XGB folds"),
                ("Training OOF AUROC LR/XGBoost", "LR/XGB OOF AUROC"),
                ("Training-only recommendation", "Recommendation"),
                ("Previously reported", "Reported"),
                ("Concordant", "Concordant"),
            ],
            [800, 1200, 900, 1900, 1100, 1200, 1500, 1500, 1400, 800],
            True,
            6.3,
            "Only the original 80% training partition was used. Five outer and four inner stratified subject-grouped folds were used. Candidate preprocessing and hyperparameters were fixed; inner pooled OOF AUROC selected the family. CIs use 1,000 subject-cluster bootstrap resamples of nested OOF predictions. LR, logistic regression; XGB, XGBoost.",
        ),
        "S11": (
            "eICU outcome observability, hospital variation, and evaluable-versus-unevaluable comparison",
            "tab:s11",
            [
                ("Section", "Section"),
                ("Variable/model", "Variable/model"),
                ("Level/scenario", "Level"),
                ("Evaluable/result", "Evaluable"),
                ("Unevaluable/comparator", "Unevaluable"),
                ("SMD/AUROC", "Effect"),
                ("AUPRC/Brier", "AUPRC/Brier"),
                ("Note", "Note"),
            ],
            [1300, 1800, 1500, 1700, 1800, 1200, 1200, 2200],
            True,
            6.1,
            "This table concerns the any-stage incident SCr-AKI observability audit, not the severe-AKI endpoint. SMD is standardized mean difference; positive values indicate higher prevalence or mean in evaluable stays. Unevaluable stays are not assumed to be AKI-free.",
        ),
        "S12": (
            "eICU IPW, pattern-mixture, and strict pre-ICU baseline sensitivities for any-stage SCr-AKI",
            "tab:s12",
            [
                ("Section", "Section"),
                ("Landmark/scenario", "Time/scenario"),
                ("Analysis/cohort", "Analysis"),
                ("n/events", "n/events"),
                ("Event/incidence", "Event/risk"),
                ("AUROC", "AUROC"),
                ("AUPRC", "AUPRC"),
                ("Brier", "Brier"),
                ("Additional", "Additional"),
            ],
            [1200, 1800, 1700, 1300, 1200, 900, 900, 900, 2500],
            True,
            6.3,
            "These results concern the any-stage incident SCr-AKI endpoint. IPW assumes missing at random conditional on measured covariates; stabilized weights were cross-fitted and percentile-truncated. Pattern-mixture rows are assumption-driven incidence bounds and do not recover unobserved outcomes.",
        ),
    }
)
b.TABLE_SPECS = TABLE_SPECS
b.m.TABLE_SPECS = TABLE_SPECS


def word_table(doc: Document, key: str, rows: list[dict[str, object]], number: str) -> None:
    b.word_table(doc, key, rows, number)


def build_word(tables: dict[str, list[dict[str, object]]]) -> tuple[Path, Path]:
    doc = Document()
    b.m.v8.configure_word(doc, RUNNING_TITLE)
    b.m.add_cover(doc, TITLE, "Original research | Secondary analysis")
    for text in [
        f"Abstract word count: {abstract_word_count()}",
        f"Main-text word count: {main_word_count()}",
        "Tables: 4; Figures: 4; Additional files: 2",
    ]:
        p = doc.add_paragraph()
        b.m.v8.set_font(p.add_run(text), 10)
    doc.add_heading("Abstract", level=1)
    for heading, text in ABSTRACT.items():
        p = doc.add_paragraph()
        b.m.v8.set_font(p.add_run(heading + ": "), 12, bold=True)
        b.m.v8.set_font(p.add_run(text), 12)
    p = doc.add_paragraph()
    b.m.v8.set_font(p.add_run("Keywords: "), 12, bold=True)
    b.m.v8.set_font(p.add_run("; ".join(KEYWORDS)), 12)
    doc.add_heading("Background", level=1)
    for paragraph in INTRODUCTION:
        b.m.v8.add_para(doc, paragraph)
    for section, blocks in [("Methods", METHODS), ("Results", RESULTS)]:
        doc.add_heading(section, level=1)
        for heading, paragraphs in blocks.items():
            doc.add_heading(heading, level=2)
            for paragraph in paragraphs:
                b.m.v8.add_para(doc, paragraph)
    doc.add_heading("Discussion", level=1)
    for paragraph in DISCUSSION:
        b.m.v8.add_para(doc, paragraph)
    doc.add_heading("Conclusions", level=1)
    b.m.v8.add_para(doc, CONCLUSION)
    doc.add_heading("List of abbreviations", level=1)
    b.m.v8.add_para(
        doc,
        "ADQI, Acute Disease Quality Initiative; AKI, acute kidney injury; AUPRC, area under the precision-recall curve; AUROC, area under the receiver-operating-characteristic curve; CI, confidence interval; CIF, cumulative incidence function; ICU, intensive care unit; IPW, inverse-probability weighting; IQR, interquartile range; KDIGO, Kidney Disease: Improving Global Outcomes; OOF, out-of-fold; RRT, renal replacement therapy; SCr, serum creatinine.",
        indent=False,
    )
    doc.add_heading("Declarations", level=1)
    declarations = [
        ("Ethics approval and consent to participate", ETHICS),
        ("Consent for publication", "Not applicable."),
        ("Availability of data and materials", AVAILABILITY + f" [{b.m.v8.REF_INDEX['MIMIC2024']},{b.m.v8.REF_INDEX['Pollard2018']}]."),
        ("Competing interests", COMPETING),
        ("Funding", FUNDING),
        ("Author's contributions", CONTRIBUTIONS),
        ("Acknowledgements", ACKNOWLEDGEMENTS),
        ("AI-assisted editing disclosure", AI_DISCLOSURE),
    ]
    for heading, text in declarations:
        doc.add_heading(heading, level=2)
        b.m.v8.add_para(doc, text, indent=False)
    doc.add_heading("Additional files", level=1)
    b.m.v8.add_para(
        doc,
        "Additional file 1 (.docx and .pdf): Supplementary Tables S1–S12 and Supplementary Figures S1–S7.",
        indent=False,
    )
    b.m.v8.add_para(
        doc,
        "Additional file 2 (.docx and .csv): TRIPOD+AI checklist for the prediction components.",
        indent=False,
    )
    doc.add_heading("References", level=1)
    for i, (_key, reference, *_rest) in enumerate(b.m.REFS, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        b.m.v8.set_font(p.add_run(f"{i}. {reference}"), 10)
    b.m.add_landscape(doc)
    for key in ["T1", "T2", "T3", "T4"]:
        word_table(doc, key, tables[key], f"Table {key[1:]}")
    b.m.add_portrait(doc)
    doc.add_heading("Figure legends", level=1)
    for number, title, legend in FIGURE_LEGENDS:
        b.m.v8.add_para(doc, f"{number}. {title}. {legend}", indent=False)
    main_path = OUT / "secondary_manuscript_v31_en.docx"
    doc.save(main_path)

    supp = Document()
    b.m.v8.configure_word(supp, "Supplement | " + RUNNING_TITLE)
    b.m.add_cover(supp, "Additional file 1: Supplementary material", TITLE)
    b.m.add_landscape(supp)
    for i in range(1, 13):
        word_table(supp, f"S{i}", tables[f"S{i}"], f"Table S{i}")
    b.m.add_portrait(supp)
    supp.add_heading("Supplementary figures", level=1)
    for i, (number, title, legend) in enumerate(SUPP_FIGURE_LEGENDS, 1):
        p = supp.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(FIGURES / f"FigS{i}.png"), width=Inches(6.3))
        b.m.v8.add_para(supp, f"{number}. {title}. {legend}", indent=False)
    supp_path = OUT / "additional_file_1_secondary_supplement_en.docx"
    supp.save(supp_path)
    return main_path, supp_path


def latex_table(key: str, rows: list[dict[str, object]]) -> None:
    title, label, columns, _widths, landscape, _font, _legend = TABLE_SPECS[key]
    b.m.v8.latex_table(
        rows,
        columns,
        title,
        label,
        f"table{key}.tex",
        landscape=landscape,
        long=(key in {"T1", "S3", "S8", "S9", "S11", "S12"}),
    )


def write_latex(tables: dict[str, list[dict[str, object]]]) -> None:
    b.METHODS = METHODS
    b.RESULTS = RESULTS
    b.DISCUSSION = DISCUSSION
    b.write_latex(tables)
    for key in ["S10", "S11", "S12"]:
        latex_table(key, tables[key])
    main_path = LATEX / "main.tex"
    text = main_path.read_text(encoding="utf-8").replace(
        "Supplementary Tables S1--S9", "Supplementary Tables S1--S12"
    )
    main_path.write_text(text, encoding="utf-8")
    supp_path = LATEX / "supplement.tex"
    text = supp_path.read_text(encoding="utf-8")
    anchor = r"\input{tables/tableS9.tex}\clearpage"
    extra = anchor + "\n" + "\n".join(
        rf"\input{{tables/tableS{i}.tex}}\clearpage" for i in range(10, 13)
    )
    if anchor not in text:
        raise AssertionError("Could not locate supplementary-table insertion anchor")
    supp_path.write_text(text.replace(anchor, extra), encoding="utf-8")


def write_markdown() -> None:
    b.METHODS = METHODS
    b.RESULTS = RESULTS
    b.DISCUSSION = DISCUSSION
    b.write_markdown()
    old = OUT / "secondary_manuscript_v30_en.md"
    old.replace(OUT / "secondary_manuscript_v31_en.md")


def build_checklist() -> tuple[Path, Path]:
    asset = ROOT / "scripts" / "assets" / "tripod_ai_checklist_v8.tsv"
    locations = {
        "1": "Title",
        "2": "Abstract",
        "3a": "Background",
        "3b": "Background, final paragraph",
        "4": "Background, final paragraph",
        "5a": "Methods—Study design",
        "5b": "Methods—Study design",
        "6a": "Methods—Study design",
        "6b": "Methods—Study design and Figure 1",
        "7": "Methods—Landmark models",
        "8a": "Methods—Outcomes",
        "9a": "Methods—Landmark models",
        "10": "Methods—Landmarks; Figure 1; Table S1",
        "11": "Methods—eICU observability; Tables S11–S12",
        "12a": "Methods—Prediction models",
        "12b": "Methods—Nested grouped CV; Table S10",
        "12c": "Methods—Nested grouped CV and same-risk analysis; Tables S2 and S10",
        "12d": "Methods—Temporal/external validation; Table S9",
        "12e": "Methods—Prediction models",
        "12f": "Methods—External validation",
        "12g": "Methods—External validation and selection sensitivity; Tables S11–S12",
        "13": "Methods—Prediction models and nested grouped CV; Table S10",
        "15": "Methods—Prediction models; Table S10",
        "16": "Methods—Study design",
        "17": "Declarations—Ethics",
        "18a": "Declarations—Funding",
        "18b": "Declarations—Competing interests",
        "18c": "Availability of data and materials",
        "18e": "Availability of data and materials",
        "18f": "Availability of data and materials",
        "19": "Methods—Software and reporting",
        "20a": "Results—Cohort; Figure 1",
        "20b": "Table 1; Table S1",
        "20c": "Tables S2–S12",
        "21": "Results; Tables 2–4 and S10–S12",
        "22": "Availability of data and materials",
        "23a": "Results—Model performance; Figures 2–3; Tables 2–3 and S10",
        "23b": "Results—External validation and observability; Tables S5–S6 and S11–S12",
        "24": "Methods—External validation; Table 3",
        "25": "Discussion",
        "26": "Discussion—Limitations",
        "27a": "Methods—Observability and selection sensitivity; Tables S11–S12",
        "27b": "Methods—Observability",
        "27c": "Discussion; Conclusions",
    }
    rows = []
    with asset.open(encoding="utf-8") as handle:
        for row in csv.DictReader(handle, delimiter="\t"):
            row["location"] = locations.get(row["item"], "See manuscript or supplement")
            row["status"] = "Reported"
            rows.append(row)
    csv_path = OUT / "additional_file_2_tripod_ai_checklist.csv"
    with csv_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0]))
        writer.writeheader()
        writer.writerows(rows)
    doc = Document()
    b.m.v8.configure_word(doc, "TRIPOD+AI checklist")
    b.m.add_cover(doc, "Additional file 2: TRIPOD+AI checklist", TITLE)
    b.m.v8.add_para(
        doc,
        "This checklist applies to the severe-AKI prediction components. The multistate and competing-risk analyses are described separately.",
        indent=False,
    )
    b.m.add_landscape(doc)
    b.m.v8.add_table(
        doc,
        "",
        "",
        rows,
        [
            ("section", "Section"),
            ("item", "Item"),
            ("checklist_item", "Reporting recommendation"),
            ("location", "Location"),
            ("status", "Status"),
        ],
        [1200, 650, 4600, 4000, 1100],
        font=6.8,
    )
    docx_path = OUT / "additional_file_2_tripod_ai_checklist.docx"
    doc.save(docx_path)
    return docx_path, csv_path


def copy_figures_and_source_data() -> None:
    b.copy_figures()
    b.copy_source_data()
    sources = {
        V31 / "audit_v31_training_only_model_selection_summary.csv": "Table_S10_nested_selection_source_data.csv",
        V31 / "audit_v31_nested_outer_fold_performance.csv": "Table_S10_outer_fold_source_data.csv",
        V18 / "audit_v18_evaluable_vs_unevaluable_comparison.csv": "Table_S11_evaluable_comparison_source_data.csv",
        V18 / "model_v18_outcome_observability_performance.csv": "Table_S11_observability_model_source_data.csv",
        V21 / "analysis_v21_ipw_weighted_external_performance.csv": "Table_S12_ipw_source_data.csv",
        V21 / "analysis_v21_pattern_mixture_aki_incidence.csv": "Table_S12_pattern_mixture_source_data.csv",
        V17 / "analysis_v17_strict_pre_icu_baseline_sensitivity.csv": "Table_S12_strict_baseline_source_data.csv",
    }
    for src, name in sources.items():
        shutil.copy2(src, SOURCE_DATA / name)


def write_audits() -> None:
    b.write_audits()
    evidence = list(csv.DictReader((OUT / "claim_evidence_map.csv").open(encoding="utf-8-sig")))
    evidence.extend(
        [
            {
                "Claim": "Training-only nested grouped model-selection audit",
                "Source": "audit_v31_training_only_model_selection_summary.csv; audit_v31_nested_outer_fold_performance.csv",
                "Manuscript location": "Methods; Results; Table S10",
            },
            {
                "Claim": "eICU outcome observability and measurement selection",
                "Source": "audit_v18_evaluable_vs_unevaluable_comparison.csv; model_v18_outcome_observability_performance.csv",
                "Manuscript location": "Methods; Results; Table S11",
            },
            {
                "Claim": "eICU IPW, pattern-mixture, and strict baseline sensitivity",
                "Source": "analysis_v21_ipw_weighted_external_performance.csv; analysis_v21_pattern_mixture_aki_incidence.csv; analysis_v17_strict_pre_icu_baseline_sensitivity.csv",
                "Manuscript location": "Methods; Results; Table S12",
            },
        ]
    )
    with (OUT / "claim_evidence_map.csv").open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["Claim", "Source", "Manuscript location"])
        writer.writeheader()
        writer.writerows(evidence)


def write_readme() -> None:
    text = f"""# Revised secondary manuscript package v31

This package preserves v30 and adds two requested audit layers.

- Title: {TITLE}
- Author: Bizhi Wei
- Abstract words: {abstract_word_count()}
- Main-text words: {main_word_count()}
- Main items: Tables 1–4 and Figures 1–4
- Additional file 1: Tables S1–S12 and Figures S1–S7
- Additional file 2: TRIPOD+AI checklist
- New Table S10: training-only five-by-four nested subject-grouped model-family selection audit
- New Table S11: complete eICU outcome-observability and evaluable-versus-unevaluable audit
- New Table S12: eICU IPW, pattern-mixture, and strict pre-ICU baseline sensitivities
- External sensitivity boundary: Tables S11–S12 concern the previously defined any-stage incident SCr-AKI endpoint and are explicitly distinguished from the severe-AKI external models
- Repository release cited in the manuscript: v1.0.3
- Archived DOI: to be provided before publication

The locked test partition, scientific outcome counts, figures, and previously reported held-out/external performance estimates were not changed.
"""
    (OUT / "README.md").write_text(text, encoding="utf-8")


def write_manifest() -> None:
    rows = []
    for path in sorted(OUT.rglob("*")):
        if not path.is_file() or path.name == "file_manifest_sha256.csv":
            continue
        rows.append(
            {
                "relative_path": path.relative_to(OUT).as_posix(),
                "bytes": path.stat().st_size,
                "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
            }
        )
    with (OUT / "file_manifest_sha256.csv").open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["relative_path", "bytes", "sha256"])
        writer.writeheader()
        writer.writerows(rows)


def main() -> None:
    if OUT.exists():
        shutil.rmtree(OUT)
    for path in [OUT, LATEX, TABLES, FIGURES, QA, SOURCE_DATA]:
        path.mkdir(parents=True, exist_ok=True)
    tables = build_tables()
    copy_figures_and_source_data()
    write_latex(tables)
    write_markdown()
    build_word(tables)
    build_checklist()
    write_audits()
    write_readme()
    write_manifest()
    print(f"Built v31 manuscript package: {OUT}")
    print(
        f"Abstract words={abstract_word_count()}; main words={main_word_count()}; references={len(b.m.REFS)}"
    )


if __name__ == "__main__":
    main()
