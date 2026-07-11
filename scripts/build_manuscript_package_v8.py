from __future__ import annotations

import csv
import importlib.util
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
V6 = ROOT / "outputs" / "final_manuscript_v6"
V7 = ROOT / "outputs" / "manuscript_package_v7"
OUT = ROOT / "outputs" / "manuscript_package_v8_critical_care"
LATEX = OUT / "latex"
FIGURES = OUT / "figures"
TABLES = OUT / "tables"
QA = OUT / "qa"
ASSETS = ROOT / "scripts" / "assets"
GRAPHICAL = ROOT / "outputs" / "_v8_figure_stage"

spec = importlib.util.spec_from_file_location("v7builder", ROOT / "scripts" / "build_manuscript_package_v7.py")
v7 = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(v7)

TITLE = (
    "Development and internal validation of dynamic models for incident postoperative acute kidney injury "
    "in surgical intensive care: a retrospective MIMIC-IV cohort study"
)
RUNNING_TITLE = "Dynamic prediction of postoperative AKI"
AUTHORS_WORD = "Bizhi Wei¹*"
AFFILIATION_WORD = "¹ Pu Ai Medical School, Shaoyang University, Shaoyang 422000, Hunan, China"
CORRESPONDENCE_WORD = (
    "*Correspondence: Bizhi Wei, Pu Ai Medical School, Shaoyang University, "
    "Shaoyang 422000, Hunan, China; Email: 15619056250wbz@gmail.com"
)
ETHICS_STATEMENT = (
    "MIMIC-IV was created under institutional review board approval at Beth Israel Deaconess Medical Center "
    "and the Massachusetts Institute of Technology, with waiver of individual informed consent for the "
    "deidentified research resource. The present study used only deidentified data accessed through credentialed "
    "PhysioNet access and did not involve direct contact with human participants. No additional local ethics "
    "approval was required for this retrospective analysis of publicly available deidentified data."
)
AVAILABILITY_STATEMENT = (
    "MIMIC-IV version 3.1 is available through PhysioNet to credentialed users who complete the required training "
    "and sign the data use agreement. The authors are not permitted to redistribute patient-level MIMIC-IV data "
    "or derived patient-level analytic datasets. Analytic code is publicly available at "
    "https://github.com/Bizhi-Wei/postoperative-aki-dynamic-prediction (release v1.0.1: "
    "https://github.com/Bizhi-Wei/postoperative-aki-dynamic-prediction/releases/tag/v1.0.1). "
    "The archived software DOI will be provided before publication."
)
FUNDING_STATEMENT = (
    "This research received no specific grant from any funding agency in the public, commercial, or not-for-profit sectors."
)
COMPETING_INTERESTS_STATEMENT = "The authors declare that they have no competing interests."
CONTRIBUTIONS_STATEMENT = (
    "Bizhi Wei: Conceptualization, methodology, investigation, data curation, formal analysis, visualization, "
    "writing – original draft, writing – review and editing, project administration, and supervision. "
    "The author read and approved the final manuscript."
)
ACKNOWLEDGEMENTS_STATEMENT = "The authors thank the developers and maintainers of the MIMIC-IV database and PhysioNet."
AI_DISCLOSURE_STATEMENT = (
    "During manuscript preparation, the authors used ChatGPT for language editing, structural organization, and "
    "formatting assistance. The authors reviewed and verified all content, analyses, interpretations, and references, "
    "and take full responsibility for the final manuscript."
)

ABSTRACT = {
    "Background": (
        "Postoperative acute kidney injury (AKI) is common after major surgery, but risk evolves during early "
        "critical illness. We developed and internally validated landmark-specific models for incident AKI in "
        "a rigorously defined surgical intensive care cohort."
    ),
    "Methods": (
        "We conducted a retrospective cohort study using MIMIC-IV version 3.1. The analysis unit was the first "
        "intensive care unit (ICU) stay per hospital admission. Incident AKI within 7 days of ICU admission was "
        "recomputed using Kidney Disease: Improving Global Outcomes serum creatinine criteria; urine output was "
        "not used. Separate risk sets were constructed at ICU admission, 6 h, and 24 h, excluding patients whose "
        "AKI had occurred by each landmark. Logistic regression, random forest, XGBoost, and LightGBM were evaluated "
        "using an 80:20 patient-grouped split. Performance assessment included discrimination, calibration, "
        "decision-curve analysis, patient-level bootstrap confidence intervals, subgroup analyses, SHAP attribution, "
        "and prespecified baseline-creatinine sensitivity analyses."
    ),
    "Results": (
        "Among 11,943 qualifying admissions, 10,877 were evaluable for incident AKI and 4,531 developed AKI (41.7%). "
        "Stage 1, 2, and 3 accounted for 84.9%, 9.8%, and 5.3% of events. Risk sets included 10,877 admissions at "
        "0 h, 10,624 at 6 h, and 9,301 at 24 h. XGBoost had the highest test AUROC at 0 h (0.728, 95% CI 0.706-0.748) "
        "and 6 h (0.740, 0.719-0.762). At 24 h, logistic regression had the highest AUROC (0.754, 0.729-0.777), "
        "an area under the precision-recall curve of 0.602, and a Brier score of 0.176. Removing creatinine-derived "
        "predictors reduced 24-h AUROC by 0.009-0.025 across model families, whereas restriction to a pre-ICU "
        "7-day baseline creatinine produced no material change."
    ),
    "Conclusions": (
        "Routinely available peri-ICU data provided moderate internal discrimination for incident postoperative AKI "
        "across dynamic landmarks. Performance at 24 h partly depended on creatinine trajectories, although "
        "non-creatinine models retained discrimination. External and temporal validation are required before clinical use."
    ),
}
KEYWORDS = [
    "acute kidney injury", "postoperative complications", "intensive care", "dynamic prediction",
    "machine learning", "MIMIC-IV", "clinical prediction model",
]

INTRODUCTION = [
    "Postoperative acute kidney injury (AKI) is a frequent complication of major surgery and is associated with "
    "greater short-term morbidity, chronic kidney disease, cardiovascular events, and death [[CITE:Boyer2022,Prowle2021,Wang2017,Nadim2018,Hoste2015,Grams2016,Biteker2014]]. "
    "Its pathophysiology is heterogeneous: pre-existing kidney vulnerability may interact with hemodynamic instability, "
    "inflammation, anemia, nephrotoxic exposure, and procedure-specific insults. Consequently, clinically relevant risk "
    "can change substantially during the first postoperative day.",
    "Many perioperative studies assess risk once, before surgery or at ICU admission. Such models are useful for baseline "
    "stratification but cannot fully represent evolving physiology. Urine output, arterial pressure, and early laboratory "
    "trajectories have been associated with postoperative AKI, yet their interpretation depends on when they become "
    "available and whether the outcome has already occurred [[CITE:Shiba2018,Kheterpal2007]]. A model evaluated at 6 or 24 h "
    "should therefore estimate future risk only among patients who remain event free at that landmark.",
    "The Kidney Disease: Improving Global Outcomes (KDIGO) definition provides a standardized framework for AKI ascertainment "
    "from changes in serum creatinine and urine output [[CITE:KDIGO2012,Kellum2013]]. Serum creatinine, however, is both a delayed "
    "marker of kidney dysfunction and a dominant predictor in many AKI models. When a creatinine-defined outcome is predicted "
    "after ICU admission, early creatinine changes may increase apparent performance because they are already close to the "
    "outcome definition. Strict temporal alignment and sensitivity analyses that remove creatinine-derived predictors are "
    "therefore necessary to assess information leakage and clinical interpretability.",
    "Electronic health record studies have demonstrated the feasibility of longitudinal AKI prediction in general inpatient, "
    "critical care, and cardiac surgical populations [[CITE:Koyner2018,Tomasev2019,Tseng2020,Ryan2023,Demirjian2022]]. However, "
    "published models differ in surgical case mix, baseline kidney-function definitions, prediction horizons, and validation "
    "strategies. We therefore developed and internally validated models for incident AKI at ICU admission, 6 h, and 24 h in a "
    "strict therapeutic surgical cohort. We also evaluated calibration, clinical net benefit, model attribution, subgroup "
    "performance, a no-creatinine analysis, and restriction to patients with a pre-ICU baseline creatinine."
]

METHODS = {
    "Study design, setting, and reporting": [
        "We performed a retrospective prediction-model development and internal-validation study using MIMIC-IV version 3.1, "
        "a deidentified electronic health record database from Beth Israel Deaconess Medical Center in Boston, Massachusetts, "
        "United States [[CITE:Johnson2023,MIMIC2024]]. MIMIC-IV version 3.1 contains admissions from 2008 through 2022. The "
        "analysis unit was the first qualifying ICU stay within each hospital admission. Reporting was guided by TRIPOD+AI, "
        "and design decisions were considered in relation to PROBAST [[CITE:Collins2024,Wolff2019]]."
    ],
    "Cohort definition": [
        "Adults were eligible when an explicitly therapeutic major operation was recorded on the day of ICU admission or the "
        "preceding day and the first ICU location was surgical, cardiac vascular, trauma surgical, mixed medical-surgical, "
        "neurosurgical, or post-anesthesia care. Eligible procedures included cardiac, vascular, general gastrointestinal or "
        "hepatobiliary, major orthopedic, neurosurgical, and thoracic or respiratory operations. Diagnostic imaging, "
        "electrocardiography, vascular access, tracheal intubation, enteral nutrition, dialysis, nonoperative respiratory "
        "measurements, routine chest radiography, and obstetric procedures did not qualify as surgery.",
        "The index time was ICU admission. The recorded operation date was retained for audit because procedure coding provided "
        "date-level rather than precise operative-end timestamps. Admissions with AKI present at or before the index time, no "
        "usable baseline serum creatinine (SCr), or no post-index SCr measurement within 7 days were retained in cohort audits but "
        "excluded from incident-AKI model development."
    ],
    "Baseline kidney function and outcome": [
        "Baseline SCr was defined as the lowest value during the 7 days before ICU admission. If no such measurement was "
        "available, the earliest admission SCr within the first 24 h of hospitalization was used; the source and timestamp were "
        "retained. At a prediction landmark, a fallback baseline was exposed to a model only when its measurement preceded that landmark.",
        "Incident AKI was adjudicated from timestamped SCr values using KDIGO criteria, without urine output [[CITE:KDIGO2012]]. "
        "AKI was present when SCr increased by at least 0.3 mg/dL from a prior measurement within 48 h or reached at least 1.5 "
        "times baseline within 7 days after ICU admission. Patients meeting either criterion before or at ICU admission were "
        "classified as having prevalent AKI. Stage 1 was defined as a 1.5- to less than 2.0-fold baseline increase or an absolute "
        "increase of at least 0.3 mg/dL; stage 2 as a 2.0- to less than 3.0-fold increase; and stage 3 as at least a 3.0-fold "
        "increase or peak SCr of at least 4.0 mg/dL. The model outcome at each landmark was the probability of incident AKI after "
        "that landmark and within 7 days of ICU admission."
    ],
    "Landmark datasets and predictors": [
        "Prediction datasets were constructed at ICU admission (0 h), 6 h, and 24 h. Patients with AKI onset at or before a "
        "landmark were excluded from that landmark risk set. Static candidate predictors comprised age, sex, race, admission "
        "characteristics, comorbidities, Charlson score, surgical category, first ICU location, and baseline kidney-function variables.",
        "For the 0-h model, only variables available at or before ICU admission were retained. The most recent laboratory values "
        "during the 24 h before ICU admission were used as pre-index features. At 6 h and 24 h, minimum, maximum, most recent, and "
        "measurement-count features were recalculated from timestamped laboratory and vital-sign observations in (0, landmark]. "
        "Whole-follow-up summaries, mortality, length of stay, AKI-derived variables, and untimed post-index summaries were excluded. "
        "All preprocessing rules were prespecified and applied identically across demographic groups."
    ],
    "Sample size and missing data": [
        "All eligible admissions were included; no a priori sample-size calculation was performed. The available 0-h development "
        "cohort contained 10,877 admissions and 4,531 events. Study-size adequacy was considered using the number of outcome events, "
        "candidate predictors, and uncertainty in held-out performance estimates rather than a fixed events-per-variable rule "
        "[[CITE:Riley2020]]. Continuous predictors were median-imputed and accompanied by missingness indicators. Categorical and "
        "binary predictors were imputed using the most frequent training value; categorical variables were one-hot encoded. "
        "Imputation parameters were estimated in the training data and then applied to the test data."
    ],
    "Model development and internal validation": [
        "Logistic regression, random forest, XGBoost, and LightGBM were fitted separately at each landmark "
        "[[CITE:Breiman2001,Chen2016,Ke2017]]. Continuous predictors were standardized for logistic regression. An 80:20 "
        "patient-grouped split was selected from 500 candidate GroupShuffleSplit assignments to approximate the overall event "
        "prevalence. The same subject assignment was reused across landmarks, and no patient appeared in both development and test sets.",
        "Random forest used 500 trees and a minimum leaf size of 5. XGBoost used 500 trees, a learning rate of 0.03, and maximum "
        "depth of 4. LightGBM used 500 trees, a learning rate of 0.03, and 31 leaves. Hyperparameters were fixed before test-set "
        "evaluation. No outcome resampling, synthetic sampling, or class weighting was used. Models returned predicted probabilities. "
        "A secondary binary classification threshold maximizing the Youden index was selected in development predictions and applied unchanged to the test set."
    ],
    "Performance assessment": [
        "Discrimination was summarized using the area under the receiver-operating-characteristic curve (AUROC) and area under the "
        "precision-recall curve (AUPRC). Accuracy, sensitivity, specificity, precision, F1 score, and confusion matrices were "
        "calculated at a threshold of 0.5 and at the development-derived Youden threshold. Overall performance and calibration were "
        "assessed using the Brier score, calibration intercept, calibration slope, and ten equal-frequency calibration groups "
        "[[CITE:Steyerberg2010,VanCalster2019]]. Decision-curve analysis estimated net benefit across threshold probabilities from "
        "0.05 to 0.80 relative to treat-all and treat-none strategies [[CITE:Vickers2006]].",
        "Confidence intervals were obtained from 1,000 patient-level bootstrap resamples of the held-out test set. Prespecified "
        "subgroup analyses considered sex, age, race, baseline SCr source, first ICU location, chronic kidney disease, and surgical "
        "category. Subgroup intervals used 300 patient-level bootstrap resamples; estimates were omitted for groups with fewer than "
        "50 observations or a single outcome class. Because the data came from one center, between-center heterogeneity was not estimable."
    ],
    "Interpretability, fairness, and sensitivity analyses": [
        "Global SHAP values were calculated in up to 1,000 held-out observations for the tree model with the highest test AUROC at "
        "each landmark [[CITE:Lundberg2017]]. SHAP values were interpreted as model attribution, not as causal effects or evidence of modifiability.",
        "Fairness was explored descriptively through performance estimates by sex, age, and race; no fairness-constrained optimization "
        "was applied. Two prespecified sensitivity analyses retained the original patient assignment. First, all creatinine, baseline "
        "SCr, and baseline-to-ICU timing predictors were removed before model refitting. Second, models were refitted only among "
        "patients whose baseline SCr was measured during the 7 days before ICU admission. Paired patient-level bootstrap resampling "
        "compared predictions in identical test patients."
    ],
    "Software and reproducibility": [
        "Analyses were performed in Python 3.14 using pandas 3.0, scikit-learn 1.9, XGBoost 3.3, LightGBM 4.6, and SHAP 0.52 "
        "[[CITE:Pedregosa2011]]. Cohort construction, outcome derivation, feature engineering, model fitting, evaluation, and figure "
        "generation were implemented as versioned scripts. During manuscript preparation, ChatGPT assisted with language editing, "
        "structural organization, and formatting; it did not generate or modify study data or analyses. The authors reviewed and "
        "verified all numerical results, claims, interpretations, and references."
    ],
    "Patient and public involvement": [
        "Patients and members of the public were not involved in the design, conduct, interpretation, or reporting of this retrospective database study."
    ],
}

RESULTS = {
    "Cohort and incident AKI": [
        "The strict postoperative surgical ICU cohort included 11,943 hospital admissions (Fig. 1). We excluded 1,014 admissions "
        "with AKI present at or before ICU admission, 50 without a usable baseline SCr, and two without a post-index SCr measurement. "
        "The incident-AKI analytic cohort therefore comprised 10,877 admissions. Incident AKI occurred in 4,531 admissions (41.7%): "
        "3,847 events (84.9%) were stage 1, 445 (9.8%) were stage 2, and 239 (5.3%) were stage 3. Median time to AKI was 31.7 h "
        "after ICU admission (interquartile range [IQR], 16.9-42.5 h).",
        "Median age was 66 years (IQR, 57-74) and was higher among patients who developed AKI than among those who did not "
        "(69 vs 64 years; Table 1). Cardiac operations accounted for 74.1% of admissions, and 77.6% had a baseline SCr measured "
        "within 7 days before ICU admission. Chronic kidney disease, heart failure, hypertension, diabetes, anemia, and cardiac "
        "surgery were more common among admissions with incident AKI."
    ],
    "Dynamic risk sets": [
        "All 10,877 evaluable admissions entered the 0-h dataset, with 4,531 subsequent events (41.7%). Exclusion of 253 events "
        "occurring by 6 h yielded 10,624 admissions and 4,278 future events. Exclusion of 1,576 cumulative events occurring by 24 h "
        "yielded 9,301 admissions and 2,955 future events (Fig. 1). Because landmark populations and remaining outcome windows differed, "
        "between-landmark metric differences were not interpreted as paired within-patient improvements."
    ],
    "Model performance": [
        "At 0 h, XGBoost had the highest test discrimination, with AUROC 0.728 (95% CI, 0.706-0.748) and AUPRC 0.665 "
        "(0.633-0.696). XGBoost also had the highest AUROC at 6 h (0.740, 0.719-0.762) and AUPRC (0.675, 0.644-0.708). "
        "At 24 h, logistic regression had the highest AUROC (0.754, 0.729-0.777), AUPRC (0.602, 0.562-0.643), and a Brier "
        "score of 0.176 (Fig. 2; Table 2). Confidence intervals overlapped substantially across model families and did not "
        "establish the superiority of a single algorithm.",
        "Calibration was close to ideal for XGBoost at 0 h (intercept, -0.003; slope, 0.997) and remained acceptable at 6 h "
        "(intercept, -0.018; slope, 0.935). The 24-h logistic model had an intercept of -0.144 and a slope of 0.831, indicating "
        "moderately over-extreme predictions in the held-out sample. Selected models had greater internal net benefit than "
        "treat-all and treat-none strategies across most evaluated thresholds (Fig. 3)."
    ],
    "Model attribution and subgroup performance": [
        "At 6 h, baseline SCr, age, early minimum hemoglobin, Charlson score, and chronic kidney disease were prominent XGBoost "
        "features. At 24 h, the most recent creatinine had the largest mean absolute SHAP value, followed by age, minimum hemoglobin, "
        "cardiac vascular ICU location, cardiac surgery, minimum creatinine, white blood cell count, and blood urea nitrogen (Fig. 4).",
        "Discrimination was broadly maintained across sex and age groups, although estimates were less precise in smaller surgical "
        "subgroups (Additional file 1: Fig. S4 and Table S2). The selected 24-h logistic model had an AUROC of 0.767 (95% CI, "
        "0.733-0.807) in women and 0.745 (0.715-0.776) in men. Corresponding estimates were 0.732 (0.693-0.767) among patients "
        "younger than 65 years and 0.750 (0.719-0.776) among those aged 65 years or older."
    ],
    "Sensitivity analyses": [
        "Removing creatinine-derived predictors reduced 24-h discrimination in every model family (Additional file 1: Fig. S2; "
        "Table 3). Logistic-regression AUROC decreased from 0.754 to 0.729, a paired difference of -0.025 (95% CI, -0.036 to "
        "-0.013). Paired intervals also excluded zero for random forest, XGBoost, and LightGBM. Nevertheless, no-creatinine models "
        "retained AUROCs of approximately 0.73.",
        "Restriction to patients with a pre-index 7-day SCr baseline did not materially alter 24-h discrimination (Additional file 1: "
        "Fig. S3). The refitted logistic model achieved an AUROC of 0.755, compared with 0.757 for the full-cohort model evaluated in "
        "the same restricted test patients; the paired difference was -0.002 (95% CI, -0.007 to 0.004). Paired AUROC intervals "
        "crossed zero for all four model families."
    ],
}

DISCUSSION = [
    "In this strict postoperative surgical ICU cohort, routinely available data supported moderate prediction of incident AKI at "
    "ICU admission, 6 h, and 24 h. XGBoost had the highest AUROC at 0 h and 6 h, whereas logistic regression performed best at "
    "24 h. Differences among model families were small relative to bootstrap uncertainty, arguing against an assumption that greater "
    "algorithmic complexity necessarily improves later postoperative risk estimation.",
    "The observed discrimination was lower than that reported by some cardiac surgery-specific models, which have achieved AUROCs "
    "above 0.80 for selected outcomes or populations [[CITE:Tseng2020,Ryan2023,Demirjian2022]]. This difference is plausible because "
    "our cohort included several surgical specialties, focused on incident AKI after excluding prevalent disease, and enforced "
    "landmark-specific predictor windows. These design choices make the prediction target more clinically explicit but also remove "
    "information that can inflate apparent performance.",
    "The landmark design is central to interpretation. At 6 h and 24 h, patients whose AKI had already occurred were removed and "
    "predictors were restricted to measurements available by the corresponding time. Later models therefore estimated residual future "
    "risk among event-free patients rather than repeatedly scoring a fixed cohort. The modest increase in AUROC across landmarks cannot "
    "be interpreted as a within-patient improvement caused by accumulating data because event prevalence, case mix, and prediction horizon changed simultaneously.",
    "Creatinine was the dominant 24-h XGBoost attribution, and removing all creatinine-related predictors reduced AUROC across model "
    "families. Early postoperative kidney-function trajectories clearly contained predictive information; however, the outcome was "
    "itself defined by future creatinine change, so part of this signal reflects temporal proximity to the outcome definition. The "
    "approximately 0.73 AUROC of no-creatinine models indicates that demographic, comorbidity, surgical, hematologic, and physiologic "
    "features retained discrimination beyond direct creatinine measurements.",
    "Calibration and decision consequences deserve the same attention as discrimination. The selected 24-h logistic model had the "
    "highest AUROC but a calibration slope below one and a negative intercept, indicating that recalibration may be required in a new "
    "setting [[CITE:VanCalster2019]]. Decision curves suggested potential internal net benefit across a range of thresholds, but did "
    "not define an intervention, prove patient benefit, or quantify alert burden. A clinical implementation study would require a "
    "prespecified response pathway, threshold selection with clinicians, calibration surveillance, and prospective evaluation.",
    "Subgroup analyses did not reveal a large loss of discrimination by sex or age, but smaller racial and surgical strata produced "
    "wide intervals. These analyses should not be interpreted as evidence of algorithmic fairness. The database reflects one health "
    "system, clinical measurement intensity can differ across patient groups, and the study did not evaluate equalized error rates, "
    "calibration within all intersectional groups, or downstream treatment effects.",
    "The study has several strengths. Surgical eligibility excluded diagnostic and bedside procedures that are frequently "
    "misclassified when any procedure code is treated as surgery. AKI was recomputed from timestamped laboratory measurements after "
    "the index time was fixed. Landmark-specific features were reconstructed from their source timestamps, patient identity was "
    "respected during data splitting, and paired bootstrap analyses directly tested sensitivity to creatinine predictors and baseline source.",
    "Several limitations remain. MIMIC-IV represents one tertiary academic center, and cardiac surgery dominated the cohort, limiting "
    "transportability. The retrospective design is vulnerable to selection and measurement processes embedded in routine care. AKI "
    "ascertainment excluded urine output; an admission-based fallback SCr may not represent stable outpatient kidney function; and "
    "procedure timing was available only by date. Intraoperative exposures, fluid balance, medication timing, and potentially modifiable "
    "factors such as hypotension were not modeled comprehensively [[CITE:Park2020]]. Internal random-split validation does not assess "
    "temporal or geographic transportability. Finally, SHAP describes predictive attribution and cannot identify causal or modifiable risk factors."
]

CONCLUSION = (
    "Dynamic landmark models based on routinely available peri-ICU data achieved moderate internal discrimination for incident "
    "postoperative AKI. XGBoost performed best at ICU admission and 6 h, whereas logistic regression performed best at 24 h; no "
    "algorithm showed clear superiority. Creatinine trajectories contributed materially at 24 h, but non-creatinine information "
    "retained discrimination. External and temporal validation, urine-output outcome assessment, recalibration, and prospective "
    "workflow evaluation are required before clinical deployment."
)

FIGURE_INFO = [
    ("Fig1", "Figure_1_cohort_flowchart", "Cohort derivation and dynamic landmark risk sets",
     "The strict surgical ICU cohort was restricted to incident-AKI-evaluable admissions. The 6-h and 24-h risk sets exclude AKI with onset at or before each landmark. SCr, serum creatinine."),
    ("Fig2", "Figure_2_dynamic_ROC", "Discrimination across dynamic prediction landmarks",
     "Receiver-operating-characteristic curves for four model families at ICU admission, 6 h, and 24 h. Landmark populations and remaining outcome windows differ and should not be interpreted as paired longitudinal comparisons."),
    ("Fig3", "Figure_3_calibration_DCA", "Calibration and clinical net benefit",
     "Top row: observed versus predicted risk in ten equal-frequency groups. Bottom row: decision curves relative to treat-all and treat-none strategies. DCA, decision-curve analysis."),
    ("Fig4", "Figure_4_SHAP_6h_24h", "Global feature importance at 6 and 24 hours",
     "Mean absolute SHAP values for the 12 leading XGBoost features. Values quantify global model attribution and do not represent causal or modifiable effects."),
]
SUPP_FIGURE_INFO = [
    ("FigS1", "Figure_S1_precision_recall", "Precision-recall curves across prediction landmarks", "Curves are shown for all four model families at 0 h, 6 h, and 24 h."),
    ("FigS2", "Figure_S2_no_creatinine_sensitivity", "No-creatinine sensitivity analysis at 24 hours", "Paired AUROC differences and model performance after removing all creatinine-derived predictors."),
    ("FigS3", "Figure_S3_preindex_baseline_sensitivity", "Pre-index baseline sensitivity analysis at 24 hours", "Paired comparison after restricting model development and evaluation to patients with a baseline SCr measured within 7 days before ICU admission."),
    ("FigS4", "Figure_S4_subgroup_performance", "Performance across prespecified clinical subgroups", "AUROC estimates and patient-level bootstrap 95% confidence intervals for selected landmark models."),
    ("FigS5", "Figure_S5_predictor_missingness", "Missingness of pre-index laboratory predictors", "Percent missingness for laboratory predictors with more than 40% missing data across landmark datasets."),
]

REFS = [
    ("Boyer2022", "Boyer N, Eldridge J, Prowle JR, Forni LG. Postoperative acute kidney injury. Clin J Am Soc Nephrol. 2022;17:1535-1545. doi:10.2215/CJN.16541221.", "10.2215/CJN.16541221", "35710717", "postoperative AKI context"),
    ("Prowle2021", "Prowle JR, Forni LG, Bell M, et al. Postoperative acute kidney injury in adult non-cardiac surgery: joint consensus report. Nat Rev Nephrol. 2021;17:605-618. doi:10.1038/s41581-021-00418-2.", "10.1038/s41581-021-00418-2", "33976395", "definition and perioperative context"),
    ("Wang2017", "Wang Y, Bellomo R. Cardiac surgery-associated acute kidney injury: risk factors, pathophysiology and treatment. Nat Rev Nephrol. 2017;13:697-711. doi:10.1038/nrneph.2017.119.", "10.1038/nrneph.2017.119", "28869251", "cardiac surgical AKI"),
    ("Nadim2018", "Nadim MK, Forni LG, Bihorac A, et al. Cardiac and vascular surgery-associated acute kidney injury: the 20th ADQI consensus conference. J Am Heart Assoc. 2018;7:e008834. doi:10.1161/JAHA.118.008834.", "10.1161/JAHA.118.008834", "29858368", "cardiac and vascular surgery consensus"),
    ("Hoste2015", "Hoste EAJ, Bagshaw SM, Bellomo R, et al. Epidemiology of acute kidney injury in critically ill patients: the multinational AKI-EPI study. Intensive Care Med. 2015;41:1411-1423. doi:10.1007/s00134-015-3934-7.", "10.1007/s00134-015-3934-7", "26162677", "ICU AKI epidemiology"),
    ("Grams2016", "Grams ME, Sang Y, Coresh J, et al. Acute kidney injury after major surgery: a retrospective analysis of Veterans Health Administration data. Am J Kidney Dis. 2016;67:872-880. doi:10.1053/j.ajkd.2015.07.022.", "10.1053/j.ajkd.2015.07.022", "26337133", "major-surgery epidemiology"),
    ("Biteker2014", "Biteker M, Dayan A, Tekkesin AI, et al. Incidence, risk factors, and outcomes of perioperative acute kidney injury in noncardiac and nonvascular surgery. Am J Surg. 2014;207:53-59. doi:10.1016/j.amjsurg.2013.04.006.", "10.1016/j.amjsurg.2013.04.006", "24050540", "noncardiac surgery epidemiology"),
    ("Shiba2018", "Shiba A, Uchino S, Fujii T, Takinami M, Uezono S. Association between intraoperative oliguria and acute kidney injury after major noncardiac surgery. Anesth Analg. 2018;127:1229-1235. doi:10.1213/ANE.0000000000003576.", "10.1213/ANE.0000000000003576", "29933276", "dynamic perioperative risk"),
    ("Kheterpal2007", "Kheterpal S, Tremper KK, Englesbe MJ, et al. Predictors of postoperative acute renal failure after noncardiac surgery in patients with previously normal renal function. Anesthesiology. 2007;107:892-902. doi:10.1097/01.anes.0000290588.29668.38.", "10.1097/01.anes.0000290588.29668.38", "18043057", "preoperative risk prediction"),
    ("KDIGO2012", "Kidney Disease: Improving Global Outcomes Acute Kidney Injury Work Group. KDIGO clinical practice guideline for acute kidney injury. Kidney Int Suppl. 2012;2:1-138.", "", "", "AKI definition"),
    ("Kellum2013", "Kellum JA, Lameire N. Diagnosis, evaluation, and management of acute kidney injury: a KDIGO summary (Part 1). Crit Care. 2013;17:204. doi:10.1186/cc11454.", "10.1186/cc11454", "", "KDIGO implementation"),
    ("Koyner2018", "Koyner JL, Carey KA, Edelson DP, Churpek MM. The development of a machine learning inpatient acute kidney injury prediction model. Crit Care Med. 2018;46:1070-1077. doi:10.1097/CCM.0000000000003123.", "10.1097/CCM.0000000000003123", "29596073", "longitudinal AKI prediction"),
    ("Tomasev2019", "Tomasev N, Glorot X, Rae JW, et al. A clinically applicable approach to continuous prediction of future acute kidney injury. Nature. 2019;572:116-119. doi:10.1038/s41586-019-1390-1.", "10.1038/s41586-019-1390-1", "31367026", "continuous AKI prediction"),
    ("Tseng2020", "Tseng PY, Chen YT, Wang CH, et al. Prediction of the development of acute kidney injury following cardiac surgery by machine learning. Crit Care. 2020;24:478. doi:10.1186/s13054-020-03179-9.", "10.1186/s13054-020-03179-9", "32736589", "cardiac surgical ML"),
    ("Ryan2023", "Ryan CT, Zeng Z, Chatterjee S, et al. Machine learning for dynamic and early prediction of acute kidney injury after cardiac surgery. J Thorac Cardiovasc Surg. 2023;166:e551-e564. doi:10.1016/j.jtcvs.2022.09.045.", "10.1016/j.jtcvs.2022.09.045", "36347651", "dynamic cardiac surgical prediction"),
    ("Demirjian2022", "Demirjian S, Bashour CA, Shaw A, et al. Predictive accuracy of a perioperative laboratory test-based prediction model for moderate to severe acute kidney injury after cardiac surgery. JAMA. 2022;327:956-964. doi:10.1001/jama.2022.1751.", "10.1001/jama.2022.1751", "35258532", "perioperative laboratory model"),
    ("Johnson2023", "Johnson AEW, Bulgarelli L, Shen L, et al. MIMIC-IV, a freely accessible electronic health record dataset. Sci Data. 2023;10:1. doi:10.1038/s41597-022-01899-x.", "10.1038/s41597-022-01899-x", "36596836", "database description"),
    ("MIMIC2024", "Johnson A, Bulgarelli L, Pollard T, et al. MIMIC-IV. Version 3.1. PhysioNet. 2024. doi:10.13026/kpb9-mt58.", "10.13026/kpb9-mt58", "", "versioned dataset citation"),
    ("Collins2024", "Collins GS, Moons KGM, Dhiman P, et al. TRIPOD+AI statement: updated guidance for reporting clinical prediction models. BMJ. 2024;385:e078378. doi:10.1136/bmj-2023-078378.", "10.1136/bmj-2023-078378", "38626948", "reporting guideline"),
    ("Wolff2019", "Wolff RF, Moons KGM, Riley RD, et al. PROBAST: a tool to assess the risk of bias and applicability of prediction model studies. Ann Intern Med. 2019;170:51-58. doi:10.7326/M18-1376.", "10.7326/M18-1376", "30596875", "risk of bias"),
    ("Riley2020", "Riley RD, Ensor J, Snell KIE, et al. Calculating the sample size required for developing a clinical prediction model. BMJ. 2020;368:m441. doi:10.1136/bmj.m441.", "10.1136/bmj.m441", "32188600", "prediction-model sample size"),
    ("Breiman2001", "Breiman L. Random forests. Mach Learn. 2001;45:5-32. doi:10.1023/A:1010933404324.", "10.1023/A:1010933404324", "", "random forest method"),
    ("Chen2016", "Chen T, Guestrin C. XGBoost: a scalable tree boosting system. In: Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining. 2016:785-794. doi:10.1145/2939672.2939785.", "10.1145/2939672.2939785", "", "XGBoost method"),
    ("Ke2017", "Ke G, Meng Q, Finley T, et al. LightGBM: a highly efficient gradient boosting decision tree. Adv Neural Inf Process Syst. 2017;30:3146-3154.", "", "", "LightGBM method"),
    ("Steyerberg2010", "Steyerberg EW, Vickers AJ, Cook NR, et al. Assessing the performance of prediction models: a framework for traditional and novel measures. Epidemiology. 2010;21:128-138. doi:10.1097/EDE.0b013e3181c30fb2.", "10.1097/EDE.0b013e3181c30fb2", "20010215", "model-performance framework"),
    ("VanCalster2019", "Van Calster B, McLernon DJ, van Smeden M, Wynants L, Steyerberg EW. Calibration: the Achilles heel of predictive analytics. BMC Med. 2019;17:230. doi:10.1186/s12916-019-1466-7.", "10.1186/s12916-019-1466-7", "31842878", "calibration"),
    ("Vickers2006", "Vickers AJ, Elkin EB. Decision curve analysis: a novel method for evaluating prediction models. Med Decis Making. 2006;26:565-574. doi:10.1177/0272989X06295361.", "10.1177/0272989X06295361", "17099194", "decision-curve analysis"),
    ("Lundberg2017", "Lundberg SM, Lee SI. A unified approach to interpreting model predictions. Adv Neural Inf Process Syst. 2017;30:4765-4774.", "", "", "SHAP method"),
    ("Park2020", "Park S, Lee HC, Jung CW, et al. Intraoperative arterial pressure variability and postoperative acute kidney injury. Clin J Am Soc Nephrol. 2020;15:35-46. doi:10.2215/CJN.06620619.", "10.2215/CJN.06620619", "31888922", "potentially modifiable perioperative exposure"),
    ("Pedregosa2011", "Pedregosa F, Varoquaux G, Gramfort A, et al. Scikit-learn: machine learning in Python. J Mach Learn Res. 2011;12:2825-2830.", "", "", "software"),
]
REF_INDEX = {key: i + 1 for i, (key, *_rest) in enumerate(REFS)}


def clean_model(value: str) -> str:
    return {"Logistic Regression": "Logistic regression", "Random Forest": "Random forest"}.get(value, value)


def compress_numbers(numbers: list[int]) -> str:
    numbers = sorted(set(numbers)); parts = []; start = previous = numbers[0]
    for value in numbers[1:] + [None]:
        if value is not None and value == previous + 1:
            previous = value; continue
        parts.append(str(start) if start == previous else f"{start}-{previous}")
        if value is not None: start = previous = value
    return ",".join(parts)


def render_text(text: str, target: str) -> str:
    pieces = re.split(r"(\[\[CITE:[^\]]+\]\])", text)
    output = []
    for piece in pieces:
        if piece.startswith("[[CITE:"):
            keys = piece[7:-2].split(",")
            if target == "tex": output.append(r"\cite{" + ",".join(keys) + "}")
            else: output.append("[" + compress_numbers([REF_INDEX[key] for key in keys]) + "]")
        else:
            output.append(tex_escape(piece) if target == "tex" else piece)
    return "".join(output)


def tex_escape(value: object) -> str:
    text = str(value)
    for source, target in [("\\", r"\textbackslash{}"), ("&", r"\&"), ("%", r"\%"), ("#", r"\#"),
                           ("_", r"\_"), ("–", "--"), ("—", "---"), ("≥", r"$\geq$"), ("≤", r"$\leq$")]:
        text = text.replace(source, target)
    return text


def abstract_word_count() -> int:
    return len(re.findall(r"\b[\w-]+\b", " ".join(ABSTRACT.values())))


def main_word_count() -> int:
    blocks = INTRODUCTION + [p for values in METHODS.values() for p in values] + [p for values in RESULTS.values() for p in values] + DISCUSSION + [CONCLUSION]
    return len(re.findall(r"\b[\w-]+\b", re.sub(r"\[\[CITE:[^\]]+\]\]", "", " ".join(blocks))))


def humanize_predictor(name: str) -> str:
    direct = {
        "first_careunit": "First ICU location", "gender": "Sex", "anchor_age": "Age", "race": "Race",
        "admission_type": "Admission type", "insurance": "Insurance", "marital_status": "Marital status",
        "chf": "Congestive heart failure", "hypertension": "Hypertension", "dm": "Diabetes mellitus",
        "dm_comp": "Diabetes with complications", "ckd": "Chronic kidney disease", "copd": "Chronic pulmonary disease",
        "liver": "Liver disease", "cancer": "Cancer", "pvd": "Peripheral vascular disease", "stroke": "Stroke",
        "mi": "Myocardial infarction", "obesity": "Obesity", "anemia": "Anemia", "charlson_score": "Charlson comorbidity score",
        "n_qualifying_codes": "Number of qualifying procedure codes", "days_from_procedure_to_icu": "Days from procedure to ICU admission",
        "baseline_scr_available_at_landmark": "Baseline SCr available", "baseline_scr_at_landmark": "Baseline SCr",
        "baseline_scr_source_at_landmark": "Baseline SCr source", "baseline_to_icu_hours_at_landmark": "Baseline-to-ICU interval (h)",
    }
    surgery = {
        "cardiac_surgery": "Cardiac surgery", "non_cardiac_surgery": "Noncardiac surgery", "vascular_surgery": "Vascular surgery",
        "general_gi_hepatobiliary_surgery": "General/GI/hepatobiliary surgery", "orthopedic_major_surgery": "Major orthopedic surgery",
        "neurosurgery": "Neurosurgery", "thoracic_respiratory_surgery": "Thoracic/respiratory surgery",
    }
    if name in direct: return direct[name]
    if name in surgery: return surgery[name]
    match = re.match(r"(lab|vital)_(pre24h|0_6h|0_24h)_([a-z0-9]+(?:_[a-z0-9]+)*)_(min|max|last|count)$", name)
    if match:
        kind, window, variable, statistic = match.groups()
        variables = {"bun": "blood urea nitrogen", "wbc": "white blood cell count", "inr": "INR", "paco2": "PaCO2",
                     "pao2": "PaO2", "ph": "pH", "spo2": "SpO2", "map": "mean arterial pressure", "sbp": "systolic blood pressure",
                     "dbp": "diastolic blood pressure", "temperature_c": "temperature", "heart_rate": "heart rate",
                     "respiratory_rate": "respiratory rate", "platelet": "platelet count"}
        label = variables.get(variable, variable.replace("_", " "))
        prefix = {"pre24h": "Pre-ICU 24-h", "0_6h": "0-6 h", "0_24h": "0-24 h"}[window]
        suffix = {"min": "minimum", "max": "maximum", "last": "most recent", "count": "measurement count"}[statistic]
        return f"{prefix} {label} ({suffix})"
    return name.replace("_", " ").capitalize()


def standardize_tables():
    TABLES.mkdir(parents=True, exist_ok=True)
    t1 = v7.read_csv("Table_1_baseline_characteristics.csv")
    replacements = {"Anaemia": "Anemia", "Orthopaedic": "Orthopedic", "No AKI": "No incident AKI"}
    for row in t1:
        for old, new in replacements.items(): row["Characteristic"] = row["Characteristic"].replace(old, new)
    t2 = v7.read_csv("Table_2_selected_model_performance.csv")
    for row in t2: row["Selected model"] = clean_model(row["Selected model"])
    t3 = v7.read_csv("Table_3_sensitivity_analyses_24h.csv")
    for row in t3: row["Model"] = clean_model(row["Model"])
    s1 = v7.combine_s1(); s2 = v7.combine_s2(); s3_display = v7.high_missingness(); s4 = v7.combine_s4()
    for row in s3_display:
        row["predictor"] = row["predictor"].replace("haemoglobin", "hemoglobin").replace("Haemoglobin", "Hemoglobin")
    for rows in [s1, s2, s4]:
        for row in rows:
            if "model" in row: row["model"] = clean_model(row["model"])
    # Enforce American medical English consistently in every displayed table.
    for rows in [t1, t2, t3, s1, s2, s3_display, s4]:
        for row in rows:
            for key, value in row.items():
                if isinstance(value, str):
                    row[key] = (value.replace("orthopaedic", "orthopedic")
                                      .replace("Orthopaedic", "Orthopedic")
                                      .replace("haemoglobin", "hemoglobin")
                                      .replace("Haemoglobin", "Hemoglobin")
                                      .replace("anaemia", "anemia")
                                      .replace("Anaemia", "Anemia"))
    all_missing = v7.read_csv("Table_S3_predictor_missingness.csv")
    for row in all_missing:
        row["predictor"] = humanize_predictor(row["predictor"])
        row["variable_type"] = {"continuous_numeric": "Continuous", "binary": "Binary", "categorical": "Categorical"}.get(row["variable_type"], row["variable_type"])
    tables = {
        "Table_1_baseline_characteristics.csv": t1,
        "Table_2_selected_model_performance.csv": t2,
        "Table_3_sensitivity_analyses_24h.csv": t3,
        "Table_S1_all_model_performance.csv": s1,
        "Table_S2_subgroup_performance.csv": s2,
        "Table_S3_high_missingness.csv": s3_display,
        "Table_S3_all_predictor_missingness.csv": all_missing,
        "Table_S4_all_landmark_sensitivity_analyses.csv": s4,
    }
    for filename, rows in tables.items():
        with (TABLES / filename).open("w", newline="", encoding="utf-8-sig") as handle:
            writer = csv.DictWriter(handle, fieldnames=list(rows[0])); writer.writeheader(); writer.writerows(rows)
    return t1, t2, t3, s1, s2, s3_display, s4


def write_bibtex():
    entries = []
    for key, citation, doi, _pmid, _role in REFS:
        title = citation.split(". ", 1)[1].split(". ", 1)[0] if ". " in citation else citation
        entries.append(f"@misc{{{key},\n  title={{{title}}},\n  note={{{citation}}}" + (f",\n  doi={{{doi}}}" if doi else "") + "\n}")
    (LATEX / "references.bib").write_text("\n\n".join(entries) + "\n", encoding="utf-8")


def copy_figures():
    FIGURES.mkdir(parents=True, exist_ok=True)
    for short, source, _title, _legend in FIGURE_INFO + SUPP_FIGURE_INFO:
        for ext in ["png", "pdf", "svg"]:
            shutil.copy2(V6 / "figures" / f"{source}.{ext}", FIGURES / f"{short}.{ext}")
    for ext in ["png", "pdf", "svg"]:
        shutil.copy2(GRAPHICAL / f"graphical_abstract.{ext}", FIGURES / f"Graphical_Abstract.{ext}")


def latex_table(rows, columns, caption, label, filename, landscape=False, long=False):
    v7.write_tex_table(rows, columns, caption, label, LATEX / "tables" / filename, landscape=landscape, long=long)


def build_latex(tables):
    LATEX.mkdir(parents=True, exist_ok=True); (LATEX / "tables").mkdir(exist_ok=True)
    t1, t2, t3, s1, s2, s3, s4 = tables
    latex_table(t1, [("Characteristic", "Characteristic"), ("Overall (N=10,877)", "Overall"), ("No AKI (N=6,346)", "No incident AKI"),
                     ("Incident AKI (N=4,531)", "Incident AKI"), ("Standardized mean difference", "SMD"), ("Missing, n", "Missing, n")],
                "Baseline characteristics by incident acute kidney injury status", "tab:baseline", "table1.tex", landscape=True)
    latex_table(t2, [("Landmark", "Landmark"), ("Selected model", "Selected model"), ("Test N", "Test n"), ("Event rate", "Event rate"),
                     ("AUROC (95% CI)", "AUROC (95% CI)"), ("AUPRC (95% CI)", "AUPRC (95% CI)"), ("Brier score", "Brier score"),
                     ("Calibration intercept", "Intercept"), ("Calibration slope", "Slope")],
                "Performance of selected models at each prediction landmark", "tab:selected", "table2.tex", landscape=True)
    latex_table(t3, [("Sensitivity analysis", "Analysis"), ("Model", "Model"), ("Full/reference AUROC", "Reference AUROC"),
                     ("Sensitivity AUROC", "Sensitivity AUROC"), ("ΔAUROC (95% paired CI)", "AUROC difference (95% CI)")],
                "Sensitivity analyses at the 24-hour landmark", "tab:sensitivity", "table3.tex", landscape=True)
    latex_table(s1, [("time", "Landmark"), ("model", "Model"), ("n", "Test n"), ("event", "Event rate"), ("auroc", "AUROC (95% CI)"),
                     ("auprc", "AUPRC (95% CI)"), ("brier", "Brier"), ("cal", "Intercept/slope"), ("youden", "Youden threshold (sensitivity/specificity)")],
                "Performance of all models across prediction landmarks", "tab:s1", "tableS1.tex", landscape=True)
    latex_table(s2, [("time", "Landmark"), ("model", "Model"), ("group", "Subgroup"), ("level", "Level"), ("n", "n/events"),
                     ("auroc", "AUROC (95% CI)"), ("auprc", "AUPRC (95% CI)"), ("brier", "Brier")],
                "Model performance across prespecified clinical subgroups", "tab:s2", "tableS2.tex", landscape=True, long=True)
    latex_table(s3, [("landmark_hours", "Landmark (h)"), ("predictor", "Predictor"), ("variable_type", "Type"), ("n_missing", "Missing n"),
                     ("missing_percent", "Missing (%)"), ("n_observed", "Observed n")],
                "Predictors with more than 40 percent missing data", "tab:s3", "tableS3.tex", landscape=True, long=True)
    latex_table(s4, [("analysis", "Analysis"), ("time", "Landmark"), ("model", "Model"), ("n", "Test n"), ("reference", "Reference AUROC"),
                     ("sensitivity", "Sensitivity AUROC"), ("delta", "AUROC difference (95% CI)")],
                "Sensitivity analyses across all prediction landmarks", "tab:s4", "tableS4.tex", landscape=True, long=True)

    abstract_tex = "\n\n".join(r"\textbf{" + key + ":} " + tex_escape(value) for key, value in ABSTRACT.items())
    methods_tex = "\n".join(r"\subsection{" + tex_escape(key) + "}\n" + "\n\n".join(render_text(p, "tex") for p in values) for key, values in METHODS.items())
    results_tex = "\n".join(r"\subsection{" + tex_escape(key) + "}\n" + "\n\n".join(render_text(p, "tex") for p in values) for key, values in RESULTS.items())
    legends_tex = []
    for short, _source, title, legend in FIGURE_INFO:
        legends_tex.append(r"\begin{figure}[p]\centering\includegraphics[width=\textwidth]{../figures/" + short + ".pdf}\n" +
                           r"\caption{" + tex_escape(title) + ". " + tex_escape(legend) + r"}\end{figure}")
    refs = "\n".join(r"\bibitem{" + key + "} " + tex_escape(citation) for key, citation, *_ in REFS)
    main = r"""\documentclass[11pt]{article}
\usepackage[letterpaper,margin=1in]{geometry}
\usepackage[T1]{fontenc}\usepackage{lmodern}\usepackage{microtype}
\usepackage{graphicx,booktabs,longtable,pdflscape,caption,setspace,lineno,hyperref}
\usepackage[numbers,sort&compress]{natbib}
\doublespacing\linenumbers\captionsetup{font=small,labelfont=bf}
\title{__TITLE__}
\author{Bizhi Wei\textsuperscript{1*}\\{}
\textsuperscript{1}Pu Ai Medical School, Shaoyang University, Shaoyang 422000, Hunan, China\\{}
\textsuperscript{*}Correspondence: Bizhi Wei\\{}
Pu Ai Medical School, Shaoyang University, Shaoyang 422000, Hunan, China\\{}
Email: 15619056250wbz@gmail.com}
\date{}
\begin{document}\maketitle
\noindent\textbf{Article type:} Research\par
\noindent\textbf{Abstract word count:} __ABSWC__\par
\noindent\textbf{Main-text word count:} __MAINWC__\par
\noindent\textbf{Tables/Figures:} 3/4\par
\section*{Abstract}
__ABSTRACT__
\noindent\textbf{Keywords:} __KEYWORDS__
\section{Background}
__INTRO__
\section{Methods}
__METHODS__
\section{Results}
__RESULTS__
\section{Discussion}
__DISCUSSION__
\section{Conclusions}
__CONCLUSION__
\section*{List of abbreviations}
AKI, acute kidney injury; AUPRC, area under the precision-recall curve; AUROC, area under the receiver-operating-characteristic curve; CI, confidence interval; DCA, decision-curve analysis; ICU, intensive care unit; IQR, interquartile range; KDIGO, Kidney Disease: Improving Global Outcomes; SCr, serum creatinine; SHAP, SHapley Additive exPlanations.
\section*{Declarations}
\subsection*{Ethics approval and consent to participate}
__ETHICS__
\subsection*{Consent for publication} Not applicable.
\subsection*{Availability of data and materials}
__AVAILABILITY__ \cite{MIMIC2024}
\subsection*{Competing interests} __COMPETING__
\subsection*{Funding} __FUNDING__
\subsection*{Authors' contributions} __CONTRIBUTIONS__
\subsection*{Acknowledgements} __ACKNOWLEDGEMENTS__
\subsection*{AI-assisted editing disclosure} __AI_DISCLOSURE__
\section*{Additional files}
Additional file 1 (.docx and .pdf): Supplementary Tables S1-S4 and Supplementary Figures S1-S5.\\
Additional file 2 (.docx and .csv): Completed TRIPOD+AI checklist.
\begin{thebibliography}{99}
__REFS__
\end{thebibliography}
\clearpage\input{tables/table1.tex}\clearpage\input{tables/table2.tex}\clearpage\input{tables/table3.tex}
\clearpage
__FIGURES__
\end{document}
"""
    replacements = {
        "__TITLE__": tex_escape(TITLE), "__ABSWC__": str(abstract_word_count()), "__MAINWC__": str(main_word_count()),
        "__ABSTRACT__": abstract_tex, "__KEYWORDS__": "; ".join(KEYWORDS),
        "__INTRO__": "\n\n".join(render_text(p, "tex") for p in INTRODUCTION), "__METHODS__": methods_tex,
        "__RESULTS__": results_tex, "__DISCUSSION__": "\n\n".join(render_text(p, "tex") for p in DISCUSSION),
        "__CONCLUSION__": render_text(CONCLUSION, "tex"), "__REFS__": refs, "__FIGURES__": "\n".join(legends_tex),
        "__ETHICS__": tex_escape(ETHICS_STATEMENT), "__AVAILABILITY__": tex_escape(AVAILABILITY_STATEMENT),
        "__COMPETING__": tex_escape(COMPETING_INTERESTS_STATEMENT), "__FUNDING__": tex_escape(FUNDING_STATEMENT),
        "__CONTRIBUTIONS__": tex_escape(CONTRIBUTIONS_STATEMENT),
        "__ACKNOWLEDGEMENTS__": tex_escape(ACKNOWLEDGEMENTS_STATEMENT),
        "__AI_DISCLOSURE__": tex_escape(AI_DISCLOSURE_STATEMENT),
    }
    for key, value in replacements.items(): main = main.replace(key, value)
    (LATEX / "main.tex").write_text(main, encoding="utf-8")

    supp_figs = []
    for short, _source, title, legend in SUPP_FIGURE_INFO:
        supp_figs.append(r"\begin{figure}[p]\centering\includegraphics[width=\textwidth]{../figures/" + short + ".pdf}\n" +
                         r"\caption{" + tex_escape(title) + ". " + tex_escape(legend) + r"}\end{figure}\clearpage")
    supplement = r"""\documentclass[11pt]{article}
\usepackage[letterpaper,margin=0.75in]{geometry}\usepackage[T1]{fontenc}\usepackage{lmodern}
\usepackage{graphicx,booktabs,longtable,pdflscape,caption,setspace,hyperref}\onehalfspacing
\begin{document}\begin{center}{\Large\bfseries Additional file 1: Supplementary material}\\[8pt]
__TITLE__\\[8pt]
Bizhi Wei\textsuperscript{1*}\\
\textsuperscript{1}Pu Ai Medical School, Shaoyang University, Shaoyang 422000, Hunan, China\\
\textsuperscript{*}Correspondence: Bizhi Wei; 15619056250wbz@gmail.com
\end{center}
\input{tables/tableS1.tex}\clearpage\input{tables/tableS2.tex}\clearpage\input{tables/tableS3.tex}\clearpage\input{tables/tableS4.tex}\clearpage
__SUPPFIGS__
\end{document}
""".replace("__TITLE__", tex_escape(TITLE)).replace("__SUPPFIGS__", "\n".join(supp_figs))
    (LATEX / "supplement.tex").write_text(supplement, encoding="utf-8")
    write_bibtex()


def set_font(run, size=12, bold=None, italic=None):
    run.font.name = "Times New Roman"; run.font.size = Pt(size)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic


def add_field(paragraph, instruction):
    run = paragraph.add_run(); begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = instruction
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = "1"; end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for element in [begin, instr, separate, text, end]: run._r.append(element)
    set_font(run, 9)


def configure_word(doc, running_title):
    section = doc.sections[0]; section.page_width = Inches(8.5); section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    sectpr = section._sectPr; ln = OxmlElement("w:lnNumType"); ln.set(qn("w:countBy"), "1"); ln.set(qn("w:restart"), "continuous"); ln.set(qn("w:distance"), "360"); sectpr.append(ln)
    styles = doc.styles
    normal = styles["Normal"]; normal.font.name = "Times New Roman"; normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing = 2.0; normal.paragraph_format.space_after = Pt(0)
    for name, size in [("Title", 16), ("Heading 1", 14), ("Heading 2", 12)]:
        style = styles[name]; style.font.name = "Times New Roman"; style.font.size = Pt(size); style.font.bold = True; style.font.color.rgb = None
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.space_before = Pt(10); style.paragraph_format.space_after = Pt(4)
    header = section.header.paragraphs[0]; header.text = running_title; header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs: set_font(run, 9, italic=True)
    footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_field(footer, "PAGE")


def add_para(doc, text, bold_prefix=None, indent=True):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Inches(0.3)
    if bold_prefix:
        r = p.add_run(bold_prefix); set_font(r, 12, bold=True); text = text[len(bold_prefix):].lstrip()
    r = p.add_run(render_text(text, "word")); set_font(r, 12)
    return p


def landscape(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE); section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = Inches(11), Inches(8.5)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.55)
    return section


def portrait(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE); section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.8)
    return section


def add_table(doc, number, title, rows, columns, widths, legend="", font=8):
    if number or title:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        caption = f"{number}. {title}" if number else title
        r = p.add_run(caption); set_font(r, 10, bold=True)
    table = doc.add_table(rows=1, cols=len(columns)); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = "Table Grid"; table.autofit = False
    v7.set_table_geometry(table, widths, indent=0); v7.set_repeat_table_header(table.rows[0])
    for cell, (_key, label) in zip(table.rows[0].cells, columns):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER; v7.set_cell_margins(cell, 55, 75, 55, 75)
        cell.text = label
        for run in cell.paragraphs[0].runs: set_font(run, font, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, (key, _label) in zip(cells, columns):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP; v7.set_cell_margins(cell, 45, 75, 45, 75); cell.text = str(row.get(key, ""))
            for run in cell.paragraphs[0].runs: set_font(run, font)
    if legend:
        p = doc.add_paragraph(); r = p.add_run(legend); set_font(r, 9, italic=True)
    return table


def build_word_main(tables):
    t1, t2, t3, *_ = tables
    doc = Document(); configure_word(doc, RUNNING_TITLE)
    # Avoid the built-in Word Title style, which can carry a theme-dependent
    # decorative bottom rule into LibreOffice/PDF exports.
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE); set_font(r, 16, bold=True)
    for text in [AUTHORS_WORD, AFFILIATION_WORD, CORRESPONDENCE_WORD]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(text); set_font(r, 11)
    for text in ["Article type: Research", f"Abstract word count: {abstract_word_count()}", f"Main-text word count: {main_word_count()}", "Tables: 3; Figures: 4; Additional files: 2"]:
        p = doc.add_paragraph(); r = p.add_run(text); set_font(r, 10)
    doc.add_heading("Abstract", level=1)
    for heading, text in ABSTRACT.items():
        p = doc.add_paragraph(); p.paragraph_format.first_line_indent = None
        r = p.add_run(heading + ": "); set_font(r, 12, bold=True); r = p.add_run(text); set_font(r, 12)
    p = doc.add_paragraph(); r = p.add_run("Keywords: "); set_font(r, 12, bold=True); r = p.add_run("; ".join(KEYWORDS)); set_font(r, 12)
    doc.add_heading("Background", level=1)
    for text in INTRODUCTION: add_para(doc, text)
    doc.add_heading("Methods", level=1)
    for heading, paragraphs in METHODS.items():
        doc.add_heading(heading, level=2)
        for text in paragraphs: add_para(doc, text)
    doc.add_heading("Results", level=1)
    for heading, paragraphs in RESULTS.items():
        doc.add_heading(heading, level=2)
        for text in paragraphs: add_para(doc, text)
    doc.add_heading("Discussion", level=1)
    for text in DISCUSSION: add_para(doc, text)
    doc.add_heading("Conclusions", level=1); add_para(doc, CONCLUSION)
    doc.add_heading("List of abbreviations", level=1)
    add_para(doc, "AKI, acute kidney injury; AUPRC, area under the precision-recall curve; AUROC, area under the receiver-operating-characteristic curve; CI, confidence interval; DCA, decision-curve analysis; ICU, intensive care unit; IQR, interquartile range; KDIGO, Kidney Disease: Improving Global Outcomes; SCr, serum creatinine; SHAP, SHapley Additive exPlanations.", indent=False)
    doc.add_heading("Declarations", level=1)
    declarations = {
        "Ethics approval and consent to participate": ETHICS_STATEMENT,
        "Consent for publication": "Not applicable.",
        "Availability of data and materials": AVAILABILITY_STATEMENT + " [18]",
        "Competing interests": COMPETING_INTERESTS_STATEMENT,
        "Funding": FUNDING_STATEMENT,
        "Authors' contributions": CONTRIBUTIONS_STATEMENT,
        "Acknowledgements": ACKNOWLEDGEMENTS_STATEMENT,
        "AI-assisted editing disclosure": AI_DISCLOSURE_STATEMENT,
    }
    for heading, text in declarations.items(): doc.add_heading(heading, level=2); add_para(doc, text, indent=False)
    doc.add_heading("Additional files", level=1)
    add_para(doc, "Additional file 1 (.docx and .pdf): Supplementary Tables S1-S4 and Supplementary Figures S1-S5.", indent=False)
    add_para(doc, "Additional file 2 (.docx and .csv): Completed TRIPOD+AI checklist.", indent=False)
    doc.add_heading("References", level=1)
    for i, (_key, citation, *_rest) in enumerate(REFS, 1):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.25); p.paragraph_format.first_line_indent = Inches(-0.25)
        r = p.add_run(f"{i}. {citation}"); set_font(r, 10)
    landscape(doc)
    add_table(doc, "Table 1", "Baseline characteristics by incident acute kidney injury status", t1,
              [("Characteristic", "Characteristic"), ("Overall (N=10,877)", "Overall"), ("No AKI (N=6,346)", "No incident AKI"),
               ("Incident AKI (N=4,531)", "Incident AKI"), ("Standardized mean difference", "SMD"), ("Missing, n", "Missing, n")],
              [2600, 1800, 1800, 1800, 1450, 1250], "Values are median [IQR] or n (%). SMD, standardized mean difference.")
    add_table(doc, "Table 2", "Performance of selected models at each prediction landmark", t2,
              [("Landmark", "Landmark"), ("Selected model", "Selected model"), ("Test N", "Test n"), ("Event rate", "Event rate"),
               ("AUROC (95% CI)", "AUROC (95% CI)"), ("AUPRC (95% CI)", "AUPRC (95% CI)"), ("Brier score", "Brier"),
               ("Calibration intercept", "Intercept"), ("Calibration slope", "Slope")],
              [850, 1500, 900, 950, 1850, 1850, 900, 1000, 900], "Confidence intervals were obtained by patient-level bootstrap resampling of the held-out test set.", font=7.5)
    add_table(doc, "Table 3", "Sensitivity analyses at the 24-hour landmark", t3,
              [("Sensitivity analysis", "Analysis"), ("Model", "Model"), ("Full/reference AUROC", "Reference AUROC"),
               ("Sensitivity AUROC", "Sensitivity AUROC"), ("ΔAUROC (95% paired CI)", "AUROC difference (95% CI)")],
              [2750, 1800, 1800, 1800, 2850], "AUROC differences were estimated in identical test patients using paired patient-level bootstrap resampling.")
    portrait(doc); doc.add_heading("Figure legends", level=1)
    for i, (_short, _source, title, legend) in enumerate(FIGURE_INFO, 1): add_para(doc, f"Figure {i}. {title}. {legend}", indent=False)
    path = OUT / "critical_care_main_manuscript_en.docx"; doc.save(path); return path


def build_word_supplement(tables):
    _t1, _t2, _t3, s1, s2, s3, s4 = tables
    doc = Document(); configure_word(doc, "Additional file 1 | " + RUNNING_TITLE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run("Additional file 1: Supplementary material"); set_font(r, 16, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(TITLE); set_font(r, 12, italic=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(AUTHORS_WORD); set_font(r, 11)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(AFFILIATION_WORD); set_font(r, 10)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(CORRESPONDENCE_WORD); set_font(r, 10)
    landscape(doc)
    add_table(doc, "Table S1", "Performance of all models across prediction landmarks", s1,
              [("time", "Landmark"), ("model", "Model"), ("n", "Test n"), ("event", "Event rate"), ("auroc", "AUROC (95% CI)"),
               ("auprc", "AUPRC (95% CI)"), ("brier", "Brier"), ("cal", "Intercept/slope"), ("youden", "Youden threshold (sensitivity/specificity)")],
              [700, 1400, 850, 950, 1700, 1700, 800, 1250, 2400], font=7)
    add_table(doc, "Table S2", "Model performance across prespecified clinical subgroups", s2,
              [("time", "Landmark"), ("model", "Model"), ("group", "Subgroup"), ("level", "Level"), ("n", "n/events"),
               ("auroc", "AUROC (95% CI)"), ("auprc", "AUPRC (95% CI)"), ("brier", "Brier")],
              [650, 1300, 1850, 1700, 900, 1850, 1850, 750], font=6.6)
    add_table(doc, "Table S3", "Predictors with more than 40% missing data", s3,
              [("landmark_hours", "Landmark (h)"), ("predictor", "Predictor"), ("variable_type", "Type"),
               ("n_missing", "Missing n"), ("missing_percent", "Missing (%)"), ("n_observed", "Observed n")],
              [900, 5000, 1500, 1300, 1400, 1300], "The complete predictor-level missingness table is provided as a machine-readable CSV file.", font=7.3)
    add_table(doc, "Table S4", "Sensitivity analyses across all prediction landmarks", s4,
              [("analysis", "Analysis"), ("time", "Landmark"), ("model", "Model"), ("n", "Test n"), ("reference", "Reference AUROC"),
               ("sensitivity", "Sensitivity AUROC"), ("delta", "AUROC difference (95% CI)")],
              [2200, 850, 1550, 900, 1550, 1700, 2600], font=7)
    portrait(doc); doc.add_heading("Supplementary figures", level=1)
    for i, (short, _source, title, legend) in enumerate(SUPP_FIGURE_INFO, 1):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(FIGURES / f"{short}.png"), width=Inches(6.3))
        add_para(doc, f"Figure S{i}. {title}. {legend}", indent=False)
    path = OUT / "additional_file_1_supplementary_material_en.docx"; doc.save(path); return path


def checklist_location(item):
    mapping = {
        "1": "Title page", "2": "Abstract", "3a": "Background, paragraphs 1-4", "3b": "Background, paragraph 4",
        "3c": "Methods—Interpretability, fairness, and sensitivity analyses; Discussion, paragraph 6", "4": "Background, paragraph 4",
        "5a": "Methods—Study design, setting, and reporting", "5b": "Methods—Study design, setting, and reporting",
        "6a": "Methods—Study design, setting, and reporting", "6b": "Methods—Cohort definition", "6c": "Methods—Cohort definition; not an intervention study",
        "7": "Methods—Landmark datasets and predictors", "8a": "Methods—Baseline kidney function and outcome", "8b": "Not applicable; algorithmic laboratory outcome",
        "8c": "Not applicable; outcome derived before model fitting", "9a": "Methods—Landmark datasets and predictors", "9b": "Methods—Landmark datasets and predictors",
        "9c": "Not applicable; structured EHR predictors", "10": "Methods—Sample size and missing data", "11": "Methods—Sample size and missing data",
        "12a": "Methods—Model development and internal validation", "12b": "Methods—Sample size and missing data; Model development",
        "12c": "Methods—Model development and internal validation", "12d": "Methods—Performance assessment; single-center data",
        "12e": "Methods—Performance assessment", "12f": "Not applicable; no model updating", "12g": "Not applicable; no external model evaluation",
        "13": "Methods—Model development; no imbalance correction", "14": "Methods—Interpretability, fairness, and sensitivity analyses",
        "15": "Methods—Model development and internal validation", "16": "Methods—Model development; same source population",
        "17": "Declarations—Ethics approval", "18a": "Declarations—Funding", "18b": "Declarations—Competing interests",
        "18c": "Code and data availability statement; no prospective protocol", "18d": "Code and data availability statement; not registered",
        "18e": "Declarations—Availability of data and materials", "18f": "Declarations—Availability; repository URL and DOI to be provided before publication",
        "19": "Methods—Patient and public involvement", "20a": "Results—Cohort and incident AKI; Figure 1", "20b": "Results—Cohort and incident AKI; Table 1; Table S3",
        "20c": "Additional file 1: subgroup and model-performance tables", "21": "Results—Dynamic risk sets; Table 2 and Table S1",
        "22": "Code availability statement; public repository release planned before publication", "23a": "Results—Model performance and subgroup performance; Figures 2-3",
        "23b": "Not applicable; one center", "24": "Not applicable; no model updating", "25": "Discussion, paragraphs 1-6",
        "26": "Discussion, paragraph 8", "27a": "Methods—Missing data; Discussion, paragraphs 5 and 8", "27b": "Discussion, paragraph 5",
        "27c": "Discussion, paragraphs 5 and 8; Conclusions",
    }
    return mapping.get(item, "See manuscript")


def build_checklist():
    rows = []
    with (ASSETS / "tripod_ai_checklist_v8.tsv").open(encoding="utf-8") as handle:
        for row in csv.DictReader(handle, delimiter="\t"):
            row["location"] = checklist_location(row["item"])
            row["status"] = ("Pending repository identifier" if "to be provided before publication" in row["location"]
                             else ("Not applicable" if row["location"].startswith("Not applicable") else "Reported"))
            rows.append(row)
    csv_path = OUT / "additional_file_2_tripod_ai_checklist.csv"
    with csv_path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0])); writer.writeheader(); writer.writerows(rows)
    doc = Document(); configure_word(doc, "TRIPOD+AI checklist")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run("Additional file 2: TRIPOD+AI checklist"); set_font(r, 16, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(TITLE); set_font(r, 11, italic=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(AUTHORS_WORD); set_font(r, 10)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(AFFILIATION_WORD); set_font(r, 9)
    add_para(doc, "Checklist wording is a concise author-prepared rendering of the TRIPOD+AI development/evaluation checklist. Source: Collins et al., BMJ 2024;385:e078378. Page numbers should be refreshed after the journal submission system generates its final PDF.", indent=False)
    landscape(doc)
    add_table(doc, "", "", rows, [("section", "Section"), ("item", "Item"), ("checklist_item", "Reporting recommendation"),
                                    ("location", "Location in manuscript"), ("status", "Status")],
              [1200, 650, 4600, 4000, 1700], font=6.8)
    path = OUT / "additional_file_2_tripod_ai_checklist.docx"; doc.save(path); return path, csv_path


def build_availability_doc():
    text = AVAILABILITY_STATEMENT
    doc = Document(); configure_word(doc, "Code and data availability")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run("Code and data availability statement"); set_font(r, 16, bold=True)
    add_para(doc, text, indent=False)
    path = OUT / "code_and_data_availability_statement.docx"; doc.save(path)
    (OUT / "code_and_data_availability_statement.txt").write_text(text + "\n", encoding="utf-8")
    return path


def write_audits():
    with (OUT / "reference_audit.csv").open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.writer(handle); writer.writerow(["Number", "Citation key", "Vancouver reference", "DOI", "PMID", "Role", "Verification"])
        for i, (key, citation, doi, pmid, role) in enumerate(REFS, 1): writer.writerow([i, key, citation, doi, pmid, role, "Verified against PubMed, Crossref, or official source"])
    ledger = [
        ("acute kidney injury (AKI)", "Spell out at first use; AKI thereafter"), ("incident postoperative AKI", "Primary outcome; avoid 'renal failure'"),
        ("serum creatinine (SCr)", "Use SCr; units mg/dL"), ("intensive care unit (ICU)", "Use ICU after first use"),
        ("0-h, 6-h, and 24-h landmarks", "Hyphenate adjectival landmark terms"), ("logistic regression", "Lowercase regression"),
        ("random forest", "Lowercase except at sentence start"), ("XGBoost", "Preserve capitalization"), ("LightGBM", "Preserve capitalization"),
        ("AUROC", "Area under the receiver-operating-characteristic curve"), ("AUPRC", "Area under the precision-recall curve"),
        ("decision-curve analysis (DCA)", "Use net benefit, not clinical benefit"), ("SHAP", "Model attribution, not causal importance"),
        ("hemoglobin; orthopedic", "American English used throughout"),
    ]
    with (OUT / "terminology_ledger.csv").open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.writer(handle); writer.writerow(["Canonical term", "Usage rule"]); writer.writerows(ledger)
    readme = f"""# Critical Care manuscript package v8

Target journal assumed: Critical Care (Research article).

Author: Bizhi Wei, Pu Ai Medical School, Shaoyang University.

- Structured abstract: {abstract_word_count()} words (journal maximum: 350)
- Main text: {main_word_count()} words
- Main display items: 3 tables and 4 figures
- Additional file 1: Tables S1-S4 and Figures S1-S5
- Additional file 2: TRIPOD+AI checklist
- References: {len(REFS)}, Vancouver style
- Graphical abstract: 920 x 300 px, PNG/SVG/PDF

The package uses double spacing, continuous line numbering, page numbering, SI-compatible terminology, standardized American medical English, and separate figure files. Figure and table titles contain no more than 15 words.
"""
    (OUT / "README.md").write_text(readme, encoding="utf-8")
    (OUT / "author_action_required.md").write_text("""# Remaining submission actions

1. Create the public analytic-code repository and add its URL and archived DOI before publication.
2. Add ORCID identifiers in the submission system if available.
3. Update TRIPOD+AI page numbers after the submission system creates the final PDF.
4. Do not describe SHAP features as causal or modifiable determinants.
5. Do not describe the models as externally validated or ready for clinical deployment.
""", encoding="utf-8")


def main():
    if OUT.exists(): shutil.rmtree(OUT)
    for path in [OUT, LATEX, FIGURES, TABLES, QA]: path.mkdir(parents=True, exist_ok=True)
    copy_figures(); tables = standardize_tables(); build_latex(tables)
    build_word_main(tables); build_word_supplement(tables); build_checklist(); build_availability_doc(); write_audits()
    print(f"Package: {OUT}")
    print(f"Abstract words: {abstract_word_count()}; main-text words: {main_word_count()}; references: {len(REFS)}")


if __name__ == "__main__":
    main()
