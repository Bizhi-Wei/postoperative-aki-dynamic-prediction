"""Build the revised v30 severe-AKI and renal-trajectory manuscript package.

The v29 package and locked primary manuscript remain unchanged.  This revision
adds a true cohort/analysis flow figure, a severe-AKI baseline Table 1, paired
same-risk-set incremental-value analysis, an explicit external observability
denominator, and transparent MIMIC-IV/eICU feature-harmonization documentation.
"""

from __future__ import annotations

import csv
import importlib.util
import re
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]

spec = importlib.util.spec_from_file_location("v29_builder", ROOT / "scripts" / "build_secondary_manuscript_v29.py")
m = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(m)

spec_t1 = importlib.util.spec_from_file_location("table1_builder", ROOT / "scripts" / "prepare_final_manuscript_v6.py")
t1mod = importlib.util.module_from_spec(spec_t1)
assert spec_t1.loader is not None
spec_t1.loader.exec_module(t1mod)


OUT = ROOT / "outputs" / "manuscript_package_v30_secondary_revised"
LATEX = OUT / "latex"
TABLES = OUT / "tables"
FIGURES = OUT / "figures"
QA = OUT / "qa"
SOURCE_DATA = OUT / "source_data"
V30I = ROOT / "outputs" / "modeling_v30_severe_same_risk_incremental"
V30F = ROOT / "outputs" / "manuscript_figure_v30_secondary_flow"

for name, value in {"OUT": OUT, "LATEX": LATEX, "TABLES": TABLES, "FIGURES": FIGURES, "QA": QA}.items():
    setattr(m, name, value)
    setattr(m.v8, name, value)


TITLE = m.TITLE
RUNNING_TITLE = m.RUNNING_TITLE
AUTHORS_WORD = m.AUTHORS_WORD
AFFILIATION_WORD = m.AFFILIATION_WORD
CORRESPONDENCE_WORD = m.CORRESPONDENCE_WORD
ETHICS = m.ETHICS
FUNDING = m.FUNDING
COMPETING = m.COMPETING
CONTRIBUTIONS = m.CONTRIBUTIONS
ACKNOWLEDGEMENTS = m.ACKNOWLEDGEMENTS
AI_DISCLOSURE = m.AI_DISCLOSURE
AVAILABILITY = (
    "MIMIC-IV version 3.1 and the eICU Collaborative Research Database are available through PhysioNet to "
    "credentialed users who complete the required training and sign the data use agreement. The author is not "
    "permitted to redistribute patient-level data or derived patient-level analytic datasets. Analytic code is "
    "publicly available at https://github.com/Bizhi-Wei/postoperative-aki-dynamic-prediction (release v1.0.2: "
    "https://github.com/Bizhi-Wei/postoperative-aki-dynamic-prediction/releases/tag/v1.0.2). Archived software DOI: "
    "to be provided before publication."
)


ABSTRACT = dict(m.ABSTRACT)
ABSTRACT["Methods"] = (
    "This retrospective secondary study used a strict surgical intensive care cohort from MIMIC-IV version 3.1 "
    "and feature-harmonized external validation in eICU. Severe AKI was a new Kidney Disease: Improving Global "
    "Outcomes serum creatinine stage 2 or 3 event after ICU admission, 6 h, or 24 h through day 7. Logistic "
    "regression and XGBoost used patient-grouped development/test splits. A paired same-risk-set analysis refitted "
    "models in identical 6-h and 24-h risk sets while changing only the information horizon. A seven-state, "
    "observation-time multistate analysis treated live discharge and death as competing absorbing events."
)
ABSTRACT["Results"] = (
    "Among 10,877 postoperative ICU admissions, 679 (6.2%) developed severe serum-creatinine-defined AKI. Held-out "
    "AUROCs were 0.698 (95% CI, 0.653–0.743) at 0 h, 0.790 (0.745–0.829) at 6 h, and 0.839 (0.800–0.879) at 24 h. "
    "Within fixed risk sets, adding 0–6-h information improved 6-h AUROC by 0.101 (0.064–0.138), and adding 0–24-h "
    "rather than 0–6-h information improved 24-h AUROC by 0.079 (0.045–0.119). Of 30,365 strict surgical eICU stays, "
    "14,229 (46.9%) were outcome evaluable; external AUROCs were 0.707, 0.761, and 0.784, and hospital-held-out "
    "recalibration brought slopes to 0.95, 0.98, and 0.95. Among 4,519 trajectory-eligible AKI admissions, the 48-h "
    "cumulative incidences of observed recovery and severe progression were 65.0% and 11.9%."
)
ABSTRACT["Conclusions"] = (
    "Postoperative information accrued through 6 h and 24 h added severe-AKI discrimination within fixed risk sets, "
    "and frozen feature-harmonized models retained useful external ranking, although local probability updating was "
    "needed. Recovery and recurrence were measurement-dependent states shaped by competing discharge and death. "
    "Prospective validation is required before clinical use."
)
KEYWORDS = m.KEYWORDS
INTRODUCTION = m.INTRODUCTION

METHODS = {key: list(value) for key, value in m.METHODS.items()}
METHODS["Landmark and onset-anchored prediction models"] = [
    METHODS["Landmark and onset-anchored prediction models"][0],
    (
        "To isolate the incremental value of accumulating information from changes in risk-set composition, we "
        "performed a paired same-risk-set analysis. Within the 6-h risk set, admission-only models were compared with "
        "models using information through 6 h. Within the 24-h risk set, admission-only, 0–6-h, and 0–24-h models were "
        "refitted and compared. Within each target risk set, patients, outcome window, subject split, and model family "
        "were identical; only the predictor information horizon changed. Paired differences used 1,000 subject-cluster "
        "bootstrap resamples."
    ),
    METHODS["Landmark and onset-anchored prediction models"][1],
    METHODS["Landmark and onset-anchored prediction models"][2],
]
METHODS["Temporal and external validation"] = [
    (
        "Rolling temporal validation used expanding earlier-year MIMIC-IV training data and three nonoverlapping later-year "
        "evaluation blocks per landmark. For eICU, 30,365 strict surgical first ICU stays from 197 hospitals were identified "
        "using operative organ-system admission diagnoses; 14,229 (46.9%) met the SCr outcome-observability definition at "
        "admission. Baseline, incident SCr-AKI, and landmark risk sets were rederived using the same temporal logic. Only the "
        "three severe SCr-AKI landmark models were externally validated; the onset-anchored persistence and nonrecovery "
        "models remained internal secondary analyses. Frozen feature-harmonized models used 30 predictors at 0 h and 72 at "
        "6 h and 24 h. The cross-database mapping is reported explicitly in Supplementary Table S9."
    ),
    METHODS["Temporal and external validation"][1],
]


RESULTS = {key: list(value) for key, value in m.RESULTS.items()}
RESULTS["Cohort and secondary outcomes"] = [
    (
        "The strict evaluable cohort comprised 10,877 postoperative ICU admissions, of which 679 (6.2%) developed severe "
        "active-episode SCr-AKI and 10,198 did not. Baseline characteristics by severe-AKI status are shown in Table 1. "
        "Incident SCr-AKI of any stage occurred in 4,531 admissions (41.7%), and the SCr-or-RRT severe-outcome sensitivity "
        "definition occurred in 718 admissions (Fig. 1)."
    ),
    (
        "The severe-AKI risk sets contained 10,877 admissions at 0 h, 10,856 at 6 h after 21 prevalent severe events were "
        "removed, and 10,736 at 24 h after 141 cumulative prevalent severe events were removed. Subsequent severe-event "
        "rates were 6.2%, 6.1%, and 5.0%, respectively (Fig. 1; Supplementary Table S1)."
    ),
]
SEVERE_MODEL_RESULTS = [
    (
        "In held-out patients, the selected XGBoost models at 0 h and 6 h and logistic-regression model at 24 h had AUROCs "
        "of 0.698 (95% CI, 0.653–0.743), 0.790 (0.745–0.829), and 0.839 (0.800–0.879), respectively (Table 2). "
        "Corresponding AUPRCs were 0.174, 0.313, and 0.340, and Brier scores were 0.057, 0.049, and 0.041. These estimates "
        "apply to different landmark-specific risk sets and remaining prediction windows."
    ),
    (
        "The paired same-risk-set analysis confirmed genuine incremental value of later information (Fig. 2; Supplementary "
        "Table S2). In the identical 6-h held-out risk set (n=2,162; 132 events), XGBoost AUROC increased from 0.692 with "
        "admission information to 0.794 with 0–6-h information (paired difference, 0.101; 95% CI, 0.064–0.138). In the "
        "identical 24-h held-out risk set (n=2,142; 112 events), logistic-regression AUROC increased from 0.760 with 0–6-h "
        "information to 0.839 with 0–24-h information (difference, 0.079; 0.045–0.119). Compared with admission information, "
        "the total 24-h gain was 0.145 (0.102–0.191)."
    ),
]
RESULTS["Temporal stability, external validation, and recalibration"] = [
    (
        "Across three expanding-window temporal evaluations, AUROCs ranged from 0.701 to 0.731 at 0 h, 0.762 to 0.803 at 6 h, "
        "and 0.807 to 0.847 at 24 h (Supplementary Table S4). Of 30,365 strict surgical eICU first ICU stays, 14,229 (46.9%) "
        "were outcome evaluable at admission and 16,136 were not (Fig. 1). Severe SCr-AKI occurred in 910 of 14,229 admissions "
        "at 0 h, 836 of 14,155 at 6 h, and 538 of 13,857 at 24 h. Frozen external AUROCs were 0.707 (95% CI, 0.688–0.725), "
        "0.761 (0.744–0.780), and 0.784 (0.760–0.808), respectively (Fig. 3; Table 3). These external results apply only to "
        "the severe-AKI landmark models."
    ),
    (
        "External calibration slopes were 0.79, 0.88, and 0.51. In hospitals held out from probability updating, logistic "
        "recalibration changed slopes from 0.76 to 0.95 at 0 h, 0.86 to 0.98 at 6 h, and 0.49 to 0.95 at 24 h; Brier scores "
        "changed from 0.057 to 0.055, 0.052 to 0.050, and 0.036 to 0.036. Across hospitals meeting minimum estimability "
        "criteria, median AUROCs were 0.741, 0.777, and 0.803 (Supplementary Table S6)."
    ),
]
RESULTS["Onset-anchored persistence and nonrecovery models"] = [
    RESULTS["Onset-anchored persistence and nonrecovery models"][0].replace("(Table 2)", "(Supplementary Table S3)"),
    RESULTS["Onset-anchored persistence and nonrecovery models"][1].replace("Supplementary Table S6", "Supplementary Table S7"),
]
RESULTS["Observation-time renal trajectories and competing events"] = [
    RESULTS["Observation-time renal trajectories and competing events"][0],
    RESULTS["Observation-time renal trajectories and competing events"][1].replace("Supplementary Figure S5", "Supplementary Figure S7"),
]
RESULTS = {
    "Cohort and secondary outcomes": RESULTS["Cohort and secondary outcomes"],
    "Severe-AKI model performance and incremental information value": SEVERE_MODEL_RESULTS,
    "Temporal stability, external validation, and recalibration": RESULTS["Temporal stability, external validation, and recalibration"],
    "Onset-anchored persistence and nonrecovery models": RESULTS["Onset-anchored persistence and nonrecovery models"],
    "Observation-time renal trajectories and competing events": RESULTS["Observation-time renal trajectories and competing events"],
}


DISCUSSION = list(m.DISCUSSION)
DISCUSSION[0] = (
    "This secondary study links three clinically distinct questions: who will progress to severe postoperative SCr-AKI, "
    "whether accumulating postoperative information adds predictive value in the same at-risk patients, and how observed "
    "renal states evolve when discharge and death are treated explicitly. The paired analysis showed that the higher 6-h "
    "and 24-h discrimination was not explained only by changing risk-set composition. Frozen feature-harmonized models "
    "retained useful ranking in eICU, but external calibration required local updating."
)
DISCUSSION[1] = (
    "The severe-AKI results complement studies that predict any AKI or moderate-to-severe AKI after surgery "
    "[[CITE:Kheterpal2007,Tseng2020,Ryan2023,Demirjian2022]]. Within a fixed 6-h risk set, measurements accrued through "
    "6 h improved AUROC by 0.101; within a fixed 24-h risk set, completing the 0–24-h information window improved AUROC "
    "by a further 0.079 relative to 0–6-h information. These paired estimates support real information gain, while the "
    "original landmark estimates retain their intended interpretation as conditional predictions in changing risk sets. "
    "The 6-h model may offer a more actionable compromise between timeliness and discrimination."
)
DISCUSSION[2] = (
    "External validation was a major strength but also exposed selection and transport limitations. Only 46.9% of strict "
    "surgical eICU stays were SCr-outcome evaluable, so the estimates apply to a measured subset rather than every surgical "
    "ICU patient. The three severe-AKI landmark models—not the onset-anchored recovery models—were externally validated. "
    "Similar eICU AUROCs suggest that the portable feature set preserved ranking, whereas calibration slopes, especially "
    "at 24 h, showed that the probability scale was not transportable without adjustment [[CITE:VanCalster2019]]."
)
DISCUSSION[-1] = (
    "Limitations remain. MIMIC-IV is a single-center development source dominated by cardiac surgery, and only 46.9% of the "
    "strict eICU surgical cohort was outcome evaluable. Cross-database feature harmonization used broad diagnosis-text and "
    "operative-system proxies and an approximate comorbidity score, so semantic equivalence was incomplete. RRT ascertainment "
    "was sensitive to documentation and chronic dialysis. Primary trajectory states used SCr without urine output, procedure "
    "timing was date-level, and follow-up ended at 7 days. Recovery definitions were measurement-dependent; intermittent "
    "observation may violate simple Markov interpretations. Model-family choices were made after internal comparisons before "
    "external evaluation, and no prospective silent validation was performed. Cluster bootstrap intervals do not remove "
    "residual selection bias, and predictive associations must not be interpreted as modifiable causal effects."
)
CONCLUSION = (
    "Information accrued through 6 h and 24 h improved severe postoperative SCr-AKI prediction within fixed risk sets, and "
    "feature-harmonized models retained useful external ranking among outcome-evaluable eICU patients, although local "
    "calibration updating remained necessary. After AKI onset, observed recovery was common but recurrence was not rare; "
    "discharge, death, and selective SCr measurement shaped what could be observed. Prospective silent validation with "
    "prespecified monitoring and response pathways is required before clinical implementation."
)


for name, value in {
    "TITLE": TITLE, "RUNNING_TITLE": RUNNING_TITLE, "AUTHORS_WORD": AUTHORS_WORD,
    "AFFILIATION_WORD": AFFILIATION_WORD, "CORRESPONDENCE_WORD": CORRESPONDENCE_WORD,
    "ABSTRACT": ABSTRACT, "KEYWORDS": KEYWORDS, "INTRODUCTION": INTRODUCTION,
    "METHODS": METHODS, "RESULTS": RESULTS, "DISCUSSION": DISCUSSION, "CONCLUSION": CONCLUSION,
    "AVAILABILITY": AVAILABILITY,
}.items():
    setattr(m, name, value)


FIGURE_SOURCES = {
    "Fig1": V30F / "figure_v30_cohort_analysis_flow",
    "Fig2": V30I / "figure_v30_paired_incremental_value",
    "Fig3": m.V28S / "figure_v28_severe_temporal_external_validation",
    "Fig4": m.V29 / "figure_v29_multistate_competing_risk",
    "FigS1": m.V26 / "figure_v26_severity_recovery_phenotypes",
    "FigS2": m.V26 / "figure_v26_daily_scr_state_trajectories",
    "FigS3": m.V27 / "figure_v27_severe_aki_dynamic_models",
    "FigS4": m.V27 / "figure_v27_onset_anchored_trajectory_models",
    "FigS5": m.V28S / "figure_v28_severe_external_decision_curve",
    "FigS6": m.V28R / "figure_v28_recovery_observability_ipw_competing",
    "FigS7": m.V29 / "figure_v29_competing_risk_subgroups",
}
FIGURE_LEGENDS = [
    ("Figure 1", "Cohort and analysis flow", "Panel a traces the MIMIC-IV severe-AKI landmark risk sets and the separate post-AKI trajectory pathway. Panel b reports the full strict eICU surgical denominator, SCr outcome-observability selection, and external landmark risk sets."),
    ("Figure 2", "Incremental predictive value of later information within fixed severe-AKI risk sets", "Points are paired changes in AUROC, AUPRC, and Brier score when later information is added while patients, outcome, split, and model family remain identical. Error bars are 95% subject-cluster bootstrap confidence intervals."),
    ("Figure 3", "Temporal and eICU external validation of severe-AKI landmark models", "Panels summarize MIMIC-IV rolling temporal validation, frozen multicenter eICU discrimination, hospital heterogeneity, and held-out-hospital calibration before and after logistic recalibration. Onset-anchored recovery models were not externally validated."),
    m.FIGURE_LEGENDS[3],
]
SUPP_FIGURE_LEGENDS = [
    ("Figure S1", "Severity and recovery phenotypes", "Recovery requires an observed non-AKI SCr; absence of measurement is not recovery."),
    ("Figure S2", "Daily observed SCr-state trajectories", "Daily state summaries include an explicit unobserved category."),
    ("Figure S3", "Held-out performance of dynamic severe-AKI models", "Landmark risk sets exclude severe SCr-AKI already present at each prediction time."),
    ("Figure S4", "Onset-anchored persistence and nonrecovery models", "Predictors are restricted to information available by observed AKI onset; these models had internal validation only."),
    ("Figure S5", "External decision-curve analysis for severe AKI", "Net benefit is exploratory and does not establish an intervention threshold."),
    ("Figure S6", "Recovery observability, IPW, and competing-event sensitivity", "IPW assumes missing at random conditional on observed onset covariates."),
    ("Figure S7", "Competing-risk estimates across prespecified subgroups", "Subgroup curves are descriptive and use subject-cluster bootstrap intervals."),
]
m.FIGURE_SOURCES = FIGURE_SOURCES
m.FIGURE_LEGENDS = FIGURE_LEGENDS
m.SUPP_FIGURE_LEGENDS = SUPP_FIGURE_LEGENDS


def words(text: str) -> int:
    return len(re.findall(r"\b[\w–-]+\b", re.sub(r"\[\[CITE:[^\]]+\]\]", "", text)))


def abstract_word_count() -> int:
    return words(" ".join(ABSTRACT.values()))


def main_word_count() -> int:
    blocks = INTRODUCTION + [p for ps in METHODS.values() for p in ps] + [p for ps in RESULTS.values() for p in ps] + DISCUSSION + [CONCLUSION]
    return words(" ".join(blocks))


def csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def write_csv(name: str, rows: list[dict[str, object]]) -> None:
    if not rows:
        raise ValueError(f"No rows for {name}")
    with (TABLES / name).open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0])); writer.writeheader(); writer.writerows(rows)


def build_tables() -> dict[str, list[dict[str, object]]]:
    cohort = pd.read_csv(m.V26 / "cohort_v26_strict_aki_severity_recovery.csv", low_memory=False)
    cohort["aki_final"] = cohort["severe_aki_scr_stage2_3"].fillna(False).astype(bool)
    t1df = t1mod.build_table1(cohort).rename(columns={
        "No AKI (N=6,346)": "No severe SCr-AKI (N=10,198)",
        "Incident AKI (N=4,531)": "Severe SCr-AKI (N=679)",
    })
    t1 = t1df.to_dict("records")

    perf = pd.read_csv(m.V27 / "model_v27_performance_summary.csv")
    selected = [(0, "XGBoost"), (6, "XGBoost"), (24, "Logistic Regression")]
    t2 = []
    for lm, model in selected:
        r = perf[(perf.task.eq("severe_scr")) & (perf.landmark.astype(str).eq(str(lm))) & perf.model.eq(model)].iloc[0]
        t2.append({
            "Landmark": f"{lm} h", "Selected model": model.replace("Regression", "regression"),
            "Test n/events": f"{int(r.test_n):,}/{int(r.test_event_n):,}",
            "AUROC (95% CI)": m.ci(r.auroc, r.auroc_ci95_low, r.auroc_ci95_high),
            "AUPRC (95% CI)": m.ci(r.auprc, r.auprc_ci95_low, r.auprc_ci95_high),
            "Brier": m.f3(r.brier_score), "Calibration intercept/slope": f"{r.calibration_intercept:.2f}/{r.calibration_slope:.2f}",
        })

    ext = pd.read_csv(m.V28S / "model_v28_eicu_frozen_severe_performance.csv")
    recal = pd.read_csv(m.V28S / "model_v28_heldout_hospital_recalibration_performance.csv")
    t3 = []
    for lm in [0, 6, 24]:
        e = ext[(ext.landmark_hours.eq(lm)) & ext.target.eq("SCr stage 2/3")].iloc[0]
        frozen = recal[(recal.landmark_hours.eq(lm)) & recal.method.eq("frozen")].iloc[0]
        updated = recal[(recal.landmark_hours.eq(lm)) & recal.method.eq("logistic recalibration")].iloc[0]
        t3.append({
            "Landmark": f"{lm} h", "External n/events": f"{int(e.n):,}/{int(e.event_n):,}",
            "External AUROC (95% CI)": m.ci(e.auroc, e.auroc_ci95_low, e.auroc_ci95_high),
            "External AUPRC/Brier": f"{e.auprc:.3f}/{e.brier_score:.3f}",
            "External intercept/slope": f"{e.calibration_intercept:.2f}/{e.calibration_slope:.2f}",
            "Held-out n": f"{int(frozen.n):,}", "Frozen Brier/slope": f"{frozen.brier_score:.3f}/{frozen.calibration_slope:.2f}",
            "Updated Brier/slope": f"{updated.brier_score:.3f}/{updated.calibration_slope:.2f}",
        })

    aki_cif = pd.read_csv(m.V29 / "analysis_v29_cif_after_aki.csv")
    rec_cif = pd.read_csv(m.V29 / "analysis_v29_cif_after_recovery.csv")
    t4 = []
    for source, origin in [(aki_cif, "Observed AKI onset"), (rec_cif, "First observed recovery")]:
        for r in source[(source.group_variable.eq("Overall")) & source.time_hours.isin([24, 48, 72])].itertuples():
            t4.append({"Time origin": origin, "Time": f"{int(r.time_hours)} h", "Event": r.cause,
                       "Cumulative incidence, % (95% CI)": f"{r.cif_percent:.1f} ({r.ci95_low_percent:.1f}–{r.ci95_high_percent:.1f})",
                       "Risk set n": f"{int(r.risk_set_weighted_n):,}"})

    risk = pd.read_csv(m.V27 / "audit_v27_severe_risk_set_summary.csv")
    s1 = [{"Landmark": f"{int(r.landmark_hours)} h", "Risk set n": f"{int(r.risk_set_n):,}",
           "Already severe excluded": f"{int(r.already_severe_excluded_n):,}",
           "Severe SCr-AKI events": f"{int(r.severe_scr_event_n):,} ({r.severe_scr_event_percent:.1f}%)",
           "SCr-or-RRT events": f"{int(r.severe_scr_or_rrt_event_n):,} ({r.severe_scr_or_rrt_event_percent:.1f}%)"} for r in risk.itertuples()]

    incperf = pd.read_csv(V30I / "model_v30_severe_same_risk_performance.csv")
    deltas = pd.read_csv(V30I / "model_v30_severe_paired_incremental_deltas.csv")
    s2 = []
    for risk_set, model in [(6, "XGBoost"), (24, "Logistic Regression")]:
        rows = incperf[(incperf.risk_set_hours.eq(risk_set)) & incperf.model.eq(model)].sort_values("information_hours")
        for r in rows.itertuples():
            delta_text = "Reference"
            if r.information_hours > 0:
                ref = 0 if r.information_hours == 6 else 6
                d = deltas[(deltas.risk_set_hours.eq(risk_set)) & deltas.model.eq(model)
                           & deltas.new_information_hours.eq(r.information_hours)
                           & deltas.reference_information_hours.eq(ref)].iloc[0]
                delta_text = f"{d.delta_auroc:+.3f} ({d.delta_auroc_ci_lower:+.3f} to {d.delta_auroc_ci_upper:+.3f})"
            s2.append({"Risk set": f"{risk_set} h", "Information available": r.information_set,
                       "Model": model.replace("Regression", "regression"), "Test n/events": f"{int(r.test_n):,}/{int(r.test_event_n):,}",
                       "AUROC (95% CI)": m.ci(r.auroc, r.auroc_ci_lower, r.auroc_ci_upper),
                       "AUPRC (95% CI)": m.ci(r.auprc, r.auprc_ci_lower, r.auprc_ci_upper),
                       "Brier (95% CI)": m.ci(r.brier_score, r.brier_score_ci_lower, r.brier_score_ci_upper),
                       "Delta AUROC vs preceding information": delta_text})

    s3 = [{"Task": r.task, "Landmark": str(r.landmark) + (" h" if str(r.landmark).isdigit() else ""),
           "Model": str(r.model), "Predictors": int(r.predictor_n), "Test n/events": f"{int(r.test_n):,}/{int(r.test_event_n):,}",
           "AUROC (95% CI)": m.ci(r.auroc, r.auroc_ci95_low, r.auroc_ci95_high),
           "AUPRC (95% CI)": m.ci(r.auprc, r.auprc_ci95_low, r.auprc_ci95_high), "Brier": m.f3(r.brier_score),
           "Intercept/slope": f"{r.calibration_intercept:.2f}/{r.calibration_slope:.2f}"} for r in perf.itertuples()]

    rolling = pd.read_csv(m.V28S / "analysis_v28_mimic_rolling_severe_validation.csv")
    s4 = [{"Landmark": f"{int(r.landmark_hours)} h", "Validation years": f"{int(r.test_year_start)}–{int(r.test_year_end)}",
           "Train n": f"{int(r.train_n):,}", "Test n/events": f"{int(r.test_n):,}/{int(r.test_event_n):,}", "Model": r.model,
           "AUROC": m.f3(r.auroc), "AUPRC": m.f3(r.auprc), "Brier": m.f3(r.brier_score),
           "Intercept/slope": f"{r.calibration_intercept:.2f}/{r.calibration_slope:.2f}"} for r in rolling.itertuples()]

    s5 = [{"Landmark": f"{int(r.landmark_hours)} h", "Target": r.target, "Model": r.model,
           "n/events": f"{int(r.n):,}/{int(r.event_n):,}", "Event rate": f"{r.event_percent:.1f}%",
           "AUROC (95% CI)": m.ci(r.auroc, r.auroc_ci95_low, r.auroc_ci95_high),
           "AUPRC (95% CI)": m.ci(r.auprc, r.auprc_ci95_low, r.auprc_ci95_high), "Brier": m.f3(r.brier_score),
           "Intercept/slope": f"{r.calibration_intercept:.2f}/{r.calibration_slope:.2f}"} for r in ext.itertuples()]

    hetero = pd.read_csv(m.V28S / "analysis_v28_hospital_heterogeneity_summary.csv")
    s6 = [{"Landmark": f"{int(r.landmark_hours)} h", "Hospitals": int(r.hospital_n),
           "AUROC-evaluable hospitals": int(r.hospital_n_auroc_evaluable),
           "Hospital AUROC median (IQR)": f"{r.hospital_auroc_median:.3f} ({r.hospital_auroc_q1:.3f}–{r.hospital_auroc_q3:.3f})",
           "Hospital AUROC range": f"{r.hospital_auroc_min:.3f}–{r.hospital_auroc_max:.3f}"} for r in hetero.itertuples()]

    obs = pd.read_csv(m.V28R / "model_v28_observability_performance.csv")
    ipw = pd.read_csv(m.V28R / "model_v28_recovery_ipw_performance.csv")
    s7 = []
    for r in obs.itertuples():
        s7.append({"Section": "Outcome evaluability", "Target": r.evaluable_definition, "Analysis": "Cross-fitted observability model",
                   "n/events": f"{int(r.n):,}/{int(r.event_n):,}", "Event/weighted rate": f"{r.event_percent_unweighted:.1f}%",
                   "AUROC": m.f3(r.auroc), "AUPRC": m.f3(r.auprc), "Brier": m.f3(r.brier_score)})
    for r in ipw.itertuples():
        s7.append({"Section": "Outcome model", "Target": r.task, "Analysis": r.analysis,
                   "n/events": f"{int(r.n):,}/{int(r.event_n):,}", "Event/weighted rate": f"{r.event_percent_weighted:.1f}%",
                   "AUROC": m.f3(r.auroc), "AUPRC": m.f3(r.auprc), "Brier": m.f3(r.brier_score)})

    s8 = []
    for source, origin in [(aki_cif, "Observed AKI onset"), (rec_cif, "First observed recovery")]:
        for r in source[(source.group_variable.ne("Overall")) & source.time_hours.isin([48, 72])].itertuples():
            s8.append({"Time origin": origin, "Subgroup variable": r.group_variable, "Subgroup": r.group, "n": f"{int(r.n):,}",
                       "Time": f"{int(r.time_hours)} h", "Event": r.cause,
                       "CIF, % (95% CI)": f"{r.cif_percent:.1f} ({r.ci95_low_percent:.1f}–{r.ci95_high_percent:.1f})"})

    mapping = pd.read_csv(m.V28S / "audit_v28_severe_feature_harmonization.csv")
    groups = [
        ("Demographics", "gender|anchor_age", "Patient demographics", "patient.gender and patient.age; age >89 coded as 90"),
        ("Comorbidities", "chf|hypertension|dm|ckd|copd|liver|cancer|pvd|stroke|mi|obesity|anemia|charlson_score", "Diagnosis-derived comorbidity flags and Charlson score", "diagnosis.diagnosisstring plus ICD-9 text-pattern flags; transparent sum of broad flags"),
        ("Surgical system", "cardiac_surgery|non_cardiac_surgery|general_gi_hepatobiliary_surgery|orthopedic_major_surgery|neurosurgery|thoracic_respiratory_surgery", "Therapeutic ICD procedure categories", "admissionDx operative organ system: Cardiovascular, Gastrointestinal, Musculoskeletal/Skin, Neurologic, or Respiratory"),
        ("Baseline SCr", "baseline_scr_at_landmark", "Lowest SCr in 7 d before ICU; index/admission fallback", "lab creatinine using offsets; lowest strictly pre-ICU value in 7 d, offset-0 fallback"),
        ("Pre-index laboratory values", "lab_pre24h_", "Last LABEVENTS result in -24 to 0 h", "Last eICU lab result in -24 to 0 h for the mapped analyte"),
        ("Landmark laboratory values", "lab_0_", "Last/minimum/maximum LABEVENTS value through landmark", "Last/minimum/maximum eICU lab value through landmark; BUN, creatinine, hemoglobin, lactate, WBC, platelets, potassium, sodium"),
        ("Landmark vital signs", "vital_0_", "Last/minimum/maximum CHARTEVENTS value through landmark", "Last/minimum/maximum vitalPeriodic value through landmark; MAP, heart rate, systolic BP, SpO2"),
        ("Time-updated SCr status", "prior_stage1_aki_by_landmark|hours_since_first_aki_at_landmark|current_scr_stage_at_landmark|current_scr_at_landmark|current_scr_ratio_at_landmark|scr_measurement_n_by_landmark", "Timestamped SCr states through landmark", "Offset-aligned SCr states through landmark using the same active-episode logic"),
    ]
    s9 = []
    for label, pattern, mimic_source, eicu_source in groups:
        part = mapping[mapping.predictor.str.contains(pattern, regex=True, na=False)]
        landmarks = ", ".join(f"{int(x)} h" for x in sorted(part.landmark_hours.unique()))
        s9.append({"Feature group": label, "Landmarks": landmarks, "Predictor entries": int(len(part)),
                   "MIMIC-IV operational source": mimic_source, "eICU operational source": eicu_source,
                   "MIMIC missingness range": f"{part.mimic_missing_percent.min():.1f}–{part.mimic_missing_percent.max():.1f}%",
                   "eICU missingness range": f"{part.eicu_missing_percent.min():.1f}–{part.eicu_missing_percent.max():.1f}%"})

    tables = {"T1": t1, "T2": t2, "T3": t3, "T4": t4, "S1": s1, "S2": s2, "S3": s3,
              "S4": s4, "S5": s5, "S6": s6, "S7": s7, "S8": s8, "S9": s9}
    filenames = {
        "T1": "Table_1_baseline_characteristics_by_severe_aki.csv", "T2": "Table_2_selected_severe_aki_models.csv",
        "T3": "Table_3_external_validation_recalibration.csv", "T4": "Table_4_multistate_competing_risk.csv",
        "S1": "Table_S1_severe_risk_sets.csv", "S2": "Table_S2_same_risk_incremental_value.csv",
        "S3": "Table_S3_all_secondary_models.csv", "S4": "Table_S4_rolling_temporal_validation.csv",
        "S5": "Table_S5_eicu_external_sensitivity.csv", "S6": "Table_S6_hospital_heterogeneity.csv",
        "S7": "Table_S7_recovery_observability_ipw.csv", "S8": "Table_S8_subgroup_competing_risk.csv",
        "S9": "Table_S9_feature_harmonization.csv",
    }
    for key, filename in filenames.items():
        write_csv(filename, tables[key])
    return tables


TABLE_SPECS = dict(m.TABLE_SPECS)
TABLE_SPECS.update({
    "T1": ("Baseline characteristics by severe SCr-AKI status", "tab:t1",
           [("Characteristic", "Characteristic"), ("Overall (N=10,877)", "Overall"),
            ("No severe SCr-AKI (N=10,198)", "No severe SCr-AKI"), ("Severe SCr-AKI (N=679)", "Severe SCr-AKI"),
            ("Standardized mean difference", "SMD"), ("Missing, n", "Missing")],
           [2700, 1700, 1900, 1800, 1100, 900], True, 7.4, "Values are median [IQR] or n (%). SMD compares severe with nonsevere SCr-AKI."),
    "T2": ("Held-out performance of selected severe-AKI landmark models", "tab:t2",
           [("Landmark", "Landmark"), ("Selected model", "Model"), ("Test n/events", "Test n/events"),
            ("AUROC (95% CI)", "AUROC (95% CI)"), ("AUPRC (95% CI)", "AUPRC (95% CI)"),
            ("Brier", "Brier"), ("Calibration intercept/slope", "Intercept/slope")],
           [1000, 1700, 1300, 2100, 2100, 900, 1600], True, 7.5, "The three severe-AKI landmark models were externally validated; confidence intervals use 1,000 subject-cluster bootstrap resamples."),
    "S2": ("Paired same-risk-set incremental value for severe-AKI prediction", "tab:s2",
           [("Risk set", "Risk set"), ("Information available", "Information"), ("Model", "Model"), ("Test n/events", "Test n/events"),
            ("AUROC (95% CI)", "AUROC (95% CI)"), ("AUPRC (95% CI)", "AUPRC (95% CI)"),
            ("Brier (95% CI)", "Brier (95% CI)"), ("Delta AUROC vs preceding information", "Paired delta AUROC")],
           [900, 1600, 1600, 1200, 1800, 1800, 1800, 2200], True, 6.7, "Each comparison refits the same model family in the identical target risk set; paired intervals use 1,000 subject-cluster bootstrap resamples."),
    "S3": ("Performance of all secondary prediction models", "tab:s3", m.TABLE_SPECS["S2"][2], m.TABLE_SPECS["S2"][3], True, 6.8, "Onset-anchored persistence and nonrecovery models received internal validation only."),
    "S4": ("Rolling temporal validation of severe-AKI models", "tab:s4", m.TABLE_SPECS["S3"][2], m.TABLE_SPECS["S3"][3], True, 7.0, m.TABLE_SPECS["S3"][6]),
    "S5": ("Full eICU external validation and SCr-or-RRT sensitivity", "tab:s5", m.TABLE_SPECS["S4"][2], m.TABLE_SPECS["S4"][3], True, 6.8, m.TABLE_SPECS["S4"][6]),
    "S6": ("Hospital-level heterogeneity of frozen eICU discrimination", "tab:s6", m.TABLE_SPECS["S5"][2], m.TABLE_SPECS["S5"][3], True, 8.0, m.TABLE_SPECS["S5"][6]),
    "S7": ("Recovery observability and inverse-probability-weighted sensitivity", "tab:s7", m.TABLE_SPECS["S6"][2], m.TABLE_SPECS["S6"][3], True, 7.0, m.TABLE_SPECS["S6"][6]),
    "S8": ("Subgroup cumulative incidence after AKI onset and recovery", "tab:s8", m.TABLE_SPECS["S8"][2], m.TABLE_SPECS["S8"][3], True, 6.8, m.TABLE_SPECS["S8"][6]),
    "S9": ("Operational feature harmonization between MIMIC-IV and eICU", "tab:s9",
           [("Feature group", "Feature group"), ("Landmarks", "Landmarks"), ("Predictor entries", "Entries"),
            ("MIMIC-IV operational source", "MIMIC-IV source"), ("eICU operational source", "eICU source"),
            ("MIMIC missingness range", "MIMIC missing"), ("eICU missingness range", "eICU missing")],
           [1500, 1000, 800, 2800, 3600, 1300, 1300], True, 6.5,
           "Mapping is operational rather than asserting exact semantic equivalence. Predictor entries count repeated landmark-specific fields."),
})
m.TABLE_SPECS = TABLE_SPECS


def copy_figures() -> None:
    for short, stem in FIGURE_SOURCES.items():
        for ext in ["png", "pdf", "svg"]:
            src = Path(f"{stem}.{ext}")
            if not src.exists():
                raise FileNotFoundError(src)
            shutil.copy2(src, FIGURES / f"{short}.{ext}")


def word_table(doc: Document, key: str, rows: list[dict[str, object]], number: str) -> None:
    m.word_table(doc, key, rows, number)


def build_word(tables: dict[str, list[dict[str, object]]]) -> tuple[Path, Path]:
    doc = Document(); m.v8.configure_word(doc, RUNNING_TITLE); m.add_cover(doc, TITLE, "Original research | Secondary analysis")
    for text in [f"Abstract word count: {abstract_word_count()}", f"Main-text word count: {main_word_count()}", "Tables: 4; Figures: 4; Additional files: 2"]:
        p = doc.add_paragraph(); m.v8.set_font(p.add_run(text), 10)
    doc.add_heading("Abstract", level=1)
    for heading, text in ABSTRACT.items():
        p = doc.add_paragraph(); m.v8.set_font(p.add_run(heading + ": "), 12, bold=True); m.v8.set_font(p.add_run(text), 12)
    p = doc.add_paragraph(); m.v8.set_font(p.add_run("Keywords: "), 12, bold=True); m.v8.set_font(p.add_run("; ".join(KEYWORDS)), 12)
    doc.add_heading("Background", level=1)
    for paragraph in INTRODUCTION: m.v8.add_para(doc, paragraph)
    for section, blocks in [("Methods", METHODS), ("Results", RESULTS)]:
        doc.add_heading(section, level=1)
        for heading, paragraphs in blocks.items():
            doc.add_heading(heading, level=2)
            for paragraph in paragraphs: m.v8.add_para(doc, paragraph)
    doc.add_heading("Discussion", level=1)
    for paragraph in DISCUSSION: m.v8.add_para(doc, paragraph)
    doc.add_heading("Conclusions", level=1); m.v8.add_para(doc, CONCLUSION)
    doc.add_heading("List of abbreviations", level=1)
    m.v8.add_para(doc, "ADQI, Acute Disease Quality Initiative; AKI, acute kidney injury; AUPRC, area under the precision-recall curve; AUROC, area under the receiver-operating-characteristic curve; CI, confidence interval; CIF, cumulative incidence function; ICU, intensive care unit; IPW, inverse-probability weighting; IQR, interquartile range; KDIGO, Kidney Disease: Improving Global Outcomes; RRT, renal replacement therapy; SCr, serum creatinine.", indent=False)
    doc.add_heading("Declarations", level=1)
    declarations = [("Ethics approval and consent to participate", ETHICS), ("Consent for publication", "Not applicable."),
                    ("Availability of data and materials", AVAILABILITY + f" [{m.v8.REF_INDEX['MIMIC2024']},{m.v8.REF_INDEX['Pollard2018']}]."),
                    ("Competing interests", COMPETING), ("Funding", FUNDING), ("Author's contributions", CONTRIBUTIONS),
                    ("Acknowledgements", ACKNOWLEDGEMENTS), ("AI-assisted editing disclosure", AI_DISCLOSURE)]
    for heading, text in declarations:
        doc.add_heading(heading, level=2); m.v8.add_para(doc, text, indent=False)
    doc.add_heading("Additional files", level=1)
    m.v8.add_para(doc, "Additional file 1 (.docx and .pdf): Supplementary Tables S1–S9 and Supplementary Figures S1–S7.", indent=False)
    m.v8.add_para(doc, "Additional file 2 (.docx and .csv): TRIPOD+AI checklist for the prediction components.", indent=False)
    doc.add_heading("References", level=1)
    for i, (_key, reference, *_rest) in enumerate(m.REFS, 1):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.25); p.paragraph_format.first_line_indent = Inches(-0.25)
        m.v8.set_font(p.add_run(f"{i}. {reference}"), 10)
    m.add_landscape(doc)
    for key in ["T1", "T2", "T3", "T4"]: word_table(doc, key, tables[key], f"Table {key[1:]}")
    m.add_portrait(doc); doc.add_heading("Figure legends", level=1)
    for number, title, legend in FIGURE_LEGENDS: m.v8.add_para(doc, f"{number}. {title}. {legend}", indent=False)
    main_path = OUT / "secondary_manuscript_v30_en.docx"; doc.save(main_path)

    supp = Document(); m.v8.configure_word(supp, "Supplement | " + RUNNING_TITLE); m.add_cover(supp, "Additional file 1: Supplementary material", TITLE)
    m.add_landscape(supp)
    for i in range(1, 10): word_table(supp, f"S{i}", tables[f"S{i}"], f"Table S{i}")
    m.add_portrait(supp); supp.add_heading("Supplementary figures", level=1)
    for i, (number, title, legend) in enumerate(SUPP_FIGURE_LEGENDS, 1):
        p = supp.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(FIGURES / f"FigS{i}.png"), width=Inches(6.3))
        m.v8.add_para(supp, f"{number}. {title}. {legend}", indent=False)
    supp_path = OUT / "additional_file_1_secondary_supplement_en.docx"; supp.save(supp_path)
    return main_path, supp_path


def latex_table(key: str, rows: list[dict[str, object]]) -> None:
    title, label, columns, _widths, landscape, _font, _legend = TABLE_SPECS[key]
    m.v8.latex_table(rows, columns, title, label, f"table{key}.tex", landscape=landscape, long=(key in {"T1", "S3", "S8", "S9"}))


def tex_section(blocks: dict[str, list[str]]) -> str:
    return "\n".join(r"\subsection{" + m.v8.tex_escape(h) + "}\n" + "\n\n".join(m.v8.render_text(p, "tex") for p in ps) for h, ps in blocks.items())


def write_latex(tables: dict[str, list[dict[str, object]]]) -> None:
    (LATEX / "tables").mkdir(parents=True, exist_ok=True)
    for key in ["T1", "T2", "T3", "T4", *[f"S{i}" for i in range(1, 10)]]: latex_table(key, tables[key])
    refs = "\n".join(r"\bibitem{" + key + "} " + m.v8.tex_escape(reference) for key, reference, *_ in m.REFS)
    figures = "\n".join(r"\begin{figure}[p]\centering\includegraphics[width=\textwidth]{../figures/Fig" + str(i) + ".pdf}\n" +
                         r"\caption{" + m.v8.tex_escape(title) + ". " + m.v8.tex_escape(legend) + r"}\end{figure}"
                         for i, (_n, title, legend) in enumerate(FIGURE_LEGENDS, 1))
    abstract = "\n\n".join(r"\textbf{" + k + ":} " + m.v8.tex_escape(v) for k, v in ABSTRACT.items())
    template = (m.ROOT / "scripts" / "build_secondary_manuscript_v29.py").read_text(encoding="utf-8")
    # Keep the LaTeX scaffold concise and independent of Python source parsing.
    main = r"""\documentclass[11pt]{article}
\usepackage[letterpaper,margin=1in]{geometry}\usepackage[T1]{fontenc}\usepackage{lmodern,microtype}
\usepackage{graphicx,booktabs,longtable,pdflscape,caption,setspace,lineno,hyperref}\usepackage[numbers,sort&compress]{natbib}
\doublespacing\linenumbers\captionsetup{font=small,labelfont=bf}
\title{__TITLE__}\author{Bizhi Wei\textsuperscript{1*}\\\textsuperscript{1}Pu Ai Medical School, Shaoyang University, Shaoyang 422000, Hunan, China\\\textsuperscript{*}Correspondence: Bizhi Wei; 15619056250wbz@gmail.com}\date{}
\begin{document}\maketitle
\noindent\textbf{Article type:} Original research\par
\noindent\textbf{Abstract word count:} __ABSWC__\par
\noindent\textbf{Main-text word count:} __MAINWC__\par
\noindent\textbf{Tables/Figures:} 4/4\par
\section*{Abstract}__ABSTRACT__
\noindent\textbf{Keywords:} __KEYWORDS__
\section{Background}__INTRO__
\section{Methods}__METHODS__
\section{Results}__RESULTS__
\section{Discussion}__DISCUSSION__
\section{Conclusions}__CONCLUSION__
\section*{Declarations}
\subsection*{Ethics approval and consent to participate}__ETHICS__
\subsection*{Consent for publication}Not applicable.
\subsection*{Availability of data and materials}__AVAILABILITY__ \cite{MIMIC2024,Pollard2018}
\subsection*{Competing interests}__COMPETING__
\subsection*{Funding}__FUNDING__
\subsection*{Author's contributions}__CONTRIBUTIONS__
\subsection*{Acknowledgements}__ACKNOWLEDGEMENTS__
\subsection*{AI-assisted editing disclosure}__AI__
\section*{Additional files}Additional file 1: Supplementary Tables S1--S9 and Supplementary Figures S1--S7.\\Additional file 2: TRIPOD+AI checklist.
\begin{thebibliography}{99}__REFS__\end{thebibliography}
\clearpage\input{tables/tableT1.tex}\clearpage\input{tables/tableT2.tex}\clearpage\input{tables/tableT3.tex}\clearpage\input{tables/tableT4.tex}\clearpage
__FIGURES__\end{document}"""
    replacements = {"__TITLE__": m.v8.tex_escape(TITLE), "__ABSWC__": str(abstract_word_count()), "__MAINWC__": str(main_word_count()),
                    "__ABSTRACT__": abstract, "__KEYWORDS__": "; ".join(KEYWORDS),
                    "__INTRO__": "\n\n".join(m.v8.render_text(p, "tex") for p in INTRODUCTION), "__METHODS__": tex_section(METHODS),
                    "__RESULTS__": tex_section(RESULTS), "__DISCUSSION__": "\n\n".join(m.v8.render_text(p, "tex") for p in DISCUSSION),
                    "__CONCLUSION__": m.v8.render_text(CONCLUSION, "tex"), "__ETHICS__": m.v8.tex_escape(ETHICS),
                    "__AVAILABILITY__": m.v8.tex_escape(AVAILABILITY), "__COMPETING__": m.v8.tex_escape(COMPETING),
                    "__FUNDING__": m.v8.tex_escape(FUNDING), "__CONTRIBUTIONS__": m.v8.tex_escape(CONTRIBUTIONS),
                    "__ACKNOWLEDGEMENTS__": m.v8.tex_escape(ACKNOWLEDGEMENTS), "__AI__": m.v8.tex_escape(AI_DISCLOSURE),
                    "__REFS__": refs, "__FIGURES__": figures}
    for key, value in replacements.items(): main = main.replace(key, value)
    (LATEX / "main.tex").write_text(main, encoding="utf-8")

    supp_figs = "\n".join(r"\begin{figure}[p]\centering\includegraphics[width=\textwidth]{../figures/FigS" + str(i) + ".pdf}\n" +
                           r"\caption{" + m.v8.tex_escape(title) + ". " + m.v8.tex_escape(legend) + r"}\end{figure}\clearpage"
                           for i, (_n, title, legend) in enumerate(SUPP_FIGURE_LEGENDS, 1))
    inputs = "\n".join(r"\input{tables/tableS" + str(i) + r".tex}\clearpage" for i in range(1, 10))
    supplement = (r"\documentclass[11pt]{article}\usepackage[letterpaper,margin=0.75in]{geometry}\usepackage[T1]{fontenc}\usepackage{lmodern}"
                  r"\usepackage{graphicx,booktabs,longtable,pdflscape,caption,setspace,hyperref}\onehalfspacing\begin{document}"
                  r"\begin{center}{\Large\bfseries Additional file 1: Supplementary material}\\[8pt]" + m.v8.tex_escape(TITLE) +
                  r"\\[8pt]Bizhi Wei\textsuperscript{1*}\\\textsuperscript{1}Pu Ai Medical School, Shaoyang University, Shaoyang 422000, Hunan, China\end{center}" +
                  inputs + supp_figs + r"\end{document}")
    (LATEX / "supplement.tex").write_text(supplement, encoding="utf-8")
    entries = []
    for key, reference, doi, _pmid, _role in m.REFS:
        title = reference.split(". ", 1)[1].split(". ", 1)[0] if ". " in reference else reference
        entries.append(f"@misc{{{key},\n  title={{{title}}},\n  note={{{reference}}}" + (f",\n  doi={{{doi}}}" if doi else "") + "\n}")
    (LATEX / "references.bib").write_text("\n\n".join(entries) + "\n", encoding="utf-8")


def write_markdown() -> None:
    lines = [f"# {TITLE}", "", AUTHORS_WORD, "", AFFILIATION_WORD, "", CORRESPONDENCE_WORD, "", "## Abstract", ""]
    for h, text in ABSTRACT.items(): lines.extend([f"**{h}:** {text}", ""])
    lines.extend(["**Keywords:** " + "; ".join(KEYWORDS), "", "## Background", ""])
    for p in INTRODUCTION: lines.extend([m.v8.render_text(p, "word"), ""])
    for section, blocks in [("Methods", METHODS), ("Results", RESULTS)]:
        lines.extend([f"## {section}", ""])
        for h, ps in blocks.items():
            lines.extend([f"### {h}", ""])
            for p in ps: lines.extend([m.v8.render_text(p, "word"), ""])
    lines.extend(["## Discussion", ""])
    for p in DISCUSSION: lines.extend([m.v8.render_text(p, "word"), ""])
    lines.extend(["## Conclusions", "", m.v8.render_text(CONCLUSION, "word"), "", "## Declarations", ""])
    for h, text in [("Ethics approval and consent to participate", ETHICS), ("Consent for publication", "Not applicable."),
                    ("Availability of data and materials", AVAILABILITY), ("Competing interests", COMPETING), ("Funding", FUNDING),
                    ("Author's contributions", CONTRIBUTIONS), ("Acknowledgements", ACKNOWLEDGEMENTS), ("AI-assisted editing disclosure", AI_DISCLOSURE)]:
        lines.extend([f"### {h}", "", text, ""])
    lines.extend(["## References", ""])
    for i, (_key, reference, *_rest) in enumerate(m.REFS, 1): lines.append(f"{i}. {reference}")
    (OUT / "secondary_manuscript_v30_en.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def build_checklist() -> tuple[Path, Path]:
    asset = ROOT / "scripts" / "assets" / "tripod_ai_checklist_v8.tsv"
    locations = {
        "1": "Title", "2": "Abstract", "3a": "Background", "3b": "Background, final paragraph", "4": "Background, final paragraph",
        "5a": "Methods—Study design", "5b": "Methods—Study design", "6a": "Methods—Study design", "6b": "Methods—Study design and Figure 1",
        "7": "Methods—Landmark models", "8a": "Methods—Outcomes", "9a": "Methods—Landmark models", "10": "Methods—Landmarks; Figure 1; Table S1",
        "11": "Methods—Landmarks and observability", "12a": "Methods—Prediction models", "12b": "Methods—Prediction models",
        "12c": "Methods—Prediction models and same-risk analysis", "12d": "Methods—Temporal and external validation; Table S9",
        "12e": "Methods—Prediction models", "12f": "Methods—External validation", "12g": "Methods—External validation", "13": "Methods—Prediction models",
        "15": "Methods—Prediction models", "16": "Methods—Study design", "17": "Declarations—Ethics", "18a": "Declarations—Funding",
        "18b": "Declarations—Competing interests", "18c": "Availability of data and materials", "18e": "Availability of data and materials",
        "18f": "Availability of data and materials", "19": "Methods—Software and reporting", "20a": "Results—Cohort; Figure 1",
        "20b": "Table 1; Table S1", "20c": "Tables S2–S9", "21": "Results; Tables 2–4", "22": "Availability of data and materials",
        "23a": "Results—Model performance; Figures 2–3; Tables 2–3", "23b": "Results—External validation; Tables S5–S6",
        "24": "Methods—External validation; Table 3", "25": "Discussion", "26": "Discussion—Limitations",
        "27a": "Methods—Observability; Discussion", "27b": "Methods—Observability", "27c": "Discussion; Conclusions",
    }
    rows = []
    with asset.open(encoding="utf-8") as handle:
        for row in csv.DictReader(handle, delimiter="\t"):
            row["location"] = locations.get(row["item"], "See manuscript or supplement"); row["status"] = "Reported"; rows.append(row)
    csv_path = OUT / "additional_file_2_tripod_ai_checklist.csv"
    with csv_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0])); writer.writeheader(); writer.writerows(rows)
    doc = Document(); m.v8.configure_word(doc, "TRIPOD+AI checklist"); m.add_cover(doc, "Additional file 2: TRIPOD+AI checklist", TITLE)
    m.v8.add_para(doc, "This checklist applies to the severe-AKI prediction components. The multistate and competing-risk analyses are described separately.", indent=False)
    m.add_landscape(doc); m.v8.add_table(doc, "", "", rows,
        [("section", "Section"), ("item", "Item"), ("checklist_item", "Reporting recommendation"), ("location", "Location"), ("status", "Status")],
        [1200, 650, 4600, 4000, 1100], font=6.8)
    docx_path = OUT / "additional_file_2_tripod_ai_checklist.docx"; doc.save(docx_path)
    return docx_path, csv_path


def write_audits() -> None:
    m.write_audits()
    evidence = [
        ("Severe SCr-AKI incidence 679/10,877", "audit_v29_multistate_summary.csv", "Abstract; Results; Figure 1"),
        ("Severe-AKI baseline characteristics", "cohort_v26_strict_aki_severity_recovery.csv", "Table 1"),
        ("Paired severe-AKI information gain", "model_v30_severe_paired_incremental_deltas.csv", "Results; Figure 2; Table S2"),
        ("Selected severe-AKI model performance", "model_v27_performance_summary.csv", "Table 2; Figure S3"),
        ("eICU denominator and external performance", "audit_v16_readme.md; model_v28_eicu_frozen_severe_performance.csv", "Figure 1; Figure 3; Table 3"),
        ("Cross-database feature mapping", "external_validation_eicu_v16.py; audit_v28_severe_feature_harmonization.csv", "Methods; Table S9"),
        ("Recovery observability/IPW", "model_v28_observability_performance.csv; model_v28_recovery_ipw_performance.csv", "Table S7; Figure S6"),
        ("State occupancy and competing risks", "analysis_v29_state_occupancy.csv; analysis_v29_cif_after_aki.csv; analysis_v29_cif_after_recovery.csv", "Table 4; Figure 4"),
    ]
    with (OUT / "claim_evidence_map.csv").open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle); writer.writerow(["Claim", "Source", "Manuscript location"]); writer.writerows(evidence)
    (OUT / "code_and_data_availability_statement.txt").write_text(AVAILABILITY + "\n", encoding="utf-8")


def write_readme() -> None:
    text = f"""# Revised secondary manuscript package v30

This package preserves `outputs/manuscript_package_v29_secondary` and the locked primary package.

- Title: {TITLE}
- Author: Bizhi Wei
- Abstract words: {abstract_word_count()}
- Main-text words: {main_word_count()}
- Main items: Tables 1–4 and Figures 1–4
- Additional file 1: Tables S1–S9 and Figures S1–S7
- Additional file 2: TRIPOD+AI checklist
- Core revision: severe-AKI Table 1, cohort/analysis flow, paired same-risk-set incremental analysis, explicit eICU observability denominator, and feature mapping
- External-validation boundary: only severe-AKI 0-h, 6-h, and 24-h landmark models were externally validated
- Repository release: v1.0.2
- Archived DOI: to be provided before publication

The 4,531 locked incident-AKI labels remain unchanged. Twelve post-disposition AKI onsets are excluded only from trajectory risk sets.
"""
    (OUT / "README.md").write_text(text, encoding="utf-8")


def copy_source_data() -> None:
    items = {
        V30F / "figure_v30_flow_source_data.csv": "Figure_1_flow_source_data.csv",
        V30I / "model_v30_severe_paired_incremental_deltas.csv": "Figure_2_paired_deltas_source_data.csv",
        V30I / "model_v30_severe_same_risk_performance.csv": "Figure_2_same_risk_performance_source_data.csv",
        m.V28S / "analysis_v28_mimic_rolling_severe_validation.csv": "Figure_3_temporal_source_data.csv",
        m.V29 / "analysis_v29_state_occupancy.csv": "Figure_4_state_occupancy_source_data.csv",
        m.V29 / "analysis_v29_cif_after_aki.csv": "Figure_4_cif_after_aki_source_data.csv",
        m.V29 / "analysis_v29_cif_after_recovery.csv": "Figure_4_cif_after_recovery_source_data.csv",
    }
    for src, name in items.items(): shutil.copy2(src, SOURCE_DATA / name)


def main() -> None:
    if OUT.exists(): shutil.rmtree(OUT)
    for path in [OUT, LATEX, TABLES, FIGURES, QA, SOURCE_DATA]: path.mkdir(parents=True, exist_ok=True)
    tables = build_tables(); copy_figures(); copy_source_data(); write_latex(tables); write_markdown(); build_word(tables)
    build_checklist(); write_audits(); write_readme(); m.write_manifest()
    print(f"Built revised secondary manuscript package: {OUT}")
    print(f"Abstract words={abstract_word_count()}; main words={main_word_count()}; references={len(m.REFS)}")


if __name__ == "__main__":
    main()
