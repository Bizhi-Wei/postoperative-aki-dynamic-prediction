from pathlib import Path
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
PACKAGE = PROJECT_ROOT / "outputs" / "manuscript_package_v23_locked"
SOURCE_DOCX = PACKAGE / "critical_care_main_manuscript_en.docx"
ZH_DOCX = PACKAGE / "critical_care_main_manuscript_zh_review.docx"
CHECKLIST_MD = PACKAGE / "submission_precheck_checklist_zh.md"
CHECKLIST_DOCX = PACKAGE / "submission_precheck_checklist_zh.docx"


TITLE_ZH = "外科重症监护术后新发急性肾损伤动态预测模型的开发与内部验证：一项基于 MIMIC-IV 的回顾性队列研究"


PARAGRAPH_TRANSLATIONS = {
    "Abstract": "摘要",
    "Background": "背景",
    "Methods": "方法",
    "Results": "结果",
    "Discussion": "讨论",
    "Conclusions": "结论",
    "List of abbreviations": "缩略语列表",
    "Declarations": "声明",
    "Ethics approval and consent to participate": "伦理批准与参与同意",
    "Consent for publication": "发表同意",
    "Availability of data and materials": "数据和材料可获得性",
    "Competing interests": "利益冲突",
    "Funding": "资助",
    "Authors' contributions": "作者贡献",
    "Acknowledgements": "致谢",
    "AI-assisted editing disclosure": "AI 辅助编辑声明",
    "Additional files": "附加文件",
    "References": "参考文献",
    "Figure legends": "图例",
    "Study design, setting, and reporting": "研究设计、场景与报告规范",
    "Cohort definition": "队列定义",
    "Baseline kidney function and outcome": "基线肾功能与结局",
    "Landmark datasets and predictors": "动态时间点数据集与预测变量",
    "Sample size and missing data": "样本量与缺失数据",
    "Model development and internal validation": "模型开发与内部验证",
    "Performance assessment": "模型性能评估",
    "Interpretability, fairness, and sensitivity analyses": "可解释性、公平性与敏感性分析",
    "Software and reproducibility": "软件与可重复性",
    "Patient and public involvement": "患者和公众参与",
    "Cohort and incident AKI": "队列与新发 AKI",
    "Dynamic risk sets": "动态风险集",
    "Model performance": "模型性能",
    "Model attribution and subgroup performance": "模型归因与亚组性能",
    "Sensitivity analyses": "敏感性分析",
}


BODY_TRANSLATIONS = {
    "Article type: Research": "文章类型：研究论文",
    "Abstract word count: 313": "英文摘要词数：313",
    "Main-text word count: 2746": "英文正文词数：2746",
    "Tables: 3; Figures: 4; Additional files: 2": "表格：3；图：4；附加文件：2",
    "Keywords: acute kidney injury; postoperative complications; intensive care; dynamic prediction; machine learning; MIMIC-IV; clinical prediction model":
        "关键词：急性肾损伤；术后并发症；重症监护；动态预测；机器学习；MIMIC-IV；临床预测模型",
    "Background: Postoperative acute kidney injury (AKI) is common after major surgery, but risk evolves during early critical illness. We developed and internally validated landmark-specific models for incident AKI in a rigorously defined surgical intensive care cohort.":
        "背景：术后急性肾损伤（AKI）是大手术后的常见并发症，但其风险会在重症监护早期持续变化。本研究在严格定义的外科 ICU 队列中，开发并内部验证不同动态预测时间点的新发 AKI 预测模型。",
    "Methods: We conducted a retrospective cohort study using MIMIC-IV version 3.1. The analysis unit was the first intensive care unit (ICU) stay per hospital admission. Incident AKI within 7 days of ICU admission was recomputed using Kidney Disease: Improving Global Outcomes serum creatinine criteria; urine output was not used. Separate risk sets were constructed at ICU admission, 6 h, and 24 h, excluding patients whose AKI had occurred by each landmark. Logistic regression, random forest, XGBoost, and LightGBM were evaluated using an 80:20 patient-grouped split. Performance assessment included discrimination, calibration, decision-curve analysis, patient-level bootstrap confidence intervals, subgroup analyses, SHAP attribution, and prespecified baseline-creatinine sensitivity analyses.":
        "方法：本研究基于 MIMIC-IV 3.1 版开展回顾性队列研究。分析单位为每次住院中的首次 ICU stay。ICU 入科后 7 天内的新发 AKI 依据 Kidney Disease: Improving Global Outcomes（KDIGO）血清肌酐标准重新计算；暂未使用尿量标准。分别在 ICU 入科、6 小时和 24 小时构建风险集，并在每个时间点排除此前已经发生 AKI 的患者。采用按患者分组的 80:20 训练/测试划分，评估逻辑回归、随机森林、XGBoost 和 LightGBM。模型性能评估包括区分度、校准度、决策曲线分析、患者层面 bootstrap 置信区间、亚组分析、SHAP 归因，以及预先设定的基线肌酐敏感性分析。",
    "Results: Among 11,943 qualifying admissions, 10,877 were evaluable for incident AKI and 4,531 developed AKI (41.7%). Stage 1, 2, and 3 accounted for 84.9%, 9.8%, and 5.3% of events. Risk sets included 10,877 admissions at 0 h, 10,624 at 6 h, and 9,301 at 24 h. XGBoost had the highest test AUROC at 0 h (0.728, 95% CI 0.706-0.748) and 6 h (0.740, 0.719-0.762). At 24 h, logistic regression had the highest AUROC (0.754, 0.729-0.777), an area under the precision-recall curve of 0.602, and a Brier score of 0.176. Removing creatinine-derived predictors reduced 24-h AUROC by 0.009-0.025 across model families, whereas restriction to a pre-ICU 7-day baseline creatinine produced no material change.":
        "结果：在 11,943 次符合条件的住院中，10,877 次可用于新发 AKI 评估，其中 4,531 次发生 AKI（41.7%）。1、2、3 期分别占 AKI 事件的 84.9%、9.8% 和 5.3%。0 小时、6 小时和 24 小时风险集分别包含 10,877、10,624 和 9,301 次住院。XGBoost 在 0 小时（AUROC 0.728，95% CI 0.706–0.748）和 6 小时（0.740，0.719–0.762）具有最高测试集 AUROC。24 小时时，逻辑回归的 AUROC 最高（0.754，0.729–0.777），AUPRC 为 0.602，Brier 评分为 0.176。去除肌酐相关预测变量后，各模型 24 小时 AUROC 下降 0.009–0.025；而限制为 ICU 前 7 天内有基线肌酐者后，模型性能未见实质性改变。",
    "Conclusions: Routinely available peri-ICU data provided moderate internal discrimination for incident postoperative AKI across dynamic landmarks. Performance at 24 h partly depended on creatinine trajectories, although non-creatinine models retained discrimination. External and temporal validation are required before clinical use.":
        "结论：常规可获得的围 ICU 数据在不同动态时间点对术后新发 AKI 具有中等程度的内部区分能力。24 小时模型性能部分依赖肌酐变化轨迹，但不含肌酐的模型仍保留一定区分能力。在临床应用前，仍需外部验证和时间验证。",
    "Postoperative acute kidney injury (AKI) is a frequent complication of major surgery and is associated with greater short-term morbidity, chronic kidney disease, cardiovascular events, and death [1-7]. Its pathophysiology is heterogeneous: pre-existing kidney vulnerability may interact with hemodynamic instability, inflammation, anemia, nephrotoxic exposure, and procedure-specific insults. Consequently, clinically relevant risk can change substantially during the first postoperative day.":
        "术后急性肾损伤（AKI）是大手术后的常见并发症，与短期并发症增加、慢性肾病、心血管事件和死亡相关 [1-7]。其病理生理机制具有异质性：既往肾脏易损性可能与血流动力学不稳定、炎症、贫血、肾毒性暴露以及手术特异性损伤相互作用。因此，具有临床意义的风险可在术后第 1 天内明显变化。",
    "Many perioperative studies assess risk once, before surgery or at ICU admission. Such models are useful for baseline stratification but cannot fully represent evolving physiology. Urine output, arterial pressure, and early laboratory trajectories have been associated with postoperative AKI, yet their interpretation depends on when they become available and whether the outcome has already occurred [8-9]. A model evaluated at 6 or 24 h should therefore estimate future risk only among patients who remain event free at that landmark.":
        "许多围手术期研究仅在术前或 ICU 入科时进行一次风险评估。这类模型有助于基线风险分层，但无法充分反映生理状态的动态变化。尿量、动脉压和早期实验室指标轨迹均与术后 AKI 相关，但其解释取决于这些信息在何时可获得，以及结局是否已经发生 [8-9]。因此，在 6 小时或 24 小时评估的模型，应仅在该时间点仍未发生事件的患者中估计后续风险。",
    "The Kidney Disease: Improving Global Outcomes (KDIGO) definition provides a standardized framework for AKI ascertainment from changes in serum creatinine and urine output [10-11]. Serum creatinine, however, is both a delayed marker of kidney dysfunction and a dominant predictor in many AKI models. When a creatinine-defined outcome is predicted after ICU admission, early creatinine changes may increase apparent performance because they are already close to the outcome definition. Strict temporal alignment and sensitivity analyses that remove creatinine-derived predictors are therefore necessary to assess information leakage and clinical interpretability.":
        "KDIGO 定义为基于血清肌酐和尿量变化判定 AKI 提供了标准化框架 [10-11]。然而，血清肌酐既是肾功能受损的延迟性标志物，也是许多 AKI 模型中的主导预测变量。当 ICU 入科后预测一个由肌酐定义的结局时，早期肌酐变化可能因已接近结局定义而提高表观性能。因此，必须进行严格的时间对齐，并通过去除肌酐相关预测变量的敏感性分析来评估信息泄漏和临床可解释性。",
    "Electronic health record studies have demonstrated the feasibility of longitudinal AKI prediction in general inpatient, critical care, and cardiac surgical populations [12-16]. However, published models differ in surgical case mix, baseline kidney-function definitions, prediction horizons, and validation strategies. We therefore developed and internally validated models for incident AKI at ICU admission, 6 h, and 24 h in a strict therapeutic surgical cohort. We also evaluated calibration, clinical net benefit, model attribution, subgroup performance, a no-creatinine analysis, and restriction to patients with a pre-ICU baseline creatinine.":
        "既往电子健康记录研究已证明，在普通住院、重症监护和心脏外科人群中进行纵向 AKI 预测具有可行性 [12-16]。然而，已发表模型在手术构成、基线肾功能定义、预测时间窗和验证策略方面差异较大。因此，本研究在严格定义的治疗性外科队列中，开发并内部验证 ICU 入科、6 小时和 24 小时的新发 AKI 预测模型。我们还评估了校准度、临床净获益、模型归因、亚组性能、不含肌酐模型，以及仅限 ICU 前有基线肌酐患者的敏感性分析。",
    "We performed a retrospective prediction-model development and internal-validation study using MIMIC-IV version 3.1, a deidentified electronic health record database from Beth Israel Deaconess Medical Center in Boston, Massachusetts, United States [17-18]. MIMIC-IV version 3.1 contains admissions from 2008 through 2022. The analysis unit was the first qualifying ICU stay within each hospital admission. Reporting was guided by TRIPOD+AI, and design decisions were considered in relation to PROBAST [19-20].":
        "本研究使用 MIMIC-IV 3.1 版开展回顾性预测模型开发与内部验证研究。MIMIC-IV 是来自美国马萨诸塞州波士顿 Beth Israel Deaconess Medical Center 的去标识化电子健康记录数据库 [17-18]。MIMIC-IV 3.1 版包含 2008 至 2022 年的住院记录。分析单位为每次住院中首次符合条件的 ICU stay。报告撰写参考 TRIPOD+AI，研究设计决策亦结合 PROBAST 进行考量 [19-20]。",
    "Adults were eligible when an explicitly therapeutic major operation was recorded on the day of ICU admission or the preceding day and the first ICU location was surgical, cardiac vascular, trauma surgical, mixed medical-surgical, neurosurgical, or post-anesthesia care. Eligible procedures included cardiac, vascular, general gastrointestinal or hepatobiliary, major orthopedic, neurosurgical, and thoracic or respiratory operations. Diagnostic imaging, electrocardiography, vascular access, tracheal intubation, enteral nutrition, dialysis, nonoperative respiratory measurements, routine chest radiography, and obstetric procedures did not qualify as surgery.":
        "纳入标准为成年人，且在 ICU 入科当天或前一天记录有明确治疗性大手术，同时首次 ICU 位置为外科、心脏血管、创伤外科、混合内外科、神经外科或麻醉后恢复相关单元。符合条件的手术包括心脏、血管、普通胃肠/肝胆、重大骨科、神经外科以及胸外/呼吸系统手术。单纯诊断影像、心电图、血管通路、气管插管、肠内营养、透析、非手术性呼吸测量、常规胸片和产科手术不作为本研究所定义的手术。",
    "The index time was ICU admission. The recorded operation date was retained for audit because procedure coding provided date-level rather than precise operative-end timestamps. Admissions with AKI present at or before the index time, no usable baseline serum creatinine (SCr), or no post-index SCr measurement within 7 days were retained in cohort audits but excluded from incident-AKI model development.":
        "索引时间定义为 ICU 入科时间。由于手术编码仅提供日期层面的信息，而非精确的手术结束时间戳，记录的手术日期被保留用于审计。在索引时间之前或当时已存在 AKI、无可用基线血清肌酐（SCr），或索引后 7 天内无 SCr 测量的住院记录，保留于队列审计，但不纳入新发 AKI 模型开发。",
    "Baseline SCr was defined as the lowest value during the 7 days before ICU admission. If no such measurement was available, the earliest admission SCr within the first 24 h of hospitalization was used; the source and timestamp were retained. At a prediction landmark, a fallback baseline was exposed to a model only when its measurement preceded that landmark.":
        "基线 SCr 定义为 ICU 入科前 7 天内的最低值。若无该时间窗内测量值，则使用住院后最早 24 小时内的入院 SCr；同时保留其来源和时间戳。在某一预测时间点，只有当备用基线肌酐测量发生在该时间点之前，才将其暴露给模型。",
    "Incident AKI was adjudicated from timestamped SCr values using KDIGO criteria, without urine output [10]. AKI was present when SCr increased by at least 0.3 mg/dL from a prior measurement within 48 h or reached at least 1.5 times baseline within 7 days after ICU admission. Patients meeting either criterion before or at ICU admission were classified as having prevalent AKI. Stage 1 was defined as a 1.5- to less than 2.0-fold baseline increase or an absolute increase of at least 0.3 mg/dL; stage 2 as a 2.0- to less than 3.0-fold increase; and stage 3 as at least a 3.0-fold increase or peak SCr of at least 4.0 mg/dL. The model outcome at each landmark was the probability of incident AKI after that landmark and within 7 days of ICU admission.":
        "新发 AKI 基于带时间戳的 SCr 数值并按照 KDIGO 标准判定，未使用尿量标准 [10]。若 SCr 在 48 小时内较既往测量值升高至少 0.3 mg/dL，或在 ICU 入科后 7 天内达到基线值的至少 1.5 倍，则判定为 AKI。若患者在 ICU 入科前或入科时已经满足任一标准，则归类为既存 AKI。1 期定义为较基线升高 1.5 倍至低于 2.0 倍，或绝对升高至少 0.3 mg/dL；2 期为升高 2.0 倍至低于 3.0 倍；3 期为升高至少 3.0 倍，或峰值 SCr 至少 4.0 mg/dL。每个动态时间点的模型结局为该时间点之后、ICU 入科后 7 天内发生新发 AKI 的概率。",
    "Prediction datasets were constructed at ICU admission (0 h), 6 h, and 24 h. Patients with AKI onset at or before a landmark were excluded from that landmark risk set. Static candidate predictors comprised age, sex, race, admission characteristics, comorbidities, Charlson score, surgical category, first ICU location, and baseline kidney-function variables.":
        "分别在 ICU 入科（0 小时）、6 小时和 24 小时构建预测数据集。若患者在某一动态时间点之前或当时已经发生 AKI，则从该时间点风险集中排除。静态候选预测变量包括年龄、性别、种族、入院特征、合并症、Charlson 合并症评分、手术类别、首次 ICU 位置和基线肾功能变量。",
    "For the 0-h model, only variables available at or before ICU admission were retained. The most recent laboratory values during the 24 h before ICU admission were used as pre-index features. At 6 h and 24 h, minimum, maximum, most recent, and measurement-count features were recalculated from timestamped laboratory and vital-sign observations in (0, landmark]. Whole-follow-up summaries, mortality, length of stay, AKI-derived variables, and untimed post-index summaries were excluded. All preprocessing rules were prespecified and applied identically across demographic groups.":
        "对于 0 小时模型，仅保留 ICU 入科时或之前可获得的变量。ICU 入科前 24 小时内最近一次实验室检查值被用作索引前特征。在 6 小时和 24 小时，基于带时间戳的实验室和生命体征观测，在 (0, landmark] 时间窗内重新计算最小值、最大值、最近值和测量次数特征。全随访期汇总变量、死亡、住院时长、AKI 衍生变量以及无时间限制的索引后汇总变量均被排除。所有预处理规则均预先设定，并在不同人口学组中一致应用。",
    "All eligible admissions were included; no a priori sample-size calculation was performed. The available 0-h development cohort contained 10,877 admissions and 4,531 events. Study-size adequacy was considered using the number of outcome events, candidate predictors, and uncertainty in held-out performance estimates rather than a fixed events-per-variable rule [21]. Continuous predictors were median-imputed and accompanied by missingness indicators. Categorical and binary predictors were imputed using the most frequent training value; categorical variables were one-hot encoded. Imputation parameters were estimated in the training data and then applied to the test data.":
        "纳入所有符合条件的住院记录；未进行事前样本量计算。可用于 0 小时模型开发的队列包含 10,877 次住院和 4,531 个事件。研究规模充分性依据结局事件数、候选预测变量数量以及保留测试集性能估计的不确定性综合判断，而非采用固定的每变量事件数规则 [21]。连续预测变量采用中位数填补，并加入缺失指示变量。分类和二元预测变量使用训练集最常见取值填补；分类变量进行 one-hot 编码。填补参数在训练集中估计，随后应用于测试集。",
    "Logistic regression, random forest, XGBoost, and LightGBM were fitted separately at each landmark [22-24]. Continuous predictors were standardized for logistic regression. An 80:20 patient-grouped split was selected from 500 candidate GroupShuffleSplit assignments to approximate the overall event prevalence. The same subject assignment was reused across landmarks, and no patient appeared in both development and test sets.":
        "在每个动态时间点分别拟合逻辑回归、随机森林、XGBoost 和 LightGBM [22-24]。对于逻辑回归，连续预测变量进行标准化。我们从 500 个候选 GroupShuffleSplit 分组划分中选择一个 80:20 的患者分组划分，使事件比例尽量接近总体事件率。相同的患者分配方案在各动态时间点重复使用，且没有患者同时出现在开发集和测试集中。",
    "Random forest used 500 trees and a minimum leaf size of 5. XGBoost used 500 trees, a learning rate of 0.03, and maximum depth of 4. LightGBM used 500 trees, a learning rate of 0.03, and 31 leaves. Hyperparameters were fixed before test-set evaluation. No outcome resampling, synthetic sampling, or class weighting was used. Models returned predicted probabilities. A secondary binary classification threshold maximizing the Youden index was selected in development predictions and applied unchanged to the test set.":
        "随机森林使用 500 棵树，最小叶节点样本数为 5。XGBoost 使用 500 棵树、学习率 0.03、最大深度 4。LightGBM 使用 500 棵树、学习率 0.03、31 个叶节点。所有超参数均在测试集评估前固定。未使用结局重采样、合成采样或类别权重。模型输出预测概率。另在开发集预测中选择使 Youden 指数最大的二分类阈值，并将该阈值不变地应用于测试集。",
    "Discrimination was summarized using the area under the receiver-operating-characteristic curve (AUROC) and area under the precision-recall curve (AUPRC). Accuracy, sensitivity, specificity, precision, F1 score, and confusion matrices were calculated at a threshold of 0.5 and at the development-derived Youden threshold. Overall performance and calibration were assessed using the Brier score, calibration intercept, calibration slope, and ten equal-frequency calibration groups [25-26]. Decision-curve analysis estimated net benefit across threshold probabilities from 0.05 to 0.80 relative to treat-all and treat-none strategies [27].":
        "模型区分度采用受试者工作特征曲线下面积（AUROC）和精确率-召回率曲线下面积（AUPRC）总结。在 0.5 阈值和开发集推导的 Youden 阈值下，计算准确率、敏感度、特异度、精确率、F1 值和混淆矩阵。总体性能和校准度采用 Brier 评分、校准截距、校准斜率以及十个等频校准组评估 [25-26]。决策曲线分析在 0.05 至 0.80 的阈值概率范围内，相对于全部干预和全部不干预策略估计净获益 [27]。",
    "Confidence intervals were obtained from 1,000 patient-level bootstrap resamples of the held-out test set. Prespecified subgroup analyses considered sex, age, race, baseline SCr source, first ICU location, chronic kidney disease, and surgical category. Subgroup intervals used 300 patient-level bootstrap resamples; estimates were omitted for groups with fewer than 50 observations or a single outcome class. Because the data came from one center, between-center heterogeneity was not estimable.":
        "置信区间基于保留测试集的 1,000 次患者层面 bootstrap 重采样获得。预先设定的亚组分析包括性别、年龄、种族、基线 SCr 来源、首次 ICU 位置、慢性肾病和手术类别。亚组置信区间使用 300 次患者层面 bootstrap 重采样；对于样本量少于 50 或仅有单一结局类别的亚组，不报告估计值。由于数据来自单中心，无法估计中心间异质性。",
    "Global SHAP values were calculated in up to 1,000 held-out observations for the tree model with the highest test AUROC at each landmark [28]. SHAP values were interpreted as model attribution, not as causal effects or evidence of modifiability.":
        "对于每个动态时间点测试集 AUROC 最高的树模型，在最多 1,000 个保留测试样本中计算全局 SHAP 值 [28]。SHAP 值被解释为模型归因，而非因果效应或可干预性的证据。",
    "Fairness was explored descriptively through performance estimates by sex, age, and race; no fairness-constrained optimization was applied. Two prespecified sensitivity analyses retained the original patient assignment. First, all creatinine, baseline SCr, and baseline-to-ICU timing predictors were removed before model refitting. Second, models were refitted only among patients whose baseline SCr was measured during the 7 days before ICU admission. Paired patient-level bootstrap resampling compared predictions in identical test patients.":
        "公平性通过按性别、年龄和种族分层的性能估计进行描述性探索；未采用公平性约束优化。两项预先设定的敏感性分析均保留原始患者分组。第一，在重新拟合模型前移除所有肌酐、基线 SCr 以及基线至 ICU 时间间隔相关预测变量。第二，仅在 ICU 入科前 7 天内测得基线 SCr 的患者中重新拟合模型。采用配对患者层面 bootstrap 重采样，在相同测试患者中比较预测结果。",
    "Analyses were performed in Python 3.14 using pandas 3.0, scikit-learn 1.9, XGBoost 3.3, LightGBM 4.6, and SHAP 0.52 [30]. Cohort construction, outcome derivation, feature engineering, model fitting, evaluation, and figure generation were implemented as versioned scripts. During manuscript preparation, ChatGPT assisted with language editing, structural organization, and formatting; it did not generate or modify study data or analyses. The authors reviewed and verified all numerical results, claims, interpretations, and references.":
        "分析使用 Python 3.14 完成，主要依赖 pandas 3.0、scikit-learn 1.9、XGBoost 3.3、LightGBM 4.6 和 SHAP 0.52 [30]。队列构建、结局推导、特征工程、模型拟合、评估和图形生成均通过版本化脚本实现。稿件准备过程中，ChatGPT 协助进行语言润色、结构组织和格式整理；其未生成或修改研究数据或分析。作者已审阅并核实所有数值结果、论断、解释和参考文献。",
    "Patients and members of the public were not involved in the design, conduct, interpretation, or reporting of this retrospective database study.":
        "患者和公众未参与本回顾性数据库研究的设计、实施、解释或报告。",
    "The strict postoperative surgical ICU cohort included 11,943 hospital admissions (Fig. 1). We excluded 1,014 admissions with AKI present at or before ICU admission, 50 without a usable baseline SCr, and two without a post-index SCr measurement. The incident-AKI analytic cohort therefore comprised 10,877 admissions. Incident AKI occurred in 4,531 admissions (41.7%): 3,847 events (84.9%) were stage 1, 445 (9.8%) were stage 2, and 239 (5.3%) were stage 3. Median time to AKI was 31.7 h after ICU admission (interquartile range [IQR], 16.9-42.5 h).":
        "严格定义的术后外科 ICU 队列包含 11,943 次住院（图 1）。我们排除了 1,014 次 ICU 入科前或入科时已存在 AKI 的住院、50 次无可用基线 SCr 的住院，以及 2 次索引后无 SCr 测量的住院。因此，新发 AKI 分析队列包含 10,877 次住院。共有 4,531 次住院发生新发 AKI（41.7%）：其中 3,847 例（84.9%）为 1 期，445 例（9.8%）为 2 期，239 例（5.3%）为 3 期。AKI 发生时间中位数为 ICU 入科后 31.7 小时（四分位距 [IQR]，16.9–42.5 小时）。",
    "Median age was 66 years (IQR, 57-74) and was higher among patients who developed AKI than among those who did not (69 vs 64 years; Table 1). Cardiac operations accounted for 74.1% of admissions, and 77.6% had a baseline SCr measured within 7 days before ICU admission. Chronic kidney disease, heart failure, hypertension, diabetes, anemia, and cardiac surgery were more common among admissions with incident AKI.":
        "总体年龄中位数为 66 岁（IQR，57–74），发生 AKI 者年龄高于未发生者（69 岁 vs 64 岁；表 1）。心脏手术占住院记录的 74.1%，77.6% 的住院在 ICU 入科前 7 天内测得基线 SCr。慢性肾病、心力衰竭、高血压、糖尿病、贫血和心脏手术在新发 AKI 住院中更常见。",
    "All 10,877 evaluable admissions entered the 0-h dataset, with 4,531 subsequent events (41.7%). Exclusion of 253 events occurring by 6 h yielded 10,624 admissions and 4,278 future events. Exclusion of 1,576 cumulative events occurring by 24 h yielded 9,301 admissions and 2,955 future events (Fig. 1). Because landmark populations and remaining outcome windows differed, between-landmark metric differences were not interpreted as paired within-patient improvements.":
        "全部 10,877 次可评估住院均进入 0 小时数据集，其中后续发生事件 4,531 例（41.7%）。排除 6 小时前或当时已发生的 253 个事件后，6 小时风险集包含 10,624 次住院和 4,278 个未来事件。排除 24 小时前或当时累计发生的 1,576 个事件后，24 小时风险集包含 9,301 次住院和 2,955 个未来事件（图 1）。由于不同动态时间点的人群和剩余结局窗口不同，时间点之间的指标差异不解释为患者内配对性能改善。",
    "At 0 h, XGBoost had the highest test discrimination, with AUROC 0.728 (95% CI, 0.706-0.748) and AUPRC 0.665 (0.633-0.696). XGBoost also had the highest AUROC at 6 h (0.740, 0.719-0.762) and AUPRC (0.675, 0.644-0.708). At 24 h, logistic regression had the highest AUROC (0.754, 0.729-0.777), AUPRC (0.602, 0.562-0.643), and a Brier score of 0.176 (Fig. 2; Table 2). Confidence intervals overlapped substantially across model families and did not establish the superiority of a single algorithm.":
        "在 0 小时，XGBoost 在测试集中区分度最高，AUROC 为 0.728（95% CI，0.706–0.748），AUPRC 为 0.665（0.633–0.696）。XGBoost 在 6 小时也具有最高 AUROC（0.740，0.719–0.762）和 AUPRC（0.675，0.644–0.708）。在 24 小时，逻辑回归具有最高 AUROC（0.754，0.729–0.777）、AUPRC（0.602，0.562–0.643），Brier 评分为 0.176（图 2；表 2）。不同模型家族的置信区间存在明显重叠，未能证明某一算法具有明确优越性。",
    "Calibration was close to ideal for XGBoost at 0 h (intercept, -0.003; slope, 0.997) and remained acceptable at 6 h (intercept, -0.018; slope, 0.935). The 24-h logistic model had an intercept of -0.144 and a slope of 0.831, indicating moderately over-extreme predictions in the held-out sample. Selected models had greater internal net benefit than treat-all and treat-none strategies across most evaluated thresholds (Fig. 3).":
        "0 小时 XGBoost 的校准接近理想状态（截距 -0.003；斜率 0.997），6 小时时仍可接受（截距 -0.018；斜率 0.935）。24 小时逻辑回归模型的截距为 -0.144，斜率为 0.831，提示其在保留测试样本中预测略偏极端。在多数评估阈值范围内，所选模型的内部净获益高于全部干预和全部不干预策略（图 3）。",
    "At 6 h, baseline SCr, age, early minimum hemoglobin, Charlson score, and chronic kidney disease were prominent XGBoost features. At 24 h, the most recent creatinine had the largest mean absolute SHAP value, followed by age, minimum hemoglobin, cardiac vascular ICU location, cardiac surgery, minimum creatinine, white blood cell count, and blood urea nitrogen (Fig. 4).":
        "在 6 小时模型中，基线 SCr、年龄、早期最低血红蛋白、Charlson 评分和慢性肾病是较突出的 XGBoost 特征。在 24 小时模型中，最近一次肌酐具有最大的平均绝对 SHAP 值，其后依次为年龄、最低血红蛋白、心脏血管 ICU 位置、心脏手术、最低肌酐、白细胞计数和血尿素氮（图 4）。",
    "Discrimination was broadly maintained across sex and age groups, although estimates were less precise in smaller surgical subgroups (Additional file 1: Fig. S4 and Table S2). The selected 24-h logistic model had an AUROC of 0.767 (95% CI, 0.733-0.807) in women and 0.745 (0.715-0.776) in men. Corresponding estimates were 0.732 (0.693-0.767) among patients younger than 65 years and 0.750 (0.719-0.776) among those aged 65 years or older.":
        "模型区分度在性别和年龄组中总体保持稳定，但在较小的手术亚组中估计值精确度较低（附加文件 1：图 S4 和表 S2）。所选 24 小时逻辑回归模型在女性中的 AUROC 为 0.767（95% CI，0.733–0.807），在男性中为 0.745（0.715–0.776）。在年龄小于 65 岁和 ≥65 岁患者中的对应估计分别为 0.732（0.693–0.767）和 0.750（0.719–0.776）。",
    "Removing creatinine-derived predictors reduced 24-h discrimination in every model family (Additional file 1: Fig. S2; Table 3). Logistic-regression AUROC decreased from 0.754 to 0.729, a paired difference of -0.025 (95% CI, -0.036 to -0.013). Paired intervals also excluded zero for random forest, XGBoost, and LightGBM. Nevertheless, no-creatinine models retained AUROCs of approximately 0.73.":
        "移除肌酐衍生预测变量后，所有模型家族的 24 小时区分度均下降（附加文件 1：图 S2；表 3）。逻辑回归 AUROC 从 0.754 降至 0.729，配对差异为 -0.025（95% CI，-0.036 至 -0.013）。随机森林、XGBoost 和 LightGBM 的配对置信区间也均不跨越 0。尽管如此，不含肌酐模型仍保留约 0.73 的 AUROC。",
    "Restriction to patients with a pre-index 7-day SCr baseline did not materially alter 24-h discrimination (Additional file 1: Fig. S3). The refitted logistic model achieved an AUROC of 0.755, compared with 0.757 for the full-cohort model evaluated in the same restricted test patients; the paired difference was -0.002 (95% CI, -0.007 to 0.004). Paired AUROC intervals crossed zero for all four model families.":
        "限制为索引前 7 天内有 SCr 基线的患者后，24 小时区分度未发生实质性改变（附加文件 1：图 S3）。重新拟合的逻辑回归模型 AUROC 为 0.755；相比之下，在相同受限测试患者中评估的全队列模型 AUROC 为 0.757，配对差异为 -0.002（95% CI，-0.007 至 0.004）。四个模型家族的配对 AUROC 置信区间均跨越 0。",
    "In this strict postoperative surgical ICU cohort, routinely available data supported moderate prediction of incident AKI at ICU admission, 6 h, and 24 h. XGBoost had the highest AUROC at 0 h and 6 h, whereas logistic regression performed best at 24 h. Differences among model families were small relative to bootstrap uncertainty, arguing against an assumption that greater algorithmic complexity necessarily improves later postoperative risk estimation.":
        "在这一严格定义的术后外科 ICU 队列中，常规可获得数据支持在 ICU 入科、6 小时和 24 小时对新发 AKI 进行中等程度预测。XGBoost 在 0 小时和 6 小时 AUROC 最高，而逻辑回归在 24 小时表现最佳。相较于 bootstrap 不确定性，不同模型家族之间差异较小，这提示不能假设更复杂的算法一定能改善术后后期风险估计。",
    "The observed discrimination was lower than that reported by some cardiac surgery-specific models, which have achieved AUROCs above 0.80 for selected outcomes or populations [14-16]. This difference is plausible because our cohort included several surgical specialties, focused on incident AKI after excluding prevalent disease, and enforced landmark-specific predictor windows. These design choices make the prediction target more clinically explicit but also remove information that can inflate apparent performance.":
        "本研究观察到的区分度低于部分心脏外科专病模型，后者在特定结局或人群中 AUROC 可超过 0.80 [14-16]。这一差异是合理的，因为本研究队列涵盖多个外科专科，聚焦于排除既存 AKI 后的新发 AKI，并严格执行动态时间点特异的预测变量时间窗。这些设计选择使预测目标更具临床明确性，但也去除了可能抬高表观性能的信息。",
    "The landmark design is central to interpretation. At 6 h and 24 h, patients whose AKI had already occurred were removed and predictors were restricted to measurements available by the corresponding time. Later models therefore estimated residual future risk among event-free patients rather than repeatedly scoring a fixed cohort. The modest increase in AUROC across landmarks cannot be interpreted as a within-patient improvement caused by accumulating data because event prevalence, case mix, and prediction horizon changed simultaneously.":
        "动态时间点设计是解释本研究结果的核心。在 6 小时和 24 小时，已经发生 AKI 的患者被移除，预测变量也限制为相应时间点前可获得的测量值。因此，较晚时间点模型估计的是仍未发生事件患者的剩余未来风险，而非对固定队列进行重复评分。由于事件发生率、病例构成和预测时间窗同时发生变化，不同时间点 AUROC 的小幅提高不能解释为数据累积导致的患者内性能改善。",
    "Creatinine was the dominant 24-h XGBoost attribution, and removing all creatinine-related predictors reduced AUROC across model families. Early postoperative kidney-function trajectories clearly contained predictive information; however, the outcome was itself defined by future creatinine change, so part of this signal reflects temporal proximity to the outcome definition. The approximately 0.73 AUROC of no-creatinine models indicates that demographic, comorbidity, surgical, hematologic, and physiologic features retained discrimination beyond direct creatinine measurements.":
        "肌酐是 24 小时 XGBoost 模型中最主要的归因特征，移除所有肌酐相关预测变量后，各模型家族 AUROC 均下降。术后早期肾功能轨迹显然包含预测信息；然而，由于结局本身由未来肌酐变化定义，因此该信号的一部分反映了其与结局定义的时间接近性。不含肌酐模型约 0.73 的 AUROC 表明，人口学、合并症、手术、血液学和生理学特征在直接肌酐测量之外仍保留区分能力。",
    "Calibration and decision consequences deserve the same attention as discrimination. The selected 24-h logistic model had the highest AUROC but a calibration slope below one and a negative intercept, indicating that recalibration may be required in a new setting [26]. Decision curves suggested potential internal net benefit across a range of thresholds, but did not define an intervention, prove patient benefit, or quantify alert burden. A clinical implementation study would require a prespecified response pathway, threshold selection with clinicians, calibration surveillance, and prospective evaluation.":
        "校准度和决策后果应与区分度受到同等重视。所选 24 小时逻辑回归模型具有最高 AUROC，但校准斜率低于 1 且截距为负，提示其在新场景中可能需要重新校准 [26]。决策曲线提示模型在一系列阈值范围内可能具有内部净获益，但并未定义具体干预措施、证明患者获益或量化警报负担。临床实施研究需要预先设定响应路径、与临床医生共同选择阈值、进行校准监测，并开展前瞻性评估。",
    "Subgroup analyses did not reveal a large loss of discrimination by sex or age, but smaller racial and surgical strata produced wide intervals. These analyses should not be interpreted as evidence of algorithmic fairness. The database reflects one health system, clinical measurement intensity can differ across patient groups, and the study did not evaluate equalized error rates, calibration within all intersectional groups, or downstream treatment effects.":
        "亚组分析未显示模型在性别或年龄组中存在明显区分度下降，但较小的种族和手术亚组置信区间较宽。这些分析不应被解释为算法公平性的证据。该数据库反映单一医疗系统，临床测量强度可能在不同患者群体间存在差异，本研究也未评估均衡错误率、所有交叉亚组内的校准度或下游治疗效应。",
    "The study has several strengths. Surgical eligibility excluded diagnostic and bedside procedures that are frequently misclassified when any procedure code is treated as surgery. AKI was recomputed from timestamped laboratory measurements after the index time was fixed. Landmark-specific features were reconstructed from their source timestamps, patient identity was respected during data splitting, and paired bootstrap analyses directly tested sensitivity to creatinine predictors and baseline source.":
        "本研究具有若干优势。手术纳入标准排除了诊断性和床旁操作，避免了将任何 procedure code 均视为手术时常见的误分类。AKI 在索引时间固定后，基于带时间戳的实验室测量重新计算。动态时间点特异特征从原始时间戳重建，数据划分时尊重患者身份，配对 bootstrap 分析直接检验了模型对肌酐预测变量和基线来源的敏感性。",
    "Several limitations remain. MIMIC-IV represents one tertiary academic center, and cardiac surgery dominated the cohort, limiting transportability. The retrospective design is vulnerable to selection and measurement processes embedded in routine care. AKI ascertainment excluded urine output; an admission-based fallback SCr may not represent stable outpatient kidney function; and procedure timing was available only by date. Intraoperative exposures, fluid balance, medication timing, and potentially modifiable factors such as hypotension were not modeled comprehensively [29]. Internal random-split validation does not assess temporal or geographic transportability. Finally, SHAP describes predictive attribution and cannot identify causal or modifiable risk factors.":
        "本研究仍存在若干局限。MIMIC-IV 代表单个三级学术医疗中心，且队列以心脏手术为主，限制了可迁移性。回顾性设计易受常规医疗中选择过程和测量过程影响。AKI 判定未纳入尿量；基于入院的备用 SCr 可能不能代表稳定的门诊肾功能；手术时间仅精确到日期。术中暴露、液体平衡、用药时间，以及低血压等潜在可干预因素未被全面建模 [29]。内部随机划分验证无法评估时间或地理可迁移性。最后，SHAP 描述的是预测归因，不能识别因果性或可干预风险因素。",
    "Dynamic landmark models based on routinely available peri-ICU data achieved moderate internal discrimination for incident postoperative AKI. XGBoost performed best at ICU admission and 6 h, whereas logistic regression performed best at 24 h; no algorithm showed clear superiority. Creatinine trajectories contributed materially at 24 h, but non-creatinine information retained discrimination. External and temporal validation, urine-output outcome assessment, recalibration, and prospective workflow evaluation are required before clinical deployment.":
        "基于常规可获得围 ICU 数据的动态时间点模型，对术后新发 AKI 实现了中等程度的内部区分能力。XGBoost 在 ICU 入科和 6 小时表现最佳，而逻辑回归在 24 小时表现最佳；没有算法显示明确优越性。24 小时时肌酐轨迹具有重要贡献，但非肌酐信息仍保留区分能力。在临床部署前，仍需外部和时间验证、纳入尿量的结局评估、重新校准以及前瞻性工作流程评估。",
    "AKI, acute kidney injury; AUPRC, area under the precision-recall curve; AUROC, area under the receiver-operating-characteristic curve; CI, confidence interval; DCA, decision-curve analysis; ICU, intensive care unit; IQR, interquartile range; KDIGO, Kidney Disease: Improving Global Outcomes; SCr, serum creatinine; SHAP, SHapley Additive exPlanations.":
        "AKI，急性肾损伤；AUPRC，精确率-召回率曲线下面积；AUROC，受试者工作特征曲线下面积；CI，置信区间；DCA，决策曲线分析；ICU，重症监护病房；IQR，四分位距；KDIGO，Kidney Disease: Improving Global Outcomes；SCr，血清肌酐；SHAP，SHapley Additive exPlanations。",
    "MIMIC-IV was created under institutional review board approval at Beth Israel Deaconess Medical Center and the Massachusetts Institute of Technology, with waiver of individual informed consent for the deidentified research resource. The present study used only deidentified data accessed through credentialed PhysioNet access and did not involve direct contact with human participants. No additional local ethics approval was required for this retrospective analysis of publicly available deidentified data.":
        "MIMIC-IV 数据库在 Beth Israel Deaconess Medical Center 和 Massachusetts Institute of Technology 的机构审查委员会批准下建立，并因其为去标识化研究资源而免除个体知情同意。本研究仅使用通过 PhysioNet 认证访问获得的去标识化数据，未直接接触人类受试者。对于这一基于公开可获得去标识化数据的回顾性分析，无需额外本地伦理批准。",
    "Not applicable.": "不适用。",
    "MIMIC-IV version 3.1 is available through PhysioNet to credentialed users who complete the required training and sign the data use agreement. The authors are not permitted to redistribute patient-level MIMIC-IV data or derived patient-level analytic datasets. The analytic code will be made available in a public repository before publication; the repository URL and DOI are to be provided before publication. [18]":
        "MIMIC-IV 3.1 版可通过 PhysioNet 提供给完成所需培训并签署数据使用协议的认证用户。作者无权再分发患者层面的 MIMIC-IV 数据或由其衍生的患者层面分析数据集。分析代码将在发表前存放于公开代码仓库；仓库 URL 和 DOI 将在发表前提供。[18]",
    "The authors declare that they have no competing interests.": "作者声明不存在利益冲突。",
    "This research received no specific grant from any funding agency in the public, commercial, or not-for-profit sectors.":
        "本研究未获得公共、商业或非营利部门任何资助机构的专项资助。",
    "Bizhi Wei: Conceptualization, methodology, investigation, data curation, formal analysis, visualization, writing – original draft, writing – review and editing, project administration, and supervision. The author read and approved the final manuscript.":
        "Bizhi Wei：概念化、方法学、调查、数据整理、正式分析、可视化、初稿撰写、审阅与编辑、项目管理和监督。作者已阅读并批准最终稿件。",
    "The authors thank the developers and maintainers of the MIMIC-IV database and PhysioNet.":
        "作者感谢 MIMIC-IV 数据库和 PhysioNet 的开发者与维护者。",
    "During manuscript preparation, the authors used ChatGPT for language editing, structural organization, and formatting assistance. The authors reviewed and verified all content, analyses, interpretations, and references, and take full responsibility for the final manuscript.":
        "稿件准备过程中，作者使用 ChatGPT 进行语言编辑、结构组织和格式协助。作者已审阅并核实所有内容、分析、解释和参考文献，并对最终稿件承担全部责任。",
    "Additional file 1 (.docx and .pdf): Supplementary Tables S1-S7 and Supplementary Figures S1-S6.":
        "附加文件 1（.docx 和 .pdf）：补充表 S1–S7 和补充图 S1–S6。",
    "Additional file 2 (.docx and .csv): Completed TRIPOD+AI checklist.":
        "附加文件 2（.docx 和 .csv）：已完成的 TRIPOD+AI 核对清单。",
    "Table 1. Baseline characteristics by incident acute kidney injury status":
        "表 1. 按新发急性肾损伤状态分层的基线特征",
    "Values are median [IQR] or n (%). SMD, standardized mean difference.":
        "数值为中位数 [IQR] 或 n（%）。SMD，标准化均数差。",
    "Table 2. Performance of selected models at each prediction landmark":
        "表 2. 各预测时间点所选模型的性能",
    "Confidence intervals were obtained by patient-level bootstrap resampling of the held-out test set.":
        "置信区间通过对保留测试集进行患者层面 bootstrap 重采样获得。",
    "Table 3. Sensitivity analyses at the 24-hour landmark":
        "表 3. 24 小时预测时间点的敏感性分析",
    "AUROC differences were estimated in identical test patients using paired patient-level bootstrap resampling.":
        "AUROC 差异在相同测试患者中通过配对患者层面 bootstrap 重采样估计。",
    "Figure 1. Cohort derivation and dynamic landmark risk sets. The strict surgical ICU cohort was restricted to incident-AKI-evaluable admissions. The 6-h and 24-h risk sets exclude AKI with onset at or before each landmark. SCr, serum creatinine.":
        "图 1. 队列构建流程与动态时间点风险集。严格外科 ICU 队列限制为可评估新发 AKI 的住院记录。6 小时和 24 小时风险集排除了在相应时间点之前或当时已经发生的 AKI。SCr，血清肌酐。",
    "Figure 2. Discrimination across dynamic prediction landmarks. Receiver-operating-characteristic curves for four model families at ICU admission, 6 h, and 24 h. Landmark populations and remaining outcome windows differ and should not be interpreted as paired longitudinal comparisons.":
        "图 2. 动态预测时间点的区分度。展示四类模型在 ICU 入科、6 小时和 24 小时的受试者工作特征曲线。不同时间点的人群和剩余结局窗口不同，不应解释为配对纵向比较。",
    "Figure 3. Calibration and clinical net benefit. Top row: observed versus predicted risk in ten equal-frequency groups. Bottom row: decision curves relative to treat-all and treat-none strategies. DCA, decision-curve analysis.":
        "图 3. 校准度与临床净获益。上排：十个等频组中的观察风险与预测风险。下排：相对于全部干预和全部不干预策略的决策曲线。DCA，决策曲线分析。",
    "Figure 4. Global feature importance at 6 and 24 hours. Mean absolute SHAP values for the 12 leading XGBoost features. Values quantify global model attribution and do not represent causal or modifiable effects.":
        "图 4. 6 小时和 24 小时全局特征重要性。展示 XGBoost 前 12 个主要特征的平均绝对 SHAP 值。该值量化全局模型归因，不代表因果效应或可干预效应。",
}


ROW_LABEL_TRANSLATIONS = {
    "Characteristic": "特征",
    "Overall": "总体",
    "No incident AKI": "未发生新发 AKI",
    "Incident AKI": "新发 AKI",
    "SMD": "SMD",
    "Missing, n": "缺失，n",
    "Age, years": "年龄，岁",
    "Female sex": "女性",
    "Race: White": "种族：白人",
    "Race: Black": "种族：黑人",
    "Race: Asian": "种族：亚洲人",
    "Race: Hispanic/Latino": "种族：西班牙裔/拉丁裔",
    "Race: Other/unknown": "种族：其他/未知",
    "Admission: Elective/same-day surgical": "入院：择期/当日手术",
    "Admission: Urgent/emergency": "入院：急诊/紧急",
    "Admission: Observation/other": "入院：观察/其他",
    "Baseline serum creatinine, mg/dL": "基线血清肌酐，mg/dL",
    "Pre-index 7-day creatinine baseline": "索引前 7 天肌酐基线",
    "Charlson comorbidity score": "Charlson 合并症评分",
    "Congestive heart failure": "充血性心力衰竭",
    "Hypertension": "高血压",
    "Diabetes mellitus": "糖尿病",
    "Chronic kidney disease": "慢性肾病",
    "Chronic pulmonary disease": "慢性肺病",
    "Liver disease": "肝病",
    "Cancer": "肿瘤",
    "Peripheral vascular disease": "外周血管疾病",
    "Stroke": "卒中",
    "Myocardial infarction": "心肌梗死",
    "Obesity": "肥胖",
    "Anemia": "贫血",
    "Cardiac surgery": "心脏手术",
    "Non-cardiac surgery": "非心脏手术",
    "Vascular surgery": "血管手术",
    "General/GI/hepatobiliary surgery": "普通外科/胃肠/肝胆手术",
    "Major orthopedic surgery": "重大骨科手术",
    "Neurosurgery": "神经外科手术",
    "Thoracic/respiratory surgery": "胸外/呼吸系统手术",
    "First ICU: CVICU": "首次 ICU：CVICU",
    "First ICU: SICU": "首次 ICU：SICU",
    "First ICU: TSICU": "首次 ICU：TSICU",
    "First ICU: Other surgical ICU": "首次 ICU：其他外科 ICU",
    "Landmark": "预测时间点",
    "Selected model": "所选模型",
    "Test n": "测试集 n",
    "Event rate": "事件率",
    "AUROC (95% CI)": "AUROC（95% CI）",
    "AUPRC (95% CI)": "AUPRC（95% CI）",
    "Brier": "Brier 评分",
    "Intercept": "截距",
    "Slope": "斜率",
    "Logistic regression": "逻辑回归",
    "Analysis": "分析",
    "Model": "模型",
    "Reference AUROC": "参考 AUROC",
    "Sensitivity AUROC": "敏感性 AUROC",
    "AUROC difference (95% CI)": "AUROC 差异（95% CI）",
    "No-creatinine model": "不含肌酐模型",
    "Random forest": "随机森林",
    "Pre-index baseline-only retraining": "仅索引前基线者重新训练",
}


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill="F2F4F7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_doc_defaults(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(5)


def add_para(doc, text, style=None, italic=False):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.italic = italic
    return p


def translate_text(text):
    if text in BODY_TRANSLATIONS:
        return BODY_TRANSLATIONS[text]
    if text in PARAGRAPH_TRANSLATIONS:
        return PARAGRAPH_TRANSLATIONS[text]
    # Keep references and author metadata mostly unchanged.
    return text


def add_translated_table(doc, source_table):
    rows, cols = len(source_table.rows), len(source_table.columns)
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            raw = cell.text.replace("\n", " ").strip()
            text = ROW_LABEL_TRANSLATIONS.get(raw, raw)
            set_cell_text(tbl.cell(i, j), text, bold=(i == 0))
            if i == 0:
                shade_cell(tbl.cell(i, j))
    return tbl


def build_chinese_manuscript():
    src = Document(SOURCE_DOCX)
    doc = Document()
    set_doc_defaults(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE_ZH)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    add_para(doc, "中文审阅稿（由英文投稿稿翻译生成；请以英文正式稿为投稿版本）", italic=True)

    # Paragraphs, excluding the original title and table captions that are re-added with tables.
    skip_indices = {0, 126, 127, 128, 129, 130, 131}
    for i, para in enumerate(src.paragraphs):
        text = para.text.strip()
        if not text or i in skip_indices:
            continue
        style = para.style.name
        zh = translate_text(text)
        if i >= 95 and text[:2].isdigit() or (i >= 95 and text[:1].isdigit()):
            # Keep references in English.
            zh = text
        if style == "Heading 1":
            add_para(doc, zh, "Heading 1")
        elif style == "Heading 2":
            add_para(doc, zh, "Heading 2")
        elif style == "Title":
            continue
        else:
            add_para(doc, zh)

    # Add tables at the end, matching source manuscript layout.
    captions = [
        ("表 1. 按新发急性肾损伤状态分层的基线特征", "数值为中位数 [IQR] 或 n（%）。SMD，标准化均数差。"),
        ("表 2. 各预测时间点所选模型的性能", "置信区间通过对保留测试集进行患者层面 bootstrap 重采样获得。"),
        ("表 3. 24 小时预测时间点的敏感性分析", "AUROC 差异在相同测试患者中通过配对患者层面 bootstrap 重采样估计。"),
    ]
    doc.add_page_break()
    add_para(doc, "中文表格", "Heading 1")
    for idx, table in enumerate(src.tables):
        add_para(doc, captions[idx][0], "Heading 2")
        add_para(doc, captions[idx][1], italic=True)
        add_translated_table(doc, table)
        doc.add_paragraph()

    doc.save(ZH_DOCX)


CHECKLIST_TEXT = """# 投稿前人工核对清单

用途：这是人工投稿前最后一轮审阅清单。建议不要在这一步重新计算模型，而是核对文件、数字、合规性和期刊格式。

## 1. 作者、单位和声明

- [ ] 作者信息是否最终确认：Bizhi Wei*（唯一作者）
- [ ] 通讯作者姓名、地址、邮箱是否正确
- [ ] 单位英文是否与学校/期刊系统一致
- [ ] Ethics、Consent、Funding、Competing interests、Authors' contributions 是否符合目标期刊格式
- [ ] AI-assisted editing disclosure 是否符合期刊政策
- [ ] 代码仓库 URL/DOI 是否需要在投稿前补充；若尚未生成，是否保留 “to be provided before publication”

## 2. 核心结果数字一致性

- [ ] 摘要、Results、Table 1、Table 2、Figure 1 中队列人数是否一致
- [ ] strict cohort：11,943；incident-AKI evaluable：10,877；AKI：4,531（41.7%）是否一致
- [ ] AKI 分期：Stage 1 84.9%，Stage 2 9.8%，Stage 3 5.3% 是否一致
- [ ] 0 h、6 h、24 h 风险集人数和事件数是否一致
- [ ] 主要模型 AUROC/AUPRC/Brier 是否与表格一致
- [ ] no-creatinine 和 pre-index baseline-only 敏感性分析数字是否与补充材料一致
- [ ] eICU 可评价样本量、0 h/6 h/24 h AUROC 与重校准指标是否与 Figure 4、Table 3 一致

## 3. 方法学风险点

- [ ] 分析单位是否始终写作“first ICU stay per hospital admission”
- [ ] Index time 是否始终为 ICU admission time
- [ ] AKI 结局是否明确只使用 KDIGO serum creatinine，不使用 urine output
- [ ] baseline SCr 定义和备用方案是否清楚
- [ ] 0 h/6 h/24 h landmark 数据集是否说明排除已发生 AKI 者
- [ ] 是否明确排除 whole-follow-up summary、mortality、length of stay、AKI-derived variables 等泄漏变量
- [ ] train/test split 是否说明按 subject_id/patient 分组，避免同一患者跨训练和测试集

## 4. 图表与补充材料

- [ ] Figure 1 队列流程是否清楚呈现纳入、排除和 0 h/6 h/24 h risk set
- [ ] Figure 2 ROC 曲线字体、图例和线条是否清晰
- [ ] Figure 3 calibration + DCA 是否能脱离正文理解
- [ ] Figure 4 eICU ROC 与 held-out hospital 重校准前后 calibration 是否清晰且数字一致
- [ ] Supplementary Figure S6 SHAP 是否标明“归因而非因果/可干预性”
- [ ] Table 1 是否包含必要临床基线特征
- [ ] 表格缩写是否全部解释
- [ ] Supplementary figures/tables 编号是否连续且与正文引用一致

## 5. MIMIC-IV 数据合规

- [ ] 投稿附件中不包含任何患者级 MIMIC 原始数据
- [ ] 不上传 derived patient-level analytic CSV
- [ ] 不包含任何可识别个体信息
- [ ] Data availability 明确说明 MIMIC-IV 需 PhysioNet credentialed access
- [ ] 公开代码仓库不包含原始数据、患者级输出或本地路径/敏感信息

## 6. 语言与结论边界

- [ ] 删除或弱化 “clinically useful”“ready for deployment”“highly accurate”等过度表述
- [ ] 避免把 SHAP 或 subgroup findings 写成因果或可干预因素证据
- [ ] 明确区分内部验证、时间验证和特征协调的 eICU 外部验证，避免暗示完整 MIMIC-IV 模型已被逐变量直接外部验证
- [ ] Discussion 中是否充分说明单中心、心脏手术占比高、无尿量、手术时间仅日期级等限制
- [ ] Conclusion 是否保持克制，强调仍需本地重校准、前瞻性静默验证和临床工作流评估

## 7. 目标期刊格式

- [ ] Abstract 字数和结构是否符合目标期刊
- [ ] 正文图表数量是否符合目标期刊
- [ ] Reference 格式是否符合 Vancouver/目标期刊要求
- [ ] 是否需要单独 title page、cover letter、highlights 或 graphical abstract
- [ ] TRIPOD+AI checklist 是否作为附加文件上传
- [ ] Supplementary material 是否按期刊命名规则上传

## 8. 最后打开文件检查

- [ ] Word 主稿打开无乱码、无表格跑版
- [ ] PDF 与 Word 内容一致
- [ ] 所有图像文件能单独打开，分辨率足够
- [ ] ZIP 包中没有临时文件、QA 截图、缓存、原始数据或脚本日志
- [ ] 文件名符合投稿系统要求，不含中文特殊符号或空格问题

## 人工备注

| 检查项 | 状态 | 备注 |
|---|---|---|
| 作者与声明 | 待检查 |  |
| 核心结果数字 | 待检查 |  |
| 方法学防泄漏 | 待检查 |  |
| 图表质量 | 待检查 |  |
| 数据合规 | 待检查 |  |
| 目标期刊格式 | 待检查 |  |
"""


def build_checklist_files():
    CHECKLIST_MD.write_text(CHECKLIST_TEXT, encoding="utf-8")
    doc = Document()
    set_doc_defaults(doc)
    in_notes_table = False
    for line in CHECKLIST_TEXT.splitlines():
        if line.startswith("| 检查项"):
            in_notes_table = True
            tbl = doc.add_table(rows=1, cols=3)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.style = "Table Grid"
            for j, h in enumerate(["检查项", "状态", "备注"]):
                set_cell_text(tbl.cell(0, j), h, bold=True)
                shade_cell(tbl.cell(0, j))
            for item in ["作者与声明", "核心结果数字", "方法学防泄漏", "图表质量", "数据合规", "目标期刊格式"]:
                cells = tbl.add_row().cells
                set_cell_text(cells[0], item)
                set_cell_text(cells[1], "待检查")
                set_cell_text(cells[2], "")
            continue
        if in_notes_table and line.startswith("|"):
            continue
        in_notes_table = False
        if line.startswith("# "):
            add_para(doc, line[2:], "Heading 1")
        elif line.startswith("## "):
            add_para(doc, line[3:], "Heading 2")
        elif line.startswith("- [ ] "):
            add_para(doc, "☐ " + line[6:])
        elif line.strip():
            add_para(doc, line)
    doc.save(CHECKLIST_DOCX)


def zip_outputs():
    zip_path = PACKAGE.parent / "manuscript_package_v8_critical_care.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for p in PACKAGE.rglob("*"):
            if not p.is_file():
                continue
            rel_parts = p.relative_to(PACKAGE).parts
            if any(part.startswith("_render") or part.startswith("_qa") for part in rel_parts):
                continue
            if p.name.startswith("_extracted"):
                continue
            if p.name == "contact_sheet.png":
                continue
            if p.suffix.lower() in {".tmp", ".log"}:
                continue
            else:
                zf.write(p, p.relative_to(PACKAGE.parent))
    return zip_path


if __name__ == "__main__":
    build_chinese_manuscript()
    build_checklist_files()
    zp = zip_outputs()
    print(f"Wrote: {ZH_DOCX}")
    print(f"Wrote: {CHECKLIST_MD}")
    print(f"Wrote: {CHECKLIST_DOCX}")
    print(f"Updated zip: {zp}")
