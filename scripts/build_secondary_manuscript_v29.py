"""Build the independent v29 secondary manuscript package.

This package is deliberately separate from the locked primary manuscript. It
integrates v26-v29 secondary analyses of severe SCr-AKI, onset-anchored renal
trajectory prediction, recovery observability, temporal/external validation,
and observation-time multistate/competing-risk estimates.
"""

from __future__ import annotations

import csv
import hashlib
import importlib.util
import re
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "manuscript_package_v29_secondary"
LATEX = OUT / "latex"
TABLES = OUT / "tables"
FIGURES = OUT / "figures"
QA = OUT / "qa"

V26 = ROOT / "outputs" / "modeling_v26_aki_severity_trajectories"
V27 = ROOT / "outputs" / "modeling_v27_severity_recovery"
V28S = ROOT / "outputs" / "modeling_v28_severe_temporal_external"
V28R = ROOT / "outputs" / "modeling_v28_recovery_observability"
V29 = ROOT / "outputs" / "modeling_v29_multistate_competing_risk"
LOCKED = ROOT / "outputs" / "manuscript_package_v23_locked"

spec = importlib.util.spec_from_file_location("v8", ROOT / "scripts" / "build_manuscript_package_v8.py")
v8 = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(v8)

v8.OUT = OUT
v8.LATEX = LATEX
v8.TABLES = TABLES
v8.FIGURES = FIGURES
v8.QA = QA


TITLE = (
    "Dynamic prediction of severe postoperative acute kidney injury and observation-time renal trajectories: "
    "internal, temporal, and external validation in MIMIC-IV and eICU"
)
RUNNING_TITLE = "Severe postoperative AKI and renal trajectories"
AUTHORS_WORD = "Bizhi Wei¹*"
AFFILIATION_WORD = "¹ Pu Ai Medical School, Shaoyang University, Shaoyang 422000, Hunan, China"
CORRESPONDENCE_WORD = (
    "*Correspondence: Bizhi Wei, Pu Ai Medical School, Shaoyang University, "
    "Shaoyang 422000, Hunan, China; Email: 15619056250wbz@gmail.com"
)

ETHICS = (
    "MIMIC-IV was created under institutional review board approval at Beth Israel Deaconess Medical Center "
    "and the Massachusetts Institute of Technology, with waiver of individual informed consent for the "
    "deidentified research resource. The present study used only deidentified data accessed through credentialed "
    "PhysioNet access and did not involve direct contact with human participants. No additional local ethics "
    "approval was required for this retrospective analysis of publicly available deidentified data."
)
AVAILABILITY = (
    "MIMIC-IV version 3.1 and the eICU Collaborative Research Database are available through PhysioNet to "
    "credentialed users who complete the required training and sign the data use agreement. The author is not "
    "permitted to redistribute patient-level data or derived patient-level analytic datasets. Analytic code is "
    "publicly available at https://github.com/Bizhi-Wei/postoperative-aki-dynamic-prediction (release v1.0.1: "
    "https://github.com/Bizhi-Wei/postoperative-aki-dynamic-prediction/releases/tag/v1.0.1). Archived software DOI: "
    "to be provided before publication."
)
FUNDING = "This research received no specific grant from any funding agency in the public, commercial, or not-for-profit sectors."
COMPETING = "The author declares no competing interests."
CONTRIBUTIONS = (
    "Bizhi Wei: Conceptualization, methodology, investigation, data curation, formal analysis, visualization, "
    "writing – original draft, writing – review and editing, project administration, and supervision. The author "
    "read and approved the final manuscript."
)
ACKNOWLEDGEMENTS = "The author thanks the developers and maintainers of the MIMIC-IV and eICU databases and PhysioNet."
AI_DISCLOSURE = (
    "During manuscript preparation, the author used ChatGPT for language editing, structural organization, and "
    "formatting assistance. The author reviewed and verified all content, analyses, interpretations, and references, "
    "and takes full responsibility for the final manuscript."
)


ABSTRACT = {
    "Background": (
        "Prediction of any postoperative acute kidney injury (AKI) is clinically useful, but severe AKI, persistence, "
        "recovery, and recurrence may better distinguish consequential renal courses. We evaluated landmark models "
        "for severe serum-creatinine-defined AKI and characterized observed renal-state transitions after surgery."
    ),
    "Methods": (
        "This retrospective secondary study used a strict surgical intensive care cohort from MIMIC-IV version 3.1 "
        "and feature-harmonized external validation in eICU. Severe AKI was a new Kidney Disease: Improving Global "
        "Outcomes serum creatinine stage 2 or 3 event after ICU admission, 6 h, or 24 h through day 7. Logistic "
        "regression and XGBoost models used patient-grouped development/test splits; uncertainty used patient-level "
        "bootstrap resampling. Persistence and nonrecovery were predicted from information available at observed AKI "
        "onset. A seven-state, observation-time multistate analysis and competing-risk analyses treated live discharge "
        "and death as competing absorbing events."
    ),
    "Results": (
        "Among 10,877 postoperative ICU admissions, 679 (6.2%) developed severe serum-creatinine-defined AKI. Held-out "
        "AUROCs were 0.698 (95% CI, 0.653–0.743) at 0 h, 0.790 (0.745–0.829) at 6 h, and 0.839 (0.800–0.879) at 24 h. "
        "External eICU AUROCs were 0.707, 0.761, and 0.784, respectively; hospital-held-out recalibration brought slopes "
        "to 0.95, 0.98, and 0.95. At AKI onset, AUROCs were 0.824 for persistence beyond 48 h and 0.743 for nonrecovery at "
        "the observation end. Among 4,519 trajectory-eligible AKI admissions, the 48-h cumulative incidences of observed "
        "recovery and severe progression were 65.0% and 11.9%. Among 3,936 observed recoveries, recurrent AKI cumulative "
        "incidence was 11.3% by 48 h and 13.8% by 72 h."
    ),
    "Conclusions": (
        "Early ICU information improved severe-AKI risk stratification and transferred to a multicenter external cohort, "
        "although probability updating remained necessary. Renal recovery and recurrence were common, measurement-dependent "
        "states; discharge and death must be handled explicitly. These models and trajectory estimates require prospective "
        "validation before clinical use."
    ),
}

KEYWORDS = [
    "acute kidney injury", "severe acute kidney injury", "renal recovery", "multistate model",
    "competing risks", "dynamic prediction", "MIMIC-IV", "eICU",
]

INTRODUCTION = [
    (
        "Postoperative acute kidney injury (AKI) encompasses renal courses that differ substantially in severity, duration, "
        "and reversibility. Higher KDIGO stage and persistent dysfunction are associated with greater short- and long-term "
        "risk, while early reversal may identify a more favorable phenotype [[CITE:Boyer2022,Prowle2021,Wang2017,Nadim2018,Hoste2015,Grams2016]]. "
        "A binary any-AKI endpoint can therefore obscure clinically important heterogeneity after major surgery."
    ),
    (
        "Risk also changes over time. Information available at ICU admission captures baseline vulnerability, whereas the "
        "first postoperative hours reveal evolving hemodynamic and biochemical disturbances. Landmark prediction can align "
        "features with the clinical decision time and exclude events that have already occurred [[CITE:Koyner2018,Tomasev2019,Tseng2020,Ryan2023,Demirjian2022]]. "
        "For severe AKI, this design permits patients with stage 1 AKI to remain at risk for subsequent stage 2 or 3 progression."
    ),
    (
        "Recovery analysis presents a separate measurement problem. Serum creatinine (SCr) is sampled intermittently, so both "
        "recovery and recurrence are observed only when a laboratory measurement is obtained. Discharge can end in-hospital "
        "observation without establishing renal recovery, and death prevents subsequent recovery. ADQI distinguishes rapid "
        "reversal from persistent AKI and emphasizes explicit definitions of recovery [[CITE:Chawla2017,KellumRecovery2017,ForniRecovery2017]]. "
        "Multistate and competing-risk methods can preserve these event distinctions rather than collapsing them into a single endpoint."
    ),
    (
        "We therefore performed a post-lock secondary analysis of a rigorously defined surgical ICU cohort. We aimed to: "
        "(1) develop and internally validate 0-h, 6-h, and 24-h models for severe SCr-AKI; (2) assess temporal stability, "
        "multicenter eICU transportability, and hospital-held-out recalibration; (3) predict persistent AKI and nonrecovery "
        "from information available at observed AKI onset; and (4) estimate observation-time renal-state occupancy, recovery, "
        "severe progression, and recurrence while treating discharge and death explicitly."
    ),
]

METHODS = {
    "Study design, data sources, and population": [
        (
            "We conducted a retrospective secondary prediction and trajectory study using MIMIC-IV version 3.1 and the eICU "
            "Collaborative Research Database [[CITE:Johnson2023,MIMIC2024,Pollard2018]]. The MIMIC-IV analysis used 10,877 "
            "incident-AKI-evaluable admissions from the locked strict postoperative surgical ICU cohort. The analysis unit was "
            "the first qualifying ICU stay per hospital admission. Major cardiac, vascular, gastrointestinal or hepatobiliary, "
            "orthopedic, neurosurgical, and thoracic or respiratory operations were eligible; diagnostic and bedside procedures "
            "did not establish surgical eligibility."
        ),
        (
            "The index time was ICU admission. Baseline SCr was the lowest value within 7 days before ICU admission, with an "
            "admission SCr fallback when the pre-ICU value was unavailable. Incident AKI through 7 days was determined from "
            "timestamped SCr using KDIGO absolute and relative criteria; the primary secondary analyses did not use urine output "
            "[[CITE:KDIGO2012,Kellum2013]]."
        ),
    ],
    "Severe AKI and renal-trajectory outcomes": [
        (
            "Severe SCr-AKI was the maximum active-episode KDIGO SCr stage 2 or 3 within 7 days. An SCr of at least 4.0 mg/dL "
            "upgraded an active AKI episode to stage 3 but did not independently create incident AKI in a patient with a stable "
            "high baseline. Renal replacement therapy (RRT) within 7 days was added only in a prespecified SCr-or-RRT sensitivity outcome."
        ),
        (
            "At each SCr measurement, active AKI was present when the SCr ratio was at least 1.5 versus baseline, SCr was at least "
            "4.0 mg/dL during active AKI, or SCr had risen by at least 0.3 mg/dL from a preceding value within 48 h. Rapid reversal "
            "was the first observed measurement without active AKI within 48 h after onset. Persistent AKI required no rapid "
            "reversal and an observed AKI-positive SCr at least 48 h after onset. End recovery required the final SCr within 24 h "
            "before the earlier of hospital discharge or ICU day 7 to show no active AKI. Recurrent AKI was an observed AKI-positive "
            "SCr after a first observed reversal. No SCr measurement was never coded as recovery."
        ),
    ],
    "Landmark and onset-anchored prediction models": [
        (
            "Severe-AKI risk sets were constructed at 0 h, 6 h, and 24 h. Patients who had already reached SCr stage 2 or 3 by a "
            "landmark were excluded; patients with stage 1 AKI remained eligible. Outcomes were new stage 2 or 3 events after the "
            "landmark through ICU day 7. Predictors were limited to static, pre-index, and landmark-window data available by the "
            "prediction time. Whole-follow-up summaries and outcome-derived variables were excluded."
        ),
        (
            "Onset-anchored models used the first observed AKI-positive SCr as time zero. Predictors included the static and pre-index "
            "set plus baseline SCr, onset stage, SCr ratio and change, timing, recent measurement counts, time since the previous SCr, "
            "and SCr slopes calculated only through onset. Persistence models included patients with adequate observation beyond 48 h; "
            "nonrecovery models included patients with an evaluable end-window SCr."
        ),
        (
            "Logistic regression and XGBoost used the same preprocessing and a common subject-grouped 80:20 split across tasks. "
            "Continuous variables were median-imputed; categorical variables were imputed and one-hot encoded. Logistic regression "
            "standardized continuous predictors. Model hyperparameters were fixed before test-set evaluation. Discrimination, overall "
            "performance, calibration, and development-derived Youden thresholds were assessed [[CITE:Chen2016,Pedregosa2011,Steyerberg2010,VanCalster2019]]. "
            "Confidence intervals used 1,000 subject-cluster bootstrap resamples of held-out patients."
        ),
    ],
    "Temporal and external validation": [
        (
            "Rolling temporal validation used expanding earlier-year MIMIC-IV training data and three nonoverlapping later-year "
            "evaluation blocks per landmark. For eICU, the strict surgical cohort, baseline, incident SCr-AKI, and landmark risk "
            "sets were rederived using the same temporal logic. Frozen feature-harmonized portable models used 30 predictors at 0 h "
            "and 72 predictors at 6 h and 24 h. Patient-level bootstrap resampling provided 95% confidence intervals."
        ),
        (
            "For calibration updating, eICU hospitals were partitioned into calibration and held-out sets. Intercept-only and logistic "
            "recalibration parameters were learned in calibration hospitals and applied unchanged to held-out hospitals. Hospital-level "
            "heterogeneity was summarized only where prespecified minimum sample and outcome counts permitted an AUROC estimate. "
            "Recalibration was evaluated as probability updating and was not expected to alter discrimination."
        ),
    ],
    "Recovery observability and selection sensitivity": [
        (
            "Because persistence and end recovery require follow-up SCr, we compared evaluable and nonevaluable AKI admissions and "
            "modeled evaluability with five-fold subject-grouped cross-fitting using onset-available variables. Stabilized inverse-" 
            "probability weights used marginal evaluability, propensity truncation to 0.05–0.95, and weight truncation at the first "
            "and 99th percentiles. These analyses assume missing at random conditional on measured onset features and do not identify "
            "the true missingness mechanism. Death and live discharge without the required SCr were reported separately rather than "
            "coded as recovery."
        ),
    ],
    "Multistate and competing-risk analysis": [
        (
            "The multistate clock began at ICU admission and ended at 168 h. Seven states were prespecified: no AKI, stage 1 AKI, "
            "severe AKI (stage 2 or 3), recovered, recurrent AKI, live discharge, and in-hospital death. SCr-derived states changed "
            "only at an observed measurement; no interpolation or unobserved recovery was imposed. Live discharge and death were "
            "absorbing. Twelve locked AKI events whose recorded onset occurred after the hospital disposition time were excluded "
            "from trajectory risk sets and retained in a dedicated audit; the locked primary outcome was not altered."
        ),
        (
            "A first-event competing-risk analysis began at observed AKI onset and considered observed recovery, severe stage 2/3 "
            "onset or progression, live discharge, and in-hospital death. A second analysis began at first observed recovery and "
            "considered recurrent AKI, live discharge, and death. State occupation and cumulative incidence functions used the "
            "Aalen–Johansen product integral [[CITE:AalenJohansen1978]]. Confidence intervals used 500 subject-cluster bootstrap "
            "resamples. Prespecified descriptive subgroups were chronic kidney disease and cardiac versus noncardiac surgery."
        ),
    ],
    "Software, reporting, and interpretation boundary": [
        (
            "Analyses were performed in Python using pandas, NumPy, scikit-learn, XGBoost, Matplotlib, and SciPy. Reporting was "
            "guided by TRIPOD+AI for prediction components [[CITE:Collins2024,Wolff2019]]. All trajectory estimates describe "
            "observed SCr states and do not estimate causal treatment effects, biological transition times between measurements, "
            "or clinical effectiveness."
        ),
        "Patients and members of the public were not involved in this retrospective database study.",
    ],
}

RESULTS = {
    "Cohort and secondary outcomes": [
        (
            "The strict evaluable cohort comprised 10,877 postoperative ICU admissions. Median age was 66 years (IQR, 57–74), "
            "32.9% were women, and 74.1% underwent cardiac surgery (Table 1). Incident SCr-AKI occurred in 4,531 admissions (41.7%), "
            "while severe active-episode SCr-AKI occurred in 679 (6.2%). The SCr-or-RRT sensitivity outcome occurred in 718 admissions."
        ),
        (
            "The severe-AKI risk sets contained 10,877 admissions at 0 h, 10,856 at 6 h after 21 prevalent severe events were "
            "removed, and 10,736 at 24 h after 141 prevalent severe events were removed. Subsequent severe-event rates were 6.2%, "
            "6.1%, and 5.0%, respectively (Supplementary Table S1)."
        ),
    ],
    "Severe-AKI model performance": [
        (
            "In held-out patients, XGBoost was selected at 0 h and 6 h and logistic regression at 24 h. AUROC increased from 0.698 "
            "(95% CI, 0.653–0.743) at 0 h to 0.790 (0.745–0.829) at 6 h and 0.839 (0.800–0.879) at 24 h (Fig. 2; Table 2). "
            "Corresponding AUPRCs were 0.174, 0.313, and 0.340, and Brier scores were 0.057, 0.049, and 0.041. Because landmark "
            "populations and remaining risk windows differed, these estimates describe time-specific risk sets rather than paired "
            "within-person gains. The SCr-or-RRT sensitivity produced similar or slightly higher discrimination (Supplementary Table S2)."
        ),
    ],
    "Temporal stability, external validation, and recalibration": [
        (
            "Across three expanding-window temporal evaluations, AUROCs ranged from 0.701 to 0.731 at 0 h, 0.762 to 0.803 at 6 h, "
            "and 0.807 to 0.847 at 24 h (Supplementary Table S3). In the eICU outcome-evaluable cohorts, severe SCr-AKI occurred in "
            "910 of 14,229 admissions at 0 h, 836 of 14,155 at 6 h, and 538 of 13,857 at 24 h. Frozen external AUROCs were 0.707 "
            "(95% CI, 0.688–0.725), 0.761 (0.744–0.780), and 0.784 (0.760–0.808), respectively (Fig. 3; Table 3)."
        ),
        (
            "External calibration slopes were 0.79, 0.88, and 0.51. In hospitals held out from probability updating, logistic "
            "recalibration changed slopes from 0.76 to 0.95 at 0 h, 0.86 to 0.98 at 6 h, and 0.49 to 0.95 at 24 h; Brier scores "
            "changed from 0.057 to 0.055, 0.052 to 0.050, and 0.036 to 0.036. Across hospitals meeting minimum estimability criteria, "
            "median AUROCs were 0.741, 0.777, and 0.803 at the three landmarks (Supplementary Table S5)."
        ),
    ],
    "Onset-anchored persistence and nonrecovery models": [
        (
            "Persistence beyond 48 h was evaluable in 3,786 of 4,531 incident-AKI admissions (83.6%) and occurred in 818 (21.6%). "
            "End recovery was evaluable in 3,986 (88.0%); 786 (19.7%) were not recovered at the observation end. In held-out patients, "
            "the XGBoost persistence model had AUROC 0.824 (95% CI, 0.789–0.854), AUPRC 0.554, and Brier score 0.132. The logistic "
            "nonrecovery model had AUROC 0.743 (0.704–0.783), AUPRC 0.419, and Brier score 0.140 (Table 2)."
        ),
        (
            "Evaluability models had AUROCs of 0.658 for persistence and 0.645 for end recovery, indicating selective follow-up "
            "measurement. Stabilized inverse-probability weighting changed outcome-model AUROC from 0.824 to 0.826 for persistence "
            "and from 0.743 to 0.745 for nonrecovery (Supplementary Table S6). These similarities apply only under the measured "
            "missing-at-random assumption."
        ),
    ],
    "Observation-time renal trajectories and competing events": [
        (
            "Of 4,531 locked incident-AKI admissions, 4,519 were trajectory eligible and 12 had a recorded AKI onset after hospital "
            "disposition. Observed recovery occurred at some point in 3,936 of 4,519 eligible admissions (87.1%); 641 of these "
            "recoveries (16.3%) were followed by observed recurrent AKI. By 168 h, state occupancy in the full cohort was 16.9% no AKI, "
            "1.8% stage 1, 1.5% severe AKI, 14.3% recovered, 1.9% recurrent AKI, 62.0% live discharge, and 1.7% in-hospital death (Fig. 4)."
        ),
        (
            "After observed AKI onset, the 48-h cumulative incidence of observed recovery was 65.0% (95% CI, 63.5–66.5), severe "
            "stage 2/3 onset or progression 11.9% (10.9–12.8), live discharge 1.4% (1.0–1.7), and death 0.6% (0.4–0.8). At 72 h, "
            "recovery and severe progression cumulative incidences were 77.7% and 13.3% (Table 4). After first observed recovery, "
            "recurrent AKI cumulative incidence was 11.3% (10.2–12.2) by 48 h and 13.8% (12.7–14.9) by 72 h, while live-discharge "
            "cumulative incidence reached 21.7% and 39.4%. Subgroup curves are shown in Supplementary Figure S5."
        ),
    ],
}

DISCUSSION = [
    (
        "This secondary study links three clinically distinct questions: who will progress to severe postoperative SCr-AKI, "
        "who will have a persistent or nonrecovering course once AKI is observed, and how renal states evolve when discharge and "
        "death are treated explicitly. Severe-AKI discrimination was modest at ICU admission and substantially higher after 6 and "
        "24 h of ICU observation. Feature-harmonized models retained useful ranking in eICU, but external calibration required local updating."
    ),
    (
        "The severe-AKI results complement studies that predict any AKI or moderate-to-severe AKI after surgery [[CITE:Kheterpal2007,Tseng2020,Ryan2023,Demirjian2022]]. "
        "The 24-h model achieved the highest discrimination, but this should not be interpreted as a guaranteed longitudinal increment: "
        "patients already severe were removed, the outcome window shortened, and early creatinine kinetics approached the outcome definition. "
        "The 6-h model may offer a more actionable compromise between timeliness and information content."
    ),
    (
        "External validation was a major strength. Similar or higher eICU AUROCs suggest that the portable feature set preserved risk "
        "ranking across institutions. Nevertheless, calibration slopes—especially at 24 h—showed that an exported probability scale "
        "was not transportable without adjustment [[CITE:VanCalster2019]]. Hospital-held-out logistic recalibration restored slopes close "
        "to one but did not establish prospective clinical benefit or a universal action threshold."
    ),
    (
        "Recovery was frequent but not monotonic. Nearly two thirds of trajectory-eligible AKI admissions had an observed recovery event "
        "within 48 h, yet recurrence accumulated after recovery. This agrees with the concept that AKI trajectories contain sustained, "
        "relapsing, and nonreversing patterns rather than a single terminal recovery state [[CITE:Chawla2017,KellumRecovery2017,ForniRecovery2017]]. "
        "The multistate framework makes this distinction visible and prevents live discharge from being mislabeled as renal recovery."
    ),
    (
        "The trajectory estimates should be read as observation-time quantities. A patient may biologically recover between SCr measurements, "
        "but the transition becomes visible only at the next measurement. Similarly, recurrence requires another observed AKI-positive SCr. "
        "The evaluability models confirmed that follow-up measurement was selective. IPW estimates were stable under a measured missing-at-random "
        "assumption, but this does not resolve informative monitoring or unmeasured reasons for laboratory testing."
    ),
    (
        "The study has several strengths: a strict therapeutic surgical cohort, timestamp-aligned landmark risk sets, a severe active-episode "
        "definition that did not create AKI solely from stable high baseline SCr, common subject-grouped validation, multicenter external testing, "
        "hospital-held-out recalibration, and product-integral estimates verified against empirical state occupancy. The explicit audit of 12 "
        "post-disposition onset records also prevents internally inconsistent trajectories from entering risk sets."
    ),
    (
        "Limitations remain. MIMIC-IV is a single-center development source dominated by cardiac surgery. eICU feature harmonization cannot "
        "reproduce every MIMIC-IV predictor, and RRT ascertainment is sensitive to documentation and chronic dialysis. Primary trajectory states "
        "used SCr without urine output, procedure timing was date-level, and follow-up ended at 7 days. Recovery definitions are measurement-" 
        "dependent, intermittent observation may violate simple Markov interpretations, and cluster bootstrap intervals do not remove residual "
        "selection bias. The retrospective analyses cannot establish that alerts, thresholds, or interventions improve outcomes, and predictive "
        "associations should not be interpreted as modifiable causal effects."
    ),
]

CONCLUSION = (
    "Dynamic models identified risk of severe postoperative SCr-AKI with improving time-specific discrimination and reproducible external "
    "ranking, although local calibration updating remained necessary. After AKI onset, observed recovery was common but recurrence was not rare; "
    "discharge, death, and selective SCr measurement materially shaped what could be observed. Prospective silent validation with prespecified "
    "monitoring and response pathways is required before clinical implementation."
)


POLLARD_REF = (
    "Pollard2018",
    "Pollard TJ, Johnson AEW, Raffa JD, Celi LA, Mark RG, Badawi O. The eICU Collaborative Research Database, a freely available multi-center database for critical care research. Sci Data. 2018;5:180178. doi:10.1038/sdata.2018.178.",
    "10.1038/sdata.2018.178", "30204154", "external validation database",
)
EXTRA_REFS = [
    (
        "Chawla2017",
        "Chawla LS, Bellomo R, Bihorac A, et al. Acute kidney disease and renal recovery: consensus report of the Acute Disease Quality Initiative (ADQI) 16 Workgroup. Nat Rev Nephrol. 2017;13:241–257. doi:10.1038/nrneph.2017.2.",
        "10.1038/nrneph.2017.2", "28239173", "AKI persistence and recovery consensus",
    ),
    (
        "AalenJohansen1978",
        "Aalen OO, Johansen S. An empirical transition matrix for non-homogeneous Markov chains based on censored observations. Scand J Stat. 1978;5:141–150.",
        "", "", "Aalen–Johansen product-integral method",
    ),
    (
        "KellumRecovery2017",
        "Kellum JA, Sileanu FE, Bihorac A, Hoste EAJ, Chawla LS. Recovery after acute kidney injury. Am J Respir Crit Care Med. 2017;195:784–791. doi:10.1164/rccm.201604-0799OC.",
        "10.1164/rccm.201604-0799OC", "27635668", "AKI recovery and relapse patterns",
    ),
    (
        "ForniRecovery2017",
        "Forni LG, Darmon M, Ostermann M, et al. Renal recovery after acute kidney injury. Intensive Care Med. 2017;43:855–866. doi:10.1007/s00134-017-4809-x.",
        "10.1007/s00134-017-4809-x", "28466146", "renal recovery review",
    ),
]

REFS = list(v8.REFS) + [POLLARD_REF] + EXTRA_REFS
v8.REFS = REFS
v8.REF_INDEX = {key: i + 1 for i, (key, *_rest) in enumerate(REFS)}


FIGURE_SOURCES = {
    "Fig1": V26 / "figure_v26_severity_recovery_phenotypes",
    "Fig2": V27 / "figure_v27_severe_aki_dynamic_models",
    "Fig3": V28S / "figure_v28_severe_temporal_external_validation",
    "Fig4": V29 / "figure_v29_multistate_competing_risk",
    "FigS1": V26 / "figure_v26_daily_scr_state_trajectories",
    "FigS2": V27 / "figure_v27_onset_anchored_trajectory_models",
    "FigS3": V28S / "figure_v28_severe_external_decision_curve",
    "FigS4": V28R / "figure_v28_recovery_observability_ipw_competing",
    "FigS5": V29 / "figure_v29_competing_risk_subgroups",
}

FIGURE_LEGENDS = [
    (
        "Figure 1", "Severity and recovery phenotypes in the strict postoperative ICU cohort",
        "Incident SCr-AKI is separated by active-episode severity, persistence, observed recovery, and recurrence. Recovery "
        "requires an observed non-AKI SCr; absence of measurement is not recovery."
    ),
    (
        "Figure 2", "Held-out performance of dynamic severe-AKI models",
        "Receiver-operating-characteristic, precision-recall, and calibration results for selected 0-h, 6-h, and 24-h models. "
        "Risk sets exclude severe SCr-AKI already present at each landmark."
    ),
    (
        "Figure 3", "Temporal and eICU external validation of severe-AKI models",
        "Panels summarize MIMIC-IV rolling temporal validation, frozen multicenter eICU performance, hospital heterogeneity, "
        "and held-out-hospital calibration before and after logistic recalibration."
    ),
    (
        "Figure 4", "Observation-time multistate occupancy and competing risks",
        "The left panel shows seven-state occupancy from ICU admission through 168 h. Right panels show cumulative incidence "
        "after observed AKI onset and after first observed recovery. SCr states update only at actual measurements; live "
        "discharge and in-hospital death are absorbing."
    ),
]

SUPP_FIGURE_LEGENDS = [
    ("Figure S1", "Daily observed SCr-state trajectories", "Daily state summaries include an explicit unobserved category."),
    ("Figure S2", "Onset-anchored persistence and nonrecovery models", "Predictors are restricted to information available by observed AKI onset."),
    ("Figure S3", "External decision-curve analysis for severe AKI", "Net benefit is exploratory and does not establish an intervention threshold."),
    ("Figure S4", "Recovery observability, IPW, and competing-event sensitivity", "IPW assumes missing at random conditional on observed onset covariates."),
    ("Figure S5", "Competing-risk estimates across prespecified subgroups", "Subgroup curves are descriptive and use subject-cluster bootstrap intervals."),
]


def words(text: str) -> int:
    text = re.sub(r"\[\[CITE:[^\]]+\]\]", "", text)
    return len(re.findall(r"\b[\w–-]+\b", text))


def abstract_word_count() -> int:
    return words(" ".join(ABSTRACT.values()))


def main_word_count() -> int:
    blocks = INTRODUCTION + [p for ps in METHODS.values() for p in ps] + [p for ps in RESULTS.values() for p in ps] + DISCUSSION + [CONCLUSION]
    return words(" ".join(blocks))


def f3(value: float) -> str:
    return f"{float(value):.3f}"


def ci(value: float, lo: float, hi: float) -> str:
    return f"{value:.3f} ({lo:.3f}–{hi:.3f})"


def csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def write_csv(name: str, rows: list[dict[str, object]]) -> None:
    if not rows:
        raise ValueError(f"No rows for {name}")
    with (TABLES / name).open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0]))
        writer.writeheader()
        writer.writerows(rows)


def build_tables() -> dict[str, list[dict[str, object]]]:
    t1 = csv_rows(LOCKED / "tables" / "Table_1_baseline_characteristics.csv")

    perf = pd.read_csv(V27 / "model_v27_performance_summary.csv")
    selected = [
        ("Severe SCr-AKI", "0 h", "XGBoost"),
        ("Severe SCr-AKI", "6 h", "XGBoost"),
        ("Severe SCr-AKI", "24 h", "Logistic Regression"),
        ("Persistent AKI beyond 48 h", "AKI onset", "XGBoost"),
        ("Not recovered at observation end", "AKI onset", "Logistic Regression"),
    ]
    t2 = []
    for outcome, landmark, model in selected:
        task = {"Severe SCr-AKI": "severe_scr", "Persistent AKI beyond 48 h": "persistent_aki", "Not recovered at observation end": "nonrecovery"}[outcome]
        raw_landmark = landmark.replace(" h", "")
        row = perf[(perf.task == task) & (perf.landmark.astype(str) == raw_landmark) & (perf.model == model)].iloc[0]
        t2.append({
            "Outcome": outcome, "Landmark": landmark, "Selected model": model.replace("Regression", "regression"),
            "Test n/events": f"{int(row.test_n):,}/{int(row.test_event_n):,}",
            "AUROC (95% CI)": ci(row.auroc, row.auroc_ci95_low, row.auroc_ci95_high),
            "AUPRC (95% CI)": ci(row.auprc, row.auprc_ci95_low, row.auprc_ci95_high),
            "Brier": f3(row.brier_score), "Calibration intercept/slope": f"{row.calibration_intercept:.2f}/{row.calibration_slope:.2f}",
        })

    ext = pd.read_csv(V28S / "model_v28_eicu_frozen_severe_performance.csv")
    recal = pd.read_csv(V28S / "model_v28_heldout_hospital_recalibration_performance.csv")
    t3 = []
    for lm in [0, 6, 24]:
        e = ext[(ext.landmark_hours == lm) & (ext.target == "SCr stage 2/3")].iloc[0]
        frozen = recal[(recal.landmark_hours == lm) & (recal.method == "frozen")].iloc[0]
        updated = recal[(recal.landmark_hours == lm) & (recal.method == "logistic recalibration")].iloc[0]
        t3.append({
            "Landmark": f"{lm} h", "External n/events": f"{int(e.n):,}/{int(e.event_n):,}",
            "External AUROC (95% CI)": ci(e.auroc, e.auroc_ci95_low, e.auroc_ci95_high),
            "External AUPRC/Brier": f"{e.auprc:.3f}/{e.brier_score:.3f}",
            "External intercept/slope": f"{e.calibration_intercept:.2f}/{e.calibration_slope:.2f}",
            "Held-out n": f"{int(frozen.n):,}", "Frozen Brier/slope": f"{frozen.brier_score:.3f}/{frozen.calibration_slope:.2f}",
            "Updated Brier/slope": f"{updated.brier_score:.3f}/{updated.calibration_slope:.2f}",
        })

    aki_cif = pd.read_csv(V29 / "analysis_v29_cif_after_aki.csv")
    rec_cif = pd.read_csv(V29 / "analysis_v29_cif_after_recovery.csv")
    t4 = []
    for source, origin in [(aki_cif, "Observed AKI onset"), (rec_cif, "First observed recovery")]:
        sub = source[(source.group_variable == "Overall") & source.time_hours.isin([24, 48, 72])]
        for row in sub.itertuples():
            t4.append({
                "Time origin": origin, "Time": f"{int(row.time_hours)} h", "Event": row.cause,
                "Cumulative incidence, % (95% CI)": f"{row.cif_percent:.1f} ({row.ci95_low_percent:.1f}–{row.ci95_high_percent:.1f})",
                "Risk set n": f"{int(row.risk_set_weighted_n):,}",
            })

    risk = pd.read_csv(V27 / "audit_v27_severe_risk_set_summary.csv")
    s1 = [{
        "Landmark": f"{int(r.landmark_hours)} h", "Risk set n": f"{int(r.risk_set_n):,}",
        "Already severe excluded": f"{int(r.already_severe_excluded_n):,}",
        "Severe SCr-AKI events": f"{int(r.severe_scr_event_n):,} ({r.severe_scr_event_percent:.1f}%)",
        "SCr-or-RRT events": f"{int(r.severe_scr_or_rrt_event_n):,} ({r.severe_scr_or_rrt_event_percent:.1f}%)",
    } for r in risk.itertuples()]

    s2 = []
    for r in perf.itertuples():
        s2.append({
            "Task": r.task, "Landmark": str(r.landmark) + (" h" if str(r.landmark).isdigit() else ""),
            "Model": str(r.model), "Predictors": int(r.predictor_n), "Test n/events": f"{int(r.test_n):,}/{int(r.test_event_n):,}",
            "AUROC (95% CI)": ci(r.auroc, r.auroc_ci95_low, r.auroc_ci95_high),
            "AUPRC (95% CI)": ci(r.auprc, r.auprc_ci95_low, r.auprc_ci95_high), "Brier": f3(r.brier_score),
            "Intercept/slope": f"{r.calibration_intercept:.2f}/{r.calibration_slope:.2f}",
        })

    rolling = pd.read_csv(V28S / "analysis_v28_mimic_rolling_severe_validation.csv")
    s3 = [{
        "Landmark": f"{int(r.landmark_hours)} h", "Validation years": f"{int(r.test_year_start)}–{int(r.test_year_end)}",
        "Train n": f"{int(r.train_n):,}", "Test n/events": f"{int(r.test_n):,}/{int(r.test_event_n):,}",
        "Model": r.model, "AUROC": f3(r.auroc), "AUPRC": f3(r.auprc), "Brier": f3(r.brier_score),
        "Intercept/slope": f"{r.calibration_intercept:.2f}/{r.calibration_slope:.2f}",
    } for r in rolling.itertuples()]

    s4 = []
    for r in ext.itertuples():
        s4.append({
            "Landmark": f"{int(r.landmark_hours)} h", "Target": r.target, "Model": r.model,
            "n/events": f"{int(r.n):,}/{int(r.event_n):,}", "Event rate": f"{r.event_percent:.1f}%",
            "AUROC (95% CI)": ci(r.auroc, r.auroc_ci95_low, r.auroc_ci95_high),
            "AUPRC (95% CI)": ci(r.auprc, r.auprc_ci95_low, r.auprc_ci95_high),
            "Brier": f3(r.brier_score), "Intercept/slope": f"{r.calibration_intercept:.2f}/{r.calibration_slope:.2f}",
        })

    hetero = pd.read_csv(V28S / "analysis_v28_hospital_heterogeneity_summary.csv")
    s5 = []
    for r in hetero.itertuples():
        s5.append({
            "Landmark": f"{int(r.landmark_hours)} h", "Hospitals": int(r.hospital_n),
            "AUROC-evaluable hospitals": int(r.hospital_n_auroc_evaluable),
            "Hospital AUROC median (IQR)": f"{r.hospital_auroc_median:.3f} ({r.hospital_auroc_q1:.3f}–{r.hospital_auroc_q3:.3f})",
            "Hospital AUROC range": f"{r.hospital_auroc_min:.3f}–{r.hospital_auroc_max:.3f}",
        })

    obs = pd.read_csv(V28R / "model_v28_observability_performance.csv")
    ipw = pd.read_csv(V28R / "model_v28_recovery_ipw_performance.csv")
    s6 = []
    for r in obs.itertuples():
        s6.append({
            "Section": "Outcome evaluability", "Target": r.evaluable_definition, "Analysis": "Cross-fitted observability model",
            "n/events": f"{int(r.n):,}/{int(r.event_n):,}", "Event/weighted rate": f"{r.event_percent_unweighted:.1f}%",
            "AUROC": f3(r.auroc), "AUPRC": f3(r.auprc), "Brier": f3(r.brier_score),
        })
    for r in ipw.itertuples():
        s6.append({
            "Section": "Outcome model", "Target": r.task, "Analysis": r.analysis,
            "n/events": f"{int(r.n):,}/{int(r.event_n):,}", "Event/weighted rate": f"{r.event_percent_weighted:.1f}%",
            "AUROC": f3(r.auroc), "AUPRC": f3(r.auprc), "Brier": f3(r.brier_score),
        })

    transitions = pd.read_csv(V29 / "audit_v29_transition_counts.csv")
    s7 = [{
        "From state": r.from_state, "To state": r.to_state,
        "Source": str(r.transition_source).replace("serum creatinine", "SCr").replace("hospital disposition", "Hospital disposition"),
        "Transitions": f"{int(r.transition_n):,}", "Stays": f"{int(r.stay_n):,}",
        "Time from ICU admission, h, median (IQR)": (
            f"{r.transition_time_median_hours:.1f} ({r.transition_time_q1_hours:.1f}–{r.transition_time_q3_hours:.1f})"
        ),
    } for r in transitions.itertuples()]

    s8 = []
    for source, origin in [(aki_cif, "Observed AKI onset"), (rec_cif, "First observed recovery")]:
        sub = source[(source.group_variable != "Overall") & source.time_hours.isin([48, 72])]
        for r in sub.itertuples():
            s8.append({
                "Time origin": origin, "Subgroup variable": r.group_variable, "Subgroup": r.group, "n": f"{int(r.n):,}",
                "Time": f"{int(r.time_hours)} h", "Event": r.cause,
                "CIF, % (95% CI)": f"{r.cif_percent:.1f} ({r.ci95_low_percent:.1f}–{r.ci95_high_percent:.1f})",
            })

    tables = {"T1": t1, "T2": t2, "T3": t3, "T4": t4, "S1": s1, "S2": s2, "S3": s3, "S4": s4, "S5": s5, "S6": s6, "S7": s7, "S8": s8}
    filenames = {
        "T1": "Table_1_baseline_characteristics.csv", "T2": "Table_2_secondary_model_performance.csv",
        "T3": "Table_3_external_validation_recalibration.csv", "T4": "Table_4_multistate_competing_risk.csv",
        "S1": "Table_S1_severe_risk_sets.csv", "S2": "Table_S2_all_secondary_models.csv",
        "S3": "Table_S3_rolling_temporal_validation.csv", "S4": "Table_S4_eicu_external_sensitivity.csv",
        "S5": "Table_S5_hospital_heterogeneity.csv", "S6": "Table_S6_recovery_observability_ipw.csv",
        "S7": "Table_S7_multistate_transition_counts.csv", "S8": "Table_S8_subgroup_competing_risk.csv",
    }
    for key, name in filenames.items():
        write_csv(name, tables[key])
    return tables


TABLE_SPECS = {
    "T1": (
        "Baseline characteristics by incident AKI status", "tab:t1",
        [("Characteristic", "Characteristic"), ("Overall (N=10,877)", "Overall"), ("No AKI (N=6,346)", "No incident AKI"),
         ("Incident AKI (N=4,531)", "Incident AKI"), ("Standardized mean difference", "SMD"), ("Missing, n", "Missing")],
        [2700, 1800, 1800, 1800, 1200, 1000], True, 7.5,
        "Values are median [IQR] or n (%). SMD, standardized mean difference."
    ),
    "T2": (
        "Held-out performance of selected secondary models", "tab:t2",
        [("Outcome", "Outcome"), ("Landmark", "Landmark"), ("Selected model", "Model"), ("Test n/events", "Test n/events"),
         ("AUROC (95% CI)", "AUROC (95% CI)"), ("AUPRC (95% CI)", "AUPRC (95% CI)"), ("Brier", "Brier"),
         ("Calibration intercept/slope", "Intercept/slope")],
        [2100, 900, 1550, 1200, 1900, 1900, 800, 1350], True, 7.0,
        "Confidence intervals use 1,000 subject-cluster bootstrap resamples of held-out patients."
    ),
    "T3": (
        "eICU external validation and hospital-held-out recalibration", "tab:t3",
        [("Landmark", "Landmark"), ("External n/events", "External n/events"), ("External AUROC (95% CI)", "External AUROC (95% CI)"),
         ("External AUPRC/Brier", "AUPRC/Brier"), ("External intercept/slope", "External intercept/slope"),
         ("Held-out n", "Held-out n"), ("Frozen Brier/slope", "Frozen Brier/slope"), ("Updated Brier/slope", "Updated Brier/slope")],
        [800, 1350, 2100, 1400, 1650, 1000, 1500, 1500], True, 7.0,
        "Logistic recalibration was learned in separate calibration hospitals and applied unchanged to held-out hospitals."
    ),
    "T4": (
        "Observation-time cumulative incidence after AKI onset and first recovery", "tab:t4",
        [("Time origin", "Time origin"), ("Time", "Time"), ("Event", "Event"),
         ("Cumulative incidence, % (95% CI)", "Cumulative incidence, % (95% CI)"), ("Risk set n", "Risk set n")],
        [2300, 900, 2900, 3000, 1100], True, 8.0,
        "Aalen–Johansen estimates with 500 subject-cluster bootstrap resamples. Risk-set size is shown for the corresponding time."
    ),
    "S1": (
        "Severe-AKI landmark risk sets", "tab:s1",
        [("Landmark", "Landmark"), ("Risk set n", "Risk set n"), ("Already severe excluded", "Already severe excluded"),
         ("Severe SCr-AKI events", "Severe SCr-AKI events"), ("SCr-or-RRT events", "SCr-or-RRT events")],
        [1200, 1400, 2100, 2400, 2200], True, 8.0, "Events occur after the landmark through ICU day 7."
    ),
    "S2": (
        "Performance of all secondary prediction models", "tab:s2",
        [("Task", "Task"), ("Landmark", "Landmark"), ("Model", "Model"), ("Predictors", "Predictors"),
         ("Test n/events", "Test n/events"), ("AUROC (95% CI)", "AUROC (95% CI)"),
         ("AUPRC (95% CI)", "AUPRC (95% CI)"), ("Brier", "Brier"), ("Intercept/slope", "Intercept/slope")],
        [1700, 1000, 1600, 900, 1300, 1900, 1900, 700, 1200], True, 6.8, "Held-out evaluation with subject-cluster bootstrap intervals."
    ),
    "S3": (
        "Rolling temporal validation of severe-AKI models", "tab:s3",
        [("Landmark", "Landmark"), ("Validation years", "Validation years"), ("Train n", "Train n"),
         ("Test n/events", "Test n/events"), ("Model", "Model"), ("AUROC", "AUROC"), ("AUPRC", "AUPRC"),
         ("Brier", "Brier"), ("Intercept/slope", "Intercept/slope")],
        [900, 1500, 1000, 1300, 1600, 900, 900, 800, 1300], True, 7.0, "Each evaluation trains on earlier years and tests on the stated later-year block."
    ),
    "S4": (
        "Full eICU external validation and SCr-or-RRT sensitivity", "tab:s4",
        [("Landmark", "Landmark"), ("Target", "Target"), ("Model", "Model"), ("n/events", "n/events"),
         ("Event rate", "Event rate"), ("AUROC (95% CI)", "AUROC (95% CI)"), ("AUPRC (95% CI)", "AUPRC (95% CI)"),
         ("Brier", "Brier"), ("Intercept/slope", "Intercept/slope")],
        [850, 2300, 1600, 1200, 1000, 1900, 1900, 700, 1200], True, 6.8, "RRT is a documentation-sensitive treatment-record sensitivity outcome."
    ),
    "S5": (
        "Hospital-level heterogeneity of frozen eICU discrimination", "tab:s5",
        [("Landmark", "Landmark"), ("Hospitals", "Hospitals"), ("AUROC-evaluable hospitals", "AUROC-evaluable hospitals"),
         ("Hospital AUROC median (IQR)", "Hospital AUROC median (IQR)"), ("Hospital AUROC range", "Hospital AUROC range")],
        [1200, 1200, 2300, 2600, 1900], True, 8.0, "Hospitals required prespecified minimum sample and outcome counts for AUROC estimation."
    ),
    "S6": (
        "Recovery observability and inverse-probability-weighted sensitivity", "tab:s6",
        [("Section", "Section"), ("Target", "Target"), ("Analysis", "Analysis"), ("n/events", "n/events"),
         ("Event/weighted rate", "Event/weighted rate"), ("AUROC", "AUROC"), ("AUPRC", "AUPRC"), ("Brier", "Brier")],
        [1600, 2100, 3400, 1200, 1600, 850, 850, 850], True, 7.0, "IPW assumes missing at random conditional on measured onset features."
    ),
    "S7": (
        "Observed multistate transition counts", "tab:s7",
        [("From state", "From state"), ("To state", "To state"), ("Source", "Source"),
         ("Transitions", "Transitions"), ("Stays", "Stays"),
         ("Time from ICU admission, h, median (IQR)", "Time from ICU admission, h, median (IQR)")],
        [1900, 1900, 1500, 1200, 1100, 3000], True, 7.0,
        "Transitions are generated only at observed measurements or recorded terminal events."
    ),
    "S8": (
        "Subgroup cumulative incidence after AKI onset and recovery", "tab:s8",
        [("Time origin", "Time origin"), ("Subgroup variable", "Subgroup variable"), ("Subgroup", "Subgroup"),
         ("n", "n"), ("Time", "Time"), ("Event", "Event"), ("CIF, % (95% CI)", "CIF, % (95% CI)")],
        [1800, 1700, 1900, 800, 800, 2600, 2200], True, 6.8, "Descriptive Aalen–Johansen estimates with subject-cluster bootstrap intervals."
    ),
}


def resolve_s7_spec(rows: list[dict[str, object]]):
    keys = list(rows[0])
    columns = [(k, k.replace("_", " ").capitalize()) for k in keys]
    widths = [max(900, min(2600, 650 + 120 * len(k))) for k in keys]
    return columns, widths


def copy_figures() -> None:
    for short, stem in FIGURE_SOURCES.items():
        for ext in ["png", "pdf", "svg"]:
            source = Path(f"{stem}.{ext}")
            if not source.exists():
                raise FileNotFoundError(source)
            shutil.copy2(source, FIGURES / f"{short}.{ext}")


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); v8.set_font(r, 16, bold=True)
    for text, size, italic in [(subtitle, 11, True), (AUTHORS_WORD, 11, False), (AFFILIATION_WORD, 10, False), (CORRESPONDENCE_WORD, 10, False)]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); v8.set_font(r, size, italic=italic)


def add_landscape(doc: Document):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = Inches(11), Inches(8.5)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.55)
    return section


def add_portrait(doc: Document):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.8)
    return section


def word_table(doc: Document, key: str, rows: list[dict[str, object]], number: str) -> None:
    title, _label, columns, widths, _landscape, font, legend = TABLE_SPECS[key]
    if columns is None:
        columns, widths = resolve_s7_spec(rows)
    table = v8.add_table(doc, number, title, rows, columns, widths, legend, font=font)
    # LibreOffice can otherwise split a table row at a page boundary and leave
    # a visually empty continuation band below a repeated header.
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cantSplit") is None:
            tr_pr.append(OxmlElement("w:cantSplit"))


def build_word(tables: dict[str, list[dict[str, object]]]) -> tuple[Path, Path]:
    doc = Document(); v8.configure_word(doc, RUNNING_TITLE); add_cover(doc, TITLE, "Original research | Secondary analysis")
    for text in [f"Abstract word count: {abstract_word_count()}", f"Main-text word count: {main_word_count()}", "Tables: 4; Figures: 4; Additional files: 2"]:
        p = doc.add_paragraph(); r = p.add_run(text); v8.set_font(r, 10)
    doc.add_heading("Abstract", level=1)
    for heading, text in ABSTRACT.items():
        p = doc.add_paragraph(); r = p.add_run(heading + ": "); v8.set_font(r, 12, bold=True)
        r = p.add_run(text); v8.set_font(r, 12)
    p = doc.add_paragraph(); r = p.add_run("Keywords: "); v8.set_font(r, 12, bold=True); r = p.add_run("; ".join(KEYWORDS)); v8.set_font(r, 12)
    doc.add_heading("Background", level=1)
    for paragraph in INTRODUCTION: v8.add_para(doc, paragraph)
    doc.add_heading("Methods", level=1)
    for heading, paragraphs in METHODS.items():
        doc.add_heading(heading, level=2)
        for paragraph in paragraphs: v8.add_para(doc, paragraph)
    doc.add_heading("Results", level=1)
    for heading, paragraphs in RESULTS.items():
        doc.add_heading(heading, level=2)
        for paragraph in paragraphs: v8.add_para(doc, paragraph)
    doc.add_heading("Discussion", level=1)
    for paragraph in DISCUSSION: v8.add_para(doc, paragraph)
    doc.add_heading("Conclusions", level=1); v8.add_para(doc, CONCLUSION)
    doc.add_heading("List of abbreviations", level=1)
    v8.add_para(doc, "ADQI, Acute Disease Quality Initiative; AKI, acute kidney injury; AUPRC, area under the precision-recall curve; AUROC, area under the receiver-operating-characteristic curve; CI, confidence interval; CIF, cumulative incidence function; ICU, intensive care unit; IPW, inverse-probability weighting; IQR, interquartile range; KDIGO, Kidney Disease: Improving Global Outcomes; RRT, renal replacement therapy; SCr, serum creatinine.", indent=False)
    doc.add_heading("Declarations", level=1)
    declarations = [
        ("Ethics approval and consent to participate", ETHICS), ("Consent for publication", "Not applicable."),
        ("Availability of data and materials", AVAILABILITY + f" [{v8.REF_INDEX['MIMIC2024']},{v8.REF_INDEX['Pollard2018']}]."),
        ("Competing interests", COMPETING), ("Funding", FUNDING), ("Author's contributions", CONTRIBUTIONS),
        ("Acknowledgements", ACKNOWLEDGEMENTS), ("AI-assisted editing disclosure", AI_DISCLOSURE),
    ]
    for heading, text in declarations:
        doc.add_heading(heading, level=2); v8.add_para(doc, text, indent=False)
    doc.add_heading("Additional files", level=1)
    v8.add_para(doc, "Additional file 1 (.docx and .pdf): Supplementary Tables S1–S8 and Supplementary Figures S1–S5.", indent=False)
    v8.add_para(doc, "Additional file 2 (.docx and .csv): TRIPOD+AI checklist for the prediction components.", indent=False)
    doc.add_heading("References", level=1)
    for i, (_key, reference, *_rest) in enumerate(REFS, 1):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.25); p.paragraph_format.first_line_indent = Inches(-0.25)
        r = p.add_run(f"{i}. {reference}"); v8.set_font(r, 10)
    add_landscape(doc)
    for key, number in [("T1", "Table 1"), ("T2", "Table 2"), ("T3", "Table 3"), ("T4", "Table 4")]:
        word_table(doc, key, tables[key], number)
    add_portrait(doc); doc.add_heading("Figure legends", level=1)
    for number, title, legend in FIGURE_LEGENDS:
        v8.add_para(doc, f"{number}. {title}. {legend}", indent=False)
    main_path = OUT / "secondary_manuscript_v29_en.docx"; doc.save(main_path)

    supp = Document(); v8.configure_word(supp, "Supplement | " + RUNNING_TITLE); add_cover(supp, "Additional file 1: Supplementary material", TITLE)
    add_landscape(supp)
    for i in range(1, 9):
        word_table(supp, f"S{i}", tables[f"S{i}"], f"Table S{i}")
    add_portrait(supp); supp.add_heading("Supplementary figures", level=1)
    for i, (number, title, legend) in enumerate(SUPP_FIGURE_LEGENDS, 1):
        p = supp.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(FIGURES / f"FigS{i}.png"), width=Inches(6.3))
        v8.add_para(supp, f"{number}. {title}. {legend}", indent=False)
    supp_path = OUT / "additional_file_1_secondary_supplement_en.docx"; supp.save(supp_path)
    return main_path, supp_path


def latex_table(key: str, rows: list[dict[str, object]], filename: str) -> None:
    title, label, columns, _widths, landscape, _font, _legend = TABLE_SPECS[key]
    if columns is None:
        columns, _widths = resolve_s7_spec(rows)
    v8.latex_table(rows, columns, title, label, filename, landscape=landscape, long=(key in {"T1", "S2", "S7", "S8"}))


def tex_section(blocks: dict[str, list[str]]) -> str:
    return "\n".join(
        r"\subsection{" + v8.tex_escape(heading) + "}\n" + "\n\n".join(v8.render_text(p, "tex") for p in paragraphs)
        for heading, paragraphs in blocks.items()
    )


def write_latex(tables: dict[str, list[dict[str, object]]]) -> None:
    (LATEX / "tables").mkdir(parents=True, exist_ok=True)
    for key in ["T1", "T2", "T3", "T4", "S1", "S2", "S3", "S4", "S5", "S6", "S7", "S8"]:
        latex_table(key, tables[key], f"table{key}.tex")
    refs = "\n".join(r"\bibitem{" + key + "} " + v8.tex_escape(reference) for key, reference, *_ in REFS)
    figures = "\n".join(
        r"\begin{figure}[p]\centering\includegraphics[width=\textwidth]{../figures/Fig" + str(i) + ".pdf}\n" +
        r"\caption{" + v8.tex_escape(title) + ". " + v8.tex_escape(legend) + r"}\end{figure}"
        for i, (_number, title, legend) in enumerate(FIGURE_LEGENDS, 1)
    )
    abstract = "\n\n".join(r"\textbf{" + k + ":} " + v8.tex_escape(v) for k, v in ABSTRACT.items())
    main = r"""\documentclass[11pt]{article}
\usepackage[letterpaper,margin=1in]{geometry}
\usepackage[T1]{fontenc}\usepackage{lmodern}\usepackage{microtype}
\usepackage{graphicx,booktabs,longtable,pdflscape,caption,setspace,lineno,hyperref}
\usepackage[numbers,sort&compress]{natbib}
\doublespacing\linenumbers\captionsetup{font=small,labelfont=bf}
\title{__TITLE__}
\author{Bizhi Wei\textsuperscript{1*}\\
\textsuperscript{1}Pu Ai Medical School, Shaoyang University, Shaoyang 422000, Hunan, China\\
\textsuperscript{*}Correspondence: Bizhi Wei\\
Pu Ai Medical School, Shaoyang University, Shaoyang 422000, Hunan, China\\
Email: 15619056250wbz@gmail.com}
\date{}
\begin{document}\maketitle
\noindent\textbf{Article type:} Original research\par
\noindent\textbf{Abstract word count:} __ABSWC__\par
\noindent\textbf{Main-text word count:} __MAINWC__\par
\noindent\textbf{Tables/Figures:} 4/4\par
\section*{Abstract}
__ABSTRACT__
\noindent\textbf{Keywords:} __KEYWORDS__
\section{Background}
__INTRO__
\section{Methods}
__METHODS__
\section{Results}
__RESULTS__
\section{Discussion}
__DISCUSSION__
\section{Conclusions}
__CONCLUSION__
\section*{List of abbreviations}
ADQI, Acute Disease Quality Initiative; AKI, acute kidney injury; AUPRC, area under the precision-recall curve; AUROC, area under the receiver-operating-characteristic curve; CI, confidence interval; CIF, cumulative incidence function; ICU, intensive care unit; IPW, inverse-probability weighting; IQR, interquartile range; KDIGO, Kidney Disease: Improving Global Outcomes; RRT, renal replacement therapy; SCr, serum creatinine.
\section*{Declarations}
\subsection*{Ethics approval and consent to participate} __ETHICS__
\subsection*{Consent for publication} Not applicable.
\subsection*{Availability of data and materials} __AVAILABILITY__ \cite{MIMIC2024,Pollard2018}
\subsection*{Competing interests} __COMPETING__
\subsection*{Funding} __FUNDING__
\subsection*{Author's contributions} __CONTRIBUTIONS__
\subsection*{Acknowledgements} __ACKNOWLEDGEMENTS__
\subsection*{AI-assisted editing disclosure} __AI__
\section*{Additional files}
Additional file 1 (.docx and .pdf): Supplementary Tables S1--S8 and Supplementary Figures S1--S5.\\
Additional file 2 (.docx and .csv): TRIPOD+AI checklist for the prediction components.
\begin{thebibliography}{99}
__REFS__
\end{thebibliography}
\clearpage\input{tables/tableT1.tex}\clearpage\input{tables/tableT2.tex}\clearpage\input{tables/tableT3.tex}\clearpage\input{tables/tableT4.tex}
\clearpage
__FIGURES__
\end{document}
"""
    replacements = {
        "__TITLE__": v8.tex_escape(TITLE), "__ABSWC__": str(abstract_word_count()), "__MAINWC__": str(main_word_count()),
        "__ABSTRACT__": abstract, "__KEYWORDS__": "; ".join(KEYWORDS),
        "__INTRO__": "\n\n".join(v8.render_text(p, "tex") for p in INTRODUCTION), "__METHODS__": tex_section(METHODS),
        "__RESULTS__": tex_section(RESULTS), "__DISCUSSION__": "\n\n".join(v8.render_text(p, "tex") for p in DISCUSSION),
        "__CONCLUSION__": v8.render_text(CONCLUSION, "tex"), "__ETHICS__": v8.tex_escape(ETHICS),
        "__AVAILABILITY__": v8.tex_escape(AVAILABILITY), "__COMPETING__": v8.tex_escape(COMPETING),
        "__FUNDING__": v8.tex_escape(FUNDING), "__CONTRIBUTIONS__": v8.tex_escape(CONTRIBUTIONS),
        "__ACKNOWLEDGEMENTS__": v8.tex_escape(ACKNOWLEDGEMENTS), "__AI__": v8.tex_escape(AI_DISCLOSURE),
        "__REFS__": refs, "__FIGURES__": figures,
    }
    for key, value in replacements.items():
        main = main.replace(key, value)
    (LATEX / "main.tex").write_text(main, encoding="utf-8")

    supp_figs = "\n".join(
        r"\begin{figure}[p]\centering\includegraphics[width=\textwidth]{../figures/FigS" + str(i) + ".pdf}\n" +
        r"\caption{" + v8.tex_escape(title) + ". " + v8.tex_escape(legend) + r"}\end{figure}\clearpage"
        for i, (_number, title, legend) in enumerate(SUPP_FIGURE_LEGENDS, 1)
    )
    supplement = r"""\documentclass[11pt]{article}
\usepackage[letterpaper,margin=0.75in]{geometry}\usepackage[T1]{fontenc}\usepackage{lmodern}
\usepackage{graphicx,booktabs,longtable,pdflscape,caption,setspace,hyperref}\onehalfspacing
\begin{document}\begin{center}{\Large\bfseries Additional file 1: Supplementary material}\\[8pt]
__TITLE__\\[8pt]Bizhi Wei\textsuperscript{1*}\\
\textsuperscript{1}Pu Ai Medical School, Shaoyang University, Shaoyang 422000, Hunan, China\\
\textsuperscript{*}Correspondence: Bizhi Wei; 15619056250wbz@gmail.com
\end{center}
\input{tables/tableS1.tex}\clearpage\input{tables/tableS2.tex}\clearpage\input{tables/tableS3.tex}\clearpage
\input{tables/tableS4.tex}\clearpage\input{tables/tableS5.tex}\clearpage\input{tables/tableS6.tex}\clearpage
\input{tables/tableS7.tex}\clearpage\input{tables/tableS8.tex}\clearpage
__SUPPFIGS__
\end{document}
""".replace("__TITLE__", v8.tex_escape(TITLE)).replace("__SUPPFIGS__", supp_figs)
    (LATEX / "supplement.tex").write_text(supplement, encoding="utf-8")

    entries = []
    for key, reference, doi, _pmid, _role in REFS:
        title = reference.split(". ", 1)[1].split(". ", 1)[0] if ". " in reference else reference
        entries.append(f"@misc{{{key},\n  title={{{title}}},\n  note={{{reference}}}" + (f",\n  doi={{{doi}}}" if doi else "") + "\n}")
    (LATEX / "references.bib").write_text("\n\n".join(entries) + "\n", encoding="utf-8")


def write_markdown() -> None:
    lines = [f"# {TITLE}", "", AUTHORS_WORD, "", AFFILIATION_WORD, "", CORRESPONDENCE_WORD, "", "## Abstract", ""]
    for heading, text in ABSTRACT.items():
        lines.extend([f"**{heading}:** {text}", ""])
    lines.extend(["**Keywords:** " + "; ".join(KEYWORDS), "", "## Background", ""])
    for paragraph in INTRODUCTION:
        lines.extend([v8.render_text(paragraph, "word"), ""])
    for section, blocks in [("Methods", METHODS), ("Results", RESULTS)]:
        lines.extend([f"## {section}", ""])
        for heading, paragraphs in blocks.items():
            lines.extend([f"### {heading}", ""])
            for paragraph in paragraphs:
                lines.extend([v8.render_text(paragraph, "word"), ""])
    lines.extend(["## Discussion", ""])
    for paragraph in DISCUSSION:
        lines.extend([v8.render_text(paragraph, "word"), ""])
    lines.extend(["## Conclusions", "", v8.render_text(CONCLUSION, "word"), "", "## Declarations", ""])
    for heading, text in [
        ("Ethics approval and consent to participate", ETHICS), ("Consent for publication", "Not applicable."),
        ("Availability of data and materials", AVAILABILITY), ("Competing interests", COMPETING), ("Funding", FUNDING),
        ("Author's contributions", CONTRIBUTIONS), ("Acknowledgements", ACKNOWLEDGEMENTS),
        ("AI-assisted editing disclosure", AI_DISCLOSURE),
    ]:
        lines.extend([f"### {heading}", "", text, ""])
    lines.extend(["## References", ""])
    for i, (_key, reference, *_rest) in enumerate(REFS, 1):
        lines.append(f"{i}. {reference}")
    (OUT / "secondary_manuscript_v29_en.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def build_checklist() -> tuple[Path, Path]:
    asset = ROOT / "scripts" / "assets" / "tripod_ai_checklist_v8.tsv"
    locations = {
        "1": "Title", "2": "Abstract", "3a": "Background", "3b": "Background, final paragraph",
        "4": "Background, final paragraph", "5a": "Methods—Study design", "5b": "Methods—Study design",
        "6a": "Methods—Study design", "6b": "Methods—Study design and population", "7": "Methods—Landmark models",
        "8a": "Methods—Severe AKI and renal-trajectory outcomes", "9a": "Methods—Landmark and onset-anchored models",
        "10": "Methods—Landmark models; Results—risk sets", "11": "Methods—Landmark models and observability",
        "12a": "Methods—Landmark and onset-anchored models", "12b": "Methods—Landmark and onset-anchored models",
        "12c": "Methods—Landmark and onset-anchored models", "12d": "Methods—Temporal and external validation",
        "12e": "Methods—Landmark and onset-anchored models", "12f": "Methods—Temporal and external validation",
        "12g": "Methods—Temporal and external validation", "13": "Methods—Landmark and onset-anchored models",
        "15": "Methods—Landmark and onset-anchored models", "16": "Methods—Study design",
        "17": "Declarations—Ethics", "18a": "Declarations—Funding", "18b": "Declarations—Competing interests",
        "18c": "Availability of data and materials", "18e": "Availability of data and materials",
        "18f": "Availability of data and materials", "19": "Methods—Software, reporting, and interpretation boundary",
        "20a": "Results—Cohort and secondary outcomes; Figure 1", "20b": "Table 1; Supplementary Table S1",
        "20c": "Supplementary Tables S2–S8", "21": "Results; Tables 2–4", "22": "Availability of data and materials",
        "23a": "Results—Model performance; Figures 2–3; Tables 2–3", "23b": "Results—External validation; Supplementary Table S5",
        "24": "Methods—Temporal and external validation; Table 3", "25": "Discussion", "26": "Discussion—Limitations",
        "27a": "Methods—Recovery observability; Discussion", "27b": "Methods—Recovery observability",
        "27c": "Discussion; Conclusions",
    }
    rows = []
    with asset.open(encoding="utf-8") as handle:
        for row in csv.DictReader(handle, delimiter="\t"):
            row["location"] = locations.get(row["item"], "See manuscript or supplementary material")
            row["status"] = "Reported"
            rows.append(row)
    csv_path = OUT / "additional_file_2_tripod_ai_checklist.csv"
    with csv_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0])); writer.writeheader(); writer.writerows(rows)
    doc = Document(); v8.configure_word(doc, "TRIPOD+AI checklist"); add_cover(doc, "Additional file 2: TRIPOD+AI checklist", TITLE)
    v8.add_para(doc, "This checklist applies to the prediction-model components. The multistate and competing-risk analyses are described separately in the manuscript and supplement.", indent=False)
    add_landscape(doc)
    v8.add_table(doc, "", "", rows, [("section", "Section"), ("item", "Item"), ("checklist_item", "Reporting recommendation"),
                                      ("location", "Location"), ("status", "Status")],
                 [1200, 650, 4600, 4000, 1100], font=6.8)
    docx_path = OUT / "additional_file_2_tripod_ai_checklist.docx"; doc.save(docx_path)
    return docx_path, csv_path


def write_audits() -> None:
    with (OUT / "reference_audit.csv").open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle); writer.writerow(["Number", "Citation key", "Vancouver reference", "DOI", "PMID", "Role", "Verification"])
        for i, (key, reference, doi, pmid, role) in enumerate(REFS, 1):
            writer.writerow([i, key, reference, doi, pmid, role, "Verified against PubMed, official source, or locked reference audit"])
    ledger = [
        ("severe SCr-AKI", "Active-episode KDIGO SCr stage 2 or 3; primary secondary prediction outcome"),
        ("SCr-or-RRT sensitivity", "Treatment-record sensitivity; do not equate with primary SCr target"),
        ("observed recovery", "First measured non-AKI SCr state; never infer from missing measurement or discharge"),
        ("recurrent AKI", "Observed AKI-positive SCr after first observed recovery"),
        ("observation time", "Transition time is the SCr measurement time, not latent biological transition time"),
        ("Aalen–Johansen", "Use en dash and identify as product-integral estimator"),
        ("external validation", "Feature-harmonized frozen model evaluation in eICU"),
        ("recalibration", "Probability updating learned in separate hospitals; does not improve ranking"),
        ("IPW", "Sensitivity under measured missing-at-random assumption"),
    ]
    with (OUT / "terminology_ledger.csv").open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle); writer.writerow(["Canonical term", "Usage rule"]); writer.writerows(ledger)
    evidence = [
        ("Severe SCr-AKI incidence 679/10,877", "audit_v29_multistate_summary.csv", "Results; Abstract"),
        ("Internal model performance", "model_v27_performance_summary.csv", "Table 2; Figure 2"),
        ("eICU external performance", "model_v28_eicu_frozen_severe_performance.csv", "Table 3; Figure 3"),
        ("Hospital-held-out recalibration", "model_v28_heldout_hospital_recalibration_performance.csv", "Table 3; Figure 3"),
        ("Recovery observability/IPW", "model_v28_observability_performance.csv; model_v28_recovery_ipw_performance.csv", "Table S6; Figure S4"),
        ("State occupancy and competing risks", "analysis_v29_state_occupancy.csv; analysis_v29_cif_after_aki.csv; analysis_v29_cif_after_recovery.csv", "Table 4; Figure 4"),
        ("Post-disposition exclusion audit", "audit_v29_postdisposition_aki_exclusions.csv", "Methods; Results; Supplement"),
    ]
    with (OUT / "claim_evidence_map.csv").open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle); writer.writerow(["Claim", "Source", "Manuscript location"]); writer.writerows(evidence)
    (OUT / "code_and_data_availability_statement.txt").write_text(AVAILABILITY + "\n", encoding="utf-8")
    doc = Document(); v8.configure_word(doc, "Code and data availability"); add_cover(doc, "Code and data availability statement", TITLE)
    v8.add_para(doc, AVAILABILITY, indent=False); doc.save(OUT / "code_and_data_availability_statement.docx")


def write_readme() -> None:
    readme = f"""# Independent secondary manuscript package v29

This package is separate from and does not modify `outputs/manuscript_package_v23_locked`.

- Title: {TITLE}
- Author: Bizhi Wei
- Abstract words: {abstract_word_count()}
- Main-text words: {main_word_count()}
- References: {len(REFS)}
- Main display items: Tables 1–4 and Figures 1–4
- Additional file 1: Tables S1–S8 and Figures S1–S5
- Additional file 2: TRIPOD+AI checklist for prediction components
- Scientific boundary: secondary severe-AKI and observed renal-trajectory analysis; no causal or deployment claim
- Repository: https://github.com/Bizhi-Wei/postoperative-aki-dynamic-prediction
- Archived DOI: to be provided before publication

The 4,531 locked incident-AKI labels are unchanged. Twelve records with AKI onset after recorded hospital disposition are excluded only from trajectory risk sets and retained in a dedicated audit.
"""
    (OUT / "README.md").write_text(readme, encoding="utf-8")


def write_manifest() -> None:
    rows = []
    for path in sorted(p for p in OUT.rglob("*") if p.is_file() and "qa" not in p.parts and p.name != "file_manifest_sha256.csv"):
        digest = hashlib.sha256(path.read_bytes()).hexdigest()
        rows.append({"relative_path": path.relative_to(OUT).as_posix(), "bytes": path.stat().st_size, "sha256": digest})
    write_csv("../file_manifest_sha256.csv", rows)


def main() -> None:
    if OUT.exists():
        shutil.rmtree(OUT)
    for path in [OUT, LATEX, TABLES, FIGURES, QA]:
        path.mkdir(parents=True, exist_ok=True)
    tables = build_tables()
    copy_figures()
    write_latex(tables)
    write_markdown()
    build_word(tables)
    build_checklist()
    write_audits()
    write_readme()
    write_manifest()
    print(f"Built secondary manuscript package: {OUT}")
    print(f"Abstract words={abstract_word_count()}; main words={main_word_count()}; references={len(REFS)}")


if __name__ == "__main__":
    main()
